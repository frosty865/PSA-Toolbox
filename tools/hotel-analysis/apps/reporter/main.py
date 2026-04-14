#!/usr/bin/env python3
"""
Reporter CLI: read Assessment JSON from stdin, produce DOCX with charts and summary table.
Template-first: load anchor template (ADA/report template.docx only), insert at anchors only.
Uses WORK_DIR env for working directory; prints only output path to stdout.
"""
from __future__ import annotations

import json
import os
import re
import sys
from pathlib import Path

# Startup dependency check: fail early with clear message (exit 2)
try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph
except Exception as e:
    print("ERROR: python-docx is not available", file=sys.stderr)
    print(str(e), file=sys.stderr)
    sys.exit(2)
try:
    import matplotlib
    import matplotlib.patheffects as pe
    import matplotlib.pyplot as plt
    matplotlib.use("Agg")
except ImportError:
    print("Install matplotlib: pip install matplotlib", file=sys.stderr)
    sys.exit(1)

from report_humanize import (
    assert_no_backend_leak,
    debot_vulnerability_narrative,
    dedupe_sentences,
    ensure_synthesis_formatting,
    expand_acronym,
    expand_acronym_in_text,
    normalize_spacing,
    sanitize_backend_evidence,
    strip_it_transport_mitigation_when_unconfirmed,
)
from sanitize import sanitize_text, sanitize_vulnerability_text
from docx_ops import replace_anchor_in_doc

# Category codes and display names for chart titles (charts only for these five)
CHART_CATEGORIES = (
    "ELECTRIC_POWER",
    "COMMUNICATIONS",
    "INFORMATION_TECHNOLOGY",
    "WATER",
    "WASTEWATER",
)
CATEGORY_DISPLAY = {
    "ELECTRIC_POWER": "Electric Power",
    "COMMUNICATIONS": "Communications",
    "INFORMATION_TECHNOLOGY": "Information Technology",
    "WATER": "Water",
    "WASTEWATER": "Wastewater",
    "CRITICAL_PRODUCTS": "Critical Products",  # no chart; contributes to VOFC only
}

# Canonical sector order for Part II (Technical Annex) — deterministic page breaks
SECTOR_ORDER = [
    "Electric Power",
    "Communications",
    "Information Technology",
    "Water",
    "Wastewater",
]
SECTOR_DISPLAY_TO_CODE = {CATEGORY_DISPLAY[c]: c for c in CHART_CATEGORIES}
# Template contract: narrative-only anchors (single template ADA/report template.docx)
CHART_ANCHORS = [f"[[CHART_{c}]]" for c in CHART_CATEGORIES]
TABLE_ANCHOR = "[[TABLE_SUMMARY]]"
DEP_SUMMARY_TABLE_ANCHOR = "[[DEP_SUMMARY_TABLE]]"
TABLE_DEPENDENCY_SUMMARY_ANCHOR = "[[TABLE_DEPENDENCY_SUMMARY]]"
IT_TRANSPORT_SECTION_ANCHOR = "[[IT_TRANSPORT_SECTION]]"
IT_HOSTED_SECTION_ANCHOR = "[[IT_HOSTED_SECTION]]"
VULN_NARRATIVE_ANCHOR = "[[VULN_NARRATIVE]]"
STRUCTURAL_PROFILE_SUMMARY_ANCHOR = "[[STRUCTURAL_PROFILE_SUMMARY]]"
DESIGNATION_SERVICES_ANCHOR = "[[DESIGNATION_SERVICES]]"
VULNERABILITY_COUNT_SUMMARY_ANCHOR = "[[VULNERABILITY_COUNT_SUMMARY]]"
VULNERABILITY_BLOCKS_ANCHOR = "[[VULNERABILITY_BLOCKS]]"
CROSS_INFRA_ANALYSIS_ANCHOR = "[[CROSS_INFRA_ANALYSIS]]"
SLA_PRA_SUMMARY_ANCHOR = "[[SLA_PRA_SUMMARY]]"
CROSS_DEPENDENCY_SUMMARY_ANCHOR = "[[CROSS_DEPENDENCY_SUMMARY]]"
NARRATIVE_SOURCES_ANCHOR = "[[NARRATIVE_SOURCES]]"
EXECUTIVE_SUMMARY_START_ANCHOR = "[[EXECUTIVE_SUMMARY_START]]"
VISUALIZATION_START_ANCHOR = "[[VISUALIZATION_START]]"

ANCHORS = {
    "CHART_ELECTRIC_POWER": "[[CHART_ELECTRIC_POWER]]",
    "CHART_COMMUNICATIONS": "[[CHART_COMMUNICATIONS]]",
    "CHART_INFORMATION_TECHNOLOGY": "[[CHART_INFORMATION_TECHNOLOGY]]",
    "CHART_WATER": "[[CHART_WATER]]",
    "CHART_WASTEWATER": "[[CHART_WASTEWATER]]",
    "VULN_NARRATIVE": VULN_NARRATIVE_ANCHOR,
    "TABLE_SUMMARY": "[[TABLE_SUMMARY]]",
    "SLA_PRA_SUMMARY": SLA_PRA_SUMMARY_ANCHOR,
}

# Headers that must NOT appear in output (export-style tables)
EXPORT_TABLE_BAD_HEADERS = {
    "Requires Service",
    "Time to Impact",
    "Time to Impact (hrs)",
    "Loss of Function",
    "Recovery Time",
    "Percent",
    "Capacity After Impact (No Backup)",
}

# Narrative blanks: "____" or longer underscores
UNDERSCORE_RE = re.compile(r"_{3,}")

# Canonical mapping: internal IDs / synonyms -> report display category
REPORT_CATEGORY_CANONICAL = {
    "ENERGY": "Energy",
    "ELECTRIC_POWER": "Energy",
    "ELECTRIC POWER": "Energy",
    "ELECTRIC_POWER_SUPPLY": "Energy",
    "POWER": "Energy",
    "COMMUNICATIONS": "Communications",
    "COMMS": "Communications",
    "INFORMATION_TECHNOLOGY": "Information Technology",
    "INFORMATION TECHNOLOGY": "Information Technology",
    "IT": "Information Technology",
    "INTERNET_PROVIDER": "Information Technology",
    "INTERNET PROVIDER": "Information Technology",
    "WATER": "Water",
    "WASTEWATER": "Wastewater",
    "SEWER": "Wastewater",
    "CRITICAL_PRODUCTS": "Critical Products",
    "CRITICAL PRODUCTS": "Critical Products",
}

# Dependency-only categories: fail fast if VOFC row is out-of-scope (e.g. physical security)
ALLOWED_DEP_CATEGORIES = {
    "Energy", "Communications", "Information Technology",
    "Water", "Wastewater", "Critical Products",
}


def _canon_category(raw: str | None) -> str:
    """Normalize upstream category to canonical report display name."""
    s = (raw or "").strip()
    if not s:
        return ""
    key = s.upper().replace("-", "_")
    key = " ".join(key.split())
    key = key.replace(" ", "_")
    if key in REPORT_CATEGORY_CANONICAL:
        return REPORT_CATEGORY_CANONICAL[key]
    key2 = s.upper().strip()
    if key2 in REPORT_CATEGORY_CANONICAL:
        return REPORT_CATEGORY_CANONICAL[key2]
    return s

# Physical security keywords: reject from dependency VOFC table (wrong domain)
PHYSICAL_SECURITY_KEYWORDS = (
    "cctv", "ids", "badging", "keycard", "access levels", "terminated personnel",
    "exterior ids", "interior ids", "surveillance", "access control",
)

# Placeholder patterns that must not appear in output (Gate A)
FORBIDDEN_PLACEHOLDERS = ("TBD", "Insert ", "Region__", "Insert City", "Insert PSA")
# Gate A: Placeholder phrases (case-insensitive match)
GATE_A_PLACEHOLDER_PHRASES = ("choose an item", "insert", "tbd", "lorem ipsum", "trigger conditions met")
# Gate A: Unresolved anchor pattern
GATE_A_ANCHOR_RE = re.compile(r"\[\[[A-Z0-9_]+\]\]")
# Gate A: Deprecated terms (case-insensitive). "safe" = whole-word only (block SAFE acronym, allow "unsafe"/"safety")
GATE_A_DEPRECATED_TERMS = ("safe", "security assessment at first entry")
GATE_A_SAFE_WORD_BOUNDARY_RE = re.compile(r"\bsafe\b")


def category_no_break(s: str) -> str:
    """Replace spaces with non-breaking spaces so category labels don't split mid-word."""
    return (s or "").replace(" ", "\u00A0").strip()


def render_vulnerability_block(doc: Document, block: dict) -> None:
    """
    Render a single RenderedVulnerabilityBlock with strict paragraph boundaries.
    Order: Title (bold) -> [Category] -> Condition Identified -> Operational Exposure -> Why This Matters
          -> OFC heading -> numbered OFCs -> References heading -> refs.
    When condition_identified/operational_exposure/why_this_matters present, use those; else narrative.
    """
    title = sanitize_vulnerability_text(block.get("title") or "")
    narrative = sanitize_vulnerability_text(block.get("narrative") or "")
    cond_id = sanitize_vulnerability_text(block.get("condition_identified") or "")
    op_exp = sanitize_vulnerability_text(block.get("operational_exposure") or "")
    why = sanitize_vulnerability_text(block.get("why_this_matters") or "")
    driver_cat = block.get("driverCategory") or ""
    ofcs = block.get("ofcs") or []
    references = block.get("references") or []

    # 1) Title paragraph (bold)
    p_title = doc.add_paragraph()
    r = p_title.add_run(title)
    r.bold = True
    set_paragraph_keep_with_next(p_title)

    # 2) Category (when present)
    if driver_cat:
        doc.add_paragraph(f"Category: {driver_cat}")

    # 3) Condition Identified / Operational Exposure / Why This Matters (canonical) or fallback narrative
    if cond_id or op_exp or why:
        if cond_id:
            p_h = doc.add_paragraph()
            r = p_h.add_run("Condition Identified")
            r.bold = True
            doc.add_paragraph(cond_id)
        if op_exp:
            p_h = doc.add_paragraph()
            r = p_h.add_run("Operational Exposure")
            r.bold = True
            doc.add_paragraph(op_exp)
        if why:
            p_h = doc.add_paragraph()
            r = p_h.add_run("Why This Matters")
            r.bold = True
            doc.add_paragraph(why)
    elif narrative:
        doc.add_paragraph(narrative)

    # 4) Options for Consideration heading + numbered OFCs (only when at least one non-empty OFC; max 4)
    ofcs_list = (ofcs[:4] if isinstance(ofcs, list) else [])[:4]
    ofcs_clean = [_normalize_ofc_item(str(x)) for x in ofcs_list if x and str(x).strip()]
    ofcs_clean = [x for x in ofcs_clean if x]
    if ofcs_clean:
        p_ofc_h = doc.add_paragraph()
        r = p_ofc_h.add_run("Options for Consideration")
        r.bold = True
        for i, ofc in enumerate(ofcs_clean, 1):
            doc.add_paragraph(sanitize_vulnerability_text(f"{i}. {ofc}") or ofc)

    # 5) References heading + bullet refs (always after OFCs)
    if references:
        p_ref_h = doc.add_paragraph()
        r = p_ref_h.add_run("References")
        r.bold = True
        for ref in references:
            doc.add_paragraph(sanitize_text(f"\u2022 {ref}"))


def render_vulnerability_block_after(
    doc: Document, block: dict, after_paragraph, add_spacer_after: bool = True
) -> DocxParagraph:
    """
    Strict narrative vulnerability block: Heading 3 (title), Normal (category/body), Heading 4 (section labels),
    Normal (body), List Number (OFCs, max 4), List Bullet (references). No manual "1." or "•" numbering.
    """
    title = sanitize_vulnerability_text(block.get("title") or "")
    narrative = sanitize_vulnerability_text(block.get("narrative") or "")
    cond_id = sanitize_vulnerability_text(block.get("condition_identified") or "")
    op_exp = sanitize_vulnerability_text(block.get("operational_exposure") or "")
    why = sanitize_vulnerability_text(block.get("why_this_matters") or "")
    driver_cat = block.get("driverCategory") or ""
    ofcs = block.get("ofcs") or []
    references = block.get("references") or []
    last = after_paragraph
    p_title = insert_paragraph_after(last, title, style="Heading 3")
    set_paragraph_keep_with_next(p_title)
    last = p_title
    if driver_cat:
        last = insert_paragraph_after(last, f"Category: {driver_cat}", style="Normal")
    if cond_id or op_exp or why:
        if cond_id:
            p_h = insert_paragraph_after(last, "Condition Identified", style="Heading 4")
            set_paragraph_keep_with_next(p_h)
            last = insert_paragraph_after(p_h, cond_id, style="Normal")
        if op_exp:
            p_h = insert_paragraph_after(last, "Operational Exposure", style="Heading 4")
            set_paragraph_keep_with_next(p_h)
            last = insert_paragraph_after(p_h, op_exp, style="Normal")
        if why:
            p_h = insert_paragraph_after(last, "Why This Matters", style="Heading 4")
            set_paragraph_keep_with_next(p_h)
            last = insert_paragraph_after(p_h, why, style="Normal")
    elif narrative:
        last = insert_paragraph_after(last, narrative, style="Normal")
    try:
        if last:
            last.paragraph_format.space_after = Pt(VULN_NARRATIVE_SPACING_PT)
    except Exception:
        pass
    ofcs_list = (ofcs[:4] if isinstance(ofcs, list) else [])[:4]
    ofcs_clean = [_normalize_ofc_item(str(x)) for x in ofcs_list if x and str(x).strip()]
    ofcs_clean = [x for x in ofcs_clean if x]
    if ofcs_clean:
        p_ofc_h = insert_paragraph_after(last, "Options for Consideration", style="Heading 4")
        set_paragraph_keep_with_next(p_ofc_h)
        last = p_ofc_h
        for ofc in ofcs_clean:
            last = insert_paragraph_after(last, ofc, style="List Number")
            try:
                last.paragraph_format.space_after = Pt(VULN_OFC_ITEM_SPACING_PT)
            except Exception:
                pass
    if references:
        p_ref_h = insert_paragraph_after(last, "References", style="Heading 4")
        set_paragraph_keep_with_next(p_ref_h)
        last = p_ref_h
        refs_deduped = sorted(set(str(r).strip() for r in references if r), key=str.lower)
        for ref in refs_deduped:
            last = insert_paragraph_after(last, sanitize_text(ref), style="List Bullet")
    if add_spacer_after:
        try:
            last.paragraph_format.space_after = Pt(VULN_DIVIDER_SPACING_PT)
        except Exception:
            pass
    return last


NOT_IDENTIFIED = "Not identified"  # Only when user explicitly answered NO
NOT_CONFIRMED = "not confirmed"   # When unknown, unset, or no data
SUMMARY_NOT_CONFIRMED_TEXT = "not confirmed"  # Notes when no explicit source data

# --- A) Body iteration and element insertion helpers ---


def _iter_paragraphs(doc: Document):
    """Yield every paragraph in doc (body and inside all table cells)."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def iter_block_items(doc: Document):
    """
    Yield Paragraph or Table in true document order (body block items only).
    Uses doc._element.body.iterchildren(); wraps w:p -> Paragraph, w:tbl -> Table.
    """
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, doc)


def _count_anchor_occurrences(doc: Document, token: str, body_only: bool = True) -> int:
    """Return number of paragraphs (or cells) whose trimmed text equals token. Used to assert uniqueness."""
    needle = (token or "").strip()
    count = 0
    if body_only:
        for p in doc.paragraphs:
            if (p.text or "").strip() == needle:
                count += 1
        return count
    for p in _iter_paragraphs(doc):
        if (p.text or "").strip() == needle:
            count += 1
    return count


def _doc_contains_anchor(doc: Document, token: str, body_only: bool = True) -> bool:
    """Return True if any paragraph (or cell) contains exact text token."""
    needle = (token or "").strip()
    if body_only:
        for p in doc.paragraphs:
            if (p.text or "").strip() == needle:
                return True
        return False
    for p in _iter_paragraphs(doc):
        if (p.text or "").strip() == needle:
            return True
    return False


def _verify_vulnerability_section_boundary(doc: Document) -> None:
    """
    After injecting [[VULNERABILITY_BLOCKS]], verify the next section heading (e.g. CROSS-INFRASTRUCTURE ANALYSIS)
    still appears after the last vulnerability content. Raises RuntimeError with context if section bleed detected.
    Only considers "next section" as a heading that appears *after* the last vuln content (avoids false fail when
    the same heading text appears earlier in the doc, e.g. Part I). Accepts optional numbering prefix (e.g. "C. ").
    """
    paras = doc.paragraphs
    if not paras:
        return
    # Markers: match as substring so "C. CROSS-INFRASTRUCTURE SYNTHESIS" and "CROSS-INFRASTRUCTURE" both match
    NEXT_SECTION_MARKERS = ("CROSS-INFRASTRUCTURE", "CROSS_INFRA", "[[CROSS_INFRA_ANALYSIS]]", "PRIORITY ACTIONS")
    last_vuln_idx = None
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if t.startswith("VULNERABILITY "):
            last_vuln_idx = i
    # First "next section" heading *after* last vuln content (ignore earlier occurrences e.g. in TOC/Part I)
    first_next_section_idx = None
    if last_vuln_idx is not None:
        for i in range(last_vuln_idx + 1, len(paras)):
            t = (paras[i].text or "").strip().upper()
            if any(m.upper() in t for m in NEXT_SECTION_MARKERS):
                first_next_section_idx = i
                break
    if last_vuln_idx is not None and (first_next_section_idx is None or first_next_section_idx <= last_vuln_idx):
        ctx = []
        start = max(0, last_vuln_idx - 2)
        end = min(len(paras), last_vuln_idx + 31)
        for j in range(start, end):
            ctx.append(f"  [{j}] {(paras[j].text or '').strip()[:80]!r}")
        raise RuntimeError(
            "Vulnerability section boundary error: next section heading (CROSS-INFRASTRUCTURE or similar) "
            "does not appear after the last vulnerability content. Possible content bleed or anchor replacement order. "
            "Nearest paragraphs:\n" + "\n".join(ctx)
        )


def _normalize_paragraph_text(t: str) -> str:
    """Collapse any run of whitespace to a single space and strip. Helps match anchors split across runs."""
    if not t:
        return ""
    return re.sub(r"\s+", " ", (t or "").strip())


def find_paragraph_by_exact_text(doc: Document, needle: str, body_only: bool = False):
    """
    Find paragraph whose trimmed text equals needle. If body_only=True, search only body paragraphs (not table cells).
    Tries exact match first, then normalized (whitespace collapsed) so anchors split across runs still match.
    Returns the first matching Paragraph or None.
    """
    needle = (needle or "").strip()
    needle_norm = _normalize_paragraph_text(needle)
    if body_only:
        for p in doc.paragraphs:
            raw = (p.text or "").strip()
            if raw == needle or _normalize_paragraph_text(raw) == needle_norm:
                return p
        return None
    for p in _iter_paragraphs(doc):
        raw = (p.text or "").strip()
        if raw == needle or _normalize_paragraph_text(raw) == needle_norm:
            return p
    return None


def _element_has_drawing_or_pict_child(el) -> bool:
    """True if element or any descendant has local tag name 'drawing', 'pict', or 'blip' (OOXML inline/floating image)."""
    for child in el.iter():
        local = (child.tag or "").split("}")[-1] if "}" in str(child.tag) else (child.tag or "")
        if local in ("drawing", "pict", "blip"):
            return True
    return False


def _run_has_drawing(run) -> bool:
    """True if run contains inline picture or drawing (w:drawing / w:pict)."""
    return _element_has_drawing_or_pict_child(run._element)


def paragraph_has_drawing(p) -> bool:
    """Detect inline pictures and floating drawings anchored to this paragraph/run."""
    return _element_has_drawing_or_pict_child(p._element)


def _clear_paragraph_text_preserve_drawings(p) -> None:
    """Clear text only in runs that do not contain drawings. Leaves drawing runs untouched."""
    for r in p.runs:
        if not _run_has_drawing(r):
            r.text = ""


def remove_paragraph(paragraph) -> None:
    """Remove paragraph from document (safe removal from XML tree). Never remove paragraphs that contain drawings."""
    if paragraph_has_drawing(paragraph):
        _clear_paragraph_text_preserve_drawings(paragraph)
        return
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def remove_table(table) -> None:
    """Remove table from document."""
    parent = table._tbl.getparent()
    if parent is not None:
        parent.remove(table._tbl)


def insert_paragraph_after(paragraph, text: str = "", style=None):
    """
    Insert a new paragraph immediately after the given paragraph (body or cell).
    Returns the new Paragraph. Uses add_paragraph then repositions element.
    Style is applied when the document has that style (e.g. List Bullet, Normal); otherwise skipped.
    """
    parent = paragraph._parent
    new_p = parent.add_paragraph(text)
    new_el = new_p._element
    new_el.getparent().remove(new_el)
    paragraph._element.addnext(new_el)
    if style is not None:
        try:
            new_p.style = style
        except (KeyError, ValueError, AttributeError):
            pass  # Template may not define List Bullet / List Number / etc.
    return new_p


def insert_paragraph_before(paragraph, text: str = "", style=None):
    """
    Insert a new paragraph immediately before the given paragraph (body or cell).
    Returns the new Paragraph. Style applied if present in document; otherwise skipped.
    """
    parent = paragraph._parent
    new_p = parent.add_paragraph(text)
    new_el = new_p._element
    new_el.getparent().remove(new_el)
    paragraph._element.addprevious(new_el)
    if style is not None:
        try:
            new_p.style = style
        except (KeyError, ValueError, AttributeError):
            pass
    return new_p


def add_page_break_paragraph(doc: Document, after_paragraph=None):
    """Insert a page break paragraph. If after_paragraph given, insert after it; else append to body."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    if after_paragraph is not None:
        new_el = p._element
        new_el.getparent().remove(new_el)
        after_paragraph._element.addnext(new_el)
    return p


def _paragraph_has_page_break(para) -> bool:
    """Check if paragraph contains a hard page break (w:br w:type=page)."""
    for el in para._element.iter():
        if el.tag == qn("w:br") and el.get(qn("w:type")) == "page":
            return True
    return False


def _is_page_break_paragraph(p) -> bool:
    """True if paragraph contains a hard page break."""
    if p is None:
        return False
    return _paragraph_has_page_break(p)


def _is_effectively_empty_paragraph(p) -> bool:
    """True if paragraph has no meaningful content (empty text, no images/tables)."""
    if p is None:
        return True
    t = (p.text or "").strip()
    if t:
        return False
    if paragraph_has_drawing(p):
        return False
    return True


def _replace_encoding_artifacts_in_doc(doc: Document) -> None:
    """Replace encoding artifacts (U+FFFD, mojibake) with em dash so QC passes.
    Uses high-level Run API (run.text is writable); avoids raw XML where .text can be read-only (CT_P)."""
    from qc_export import ENCODING_ARTIFACT_PATTERNS
    repl = "\u2014"  # em dash
    for pat in ENCODING_ARTIFACT_PATTERNS:
        for p in doc.paragraphs:
            for run in p.runs:
                if run.text and pat in run.text:
                    run.text = run.text.replace(pat, repl)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if run.text and pat in run.text:
                                run.text = run.text.replace(pat, repl)


def _trim_trailing_empty_paragraphs(doc: Document, max_trim: int = 25) -> None:
    """Remove trailing empty paragraphs to prevent Word from creating blank pages. Never remove drawing paragraphs."""
    trimmed = 0
    while trimmed < max_trim and doc.paragraphs:
        p = doc.paragraphs[-1]
        if _is_page_break_paragraph(p) or paragraph_has_drawing(p):
            break
        if not _is_effectively_empty_paragraph(p):
            break
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)
            trimmed += 1
        else:
            break


def _trim_empty_paragraphs_before(para, max_trim: int = 25) -> None:
    """Remove empty paragraphs immediately before the given paragraph. Never remove drawing paragraphs."""
    trimmed = 0
    while trimmed < max_trim:
        prev = para._element.getprevious()
        if prev is None or prev.tag != qn("w:p"):
            break
        prev_para = DocxParagraph(prev, para._parent)
        if _is_page_break_paragraph(prev_para) or paragraph_has_drawing(prev_para):
            break
        if not _is_effectively_empty_paragraph(prev_para):
            break
        parent = prev.getparent()
        if parent is not None:
            parent.remove(prev)
            trimmed += 1
        else:
            break


def remove_blank_pages_after_section_b(doc: Document) -> None:
    """
    Remove hard-coded blank pages after Section B (DEPENDENCY SNAPSHOT TABLE).
    Strips empty paragraphs and page-break paragraphs between SNAPSHOT_CASCADE
    and the next content (C. SECTOR ANALYSIS, D. CROSS, etc.), up to 5 paras.
    """
    p_cascade = find_paragraph_by_exact_text(doc, "[[SNAPSHOT_CASCADE]]", body_only=True)
    if p_cascade is None:
        return
    body = doc.element.body
    el = p_cascade._element
    removed = 0
    for _ in range(5):
        nxt = el.getnext()
        if nxt is None or nxt.tag != qn("w:p"):
            break
        para = DocxParagraph(nxt, doc)
        text = (para.text or "").strip()
        # Stop at next section or anchor
        if any(
            x in text.upper()
            for x in (
                "C. SECTOR",
                "D. CROSS",
                "E. PRIORITY",
                "PART II",
                "[[",
            )
        ):
            break
        if _is_page_break_paragraph(para) or _is_effectively_empty_paragraph(para):
            body.remove(nxt)
            removed += 1
            # el unchanged; getnext() will now return the element after removed nxt
        else:
            break
    if removed:
        pass  # silent; no logging needed


def remove_page_breaks_between_annex_table_and_sector_reports(doc: Document) -> None:
    """
    Remove template page break between Annex Overview table and Sector Reports.
    Prevents blank page when table ends and only 'Sector Reports' heading would appear.
    """
    # Find Sector Reports / Infrastructure Dependency heading (immediately before VULN_NARRATIVE)
    p_heading = None
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if SECTOR_REPORTS_HEADING in t or "Infrastructure Dependency" in t:
            p_heading = p
            break
    if p_heading is None:
        return
    body = doc.element.body
    for _ in range(5):
        prev = p_heading._element.getprevious()
        if prev is None:
            break
        if prev.tag == qn("w:tbl"):
            break
        if prev.tag != qn("w:p"):
            continue
        para = DocxParagraph(prev, doc)
        text = (para.text or "").strip()
        if text and (ANNEX_OVERVIEW_HEADING in text or "Dependency Summary" in text):
            break
        if _is_page_break_paragraph(para) or _is_effectively_empty_paragraph(para):
            body.remove(prev)
        else:
            break


def remove_orphaned_page_breaks_before_section_d(doc: Document) -> None:
    """
    Remove orphaned page-break paragraphs left after removing C. SECTOR ANALYSIS
    and sector content. Each sector had a page break before it; removing the
    sector headings/charts leaves 5 orphaned page breaks before D. CROSS.
    """
    p_d = None
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if "D. CROSS" in t.upper() and "SYNTHESIS" in t.upper():
            p_d = p
            break
    if p_d is None:
        return
    body = doc.element.body
    for _ in range(10):
        prev = p_d._element.getprevious()
        if prev is None:
            break
        if prev.tag != qn("w:p"):
            break
        para = DocxParagraph(prev, doc)
        if _is_page_break_paragraph(para) or _is_effectively_empty_paragraph(para):
            body.remove(prev)
        else:
            break


def collapse_consecutive_pagebreaks(doc: Document) -> None:
    """Remove page break paragraphs when two occur back-to-back."""
    body = doc.element.body
    children = list(body)
    i = 1
    while i < len(children):
        prev_el = children[i - 1]
        cur_el = children[i]
        if prev_el.tag != qn("w:p") or cur_el.tag != qn("w:p"):
            i += 1
            continue
        prev_para = DocxParagraph(prev_el, doc)
        cur_para = DocxParagraph(cur_el, doc)
        if _paragraph_has_page_break(prev_para) and _paragraph_has_page_break(cur_para):
            body.remove(cur_el)
            children = list(body)
            continue
        i += 1
    # Some templates have non-paragraph XML nodes between page-break paragraphs; clean those too.
    paras = list(doc.paragraphs)
    for i in range(1, len(paras)):
        prev_para = paras[i - 1]
        cur_para = paras[i]
        if _paragraph_has_page_break(prev_para) and _paragraph_has_page_break(cur_para):
            if not paragraph_has_drawing(cur_para):
                remove_paragraph(cur_para)


def ensure_single_page_break_before(doc: Document, para) -> None:
    """
    Ensure exactly one page break exists immediately before the given paragraph.
    Trims trailing empty paragraphs, avoids double breaks.
    """
    _trim_empty_paragraphs_before(para)
    prev = para._element.getprevious()
    if prev is not None and prev.tag == qn("w:p") and _element_has_page_break(prev):
        return
    new_p = insert_paragraph_before(para, "")
    run = new_p.add_run()
    run.add_break(WD_BREAK.PAGE)


def ensure_single_page_break_after(doc: Document, after_para):
    """
    Ensure exactly one page break exists after the given paragraph.
    Returns the paragraph to insert after (either after_para or the new break para).
    """
    if _is_page_break_paragraph(after_para):
        return after_para
    return add_page_break_paragraph(doc, after_para)


def _element_has_page_break(elm) -> bool:
    """Check if an XML element (e.g. w:p) contains a hard page break."""
    for child in elm.iter():
        if child.tag == qn("w:br") and child.get(qn("w:type")) == "page":
            return True
    return False


def ensure_part2_starts_new_page(doc: Document) -> None:
    """
    Force exactly one page break before PART II – TECHNICAL ANNEX.
    Uses page_break_before on the paragraph (no extra paragraph = no blank page).
    """
    part2_para = None
    for p in doc.paragraphs:
        t = (p.text or "").strip().upper()
        if "PART II" in t and "TECHNICAL ANNEX" in t:
            part2_para = p
            break
    if part2_para is None:
        return
    try:
        part2_para.paragraph_format.page_break_before = True
    except Exception:
        ensure_single_page_break_before(doc, part2_para)
    set_paragraph_keep_with_next(part2_para)


def _safe_get(d, *keys, default=None):
    """Navigate nested dict by keys; return default if any key missing."""
    cur = d
    for k in keys:
        if not isinstance(cur, dict) or k not in cur:
            return default
        cur = cur[k]
    return cur


def _dedupe_rows(rows):
    """Deduplicate rows by normalized tuple (lowercase, stripped); preserve first occurrence order."""
    seen = set()
    out = []
    for r in rows:
        key = tuple((x or "").strip().lower() for x in r)
        if key in seen:
            continue
        seen.add(key)
        out.append(r)
    return out


# Known ISP/transport provider names (normalized). Used to route IT-1 providers to Internet Transport only; never in Critical Hosted Services.
# Mirrors hosted_service_registry TRANSPORT_ISP logic; reporter has no registry so we use name matching.
_TRANSPORT_ISP_NORMALIZED = frozenset({
    "comcast", "xfinity", "at&t", "att", "at&t internet", "at&t fiber", "att internet", "att fiber",
    "verizon", "verizon fios", "verizon fios internet", "spectrum", "charter", "cox", "centurylink",
    "lumen", "frontier", "windstream", "earthlink", "optimum", "altice", "cogent", "zayo", "level3",
    "l3", "crown castle", "lumen technologies",
})


def _normalize_provider_name(name: str) -> str:
    """Lowercase, collapse spaces, remove common punctuation for lookup."""
    if not name or not isinstance(name, str):
        return ""
    return re.sub(r"\s+", " ", name.lower().strip().replace("&", "").replace(".", ""))


def _is_transport_provider(provider_name: str) -> bool:
    """True if provider is a known transport/ISP (include only in Internet Transport, never in Critical Hosted Services)."""
    n = _normalize_provider_name(provider_name)
    if not n:
        return False
    if n in _TRANSPORT_ISP_NORMALIZED:
        return True
    for isp in _TRANSPORT_ISP_NORMALIZED:
        if isp in n or n in isp:
            return True
    return False


def _format_path_diversity(path: dict) -> str:
    """Format physical_path_diversity dict to comma-separated labels."""
    if not path or not isinstance(path, dict):
        return ""
    parts = []
    if path.get("same_conduit"):
        parts.append("Same conduit")
    if path.get("separate_conduits"):
        parts.append("Separate conduits")
    if path.get("separate_street_approach"):
        parts.append("Separate street approach")
    if path.get("unknown"):
        parts.append("Unknown")
    return ", ".join(parts) if parts else ""


def _it_transport_table(assessment_json: dict):
    """Build IT Internet Transport Resilience table from it_transport_resilience, or legacy supply.sources."""
    assessment = assessment_json.get("assessment") or assessment_json
    it_cat = _safe_get(assessment, "categories", "INFORMATION_TECHNOLOGY", default={}) or {}
    transport = it_cat.get("it_transport_resilience")
    if transport and isinstance(transport, dict):
        path = transport.get("physical_path_diversity") or {}
        path_str = _format_path_diversity(path)
        row = [
            (transport.get("circuit_count") or "").replace("_", " ").strip() or "—",
            (transport.get("carrier_diversity") or "").replace("_", " ").strip() or "—",
            path_str or "—",
            (transport.get("building_entry_diversity") or "").replace("_", " ").strip() or "—",
            (transport.get("upstream_pop_diversity") or "").replace("_", " ").strip() or "—",
            (transport.get("notes") or "").strip() or "—",
        ]
        return {
            "type": "table",
            "title": "Information Technology – Internet Transport Resilience",
            "headers": ["Circuit Count", "Carrier Diversity", "Path Diversity", "Building Entry", "Upstream POP", "Notes"],
            "rows": [row],
        }
    # Legacy: from supply.sources
    sources = _safe_get(assessment, "categories", "INFORMATION_TECHNOLOGY", "supply", "sources", default=[]) or []
    if not sources:
        return None
    rows = []
    for s in sources:
        provider = (s.get("provider_name") or "").strip() or "Unknown"
        demarc = (s.get("demarcation_description") or "").strip()
        independence = (s.get("independence") or "").strip() or "UNKNOWN"
        notes = (s.get("notes") or "").strip()
        if provider or demarc or independence or notes:
            rows.append([provider, demarc, independence, notes])
    rows = _dedupe_rows(rows)
    if not rows:
        return None
    return {
        "type": "table",
        "title": "Information Technology – Transport Providers (Internet / Circuits)",
        "headers": ["Provider", "Demarcation / Termination", "Independence", "Notes"],
        "rows": rows,
    }


def _diversity_status_from_transport(transport: dict) -> str:
    """Derive diversity label from it_transport_resilience. Single-path / Partial / Diverse."""
    if not transport or not isinstance(transport, dict):
        return "Single-path exposure"
    circuit = (transport.get("circuit_count") or "").strip().upper().replace("-", "_")
    if circuit in ("ONE", "1"):
        return "Single-path exposure"
    carriers = (transport.get("carrier_diversity") or "").strip().upper()
    entry = (transport.get("building_entry_diversity") or "").strip().upper()
    if (carriers == "DIFFERENT_CARRIERS" and entry == "SEPARATE_ENTRIES") or carriers == "DIFFERENT_CARRIERS":
        return "Diverse transport"
    if circuit in ("TWO", "THREE_PLUS", "2", "3"):
        return "Partial diversity"
    return "Single-path exposure"


def _format_tti_for_table(val) -> str:
    """Format time-to-severe-impact for table cell (hours number or category string)."""
    if val is None:
        return "—"
    if isinstance(val, (int, float)) and val == val:
        return str(int(val)) if val == int(val) else str(val)
    s = (val or "").strip()
    return s if s else "—"


def _get_it_isp_names_from_curve(assessment: dict) -> list[str]:
    """
    Source of truth for ISP provider names: Internet/Data Connectivity curve fields only.
    Returns [primary, secondary] filtered to non-empty strings. Do NOT infer from supply.sources,
    IT-1_service_providers, hosted_services, or dependency_providers.
    """
    it_cat = _safe_get(assessment, "categories", "INFORMATION_TECHNOLOGY", default={}) or {}
    primary = (it_cat.get("curve_primary_provider") or "")
    secondary = (it_cat.get("curve_secondary_provider") or "")
    if isinstance(primary, str):
        primary = primary.strip()
    else:
        primary = str(primary).strip() if primary is not None else ""
    if isinstance(secondary, str):
        secondary = secondary.strip()
    else:
        secondary = str(secondary).strip() if secondary is not None else ""
    return [p for p in [primary, secondary] if p]


def _format_it_internet_connectivity_narrative(assessment: dict) -> str:
    """
    Narrative sentence using ISP names from curve only. Example: "The facility receives primary
    internet connectivity from Verizon with a secondary connection provided by Xfinity."
    """
    isp_names = _get_it_isp_names_from_curve(assessment)
    if not isp_names:
        return "The facility did not report primary or secondary internet connectivity providers in the assessment."
    if len(isp_names) == 1:
        return f"The facility receives primary internet connectivity from {isp_names[0]}."
    return (
        f"The facility receives primary internet connectivity from {isp_names[0]} "
        f"with a secondary connection provided by {isp_names[1]}."
    )


def _normalize_provider_name(s: str) -> str:
    """Normalize for matching provider names (strip, lowercase)."""
    return (s or "").strip().lower()


def _compact_provider_name(s: str) -> str:
    """Provider normalization for fuzzy matches (alnum only)."""
    return re.sub(r"[^a-z0-9]", "", _normalize_provider_name(s))


def _humanize_transport_independence(value) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    upper = raw.upper()
    if upper in INDEPENDENCE_LABELS:
        return INDEPENDENCE_LABELS[upper]
    if upper == "CONFIRMED":
        return "Confirmed"
    if upper == "NOT_CONFIRMED":
        return "Not confirmed"
    return raw


def _it_supply_source_lookup(it_cat: dict) -> dict:
    """
    Build lookup from normalized provider_name to supply source (demarcation, independence, notes).
    Uses categories.INFORMATION_TECHNOLOGY.supply.sources; fallback to .answers if needed.
    """
    out = {"by_name": {}, "by_compact": {}, "sources": []}
    supply = it_cat.get("supply")
    if not supply or not isinstance(supply, dict):
        answers = it_cat.get("answers") or {}
        if isinstance(answers, dict) and "supply" in answers:
            supply = answers.get("supply")
    sources = (supply or {}).get("sources") if isinstance(supply, dict) else []
    if not isinstance(sources, list):
        return out
    out["sources"] = sources
    for src in sources:
        if not src or not isinstance(src, dict):
            continue
        name = (src.get("provider_name") or "").strip()
        if not name:
            continue
        key_name = _normalize_provider_name(name)
        key_compact = _compact_provider_name(name)
        out["by_name"][key_name] = src
        out["by_compact"][key_compact] = src
    return out


def _format_demarcation(src: dict) -> str:
    """Format demarcation from source: description, or 'lat, lon', or blank when missing."""
    if not src or not isinstance(src, dict):
        return ""
    desc = (src.get("demarcation_description") or "").strip()
    if desc:
        return desc
    lat = src.get("demarcation_lat")
    lon = src.get("demarcation_lon")
    if lat is not None and lon is not None:
        return f"{lat}, {lon}"
    return ""


def _internet_transport_table(assessment_json: dict) -> dict:
    """
    BLOCK 2: INTERNET TRANSPORT. Single consolidated table. Always returns a block (never None).
    ISP provider names from curve_primary_provider and curve_secondary_provider; Demarcation,
    Independence, and Notes joined from IT supply.sources where provider_name matches.
    """
    assessment = assessment_json.get("assessment") or assessment_json
    it_cat = _safe_get(assessment, "categories", "INFORMATION_TECHNOLOGY", default={}) or {}
    isp_names = _get_it_isp_names_from_curve(assessment)
    primary_raw = (it_cat.get("curve_primary_provider") or "").strip() if isinstance(it_cat.get("curve_primary_provider"), str) else ""
    secondary_raw = (it_cat.get("curve_secondary_provider") or "").strip() if isinstance(it_cat.get("curve_secondary_provider"), str) else ""
    if (primary_raw or secondary_raw) and not isp_names:
        raise RuntimeError(
            "ISP provider names not mapped from curve inputs. Ensure assessment.categories.INFORMATION_TECHNOLOGY "
            "contains curve_primary_provider and curve_secondary_provider."
        )
    source_lookup = _it_supply_source_lookup(it_cat)
    source_by_name = source_lookup.get("by_name") or {}
    source_by_compact = source_lookup.get("by_compact") or {}
    source_list = source_lookup.get("sources") or []
    answers = it_cat.get("answers") if isinstance(it_cat.get("answers"), dict) else {}
    it_connections = it_cat.get("IT-4_service_connections")
    if not isinstance(it_connections, list):
        it_connections = answers.get("IT-4_service_connections") if isinstance(answers, dict) else []
    if not isinstance(it_connections, list):
        it_connections = []
    transport = it_cat.get("it_transport_resilience")
    if not isinstance(transport, dict) and isinstance(answers, dict):
        transport = answers.get("it_transport_resilience")
    if not isinstance(transport, dict):
        transport = {}

    def _empty_to_not_provided(v: str) -> str:
        return v.strip() if isinstance(v, str) and v.strip() else "Not provided"

    def _find_source_for_provider(provider: str, idx: int):
        n = _normalize_provider_name(provider)
        c = _compact_provider_name(provider)
        src = source_by_name.get(n)
        if src:
            return src
        src = source_by_compact.get(c)
        if src:
            return src
        for s in source_list:
            candidate = str((s or {}).get("provider_name") or "")
            nc = _normalize_provider_name(candidate)
            cc = _compact_provider_name(candidate)
            if (nc and (nc in n or n in nc)) or (cc and (cc in c or c in cc)):
                return s
        if 0 <= idx < len(source_list):
            return source_list[idx]
        return None

    def _find_connection_for_provider(provider: str, idx: int):
        n = _normalize_provider_name(provider)
        c = _compact_provider_name(provider)
        for conn in it_connections:
            assoc = str((conn or {}).get("associated_provider") or "")
            if _normalize_provider_name(assoc) == n:
                return conn
        for conn in it_connections:
            assoc = str((conn or {}).get("associated_provider") or "")
            nc = _normalize_provider_name(assoc)
            cc = _compact_provider_name(assoc)
            if (nc and (nc in n or n in nc)) or (cc and (cc in c or c in cc)):
                return conn
        if 0 <= idx < len(it_connections):
            return it_connections[idx]
        return None

    headers = ["Role", "Provider", "Demarcation", "Independence", "Notes"]
    rows = []
    if isp_names:
        roles = ["Primary Internet Provider", "Secondary Internet Provider"]
        for i, provider in enumerate(isp_names[:2]):
            role = roles[i]
            src = _find_source_for_provider(provider, i)
            conn = _find_connection_for_provider(provider, i)
            demarc_from_source = _format_demarcation(src) if isinstance(src, dict) else ""
            demarc_from_conn = str((conn or {}).get("facility_entry_location") or "").strip()
            demarc = _empty_to_not_provided(demarc_from_source or demarc_from_conn)

            indep_source = _humanize_transport_independence((src or {}).get("independence"))
            route_ind = _humanize_transport_independence((transport or {}).get("transport_route_independence"))
            phys_sep_raw = str(it_cat.get("IT-4_physically_separated") or "").strip().lower()
            if not phys_sep_raw and isinstance(answers, dict):
                phys_sep_raw = str(answers.get("IT-4_physically_separated") or "").strip().lower()
            phys_sep = "Physically separated" if phys_sep_raw == "yes" else "Not physically separated" if phys_sep_raw == "no" else ""
            indep = _empty_to_not_provided(indep_source or route_ind or phys_sep)

            notes_raw = str((src or {}).get("notes") or "").strip()
            if not notes_raw and isinstance(src, dict):
                notes_raw = str(src.get("demarcation_description") or "").strip()
            notes = notes_raw if notes_raw else ("Reported sources: 1" if (demarc != "Not provided" or indep != "Not provided") else "—")
            rows.append([role, provider, demarc, indep, notes])
    if not rows:
        rows.append(["Primary Internet Provider", "Not provided", "Not provided", "Not provided", "—"])
    return {
        "type": "table",
        "title": "INTERNET TRANSPORT",
        "headers": headers,
        "rows": rows,
    }


def _hosted_continuity_label_for_summary(entry: dict) -> str:
    """4-state continuity for Dependency Summary. Undefined = 'Not assessed'. New enum + legacy NONE/MANUAL_FALLBACK/LOCAL_MIRROR."""
    if not entry or not isinstance(entry, dict):
        return "Not assessed"
    survivability = entry.get("survivability")
    if survivability == "NO_CONTINUITY" or survivability == "NONE":
        return "No continuity"
    if survivability == "LOCAL_MIRROR_OR_CACHE" or survivability == "LOCAL_MIRROR":
        return "Local mirror/cache"
    if survivability == "ALTERNATE_PLATFORM_OR_PROVIDER" or survivability == "MANUAL_FALLBACK":
        return "Alternate platform/provider"
    if survivability == "UNKNOWN":
        return "Unknown"
    if entry.get("local_mirror_or_offline_fallback") or entry.get("local_data_export"):
        return "Local mirror/cache"
    if (
        entry.get("continuity_mechanism_in_place")
        or entry.get("offline_fallback")
        or entry.get("origin_failover")
        or entry.get("multi_pop")
        or entry.get("secondary_dns")
    ):
        return "Alternate platform/provider"
    return "Not assessed"


# Service Loss: plain-language description of what is lost if the hosted service is unreachable.
# Key = service_id (lowercase). No vendor names. Default: "Hosted Application Service".
SERVICE_LOSS_DESCRIPTIONS = {
    "aws": "Loss of hosted compute used to run business applications and services.",
    "azure": "Loss of hosted compute used to run business applications and services.",
    "gcp": "Loss of hosted compute used to run business applications and services.",
    "oracle_cloud": "Loss of hosted compute used to run business applications and services.",
    "cloudflare": "Loss of name resolution and/or content delivery that supports public and internal services.",
    "cloudflare_zero_trust": "Loss of controlled access to external internet and remote access policy enforcement.",
    "m365": "Loss of email, calendaring, and collaboration tooling used for coordination.",
    "office_365": "Loss of email, calendaring, and collaboration tooling used for coordination.",
    "teams": "Loss of video and chat collaboration used for coordination.",
    "google_workspace": "Loss of email, calendaring, and collaboration tooling used for coordination.",
    "entra_id": "Loss of centralized authentication/authorization used for application access.",
    "okta": "Loss of centralized authentication/authorization used for application access.",
    "ping": "Loss of centralized authentication/authorization used for application access.",
    "zscaler": "Loss of controlled access to external internet and remote access policy enforcement.",
    "prisma_access": "Loss of controlled access to external internet and remote access policy enforcement.",
    "cisco_secure_client": "Loss of secure remote access used for workforce connectivity.",
    "fortinet_sase": "Loss of controlled access to external internet and remote access policy enforcement.",
    "sap_erp": "Loss of ERP and core business operations systems.",
    "oracle_erp": "Loss of ERP and core business operations systems.",
    "workday_hris": "Loss of HR and payroll management systems.",
    "adp_hris": "Loss of HR and payroll management systems.",
    "salesforce_crm": "Loss of customer relationship management and customer operations.",
    "stripe": "Loss of payments and e-commerce processing.",
    "paypal": "Loss of payments and e-commerce processing.",
    "shopify": "Loss of e-commerce platform operations.",
    "zoom": "Loss of video and voice meetings.",
    "ringcentral": "Loss of voice, video, and messaging.",
    "webex": "Loss of video and voice meetings.",
    "genesys_cloud": "Loss of contact center and customer engagement.",
    "twilio": "Loss of voice, SMS, and communications API capabilities.",
    "servicenow": "Loss of IT service management and ticketing.",
    "jira_confluence": "Loss of project tracking and documentation.",
    "datadog": "Loss of infrastructure and application monitoring.",
    "onedrive_sharepoint": "Loss of access to stored files/objects used for operations and applications.",
    "google_drive": "Loss of access to stored files/objects used for operations and applications.",
    "dropbox_business": "Loss of access to stored files/objects used for operations and applications.",
    "veeam_cloud_connect": "Loss of backup and disaster recovery capabilities.",
    "web_eoc": "Loss of emergency operations and crisis management tooling.",
    "physical_security_systems": "Loss of access control and video surveillance capabilities.",
}

# Legacy map (service type labels); used only when SERVICE_LOSS_DESCRIPTIONS has no match for backward compat.
SERVICE_LOSS_MAP = {
    "aws": "Virtual Servers / Cloud Compute Infrastructure",
    "azure": "Virtual Servers / Cloud Compute Infrastructure",
    "gcp": "Virtual Servers / Cloud Compute Infrastructure",
    "oracle_cloud": "Virtual Servers / Cloud Compute Infrastructure",
    "cloudflare": "DNS / CDN / Edge Security Services",
    "cloudflare_zero_trust": "Secure Internet Gateway / Zero Trust Network Access",
    "m365": "Email / Collaboration / Document Storage",
    "office_365": "Email / Collaboration / Document Storage",
    "teams": "Video and Chat Collaboration",
    "google_workspace": "Email / Collaboration / Document Storage",
    "entra_id": "Identity / Authentication / SSO",
    "okta": "Identity / Authentication / SSO",
    "ping": "Identity / Authentication / SSO",
    "zscaler": "Secure Internet Gateway / Zero Trust Network Access",
    "prisma_access": "Secure Internet Gateway / Remote Access",
    "cisco_secure_client": "Secure Remote Access / VPN",
    "fortinet_sase": "Secure Internet Gateway / Remote Access",
    "sap_erp": "ERP / Core Business Operations",
    "oracle_erp": "ERP / Core Business Operations",
    "workday_hris": "HR / Payroll Management System",
    "adp_hris": "HR / Payroll Management System",
    "salesforce_crm": "CRM / Customer Operations",
    "stripe": "Payments / E-commerce",
    "paypal": "Payments / E-commerce",
    "shopify": "E-commerce Platform",
    "zoom": "Video / Voice Meetings",
    "ringcentral": "Voice / Video / Messaging",
    "webex": "Video / Voice Meetings",
    "genesys_cloud": "Contact Center / Customer Engagement",
    "twilio": "Voice / SMS / Communications API",
    "servicenow": "IT Service Management / Ticketing",
    "jira_confluence": "Project Tracking / Documentation",
    "datadog": "Infrastructure / Application Monitoring",
    "onedrive_sharepoint": "File Storage / Document Collaboration",
    "google_drive": "File Storage / Document Collaboration",
    "dropbox_business": "File Storage / Sync",
    "veeam_cloud_connect": "Backup / Disaster Recovery",
    "web_eoc": "Emergency Operations / Crisis Management",
    "physical_security_systems": "Access Control / Video Surveillance",
}


def _service_loss_description(service_id: str) -> str:
    """Return Service Loss column value: plain-language description of what is lost if service is unreachable."""
    key = (service_id or "").strip().lower()
    if not key or key == "other":
        return "Hosted Application Service"
    return SERVICE_LOSS_DESCRIPTIONS.get(key, SERVICE_LOSS_MAP.get(key, "Hosted Application Service"))


def _it_critical_hosted_table(assessment_json: dict) -> dict:
    """
    BLOCK 3: CRITICAL HOSTED SERVICES. Single table; only hosted dependencies (no IT-1/ISP rows).
    Data: IT-2_upstream_assets only; filter OUT any transport provider (routing guard).
    Column 4 = Service Loss (type of capability lost if unreachable). Operational Impact unchanged (curve).
    """
    assessment = assessment_json.get("assessment") or assessment_json
    it_cat = _safe_get(assessment, "categories", "INFORMATION_TECHNOLOGY", default={}) or {}
    it_answers = _safe_get(assessment_json, "sessions", "INFORMATION_TECHNOLOGY", "answers", default={}) or {}
    upstream_raw = it_cat.get("IT-2_upstream_assets") or it_answers.get("IT-2_upstream_assets") or []
    upstream = [u for u in (upstream_raw if isinstance(upstream_raw, list) else list((upstream_raw or {}).values())) if isinstance(u, dict)]
    tti = it_cat.get("time_to_impact_hours") or it_cat.get("curve_time_to_impact_hours")
    tti_str = _format_tti_for_table(tti)
    headers = ["Service", "Provider", "Operational Impact (Time to Severe Impact)", "Service Loss"]
    rows = []
    for u in upstream:
        provider = (u.get("service_provider") or "").strip() or "Not provided"
        if _is_transport_provider(provider):
            if os.environ.get("ADA_REPORTER_DEBUG", "").strip().lower() in ("1", "true", "yes"):
                print(f"[reporter] Routing guard: excluding transport provider from Critical Hosted Services: {provider!r}", file=sys.stderr)
            continue
        service_id = (u.get("service_id") or "").strip()
        service_other = (u.get("service_other") or "").strip()
        service_label = service_other if service_other and str(service_id).lower() == "other" else (service_id or "Unknown service")
        dep_id = (service_id or "").lower() if str(service_id).lower() != "other" else ""
        service_loss = _service_loss_description(dep_id or service_id)
        rows.append([service_label, provider, tti_str, service_loss])
    if not rows:
        rows.append(["No critical hosted services identified.", "—", "—", "—"])
    # Export QC: no ISP must appear in Critical Hosted Services
    for row in rows:
        if len(row) >= 2 and row[1] not in ("—", ""):
            if _is_transport_provider(str(row[1])):
                raise RuntimeError("ISP rendered as hosted service; routing bug.")
    return {
        "type": "table",
        "title": "CRITICAL HOSTED SERVICES",
        "headers": headers,
        "rows": rows,
    }


# Primary function: descriptive, impact-oriented text per IT-2 service_id (what it does; impact if unavailable).
IT_SERVICE_ID_TO_PRIMARY_FUNCTION = {
    "aws": "Hosts applications and data; outage affects dependent systems and operations",
    "azure": "Hosts applications and data; outage affects dependent systems and operations",
    "gcp": "Hosts applications and data; outage affects dependent systems and operations",
    "oracle_cloud": "Hosts applications and data; outage affects dependent systems and operations",
    "cloudflare": "CDN and DDoS protection; outage affects web availability and attack resilience",
    "m365": "Email, collaboration, and file storage; outage affects internal communication and document access",
    "office_365": "Email, collaboration, and file storage; outage affects internal communication and document access",
    "teams": "Video and chat collaboration; outage affects meetings and real-time communication",
    "google_workspace": "Email, collaboration, and file storage; outage affects internal communication and document access",
    "entra_id": "Authentication and single sign-on; outage can lock users out of multiple systems",
    "okta": "Authentication and single sign-on; outage can lock users out of multiple systems",
    "ping": "Authentication and single sign-on; outage can lock users out of multiple systems",
    "zscaler": "Secure remote access and web filtering; outage affects remote workforce and secure internet access",
    "prisma_access": "Secure remote access and web filtering; outage affects remote workforce and secure internet access",
    "cisco_secure_client": "Secure remote access (VPN); outage affects remote workforce connectivity",
    "fortinet_sase": "Secure remote access and web filtering; outage affects remote workforce and secure internet access",
    "cloudflare_zero_trust": "Secure remote access and web filtering; outage affects remote workforce and secure internet access",
    "sap_erp": "ERP (finance, supply chain); outage affects core business processes and reporting",
    "oracle_erp": "ERP (finance, supply chain); outage affects core business processes and reporting",
    "workday_hris": "Payroll and HR; outage affects pay runs and HR operations",
    "adp_hris": "Payroll and HR; outage affects pay runs and HR operations",
    "salesforce_crm": "Customer relationship management; outage affects sales and customer support",
    "stripe": "Payment processing; outage affects revenue and transaction completion",
    "paypal": "Payment processing; outage affects revenue and transaction completion",
    "shopify": "E-commerce platform; outage affects online sales and order fulfillment",
    "zoom": "Video and voice meetings; outage affects internal and external communication",
    "ringcentral": "Voice, video, and messaging; outage affects internal and customer communication",
    "webex": "Video and voice meetings; outage affects internal and external communication",
    "genesys_cloud": "Contact center and customer engagement; outage affects customer support and routing",
    "twilio": "Voice, SMS, and communications API; outage affects customer contact and notifications",
    "servicenow": "IT service management and ticketing; outage affects incident and change management",
    "jira_confluence": "Project tracking and documentation; outage affects development and knowledge sharing",
    "datadog": "Infrastructure and application monitoring; outage affects visibility and incident detection",
    "onedrive_sharepoint": "File storage and collaboration; outage affects document access and sharing",
    "google_drive": "File storage and collaboration; outage affects document access and sharing",
    "dropbox_business": "File storage and sync; outage affects document access and sharing",
    "veeam_cloud_connect": "Backup and disaster recovery; outage affects recovery capability and RTO",
    "web_eoc": "Emergency operations and crisis management; outage affects incident coordination and situational awareness",
    "physical_security_systems": "Access control, video surveillance, and alarms (web-dependent); outage affects site security and monitoring",
    "other": None,  # Use service_other or "Other"
}


def _it_primary_function(upstream_entry: dict) -> str:
    """Return primary function for an IT-2 upstream entry (service type / role)."""
    sid = (upstream_entry.get("service_id") or "").strip().lower()
    if sid == "other":
        other = (upstream_entry.get("service_other") or "").strip()
        return other if other else "Other"
    return IT_SERVICE_ID_TO_PRIMARY_FUNCTION.get(sid) or "Unknown"


def _normalize_provider_list(raw) -> list:
    """Ensure IT-1_service_providers is a list of dicts (payload may send array or keyed dict)."""
    if raw is None:
        return []
    if isinstance(raw, list):
        return [p for p in raw if isinstance(p, dict)]
    if isinstance(raw, dict):
        return [v for v in raw.values() if isinstance(v, dict)]
    return []


def _hosted_continuity_label(entry: dict) -> str:
    """Format it_hosted_resilience entry to 4-state continuity label. Undefined = 'Not assessed'. New enum + legacy."""
    if not entry or not isinstance(entry, dict):
        return "Not assessed"
    survivability = entry.get("survivability")
    if survivability == "NO_CONTINUITY" or survivability == "NONE":
        return "No continuity"
    if survivability == "LOCAL_MIRROR_OR_CACHE" or survivability == "LOCAL_MIRROR":
        return "Local mirror/cache"
    if survivability == "ALTERNATE_PLATFORM_OR_PROVIDER" or survivability == "MANUAL_FALLBACK":
        return "Alternate platform/provider"
    if survivability == "UNKNOWN":
        return "Unknown"
    if entry.get("local_mirror_or_offline_fallback") or entry.get("local_data_export"):
        return "Local mirror/cache"
    if (
        entry.get("continuity_mechanism_in_place")
        or entry.get("offline_fallback")
        or entry.get("origin_failover")
        or entry.get("multi_pop")
        or entry.get("secondary_dns")
    ):
        return "Alternate platform/provider"
    return "Not assessed"


def _it_hosted_table(assessment_json: dict):
    """Build IT hosted/upstream dependencies table from IT-2 and IT-1; add Continuity from it_hosted_resilience (3-state)."""
    assessment = assessment_json.get("assessment") or assessment_json
    it_cat = _safe_get(assessment, "categories", "INFORMATION_TECHNOLOGY", default={}) or {}
    it_answers = _safe_get(assessment_json, "sessions", "INFORMATION_TECHNOLOGY", "answers", default={}) or {}
    upstream_raw = it_cat.get("IT-2_upstream_assets") or it_answers.get("IT-2_upstream_assets") or []
    upstream = [u for u in (upstream_raw if isinstance(upstream_raw, list) else list((upstream_raw or {}).values())) if isinstance(u, dict)]
    providers = _normalize_provider_list(it_cat.get("IT-1_service_providers") or it_answers.get("IT-1_service_providers"))
    hosted_resilience = it_cat.get("it_hosted_resilience") or {}
    if not isinstance(hosted_resilience, dict):
        hosted_resilience = {}

    rows = []
    for u in upstream:
        service_id = (u.get("service_id") or "").strip()
        service_other = (u.get("service_other") or "").strip()
        service_label = service_other if service_other and str(service_id).lower() == "other" else (service_id or "Unknown service")
        provider = (u.get("service_provider") or "").strip() or "Unknown provider"
        primary_fn = _it_primary_function(u)
        dep_id = f"other_{service_other}" if (service_id or "").lower() == "other" else (service_id or "")
        indicators = _hosted_continuity_label(hosted_resilience.get(dep_id))
        rows.append([service_label, provider, primary_fn, indicators])

    for p in providers:
        provider_name = (p.get("provider_name") or "").strip()
        if not provider_name:
            continue
        dep_id = f"provider_{provider_name}"
        indicators = _hosted_continuity_label(hosted_resilience.get(dep_id))
        rows.append(["IT service provider", provider_name, "Internet and data connectivity; outage affects all external access and dependent services", indicators])

    rows = _dedupe_rows(rows)
    if not rows:
        return None

    return {
        "type": "table",
        "title": "Information Technology – Hosted / Upstream Dependencies",
        "headers": ["Dependency / Service", "Provider", "Primary function", "Continuity"],
        "rows": rows,
    }


def build_it_hosted_dependencies_block(assessment_json: dict):
    """
    Uses IT-2_upstream_assets and IT-1_service_providers to render hosted/SaaS dependencies.
    Third column is primary function (service type), not role designation.
    Returns table block dict or None if no rows.
    """
    assessment = assessment_json.get("assessment") or assessment_json
    it_cat = _safe_get(assessment, "categories", "INFORMATION_TECHNOLOGY", default={}) or {}
    it_answers = _safe_get(assessment_json, "sessions", "INFORMATION_TECHNOLOGY", "answers", default={}) or {}
    upstream_raw = it_cat.get("IT-2_upstream_assets") or it_answers.get("IT-2_upstream_assets") or []
    upstream = [u for u in (upstream_raw if isinstance(upstream_raw, list) else list((upstream_raw or {}).values())) if isinstance(u, dict)]
    providers = _normalize_provider_list(it_cat.get("IT-1_service_providers") or it_answers.get("IT-1_service_providers"))

    rows = []

    for u in upstream:
        service_id = (u.get("service_id") or "").strip()
        service_provider = (u.get("service_provider") or "").strip()
        service_other = (u.get("service_other") or "").strip()

        service_label = service_other if service_id in ("other", "OTHER") else service_id
        service_label = service_label or "Unknown service"

        primary_fn = _it_primary_function(u)
        rows.append([service_label, service_provider or "Unknown provider", primary_fn])

    for p in providers:
        provider_name = (p.get("provider_name") or "").strip()
        if provider_name:
            rows.append(["IT service provider", provider_name, "Internet and data connectivity; outage affects all external access and dependent services"])

    rows = _dedupe_rows(rows)

    if not rows:
        return None

    return {
        "type": "table",
        "title": "Hosted / Upstream IT Dependencies",
        "headers": ["Dependency / Service", "Provider", "Primary function"],
        "rows": rows,
    }


def _append_it_hosted_clause_to_vuln_blocks(vuln_blocks_str: str, assessment_json: dict) -> str:
    """
    IT vulnerability narratives are set per-finding in the export: transport exposure uses
    'Internet transport loss would isolate cloud-based services regardless of cloud-side redundancy.'
    Hosted-service exposure uses loss of internet connectivity: the facility relies on externally
    hosted services that require internet; loss of internet connectivity would render those systems
    inaccessible and disrupt operations. Do not reference vendor failure or platform outage.
    This function is a no-op; phrasing is applied by the web export per vuln type.
    """
    return vuln_blocks_str or ""


def build_it_transport_providers_block(assessment_json: dict):
    """
    ISP provider names from curve only: curve_primary_provider and curve_secondary_provider.
    Do not use supply.sources, IT-1, or hosted-service sections. Returns table block or None if no providers.
    """
    assessment = assessment_json.get("assessment") or assessment_json
    isp_names = _get_it_isp_names_from_curve(assessment)
    if not isp_names:
        return None
    rows = []
    if len(isp_names) >= 1:
        rows.append(["Primary Internet Provider", isp_names[0], "", "UNKNOWN", ""])
    if len(isp_names) >= 2:
        rows.append(["Secondary Internet Provider", isp_names[1], "", "UNKNOWN", ""])
    rows = _dedupe_rows(rows)
    if not rows:
        return None
    return {
        "type": "table",
        "title": "IT Transport Providers (Internet / Circuits)",
        "headers": ["Role", "Provider", "Demarcation / Termination", "Independence", "Notes"],
        "rows": rows,
    }


def insert_table_after(doc: Document, paragraph, rows: int, cols: int):
    """
    Insert a new table immediately after the given paragraph. doc.add_table() creates at end;
    we detach and addnext to place after paragraph. Returns the Table.
    """
    table = doc.add_table(rows=rows, cols=cols)
    tbl_el = table._tbl
    tbl_el.getparent().remove(tbl_el)
    paragraph._element.addnext(tbl_el)
    return table


def replace_anchor_with_table_only(doc: Document, anchor: str, block: dict) -> bool:
    """
    Replace a single anchor paragraph with a table only (no heading). Template owns the section heading.
    Finds paragraph with exact text anchor, inserts table in place, removes anchor paragraph.
    Returns True if anchor was found and replaced, False if anchor not present.
    """
    p = find_paragraph_by_exact_text(doc, anchor, body_only=True)
    if p is None:
        return False
    headers = block.get("headers") or []
    rows = block.get("rows") or []
    if not headers and not rows:
        rows = [["Not provided", "—", "—", "—"]]
        headers = ["", "", "", ""]
    num_cols = max(len(headers), max(len(r) for r in rows) if rows else 0)
    if num_cols == 0:
        remove_paragraph(p)
        return True
    tbl = insert_table_after(doc, p, 1 + len(rows), num_cols)
    apply_table_grid_style(tbl)
    _set_tbl_layout_fixed(tbl)
    widths = _infer_table_col_widths(block, num_cols)
    if widths:
        set_table_fixed_widths(tbl, widths)
    for c, h in enumerate(headers):
        if c < num_cols:
            tbl.rows[0].cells[c].text = sanitize_text(str(h))
    set_repeat_header_row(tbl.rows[0])
    set_table_rows_cant_split(tbl)
    for r_idx, row in enumerate(rows):
        for c_idx, cell_val in enumerate(row):
            if c_idx < num_cols and r_idx + 1 < len(tbl.rows):
                tbl.rows[r_idx + 1].cells[c_idx].text = sanitize_text(str(cell_val))
    _normalize_table_paragraph_spacing(tbl)
    remove_paragraph(p)
    return True


def insert_paragraph_after_block(doc: Document, block, text: str = "", style=None) -> DocxParagraph:
    """Insert a paragraph after a block (paragraph or table). Returns the new paragraph."""
    el = block._element if hasattr(block, "_element") else block._tbl
    new_p = doc.add_paragraph(text)
    new_el = new_p._element
    new_el.getparent().remove(new_el)
    el.addnext(new_el)
    if style is not None:
        try:
            new_p.style = style
        except (KeyError, ValueError, AttributeError):
            pass
    return new_p




def set_table_fixed_widths(table, col_widths_inches: list[float]) -> None:
    """Set fixed column widths (inches); disable autofit."""
    table.autofit = False
    for idx, w_in in enumerate(col_widths_inches):
        if idx >= len(table.columns):
            break
        w = Inches(w_in)
        table.columns[idx].width = w
        for cell in table.columns[idx].cells:
            cell.width = w


def set_repeat_header_row(row) -> None:
    """Mark row as repeating header on each page (w:tblHeader)."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trPr.append(tbl_header)


def apply_table_grid_style(table) -> None:
    """Set Word style to Table Grid for visible gridlines."""
    try:
        table.style = "Table Grid"
    except Exception:
        pass


def set_paragraph_keep_with_next(paragraph) -> None:
    """Keep paragraph with next (prevent heading/table split across pages)."""
    try:
        paragraph.paragraph_format.keep_with_next = True
    except Exception:
        pass


# Section-start markers: first paragraph of main content after TOC (stop collecting TOC here)
_TOC_END_MARKERS = (
    "part i",
    "part ii",
    "executive summary",
    "introduction",
    "overview",
)
_TOC_HEADING_MARKERS = ("table of contents", "contents")


def _is_toc_heading(text: str) -> bool:
    t = (text or "").strip().lower()
    return t in _TOC_HEADING_MARKERS or t.startswith("table of contents")


def _is_toc_end(text: str) -> bool:
    t = (text or "").strip().lower()
    if not t:
        return False
    if t.startswith(_TOC_END_MARKERS):
        return True
    # Numbered section heading: "1. Title" or "1 Title"
    if re.match(r"^\d+[.)]\s*\S", t):
        return True
    return False


def ensure_toc_remains_on_one_page(doc: Document) -> None:
    """
    Find the Table of Contents block (heading + following entries) and set keep_with_next
    on every paragraph so the whole block stays on one page.
    """
    toc_paras = []
    in_toc = False
    max_toc_paras = 80
    for block in iter_block_items(doc):
        if not isinstance(block, DocxParagraph):
            if in_toc:
                break
            continue
        text = (block.text or "").strip()
        if not in_toc:
            if _is_toc_heading(text):
                in_toc = True
                toc_paras.append(block)
            continue
        if len(toc_paras) >= max_toc_paras:
            break
        if _is_toc_end(text):
            break
        toc_paras.append(block)
    for p in toc_paras:
        set_paragraph_keep_with_next(p)


def set_table_rows_cant_split(table) -> None:
    """Add w:cantSplit to each row so rows don't split across pages."""
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(OxmlElement("w:cantSplit"))


def set_table_header_repeat(table, header_row_index: int = 0) -> None:
    """Mark header row as repeating on each page."""
    if 0 <= header_row_index < len(table.rows):
        set_repeat_header_row(table.rows[header_row_index])


# --- B) Hotel Fact Sheet table bug: remove unexpected table under "EXECUTIVE SUMMARY" ---


def remove_unexpected_exec_summary_table(doc: Document) -> None:
    """
    If the paragraph "EXECUTIVE SUMMARY" is immediately followed by a Table, remove that table.
    This fixes the bug where an unwanted table appears under the heading.
    """
    blocks = list(iter_block_items(doc))
    for i, block in enumerate(blocks):
        if isinstance(block, DocxParagraph) and (block.text or "").strip() == "EXECUTIVE SUMMARY":
            if i + 1 < len(blocks) and isinstance(blocks[i + 1], DocxTable):
                remove_table(blocks[i + 1])
            break


def remove_legacy_executive_duplicate_paragraphs(doc: Document) -> None:
    """
    Remove legacy executive copy that duplicates snapshot content.
    Applies only when snapshot mode is active.
    """
    duplicate_prefixes = (
        "this assessment documents the operational dependency profile",
    )
    duplicate_headings = {
        "key risk drivers",
    }
    for p in list(doc.paragraphs):
        text = (p.text or "").strip()
        if not text:
            continue
        lower = text.lower()
        if any(lower.startswith(prefix) for prefix in duplicate_prefixes):
            remove_paragraph(p)
            continue
        if lower in duplicate_headings:
            remove_paragraph(p)


# --- C) Chart insertion at anchor (body-only; remove anchor para) ---

# Word insertion: fixed width+height (2.82:1) to preserve Excel ratio
# 6.0" fits within letter (8.5") with 0.75" margins (7" usable); avoids overflow
CHART_W_INCHES = 6.0
CHART_H_INCHES = 6.0 / 2.82  # ≈ 2.13
CHART_SPACING_PT = 6
# Chart placement: Section C — Heading 3 (12pt before, 6pt after), then image only; no caption, no trailing spacer.
CHART_HEADING_SPACE_BEFORE_PT = 12
CHART_HEADING_SPACE_AFTER_PT = 6
CHART_IMAGE_SPACING_PT = 6

# Exact sector names for Section C headings (Heading 3)
SECTION_C_SECTOR_HEADINGS = {
    "ELECTRIC_POWER": "ELECTRIC POWER",
    "COMMUNICATIONS": "COMMUNICATIONS",
    "INFORMATION_TECHNOLOGY": "INFORMATION TECHNOLOGY",
    "WATER": "WATER",
    "WASTEWATER": "WASTEWATER",
}


def insert_charts_at_anchors(doc: Document, chart_paths: dict[str, Path]) -> None:
    """
    Insert chart image at each [[CHART_*]] anchor. Heading 3 (exact sector name) above image only; no caption, no spacer.
    Hard-fail if any chart file is missing or anchor count != 1.
    """
    for code in CHART_CATEGORIES:
        anchor = f"[[CHART_{code}]]"
        path = chart_paths.get(code)
        if path is None:
            raise RuntimeError(f"Chart path not set for sector: {code}")
        path = Path(path)
        if not path.exists():
            raise RuntimeError(f"Chart missing: {code} expected at {path}")
        candidates = list(find_anchor_paragraph_exact(doc, anchor))
        if len(candidates) != 1:
            raise RuntimeError(
                f"Chart anchor must appear exactly once: {anchor} (found {len(candidates)} occurrences)"
            )
        sector_heading = SECTION_C_SECTOR_HEADINGS.get(code, code)
        insert_chart_at_anchor(doc, anchor, path, sector_label=sector_heading)


def insert_chart_at_anchor(
    doc: Document,
    anchor_text: str,
    png_path: Path | str,
    width_in: float | None = None,
    height_in: float | None = None,
    caption: str | None = None,
    sector_label: str | None = None,
) -> None:
    """
    Find anchor paragraph in BODY only (exact match). Insert Heading 3 (sector label) above image, then image.
    No caption, no trailing spacer. Heading: keep_with_next=True, space_before=12pt, space_after=6pt.
    """
    w = width_in if width_in is not None else CHART_W_INCHES
    h = height_in if height_in is not None else CHART_H_INCHES
    p = find_paragraph_by_exact_text(doc, anchor_text, body_only=True)
    if p is None:
        raise RuntimeError(f"Chart anchor not found: {anchor_text}")
    insert_after = p
    if sector_label:
        label_para = insert_paragraph_after(insert_after, sanitize_text(sector_label))
        try:
            label_para.style = "Heading 3"
        except Exception:
            pass
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            label_para.paragraph_format.space_before = Pt(CHART_HEADING_SPACE_BEFORE_PT)
            label_para.paragraph_format.space_after = Pt(CHART_HEADING_SPACE_AFTER_PT)
        except Exception:
            pass
        set_paragraph_keep_with_next(label_para)
        insert_after = label_para
    img_para = insert_paragraph_after(insert_after, "")
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        img_para.paragraph_format.space_before = Pt(CHART_IMAGE_SPACING_PT)
        img_para.paragraph_format.space_after = Pt(CHART_IMAGE_SPACING_PT)
    except Exception:
        pass
    set_paragraph_keep_with_next(img_para)
    run = img_para.add_run()
    run.add_picture(str(png_path), width=Inches(w), height=Inches(h))
    if caption:
        caption_para = insert_paragraph_after(img_para, sanitize_text(caption))
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            caption_para.paragraph_format.space_before = Pt(CHART_IMAGE_SPACING_PT)
            caption_para.paragraph_format.space_after = Pt(CHART_IMAGE_SPACING_PT)
        except Exception:
            pass
        set_paragraph_keep_with_next(caption_para)
    remove_paragraph(p)


def insert_chart_after_paragraph(
    doc: Document,
    after_para,
    png_path: Path | str,
    caption: str | None = None,
) -> object:
    """
    Insert chart image (and optional caption) after the given paragraph.
    Order: image first, then caption (chart appears between title and caption).
    Returns the last inserted paragraph (for chaining).
    """
    w, h = CHART_W_INCHES, CHART_H_INCHES
    insert_after = after_para
    # Create paragraph at doc end, add picture (while in doc), then move to position
    img_para = doc.add_paragraph("")
    run = img_para.add_run()
    run.add_picture(str(png_path), width=Inches(w), height=Inches(h))
    img_para._element.getparent().remove(img_para._element)
    insert_after._element.addnext(img_para._element)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        img_para.paragraph_format.space_before = Pt(CHART_SPACING_PT)
        img_para.paragraph_format.space_after = Pt(CHART_SPACING_PT)
    except Exception:
        pass
    set_paragraph_keep_with_next(img_para)
    insert_after = img_para
    if caption:
        caption_para = insert_paragraph_after(insert_after, sanitize_text(caption))
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            caption_para.paragraph_format.space_before = Pt(CHART_SPACING_PT)
            caption_para.paragraph_format.space_after = Pt(CHART_SPACING_PT)
        except Exception:
            pass
        set_paragraph_keep_with_next(caption_para)
        insert_after = caption_para
    spacer = insert_paragraph_after(insert_after, "")
    return spacer


# --- D) Narrative blanks: fill by proximity to chart anchor (5 values per category) ---


# When mitigated loss is unknown, replace the 5th-blank clause with this (no "Not identified % loss").
MITIGATED_UNKNOWN_CLAUSE = "may reduce operational loss; the mitigated loss percentage was not confirmed."


def replace_nth_underscore_blanks(
    paragraph, replacements_list: list[str], mitigated_unknown_clause: bool = False
) -> None:
    """
    Replace underscore blanks (r"_{3,}") in paragraph sequentially with replacements_list.
    If mitigated_unknown_clause is True, the 5th blank's clause (e.g. "only suffer a  ____ % loss")
    is replaced entirely with MITIGATED_UNKNOWN_CLAUSE instead of "not confirmed".
    """
    full_text = paragraph.text or ""
    blanks = list(UNDERSCORE_RE.finditer(full_text))
    if len(blanks) != len(replacements_list):
        raise RuntimeError(
            f"Template drift: paragraph has {len(blanks)} underscore blanks but {len(replacements_list)} replacements provided."
        )
    # Build (start, end, replacement) for each; 5th blank can be whole-clause replacement
    clause_re = re.compile(r"only\s+suffer\s+a\s+_{3,}\s*%\s*loss", re.IGNORECASE)
    replacements: list[tuple[int, int, str]] = []
    for i in range(len(blanks)):
        m = blanks[i]
        if i == 4 and mitigated_unknown_clause:
            match = clause_re.search(full_text)
            if match:
                replacements.append((match.start(), match.end(), MITIGATED_UNKNOWN_CLAUSE))
            else:
                val = replacements_list[i] if i < len(replacements_list) else NOT_CONFIRMED
                replacements.append((m.start(), m.end(), val or NOT_CONFIRMED))
        else:
            val = replacements_list[i] if i < len(replacements_list) else NOT_CONFIRMED
            replacements.append((m.start(), m.end(), val or NOT_CONFIRMED))
    # Apply from end so indices stay valid
    for start, end, repl in sorted(replacements, key=lambda x: -x[0]):
        full_text = full_text[:start] + repl + full_text[end:]
    for r in paragraph.runs:
        r.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = full_text
    else:
        paragraph.add_run(full_text)


def _round_hours_for_narrative(h: float | None) -> str:
    """Round to nearest 0.25; integer if whole. Never '~0'; use '0'."""
    if h is None or (isinstance(h, (int, float)) and h != h):
        return "not confirmed"
    v = max(0, min(96, float(h)))
    if v == 0:
        return "0"
    if v == int(v):
        return str(int(v))
    return str(round(v * 4) / 4)


def _round_pct_for_narrative(p: float | None) -> str:
    """Whole percent, clamp 0-100."""
    if p is None or (isinstance(p, (int, float)) and p != p):
        return "not confirmed"
    return str(round(max(0, min(100, float(p)))))


def build_sector_narrative(sector_label: str, metrics: dict, tags: dict) -> str:
    """
    Deterministic 3-sentence sector narrative. Dependency-focused, no doctrine.
    C3: No "~not confirmed%" or disjointed alt_loss clauses. Omit alt_loss when missing;
    when alt_loss >= loss - 2, say alternate does not materially reduce loss.
    D4: Never output "~not confirmed%" or "not confirmed%" - use "loss not quantified" when missing.
    """
    requires = metrics.get("requires_service", True)
    if not requires:
        svc = sector_label.lower()
        return f"{sector_label}: This facility does not require {svc} service for normal operations."
    tti_str = _round_hours_for_narrative(metrics.get("tti"))
    loss_str = _round_pct_for_narrative(metrics.get("loss"))
    if loss_str == "not confirmed":
        loss_str = "loss not quantified"
    alt_present = metrics.get("alt_present") is True
    alt_sust = metrics.get("alt_sust")
    alt_loss = metrics.get("alt_loss")
    provider = "confirmed" if tags.get("provider_confirmed") else "not confirmed"
    redund = (tags.get("redund") or "weak").lower()
    corridor = "co-located corridor exposure" if tags.get("corridor_colocated") else ""
    priority = "established" if tags.get("priority_restore_yes") else "not established"

    loss_clause = f"~{loss_str}% functional loss" if loss_str != "loss not quantified" else "functional loss not quantified"
    s1 = f"{sector_label} reaches severe impact in {tti_str} hours with {loss_clause} when alternate capability is absent."
    ra_mode = metrics.get("redundancy_mode") or "UNKNOWN"
    ra_delay_min = metrics.get("activation_delay_min")
    ra_documented = metrics.get("documented_and_tested")

    if alt_present:
        sust_str = _round_hours_for_narrative(alt_sust)
        alt_loss_str = _round_pct_for_narrative(alt_loss) if alt_loss is not None else None

        if ra_mode == "AUTOMATIC":
            s2 = "Alternate capability activates automatically."
            if alt_sust is not None and isinstance(alt_sust, (int, float)) and float(alt_sust) < 96:
                s2 += f" Alternate sustainment is ~{sust_str} hours."
            elif alt_sust is not None:
                s2 += f" Alternate sustainment is ~{sust_str} hours."
            if alt_loss_str and alt_loss_str != "not confirmed":
                s2 += f" Modeled functional loss with backup is ~{alt_loss_str}%."
        elif ra_mode in ("MANUAL_ONSITE", "MANUAL_REMOTE", "VENDOR_REQUIRED"):
            init_label = "on-site" if ra_mode == "MANUAL_ONSITE" else "remote" if ra_mode == "MANUAL_REMOTE" else "vendor"
            s2 = f"Alternate capability uses {init_label} initiation."
            if ra_delay_min is not None and isinstance(ra_delay_min, (int, float)) and ra_delay_min >= 0:
                if ra_delay_min >= 60:
                    delay_str = f"~{int(ra_delay_min / 60)} hours"
                else:
                    delay_str = f"~{int(ra_delay_min)} minutes"
                s2 += f" Activation delay is {delay_str}."
            else:
                s2 += " Activation delay is not provided."
            s2 += f" Alternate capability sustainment is ~{sust_str} hours."
            if alt_loss_str and alt_loss_str != "not confirmed":
                s2 += f" Functional loss with backup ~{alt_loss_str}% (where modeled)."
        else:
            s2 = "Multiple connections or alternate capability are present; route independence and failover mechanism are not confirmed."
            s2 += f" Alternate sustainment is ~{sust_str} hours."
            if alt_loss_str and alt_loss_str != "not confirmed":
                s2 += f" Functional loss with backup ~{alt_loss_str}% (where modeled)."

        if ra_documented is False:
            s2 += " Activation process is not confirmed as tested."
    else:
        s2 = "No alternate capability is available for extended outage scenarios."
    parts = [f"Structural posture is {redund}, provider detail status is {provider}, restoration priority is {priority}."]
    if corridor:
        parts.insert(0, corridor)
    s3 = " ".join(parts)
    return f"{s1} {s2} {s3}"


def _validate_chart_points_for_export(assessment: dict) -> None:
    """
    B1: Fail export if any sector lacks valid curve points (0..96, capacity 0..100).
    COMMUNICATIONS uses PACE model; other sectors use build_curve.
    """
    categories = assessment.get("categories") or {}
    for code in CHART_CATEGORIES:
        inp = categories.get(code) or {}
        if code == "COMMUNICATIONS":
            pace = build_pace_model_from_comm(inp)
            if pace.get("enabled") and pace.get("layers"):
                primary = (pace.get("layers") or {}).get("PRIMARY") or {}
                curve = primary.get("curve") or []
                if not curve or (curve and curve[-1].get("t_hours", 0) < 96):
                    print("ERROR: Communications PACE curve does not cover 0..96 hours", file=sys.stderr)
                    sys.exit(1)
            continue
        points = build_curve(inp)
        if not points or len(points) < 2:
            print(f"ERROR: Missing curve points for {CATEGORY_DISPLAY.get(code, code)}", file=sys.stderr)
            sys.exit(1)
        pts = _ensure_curve_complete(points, horizon=96)
        if pts[-1]["t_hours"] < 96:
            print(f"ERROR: Curve for {CATEGORY_DISPLAY.get(code, code)} does not cover 0..96 hours", file=sys.stderr)
            sys.exit(1)
        for p in pts:
            if not (0 <= p.get("capacity_without_backup", 0) <= 100):
                print(f"ERROR: Invalid capacity in curve for {CATEGORY_DISPLAY.get(code, code)}", file=sys.stderr)
                sys.exit(1)


def _validate_sector_metrics_for_export(assessment: dict) -> None:
    """
    Part D: Fail export if any sector with requires_service is missing TTI, LOSS, REC.
    If ALT_PRESENT, ALT_SUST must be present.
    COMMUNICATIONS: when PACE data exists (comm_pace_* viable), skip legacy curve validation.
    """
    categories = assessment.get("categories") or {}
    errors = []
    for code in CHART_CATEGORIES:
        inp = categories.get(code) or {}
        req = inp.get("requires_service") or inp.get("curve_requires_service")
        if not req:
            continue
        display = CATEGORY_DISPLAY.get(code, code)
        if code == "COMMUNICATIONS":
            has_pace = any(
                _is_layer_viable(inp.get(k) or {})
                for k in ("comm_pace_P", "comm_pace_A", "comm_pace_C", "comm_pace_E")
            )
            if has_pace:
                continue
        tti = inp.get("time_to_impact_hours") if inp.get("time_to_impact_hours") is not None else inp.get("curve_time_to_impact_hours")
        loss_frac = inp.get("loss_fraction_no_backup") if inp.get("loss_fraction_no_backup") is not None else inp.get("curve_loss_fraction_no_backup")
        rec = inp.get("recovery_time_hours") if inp.get("recovery_time_hours") is not None else inp.get("curve_recovery_time_hours")
        has_backup = _effective_has_backup(inp) or inp.get("curve_backup_available") in (True, "yes", "Yes")
        alt_sust = inp.get("backup_duration_hours") if inp.get("backup_duration_hours") is not None else inp.get("curve_backup_duration_hours")
        if tti is None or (isinstance(tti, (int, float)) and tti != tti):
            errors.append(f"{display}: time_to_impact_hours missing or invalid")
        if loss_frac is None and (inp.get("loss_fraction_no_backup") is None and inp.get("curve_loss_fraction_no_backup") is None):
            errors.append(f"{display}: loss_fraction_no_backup missing")
        if rec is None or (isinstance(rec, (int, float)) and rec != rec):
            errors.append(f"{display}: recovery_time_hours missing or invalid")
        if has_backup and (alt_sust is None or (isinstance(alt_sust, (int, float)) and alt_sust != alt_sust)):
            if code != "INFORMATION_TECHNOLOGY":
                errors.append(f"{display}: backup_duration_hours missing (alternate present)")
    if errors:
        msg = "Export blocked: missing required sector metrics.\n  " + "\n  ".join(errors[:10])
        if len(errors) > 10:
            msg += f"\n  ... and {len(errors) - 10} more"
        print(f"ERROR: {msg}", file=sys.stderr)
        sys.exit(1)


def _metrics_and_tags_for_sector(code: str, inp: dict) -> tuple[dict, dict]:
    """Extract metrics and tags for build_sector_narrative from category input."""
    tti = inp.get("time_to_impact_hours") or inp.get("curve_time_to_impact_hours")
    loss_frac = inp.get("loss_fraction_no_backup") or inp.get("curve_loss_fraction_no_backup")
    loss = round(100 * (loss_frac or 0)) if loss_frac is not None else None
    has_backup = _effective_has_backup(inp) or inp.get("curve_backup_available") in (True, "yes", "Yes")
    alt_sust = inp.get("backup_duration_hours") or inp.get("curve_backup_duration_hours") if has_backup else None
    loss_with = inp.get("loss_fraction_with_backup") or inp.get("curve_loss_fraction_with_backup")
    alt_loss = (1 - (loss_with or 0)) * 100 if loss_with is not None else None
    rec = inp.get("recovery_time_hours") or inp.get("curve_recovery_time_hours")
    supply = inp.get("supply") or {}
    sources = (supply if isinstance(supply, dict) else {}).get("sources") or []
    provider_confirmed = bool(sources and any((s or {}).get("provider_name") or (s or {}).get("provider") for s in sources))
    agreements = inp.get("agreements") or {}
    priority_yes = agreements.get("has_sla") is True or agreements.get("has_pra") is True
    sources_str = (inp.get("sources") or "").lower() if isinstance(inp.get("sources"), str) else ""
    redund = "strong" if "independent" in sources_str else ("partial" if "2+" in sources_str else "weak")
    corridor_colocated = False
    if code == "WATER" and inp.get("W_Q4_collocated_corridor") == "yes":
        corridor_colocated = True
    elif code == "WASTEWATER" and inp.get("WW_Q4_collocated_corridor") == "yes":
        corridor_colocated = True
    elif code in ("ELECTRIC_POWER", "COMMUNICATIONS", "INFORMATION_TECHNOLOGY"):
        raw = inp.get("E-4_service_connections") if code == "ELECTRIC_POWER" else inp.get("CO-4_service_connections") if code == "COMMUNICATIONS" else inp.get("IT-4_service_connections")
        for s in (raw or []):
            if isinstance(s, dict) and (s.get("shared_corridor_with_other_utilities") or s.get("shared_corridor")) == "yes":
                corridor_colocated = True
                break
    req = inp.get("requires_service") or inp.get("curve_requires_service")
    redundancy_activation = inp.get("redundancy_activation") or {}
    if isinstance(redundancy_activation, dict):
        ra_mode = redundancy_activation.get("mode") or "UNKNOWN"
        ra_delay_min = redundancy_activation.get("activation_delay_min")
        ra_documented = redundancy_activation.get("documented_and_tested")
    else:
        ra_mode = "UNKNOWN"
        ra_delay_min = None
        ra_documented = None
    metrics = {
        "tti": tti,
        "loss": loss,
        "alt_present": has_backup,
        "alt_sust": alt_sust,
        "alt_loss": alt_loss,
        "rec": rec,
        "requires_service": bool(req),
        "redundancy_mode": ra_mode,
        "activation_delay_min": ra_delay_min,
        "documented_and_tested": ra_documented,
    }
    tags = {"provider_confirmed": provider_confirmed, "redund": redund, "corridor_colocated": corridor_colocated, "priority_restore_yes": priority_yes}
    return metrics, tags


def _cat_values_for_narrative(category_code: str, inp: dict) -> dict:
    """
    Build the 5 values for narrative fill from category input.
    - time_to_impact_hours -> hours (1)
    - loss_fraction_no_backup -> operational capacity % = 100 - loss_fraction_percent (2)
    - backup label (3, 4): "backup generator" for Energy, else "backup system"
    - mitigated loss % (5): loss_fraction_with_backup * 100 or backup capacity derived
    """
    hours = inp.get("time_to_impact_hours")
    hours_str = str(int(hours)) if hours is not None and isinstance(hours, (int, float)) else NOT_CONFIRMED

    loss_frac = inp.get("loss_fraction_no_backup")
    if loss_frac is not None and isinstance(loss_frac, (int, float)):
        capacity_pct = pct((1 - loss_frac) * 100)
        capacity_str = f"{capacity_pct:.1f}"
    else:
        capacity_str = NOT_CONFIRMED

    # Use backup_type verbatim when present; otherwise generic "backup system" (do not default to "generator").
    backup_label = inp.get("backup_type") or "backup system"
    # B4: non-Energy must never say "backup generator" — use neutral phrasing
    NON_ENERGY_CODES = ("COMMUNICATIONS", "INFORMATION_TECHNOLOGY", "WATER", "WASTEWATER")
    if category_code in NON_ENERGY_CODES and "generator" in (backup_label or "").lower():
        if "comms" in category_code.lower() or category_code == "COMMUNICATIONS":
            backup_label = "alternate communications capability"
        elif category_code == "INFORMATION_TECHNOLOGY":
            backup_label = "secondary connectivity or backup capability"
        elif category_code in ("WATER", "WASTEWATER"):
            backup_label = "alternate supply / alternate capability"
        else:
            backup_label = "alternate capability"

    loss_with = inp.get("loss_fraction_with_backup")
    mitigated_unknown = False
    if loss_with is not None and isinstance(loss_with, (int, float)):
        mitigated_pct = pct(loss_with * 100)
        mitigated_str = f"{mitigated_pct:.1f}"
    else:
        cap_backup = inp.get("backup_capacity_pct") or inp.get("backup_capacity_percent")
        if cap_backup is not None and isinstance(cap_backup, (int, float)):
            mitigated_str = f"{pct(100 - cap_backup):.1f}"
        else:
            mitigated_str = NOT_CONFIRMED
            mitigated_unknown = True

    return {
        "hours": sanitize_text(hours_str),
        "capacity_pct": sanitize_text(capacity_str),
        "backup_label": sanitize_text(backup_label or "backup system"),
        "mitigated_loss_pct": sanitize_text(mitigated_str),
        "mitigated_unknown": mitigated_unknown,
    }


def fill_narrative_for_anchor(
    doc: Document, anchor_text: str, cat_values: dict, sector_narrative: str | None = None
) -> None:
    """
    Find chart anchor; inject deterministic 3-sentence sector narrative (Part C).
    If sector_narrative provided, replace first narrative paragraph with it.
    Else fall back to legacy underscore-blank replacement.
    """
    anchor_para = find_paragraph_by_exact_text(doc, anchor_text, body_only=True)
    if anchor_para is None:
        raise RuntimeError(f"Chart anchor not found for narrative: {anchor_text}")
    blocks = list(iter_block_items(doc))
    anchor_el = anchor_para._element
    anchor_index = None
    for i, block in enumerate(blocks):
        if isinstance(block, DocxParagraph) and block._element is anchor_el:
            anchor_index = i
            break
    if anchor_index is None:
        raise RuntimeError(f"Chart anchor paragraph not in body order: {anchor_text}")
    next_chart_anchors = [ANCHORS[k] for k in ("CHART_ELECTRIC_POWER", "CHART_COMMUNICATIONS", "CHART_INFORMATION_TECHNOLOGY", "CHART_WATER", "CHART_WASTEWATER")]
    vofc_section_header = "Vulnerabilities and Options for Consideration"
    replacements = [
        cat_values["hours"],
        cat_values["capacity_pct"],
        cat_values["backup_label"],
        cat_values["backup_label"],
        cat_values["mitigated_loss_pct"],
    ]
    for i in range(anchor_index + 1, len(blocks)):
        block = blocks[i]
        if isinstance(block, DocxParagraph):
            text = (block.text or "").strip()
            if any(a in text for a in next_chart_anchors):
                break
            if vofc_section_header in text:
                break
            if sector_narrative:
                block.clear()
                block.add_run(sanitize_text(sector_narrative))
                return
            if "approximately" in text and UNDERSCORE_RE.search(text):
                replace_nth_underscore_blanks(block, replacements, mitigated_unknown_clause=cat_values.get("mitigated_unknown"))
                return
    if sector_narrative:
        para_after = blocks[anchor_index + 1] if anchor_index + 1 < len(blocks) and isinstance(blocks[anchor_index + 1], DocxParagraph) else None
        if para_after:
            para_after.clear()
            para_after.add_run(sanitize_text(sector_narrative))
            return
    raise RuntimeError(f"No narrative paragraph found after anchor: {anchor_text}")


# --- E) VOFC: lane classification (dependency vs cyber), OFC cleanup, two subsections ---

# Column widths in inches (match template; total ~6.5 for printable width)
VOFC_COL_WIDTHS = [1.4, 2.4, 2.7]
CHOOSE_AN_ITEM = "Choose an item."
GATE_D_MAX_FINDINGS_PER_SECTOR = 4
REPORT_MAX_FINDINGS_PER_SECTOR_MAIN = 6  # Main report (CRITICAL INFRASTRUCTURE) can show up to 6 per sector
GATE_D_TRUNCATION_NOTE = "Additional findings exist but are not displayed in this brief."
# Terms that indicate cyber program content (IT category only); do not change vulnerability text.
CYBER_TERMS_RE = re.compile(
    r"\b(cyber|cybersecurity|incident|NIST|ICS-CERT|US-CERT|ISAC|plan|training|exercise|scan|segmentation)\b",
    re.IGNORECASE,
)
DEPENDENCY_VOFC_HEADING = "Infrastructure Dependency Vulnerabilities and Options for Consideration"  # Legacy; replaced by SECTOR_REPORTS_HEADING
SECTOR_REPORTS_HEADING = "Sector Reports"
ANNEX_OVERVIEW_HEADING = "Annex Overview"
CYBER_VOFC_HEADING = "Cybersecurity Program & Resilience Vulnerabilities and Options for Consideration"


def _ofc_has_source(text: str) -> bool:
    return bool(re.search(r"\(\s*source\s*:[^)]+\)", text or "", flags=re.IGNORECASE))


def _coerce_ofc_with_source(item: dict, ofc_text: str) -> str:
    txt = (ofc_text or "").strip()
    if not txt:
        return ""
    if _ofc_has_source(txt):
        return txt
    source_candidates = [
        str(item.get("source_ref") or "").strip(),
        str(item.get("source_reference") or "").strip(),
        str(item.get("source_registry_id") or "").strip(),
    ]
    source = next((s for s in source_candidates if s), "")
    if source:
        return f"{txt} (Source: {source})"
    # Legacy vofc_collection rows may not carry per-item source fields.
    # Do not hard-fail export here; source handling is validated in canonical/derived block paths.
    return txt


def _vofc_lane(category: str, vulnerability_text: str) -> str:
    """Return 'cyber' if IT category and text contains cyber terms; else 'dependency'. Classification only."""
    vuln = (vulnerability_text or "").strip()
    cat_upper = (category or "").upper()
    if "INFORMATION_TECHNOLOGY" in cat_upper or category == "Information Technology":
        if CYBER_TERMS_RE.search(vuln):
            return "cyber"
    return "dependency"


def _is_physical_security_vofc(vuln: str) -> bool:
    """True if vulnerability text contains physical security keywords (wrong domain for dependency report)."""
    v = (vuln or "").lower()
    return any(kw in v for kw in PHYSICAL_SECURITY_KEYWORDS)


def _normalize_dependency_rows(rows: list[tuple[str, str, str]]) -> None:
    """Normalize category in each row to canonical display name (in-place)."""
    for i in range(len(rows)):
        cat, vuln, ofc = rows[i]
        rows[i] = (_canon_category(cat), vuln, ofc)


def _assert_dependency_only_vofc(rows: list[tuple[str, str, str]]) -> None:
    """Fail fast if any dependency VOFC row has out-of-scope category (e.g. physical security)."""
    bad = []
    for cat, vuln, _ in rows:
        cat_clean = (cat or "").strip()
        if not cat_clean:
            continue
        if cat_clean not in ALLOWED_DEP_CATEGORIES:
            bad.append(cat_clean)
    if bad:
        bad_u = sorted(set(bad))
        raise ValueError(
            f"Out-of-scope VOFC category: {bad_u[0]!r}. "
            f"Allowed: {sorted(ALLOWED_DEP_CATEGORIES)}. "
            "Report is dependency-only: Energy, Communications, IT, Water, Wastewater, Critical Products."
        )


def _cap_vofc_rows_per_sector(
    rows: list[tuple[str, str, str]], max_per: int = GATE_D_MAX_FINDINGS_PER_SECTOR
) -> tuple[list[tuple[str, str, str]], bool]:
    """
    Gate D: Cap at max_per rows per category. Return (capped_rows, truncated).
    """
    by_cat: dict[str, list[tuple[str, str, str]]] = {}
    for r in rows:
        cat = r[0]
        if cat not in by_cat:
            by_cat[cat] = []
        by_cat[cat].append(r)
    capped: list[tuple[str, str, str]] = []
    truncated = False
    for cat in by_cat:
        items = by_cat[cat]
        if len(items) > max_per:
            truncated = True
            items = items[:max_per]
        capped.extend(items)
    return capped, truncated


def _clean_ofcs_and_build_rows(items: list[dict]) -> tuple[list[tuple[str, str, str]], list[tuple[str, str, str]]]:
    """
    Group items by (category, vulnerability). Per group: drop OFC if empty or 'Choose an item.', cap 4;
    fail if 0 OFCs remain (no placeholder fallback). Classify lane (dependency/cyber). Return (dependency_rows, cyber_rows).
    C1: Reject physical security VOFCs (CCTV, IDS, badging, etc.) from dependency rows.
    Gate D: Output is capped at GATE_D_MAX_FINDINGS_PER_SECTOR per category by caller.
    """
    # Group by (category, vulnerability)
    groups: dict[tuple[str, str], list[str]] = {}
    for item in _vofc_items_sorted_for_report(items):
        cat = str(item.get("category") or "").strip()
        vuln = str(item.get("vulnerability") or "").strip()
        ofc = str(item.get("option_for_consideration") or "").strip()
        if not cat and not vuln:
            continue
        # C1: Skip physical security content (wrong domain)
        if _is_physical_security_vofc(vuln):
            continue
        key = (cat, vuln)
        if key not in groups:
            groups[key] = []
        if ofc and ofc != CHOOSE_AN_ITEM:
            ofc = _coerce_ofc_with_source(item, ofc)
            if len(groups[key]) < 4:
                groups[key].append(ofc)
    dependency_rows: list[tuple[str, str, str]] = []
    cyber_rows: list[tuple[str, str, str]] = []
    for (cat, vuln), ofcs in groups.items():
        if not ofcs:
            raise ValueError(
                f'No OFCs available for "{(vuln or "").strip() or "(unknown vulnerability)"}" '
                f'in category "{(cat or "").strip() or "(unknown category)"}".'
            )
        lane = _vofc_lane(cat, vuln)
        for ofc in ofcs:
            if lane == "cyber":
                cyber_rows.append((cat, vuln, ofc))
            else:
                dependency_rows.append((cat, vuln, ofc))
    return dependency_rows, cyber_rows


CATEGORY_DISPLAY_TO_CODE = {v: k for k, v in CATEGORY_DISPLAY.items() if k in CHART_CATEGORIES}


def _vofc_rows_to_energy_and_sections(
    dep_rows: list[tuple[str, str, str]],
) -> tuple[dict | None, list]:
    """
    Convert dependency VOFC rows (cat, vuln, ofc) to energy_dependency and dependency_sections.
    Used when payload has vofc_collection but not energy_dependency/dependency_sections.
    """
    by_cat: dict[str, list[tuple[str, str, str]]] = {}
    for cat, vuln, ofc in dep_rows:
        cat = (cat or "").strip()
        if not cat:
            continue
        code = CATEGORY_DISPLAY_TO_CODE.get(cat) or (
            "ELECTRIC_POWER" if "electric" in cat.lower() or "power" in cat.lower() or "energy" in cat.lower()
            else "COMMUNICATIONS" if "comm" in cat.lower()
            else "INFORMATION_TECHNOLOGY" if "information" in cat.lower() or "technology" in cat.lower() or " it " in f" {cat} "
            else "WATER" if cat.lower() == "water"
            else "WASTEWATER" if "waste" in cat.lower()
            else None
        )
        if code and code in CHART_CATEGORIES:
            key = code
        else:
            key = cat
        if key not in by_cat:
            by_cat[key] = []
        by_cat[key].append((cat, vuln, ofc))
    # Group by (cat, vuln) and collect OFCs
    blocks_by_cat: dict[str, list[dict]] = {}
    for key, rows in by_cat.items():
        vuln_groups: dict[str, list[str]] = {}
        for _, vuln, ofc in rows:
            vuln = (vuln or "").strip()
            if not vuln:
                continue
            if vuln not in vuln_groups:
                vuln_groups[vuln] = []
            if ofc and ofc != CHOOSE_AN_ITEM and ofc not in vuln_groups[vuln]:
                if len(vuln_groups[vuln]) < 4:
                    vuln_groups[vuln].append(ofc)
        for vuln, ofcs in vuln_groups.items():
            if not ofcs:
                raise ValueError(
                    f'No OFCs available for "{vuln}" while building dependency sections for "{key}".'
                )
            block = {"title": vuln, "narrative": "", "ofcs": ofcs, "references": []}
            if key not in blocks_by_cat:
                blocks_by_cat[key] = []
            blocks_by_cat[key].append(block)
    energy_dependency = None
    dependency_sections = []
    for code in CHART_CATEGORIES:
        blocks = blocks_by_cat.get(code) or []
        if code == "ELECTRIC_POWER":
            energy_dependency = {"vulnerability_blocks": blocks}
        else:
            display = CATEGORY_DISPLAY.get(code, code)
            dependency_sections.append({"name": display, "vulnerability_blocks": blocks})
    return energy_dependency, dependency_sections


def build_vofc_table_at_anchor(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    """VOFC table removed; use VULN_NARRATIVE narrative injection."""
    raise RuntimeError("VOFC table removed; use VULN_NARRATIVE narrative injection.")


def inject_text_at_anchor(doc: Document, anchor: str, text: str, body_only: bool = True) -> int:
    """
    Replace anchor with text as multiple paragraphs. Text is split on newlines (single or double).
    Returns number of paragraphs inserted. Anchor must appear exactly once.
    """
    text_trimmed = (text or "").replace("\u00a0", " ").strip()
    candidates = list(find_anchor_paragraph_exact(doc, anchor, body_only=body_only))
    if len(candidates) != 1:
        raise RuntimeError(
            f"Anchor {anchor} must appear exactly once (found {len(candidates)} occurrences)"
        )
    p, _ = candidates[0]
    if paragraph_has_drawing(p):
        _clear_paragraph_text_preserve_drawings(p)
    else:
        p.clear()
    last = p
    count = 0
    blocks = re.split(r"\n\s*\n", text_trimmed) if text_trimmed else []
    for block in blocks:
        for line in block.split("\n"):
            line = line.strip()
            if line:
                last = insert_paragraph_after(last, sanitize_text(line), style="Normal")
                count += 1
    if not paragraph_has_drawing(p):
        remove_paragraph(p)
    return count


# Franklin-based typography for Part II vulnerability blocks (government-doc look)
FRANKLIN_FONT_DEMI = "Franklin Gothic Demi"
FRANKLIN_FONT_MEDIUM = "Franklin Gothic Medium"
FRANKLIN_FONT_BOOK = "Franklin Gothic Book"
ADA_VULN_STYLE_NAMES = (
    "ADA_Vuln_Header",
    "ADA_Vuln_Severity",
    "ADA_Vuln_Meta",
    "ADA_Vuln_Label",
    "ADA_Vuln_Body",
    "ADA_Vuln_Bullets",
    "ADA_Vuln_Numbered",
)


def _get_style_safe(doc: Document, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return None


def _get_paragraph_style_names(doc: Document) -> list[str]:
    """Return list of paragraph style names in the document (exact names as in Word)."""
    names = []
    try:
        for s in doc.styles:
            if getattr(s, "type", None) == WD_STYLE_TYPE.PARAGRAPH:
                names.append(getattr(s, "name", "") or "")
    except Exception:
        pass
    return sorted(names)


REQUIRED_ADA_STYLES = [
    "ADA_Vuln_Header",
    "ADA_Vuln_Severity",
    "ADA_Vuln_Meta",
    "ADA_Vuln_Label",
    "ADA_Vuln_Body",
    "ADA_Vuln_Bullets",
    "ADA_Vuln_Numbered",
]


def _norm_style(x: str | None) -> str:
    """Normalize for style name/styleId comparison (strip only; case-sensitive)."""
    return (x or "").strip()


def preflight_assert_ada_styles(doc: Document) -> None:
    """
    Assert all required ADA_Vuln_* paragraph styles exist in the template.
    Match by style name OR style_id (normalized strip); only PARAGRAPH styles count.
    Call after ensure_ada_vuln_styles() so on-the-fly creation has run.
    """
    present = set()
    for s in doc.styles:
        try:
            if getattr(s, "type", None) != WD_STYLE_TYPE.PARAGRAPH:
                continue
            present.add(_norm_style(getattr(s, "name", None)))
            present.add(_norm_style(getattr(s, "style_id", None)))
        except Exception:
            pass
    missing = [r for r in REQUIRED_ADA_STYLES if _norm_style(r) not in present]
    if missing:
        # Debug: show exactly what python-docx sees (remove once QC passes)
        print("[QC] Available PARAGRAPH styles (name | style_id):")
        for s in doc.styles:
            try:
                if getattr(s, "type", None) == WD_STYLE_TYPE.PARAGRAPH:
                    print(f"[QC] {getattr(s,'name',None)!r} | {getattr(s,'style_id',None)!r}")
            except Exception:
                pass
        available = sorted(present)
        raise ValueError(
            f"Template is missing required paragraph styles: {missing!r}. "
            "Open report template.docx and add the styles with exact names (e.g. ADA_Vuln_Severity, ADA_Vuln_Label), "
            "or run: python apps/reporter/add_ada_vuln_styles_to_template.py. "
            f"Available paragraph styles (name/styleId): {available!r}"
        )


def ensure_ada_vuln_styles(doc: Document) -> None:
    """Ensure ADA_Vuln_* paragraph styles exist (Franklin typography). Idempotent per style."""
    styles = doc.styles
    if _get_style_safe(doc, "ADA_Vuln_Header") is None:
        s = styles.add_style("ADA_Vuln_Header", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = FRANKLIN_FONT_DEMI
        s.font.size = Pt(13)
        s.font.bold = True
        s.font.all_caps = True
        pf = s.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.0
        pf.keep_with_next = True
    if _get_style_safe(doc, "ADA_Vuln_Severity") is None:
        s = styles.add_style("ADA_Vuln_Severity", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = FRANKLIN_FONT_MEDIUM
        s.font.size = Pt(11)
        s.font.bold = True
        s.font.all_caps = True
        pf = s.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.keep_with_next = True
    if _get_style_safe(doc, "ADA_Vuln_Meta") is None:
        s = styles.add_style("ADA_Vuln_Meta", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = FRANKLIN_FONT_BOOK
        s.font.size = Pt(10.5)
        s.font.bold = False
        pf = s.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15
        pf.keep_with_next = True
    if _get_style_safe(doc, "ADA_Vuln_Label") is None:
        s = styles.add_style("ADA_Vuln_Label", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = FRANKLIN_FONT_MEDIUM
        s.font.size = Pt(10.5)
        s.font.bold = True
        pf = s.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.0
    if _get_style_safe(doc, "ADA_Vuln_Body") is None:
        s = styles.add_style("ADA_Vuln_Body", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = FRANKLIN_FONT_BOOK
        s.font.size = Pt(10.5)
        s.font.bold = False
        pf = s.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15
    if _get_style_safe(doc, "ADA_Vuln_Bullets") is None:
        s = styles.add_style("ADA_Vuln_Bullets", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = FRANKLIN_FONT_BOOK
        s.font.size = Pt(10.5)
        pf = s.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.first_line_indent = Inches(-0.15)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.15
    if _get_style_safe(doc, "ADA_Vuln_Numbered") is None:
        s = styles.add_style("ADA_Vuln_Numbered", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = FRANKLIN_FONT_BOOK
        s.font.size = Pt(10.5)
        pf = s.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.first_line_indent = Inches(-0.15)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.15


# Severity line: "HIGH – Structural Concentration" (dash can be -, –, —; case-insensitive)
_SEVERITY_LINE_PATTERN = re.compile(
    r"(?i)^(HIGH|ELEVATED|MODERATE)\s*[\-\u2013\u2014]\s*.+"
)
# Fallback: line that starts with severity word then space/dash/colon and more (ensures P2 always gets Severity style)
_SEVERITY_LINE_FALLBACK = re.compile(
    r"(?i)^(HIGH|ELEVATED|MODERATE)[\s\-–—:].+"
)

# Canonical en-dash for severity line (no smart quotes or em-dash)
_EN_DASH = "\u2013"


def _normalize_severity_line(line: str) -> str:
    """Remove stray Unicode from severity line; produce exactly 'SEVERITY – DRIVER' with U+2013 only."""
    if not line or not line.strip():
        return line
    s = (line or "").strip()
    # Strip smart quotes and stray double-quote that can appear after dash
    s = s.replace("\u201c", "").replace("\u201d", "").replace('"', "")
    # Normalize any dash (em-dash, en-dash, hyphen) to single en-dash with spaces
    s = re.sub(r"[\u2014\u2013\-]+", _EN_DASH, s)
    s = re.sub(r"\s+", " ", s).strip()
    # Ensure single " – " between severity word and rest (collapse multiple dashes/spaces)
    parts = re.split(r"\s*" + re.escape(_EN_DASH) + r"\s*", s, maxsplit=1)
    if len(parts) == 2:
        severity_part = parts[0].strip()
        driver_part = parts[1].strip()
        return f"{severity_part} {_EN_DASH} {driver_part}"
    return s


def _classify_ada_vuln_line(line: str) -> str:
    """Return ADA_Vuln_* style name for a single line of federal block content.
    Severity (P2) MUST be its own paragraph with style ADA_Vuln_Severity; we use two patterns to catch all variants.
    """
    line = (line or "").strip()
    if not line:
        return "ADA_Vuln_Body"
    if re.match(r"^VULNERABILITY \d+\s+(?:HIGH|ELEVATED|MODERATE)\s*[\u2013\-]\s*.+", line, re.IGNORECASE):
        return "ADA_Vuln_Header"
    if re.match(r"^VULNERABILITY \d+\s+(?:HIGH|ELEVATED|MODERATE)\s*$", line, re.IGNORECASE):
        return "ADA_Vuln_Header"
    if re.match(r"^VULNERABILITY \d+$", line, re.IGNORECASE):
        return "ADA_Vuln_Header"
    if _SEVERITY_LINE_PATTERN.match(line) or _SEVERITY_LINE_FALLBACK.match(line):
        return "ADA_Vuln_Severity"
    if line.startswith("Infrastructure Domain:") or line.startswith("Risk Type:"):
        return "ADA_Vuln_Meta"
    if line in (
        "Operational Consequence",
        "Exposure Description",
        "Structural Indicator Profile",
        "Standards Alignment",
        "Options for Consideration",
    ):
        return "ADA_Vuln_Label"
    if re.match(r"^\d+\.\s+", line):
        return "ADA_Vuln_Numbered"
    if line.startswith("- "):
        return "ADA_Vuln_Bullets"
    return "ADA_Vuln_Body"


def _add_severity_paragraph_after(last, raw_line: str, doc: Document) -> "DocxParagraph":
    """
    Create a dedicated paragraph for the severity line. Style is set on the paragraph (not run).
    Uses exact 'ADA_Vuln_Severity'; normalizes raw line to remove stray Unicode (smart quotes, em-dash).
    """
    severity_line = _normalize_severity_line(raw_line)
    p = insert_paragraph_after(last, "", style=None)
    p.clear()
    p.add_run(sanitize_text(severity_line))
    if _get_style_safe(doc, "ADA_Vuln_Severity") is None:
        available = [getattr(s, "name", "") or "" for s in doc.styles]
        raise ValueError(
            "Template is missing required paragraph style ADA_Vuln_Severity. "
            "Open report template.docx and add the style with exact name, or run: python apps/reporter/add_ada_vuln_styles_to_template.py. "
            f"Available styles: {available!r}"
        )
    p.style = doc.styles["ADA_Vuln_Severity"]
    return p


def _add_paragraph_with_style_after(last, text: str, style_name: str, doc: Document):
    """Insert a paragraph after last and set its style (ADA_Vuln_Header/Label/Body/etc.) so Export QC detects it."""
    p = insert_paragraph_after(last, "", style=None)
    p.clear()
    run = p.add_run(sanitize_text(text))
    if style_name == "ADA_Vuln_Header":
        run.bold = True
    if style_name in ADA_VULN_STYLE_NAMES:
        if _get_style_safe(doc, style_name) is None:
            available = _get_paragraph_style_names(doc)
            raise RuntimeError(
                f"Template is missing required paragraph style {style_name!r}. "
                "Open report template.docx and add the styles with exact names, or run: python apps/reporter/add_ada_vuln_styles_to_template.py. "
                f"Available paragraph styles: {available!r}"
            )
        p.style = doc.styles[style_name]
    else:
        try:
            p.style = doc.styles[style_name]
        except (KeyError, ValueError, AttributeError):
            try:
                p.style = "Normal"
            except Exception:
                pass
    return p


def _add_meta_paragraph_after(last, line: str, doc: Document) -> "DocxParagraph":
    """Insert ADA_Vuln_Meta paragraph with bold label run and normal value run (e.g. 'Infrastructure Domain: Electric Power')."""
    style_name = "ADA_Vuln_Meta"
    if ": " in line:
        label, _, value = line.partition(": ")
        label = label.strip() + ": "
        value = value.strip()
    else:
        label, value = "", line
    p = insert_paragraph_after(last, "", style=None)
    p.clear()
    if label:
        r1 = p.add_run(sanitize_text(label))
        r1.bold = True
    if value:
        p.add_run(sanitize_text(value))
    if _get_style_safe(doc, style_name) is None:
        available = _get_paragraph_style_names(doc)
        raise RuntimeError(
            f"Required paragraph style {style_name!r} not in template. "
            f"Available paragraph styles: {available!r}"
        )
    p.style = doc.styles[style_name]
    return p


# Canonical category order for Part II themed findings (matches export)
_PART2_CATEGORY_ORDER = (
    "ELECTRIC_POWER",
    "COMMUNICATIONS",
    "INFORMATION_TECHNOLOGY",
    "WATER",
    "WASTEWATER",
)
_SEVERITY_ORDER = ("HIGH", "ELEVATED", "MODERATE", "LOW")

# Back-compat: when assessment.sessions.*.derived is missing, show this instead of re-deriving
_DERIVED_FINDINGS_MISSING_MSG = (
    "Infrastructure vulnerabilities were not generated for this export (derived findings missing)."
)


def _has_derived_findings(data: dict) -> bool:
    """True if payload has assessment.sessions.<domain>.derived for at least one domain."""
    assessment = data.get("assessment") or {}
    sessions = assessment.get("sessions") or {}
    if not isinstance(sessions, dict):
        return False
    for code in _PART2_CATEGORY_ORDER:
        domain = sessions.get(code) or {}
        if isinstance(domain, dict) and domain.get("derived") is not None:
            return True
    return False


def load_derived_findings_from_payload(data: dict) -> list[dict]:
    """
    Load canonical vulnerabilities + OFCs from assessment.sessions.<domain>.derived.
    Single source of truth: reporter only renders, never recomputes.
    Priority per domain: derived.themedFindings -> derived.vulnerabilities -> [].
    OFC text comes only from derived.ofcs joined by vulnerability_id (order preserved; multiple OFCs joined with \\n).
    Returns list of normalized findings in domain order then array order; each item has id, title, narrative, infrastructure, evidence, ofcText.
    """
    assessment = data.get("assessment") or {}
    sessions = assessment.get("sessions") or {}
    if not isinstance(sessions, dict):
        return []
    out: list[dict] = []
    for domain_code in _PART2_CATEGORY_ORDER:
        domain = sessions.get(domain_code) or {}
        if not isinstance(domain, dict):
            continue
        derived = domain.get("derived") or {}
        if not isinstance(derived, dict):
            continue
        # A) themedFindings preferred; B) vulnerabilities fallback; C) empty
        themed = derived.get("themedFindings")
        if not isinstance(themed, list):
            themed = derived.get("vulnerabilities")
        if not isinstance(themed, list):
            themed = []
        ofcs_raw = derived.get("ofcs") or []
        if not isinstance(ofcs_raw, list):
            ofcs_raw = []
        # Build ofc_by_vuln_id: vulnerability_id -> list of text (order preserved)
        ofc_by_vuln_id: dict[str, list[str]] = {}
        for o in ofcs_raw:
            if not isinstance(o, dict):
                continue
            vid = (o.get("vulnerability_id") or "").strip()
            if not vid:
                continue
            text = (o.get("text") or "").strip()
            ofc_by_vuln_id.setdefault(vid, []).append(text)
        for item in themed:
            if not item or not isinstance(item, dict):
                continue
            fid = (item.get("id") or "").strip()
            title = (item.get("title") or "").strip()
            narrative = (item.get("narrative") or "").strip()
            if not title and "text" in item:
                # Fallback from derived.vulnerabilities: text is "Title. Narrative"
                t = (item.get("text") or "").strip()
                if ". " in t:
                    title, _, rest = t.partition(". ")
                    title = title.strip()
                    narrative = rest.strip()
                else:
                    title = t
            if not title:
                continue
            ofc_texts = ofc_by_vuln_id.get(fid, [])
            ofc_joined = "\n".join(ofc_texts) if ofc_texts else None
            out.append({
                "id": fid or (item.get("title") or "").strip(),
                "title": title,
                "narrative": narrative,
                "infrastructure": domain_code,
                "evidence": item.get("evidence"),
                "ofcText": ofc_joined,
                "severity": (item.get("severity") or "").strip(),
                "references": item.get("references"),
            })
    return out


def _normalize_ofcs(ofcs: list[str]) -> list[str]:
    out: list[str] = []
    for ofc in ofcs:
        t = _normalize_ofc_item(ofc)
        if t:
            out.append(t)
    return out


def _canonical_vuln_block_to_block(c: dict) -> dict:
    """Convert payload canonicalVulnBlocks[] item to render block (title, narrative, severity, ofcs). OFC text used exactly."""
    title = (c.get("title") or "").strip()
    if not title:
        return {}
    narrative = (c.get("narrative") or "").strip()
    ofc_text = c.get("ofcText")
    if ofc_text is None or (isinstance(ofc_text, str) and not ofc_text.strip()):
        ofcs = []
    else:
        lines = [s.strip() for s in (ofc_text.strip().split("\n")) if s.strip()]
        ofcs = _normalize_ofcs(lines if lines else [])
    severity = (c.get("severity") or "").strip()
    if severity and severity.upper() not in _SEVERITY_ORDER:
        severity = ""
    refs = c.get("references") or []
    if not isinstance(refs, list):
        refs = []
    refs = [sanitize_text(str(r).strip()) for r in refs[:10] if str(r).strip()]
    return {
        "title": title,
        "narrative": narrative,
        "severity": severity or "Structural",
        "ofcs": ofcs,
        "references": refs,
    }


def _derived_finding_to_block(f: dict) -> dict:
    """Convert a loaded derived finding to the block shape expected by _render_structured_vulnerabilities / _part2_vuln_to_block."""
    title = (f.get("title") or "").strip()
    if not title:
        return {}
    narrative = (f.get("narrative") or "").strip()
    ofc_text = f.get("ofcText")
    if ofc_text is None or (isinstance(ofc_text, str) and not ofc_text.strip()):
        ofcs = []
    else:
        lines = [s.strip() for s in (ofc_text.strip().split("\n")) if s.strip()]
        ofcs = _normalize_ofcs(lines if lines else [])
    severity = (f.get("severity") or "").strip()
    if severity and severity.upper() not in _SEVERITY_ORDER:
        severity = ""
    refs = f.get("references") or []
    if not isinstance(refs, list):
        refs = []
    refs = [sanitize_text(str(r).strip()) for r in refs[:10] if str(r).strip()]
    return {
        "title": title,
        "narrative": narrative,
        "severity": severity or "Structural",
        "ofcs": ofcs,
        "references": refs,
    }


def get_part2_findings(assessment: dict):
    """
    Part II vulnerability source when report_vm.part2.vulnerabilities is not present (caller handles that first).
    Strict precedence: report_themed_findings from assessment.categories only.
    Returns themed findings list ordered by severity/category/title.
    Raises ValueError when no truth-backed findings are present.
    """
    tf: list = []
    try:
        cats = assessment.get("categories") or {}
        for cat_id in _PART2_CATEGORY_ORDER:
            cat = cats.get(cat_id) or {}
            arr = (cat or {}).get("report_themed_findings") or []
            for item in arr:
                if not item or not isinstance(item, dict):
                    continue
                title = (item.get("title") or "").strip()
                narrative = (item.get("narrative") or "").strip()
                ofc = (item.get("ofcText") or "").strip()
                if title and (narrative or ofc):
                    tf.append({**item, "category": cat_id})
    except Exception:
        tf = []
    if tf:
        # Order: severity (HIGH first), then category order, then title
        def _key(f):
            sev = (f.get("severity") or "").upper()
            sev_idx = _SEVERITY_ORDER.index(sev) if sev in _SEVERITY_ORDER else 99
            cat_idx = _PART2_CATEGORY_ORDER.index(f.get("category") or "") if (f.get("category") or "") in _PART2_CATEGORY_ORDER else 99
            return (sev_idx, cat_idx, (f.get("title") or "").lower())
        return sorted(tf, key=_key)
    raise ValueError("Missing structured Part II findings: report_themed_findings not present.")


def _split_ofc_text(ofc_text: str) -> list[str]:
    """Split ofcText into 1–4 bullets (newline, bullet, or numbered)."""
    if not ofc_text or not isinstance(ofc_text, str):
        return []
    t = ofc_text.replace("\u00a0", " ").strip()
    if not t:
        return []
    # Split on newlines, bullet glyphs, or list markers at line starts.
    parts = re.split(r"\n+|[\u2022\u2023]+|(?:(?<=^)|(?<=\n))\s*(?:[-*]|\d+[.)])\s+", t)
    out = [_normalize_ofc_item(p) for p in parts if p and p.strip()]
    out = [p for p in out if p]
    return out[:4]


def _normalize_ofc_item(text: str) -> str:
    """Normalize a single OFC line; strip list artifacts like leading ':' or '- : '."""
    t = sanitize_vulnerability_text(str(text or "")).replace("\u00a0", " ").strip()
    if not t:
        return ""
    # Repeatedly strip list markers then any dangling leading colon.
    prior = None
    while t != prior:
        prior = t
        t = re.sub(r"^\s*(?:[\u2022\u2023\-*]+|\(?\d+[.)]\)?|[A-Za-z][.)])\s*", "", t).strip()
        t = re.sub(r"^\s*:\s*", "", t).strip()
    return t


def inject_themed_findings_at_anchor(doc: Document, anchor: str, findings: list, body_only: bool = True) -> int:
    """
    Replace anchor with Part II content from report_themed_findings (title, narrative, ofcText).
    Uses ADA_Vuln_* styles. No de-bot banlist (use themed text verbatim; light trim + punctuation).
    """
    ensure_ada_vuln_styles(doc)
    preflight_assert_ada_styles(doc)
    candidates = list(find_anchor_paragraph_exact(doc, anchor, body_only=body_only))
    if len(candidates) != 1:
        raise RuntimeError(
            f"Anchor {anchor} must appear exactly once (found {len(candidates)} occurrences)"
        )
    p, _ = candidates[0]
    if paragraph_has_drawing(p):
        _clear_paragraph_text_preserve_drawings(p)
    else:
        p.clear()
    last = p
    count = 0
    for idx, f in enumerate(findings, start=1):
        title = (f.get("title") or "").strip()
        narrative = (f.get("narrative") or "").strip()
        ofc_text = (f.get("ofcText") or "").strip()
        severity = (f.get("severity") or "").strip()
        if not title:
            continue
        # Ensure period at end of narrative (light normalization)
        if narrative and narrative[-1] not in ".!?":
            narrative = narrative + "."
        # Header: VULNERABILITY N + title (match template doctrine)
        header_text = f"VULNERABILITY {idx} {title}"
        last = _add_paragraph_with_style_after(last, header_text, "ADA_Vuln_Header", doc)
        count += 1
        if severity and severity.upper() in _SEVERITY_ORDER:
            sev_line = f"{severity.upper()} \u2013 Structural"
            last = _add_severity_paragraph_after(last, sev_line, doc)
            count += 1
        # Exposure Description
        last = _add_paragraph_with_style_after(last, "Exposure Description", "ADA_Vuln_Label", doc)
        count += 1
        if narrative:
            last = _add_paragraph_with_style_after(last, narrative, "ADA_Vuln_Body", doc)
            count += 1
        # Options for Consideration
        bullets = _split_ofc_text(ofc_text)
        if bullets:
            last = _add_paragraph_with_style_after(last, "Options for Consideration", "ADA_Vuln_Label", doc)
            count += 1
            for bullet in bullets:
                last = _add_paragraph_with_style_after(last, bullet, "ADA_Vuln_Bullets", doc)
                count += 1
        # Separation handled by paragraph style spacing; avoid explicit blank spacer paragraphs.
    if not paragraph_has_drawing(p):
        remove_paragraph(p)
    return count


def _render_structured_vulnerabilities(doc: Document, anchor: str, vulns: list, body_only: bool = True) -> int:
    """
    Replace anchor with structured Part II vulnerabilities from report_vm.part2.vulnerabilities.
    Single source: same findings as in-app Summary. Uses ADA_Vuln_* styles.
    Header format: "VULNERABILITY {n} {severity} – {title}" (severity appears once).
    """
    ensure_ada_vuln_styles(doc)
    preflight_assert_ada_styles(doc)
    candidates = list(find_anchor_paragraph_exact(doc, anchor, body_only=body_only))
    if len(candidates) != 1:
        raise RuntimeError(
            f"Anchor {anchor} must appear exactly once (found {len(candidates)} occurrences)"
        )
    p, _ = candidates[0]
    if paragraph_has_drawing(p):
        _clear_paragraph_text_preserve_drawings(p)
    else:
        p.clear()
    last = p
    count = 0
    for idx, v in enumerate(vulns or [], start=1):
        block = _part2_vuln_to_block(v)
        if not block or not block.get("title"):
            continue
        title = block["title"]
        narrative = block.get("narrative") or ""
        severity = (block.get("severity") or "").strip()
        ofcs = block.get("ofcs") or []
        refs = block.get("references") or []
        # Header: "VULNERABILITY {n} {severity} – {title}" so severity appears exactly once
        if severity and severity.upper() in _SEVERITY_ORDER:
            header_text = f"VULNERABILITY {idx} {severity.upper()} \u2013 {title}"
        else:
            header_text = f"VULNERABILITY {idx} {title}"
        last = _add_paragraph_with_style_after(last, header_text, "ADA_Vuln_Header", doc)
        count += 1
        # Severity line: risk type only (no duplicate severity word)
        last = _add_severity_paragraph_after(last, "Structural", doc)
        count += 1
        last = _add_paragraph_with_style_after(last, "Exposure Description", "ADA_Vuln_Label", doc)
        count += 1
        if narrative:
            last = _add_paragraph_with_style_after(last, narrative, "ADA_Vuln_Body", doc)
            count += 1
        ofcs_list = (ofcs if isinstance(ofcs, list) else [])
        ofcs_clean = [_normalize_ofc_item(str(x)) for x in ofcs_list if x and str(x).strip()]
        ofcs_clean = [x for x in ofcs_clean if x]
        last = _add_paragraph_with_style_after(last, "Options for Consideration", "ADA_Vuln_Label", doc)
        count += 1
        if ofcs_clean:
            for bullet in ofcs_clean:
                last = _add_paragraph_with_style_after(last, bullet, "ADA_Vuln_Bullets", doc)
                count += 1
        else:
            last = _add_paragraph_with_style_after(last, "Not provided.", "ADA_Vuln_Body", doc)
            count += 1
        if refs:
            last = _add_paragraph_with_style_after(last, "References", "ADA_Vuln_Label", doc)
            count += 1
            for r in refs:
                last = _add_paragraph_with_style_after(last, r, "ADA_Vuln_Bullets", doc)
                count += 1
        last = insert_paragraph_after(last, "", style=None)
        count += 1
    if not paragraph_has_drawing(p):
        remove_paragraph(p)
    return count


def _part2_vuln_to_block(v: dict) -> dict | None:
    """
    Adapter: convert payload.report_vm.part2.vulnerabilities[] item to reporter block dict.
    Expects title, narrative, severity; ofcs (list) or ofcText (string, split to max 4); optional references.
    Returns block with severity, title, narrative, ofcs (list, max 4), references. No inference language.
    """
    raw_title = (v.get("title") or "").strip()
    if not raw_title:
        return None
    title = sanitize_vulnerability_text(raw_title)
    raw_narrative = (v.get("narrative") or "").strip()
    narrative = debot_vulnerability_narrative(sanitize_backend_evidence(raw_narrative)) if raw_narrative else ""
    ofcs = v.get("ofcs")
    if isinstance(ofcs, list):
        ofcs = [sanitize_vulnerability_text(item if isinstance(item, str) else (item.get("text") or item.get("title") or "")) for item in ofcs[:4]]
    else:
        ofcs = _split_ofc_text((v.get("ofcText") or "").strip())
        ofcs = [sanitize_vulnerability_text(x) for x in ofcs if x][:4]
    ofcs = [x for x in ofcs if x]
    ofcs = _normalize_ofcs(ofcs)
    refs = v.get("references") or []
    if not isinstance(refs, list):
        refs = []
    refs = [sanitize_text(str(r).strip()) for r in refs[:10] if str(r).strip()]
    return {
        "severity": (v.get("severity") or "").strip(),
        "title": title,
        "narrative": narrative,
        "ofcs": ofcs,
        "references": refs,
    }


def inject_part2_vulnerabilities_at_anchor(doc: Document, anchor: str, vulns: list, body_only: bool = True) -> int:
    """
    Replace anchor with Part II vulnerabilities from report_vm.part2.vulnerabilities.
    Delegates to _render_structured_vulnerabilities (same structured findings as in-app Summary).
    """
    return _render_structured_vulnerabilities(doc, anchor, vulns, body_only)


# -----------------------------------------------------------------------------
# LEGACY: will be deleted once structured path is enforced. Used only when
# REPORTER_USE_LEGACY_VULN_STRING=1 and payload.report_vm.part2.vulnerabilities is empty.
# -----------------------------------------------------------------------------
# When vulnerability_blocks text has no newlines (flattened), split by these headers to get one paragraph per logical line.
_VULN_FLAT_SPLIT = re.compile(
    r"(?=(?:VULNERABILITY \d+|Operational Consequence|Exposure Description|"
    r"Structural Indicator Profile|Standards Alignment|Options for Consideration|"
    r"Infrastructure Domain:\s*|Risk Type:\s*))",
    re.IGNORECASE,
)


def _vulnerability_block_to_lines(text: str) -> list[str]:
    """
    Split vulnerability block string into lines for one-paragraph-per-line rendering.
    python-docx does not create new paragraphs from \\n inside a run; we must split explicitly.
    """
    if not text or not isinstance(text, str):
        return []
    t = text.replace("\u00a0", " ").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not t:
        return []
    lines = [ln.rstrip() for ln in t.split("\n")]
    # If payload has no newlines we get one long line -> split by federal section headers
    if len(lines) == 1 and len(lines[0]) > 400:
        one = lines[0]
        if "VULNERABILITY" in one and "Exposure Description" in one:
            parts = _VULN_FLAT_SPLIT.split(one)
            lines = [p.strip() for p in parts if p.strip()]
    filtered = _filter_vuln_lines_for_print(lines)
    return _merge_vuln_header_severity_lines(filtered)


def _merge_vuln_header_severity_lines(lines: list[str]) -> list[str]:
    """
    Do not merge header and severity: keep both lines so the reporter emits
    ADA_Vuln_Header (VULNERABILITY N SEVERITY) and ADA_Vuln_Severity (SEVERITY – Driver)
    as separate paragraphs for Export QC.
    """
    return lines


def _filter_vuln_lines_for_print(lines: list[str]) -> list[str]:
    """
    Remove Structural Indicator Profile and Standards Alignment sections from vulnerability
    block lines so DOCX matches in-app narrative (those sections are not user-facing in-app).
    Keeps: VULNERABILITY N, severity, domain, Risk Type, Operational Consequence, Exposure Description, Options for Consideration + OFCs.
    Handles both per-line input and flat-split chunks (chunk may start with section header and contain newlines).
    """
    out: list[str] = []
    skip = False
    for line in lines:
        s = line.strip()
        if re.match(r"^VULNERABILITY \d+", s, re.IGNORECASE):
            skip = False
            out.append(line)
            continue
        if s.startswith("Structural Indicator Profile"):
            skip = True
            continue
        if s.startswith("Standards Alignment"):
            skip = True
            continue
        if s.startswith("Options for Consideration"):
            skip = False
            out.append(line)
            continue
        if skip:
            continue
        out.append(line)
    return out


def _add_vulnerability_block_paragraphs(doc: Document, last_para, text: str):
    """
    Insert vulnerability block content as multiple paragraphs (one per line).
    Explicitly applies ADA_Vuln_* paragraph styles (Header, Severity, Label, Body, Meta, etc.)
    so Export QC detects them. Returns (last_paragraph, count).
    """
    ensure_ada_vuln_styles(doc)
    preflight_assert_ada_styles(doc)
    lines = _vulnerability_block_to_lines(text)
    count = 0
    prev_was_exposure_label = False
    for raw_line in lines:
        line_stripped = raw_line.strip()
        if not line_stripped:
            last_para = insert_paragraph_after(last_para, "", style=None)
            count += 1
            continue
        if prev_was_exposure_label and _classify_ada_vuln_line(line_stripped) == "ADA_Vuln_Body":
            line_stripped = dedupe_sentences(line_stripped)
            if not line_stripped:
                continue
        prev_was_exposure_label = line_stripped == "Exposure Description"
        style_name = _classify_ada_vuln_line(line_stripped)
        if style_name == "ADA_Vuln_Meta":
            last_para = _add_meta_paragraph_after(last_para, line_stripped, doc)
        elif style_name == "ADA_Vuln_Severity":
            last_para = _add_severity_paragraph_after(last_para, line_stripped, doc)
        else:
            last_para = _add_paragraph_with_style_after(last_para, line_stripped, style_name, doc)
        count += 1
    return last_para, count


def inject_vulnerability_blocks_at_anchor(doc: Document, anchor: str, text: str, body_only: bool = True) -> int:
    """
    LEGACY: will be deleted once structured path is enforced.
    Replace anchor with vulnerability block text, applying ADA_Vuln_* styles per line.
    Splits on newlines so each line becomes its own paragraph (python-docx does not create
    paragraphs from \\n inside a run). If text has no newlines, splits by federal section headers.
    Only used when REPORTER_USE_LEGACY_VULN_STRING=1.
    """
    text_trimmed = (text or "").replace("\u00a0", " ").strip()
    candidates = list(find_anchor_paragraph_exact(doc, anchor, body_only=body_only))
    if len(candidates) != 1:
        raise RuntimeError(
            f"Anchor {anchor} must appear exactly once (found {len(candidates)} occurrences)"
        )
    p, _ = candidates[0]
    if paragraph_has_drawing(p):
        _clear_paragraph_text_preserve_drawings(p)
    else:
        p.clear()
    last, count = _add_vulnerability_block_paragraphs(doc, p, text_trimmed)
    if not paragraph_has_drawing(p):
        remove_paragraph(p)
    return count


def inject_vuln_narrative_at_anchor(doc: Document, narrative: str) -> None:
    """
    Replace [[VULN_NARRATIVE]] with pre-rendered narrative string as multiple paragraphs.
    DEPRECATED: Prefer new Part II anchors (STRUCTURAL_PROFILE_SUMMARY, VULNERABILITY_BLOCKS, etc.).
    """
    narrative_trimmed = (narrative or "").replace("\u00a0", " ").strip()
    if not narrative_trimmed:
        raise RuntimeError("VULN_NARRATIVE narrative empty")
    candidates = list(find_anchor_paragraph_exact(doc, VULN_NARRATIVE_ANCHOR, body_only=True))
    if len(candidates) != 1:
        raise RuntimeError(
            f"VULN_NARRATIVE anchor must appear exactly once (found {len(candidates)} occurrences)"
        )
    p, _ = candidates[0]
    if paragraph_has_drawing(p):
        _clear_paragraph_text_preserve_drawings(p)
    else:
        p.clear()
    last = p
    break_para = insert_paragraph_after(last, "")
    run = break_para.add_run()
    run.add_break(WD_BREAK.PAGE)
    last = break_para
    heading_para = insert_paragraph_after(last, DEPENDENCY_VOFC_HEADING)
    try:
        heading_para.style = "Heading 2"
    except Exception:
        pass
    set_paragraph_keep_with_next(heading_para)
    last = insert_paragraph_after(heading_para, "")
    blocks = re.split(r"\n\s*\n", narrative_trimmed)
    for block in blocks:
        for line in block.split("\n"):
            line = line.strip()
            if line:
                last = insert_paragraph_after(last, sanitize_text(line), style="Normal")
    if not paragraph_has_drawing(p):
        remove_paragraph(p)


def render_vofc_docx(doc: Document, anchor: str, vofc: dict) -> None:
    """
    Replace anchor paragraph with VOFC narrative (sectors + vulnerabilities as paragraphs only; no tables).
    Anchor must appear exactly once. vofc: { "sectors": [ { "sector_id", "sector_label", "sector_doctrine", "vulnerabilities": [ { "title", "condition", "impact", "evidence": [ { "kind", "text" } ], "ofcs" } ] } ] }.
    """
    candidates = list(find_anchor_paragraph_exact(doc, anchor))
    if len(candidates) != 1:
        raise RuntimeError(
            f"VULN_NARRATIVE anchor must appear exactly once (found {len(candidates)} occurrences)"
        )
    p_anchor, _ = candidates[0]
    if paragraph_has_drawing(p_anchor):
        _clear_paragraph_text_preserve_drawings(p_anchor)
    else:
        p_anchor.clear()
    last = p_anchor
    # Start vulnerability section on a new page
    break_para = insert_paragraph_after(last, "")
    run = break_para.add_run()
    run.add_break(WD_BREAK.PAGE)
    last = break_para
    sectors = vofc.get("sectors") or []
    for sec in sectors:
        sector_label = sec.get("sector_label") or ""
        sector_doctrine = sec.get("sector_doctrine") or ""
        heading_text = f"SECTOR: {sector_label} ({sector_doctrine})"
        para = insert_paragraph_after(last, "", style="Normal")
        run = para.add_run(sanitize_text(heading_text))
        run.bold = True
        set_paragraph_keep_with_next(para)
        last = para
        last = insert_paragraph_after(last, "")
        for vuln in sec.get("vulnerabilities") or []:
            title = sanitize_text((vuln.get("title") or "").strip() or "Vulnerability")
            para_title = insert_paragraph_after(last, "", style="Normal")
            para_title.add_run(title).bold = True
            last = para_title
            cond = (vuln.get("condition") or "").strip()
            if cond:
                last = insert_paragraph_after(last, "Condition: " + sanitize_text(cond), style="Normal")
            imp = (vuln.get("impact") or "").strip()
            if imp:
                last = insert_paragraph_after(last, "Impact: " + sanitize_text(imp), style="Normal")
            evidence = vuln.get("evidence") or []
            non_ref = [e.get("text", "").strip() for e in evidence if (e.get("kind") or "").strip() != "reference" and (e.get("text") or "").strip()]
            refs = [e.get("text", "").strip() for e in evidence if (e.get("kind") or "").strip() == "reference" and (e.get("text") or "").strip()]
            last = insert_paragraph_after(last, "Evidence/Drivers:", style="Normal")
            for item in non_ref:
                if item:
                    last = insert_paragraph_after(last, sanitize_text(item), style="List Bullet")
            last = insert_paragraph_after(last, "Options for Consideration:", style="Normal")
            for ofc in (vuln.get("ofcs") or [])[:4]:
                t = (ofc or "").strip()
                if t:
                    last = insert_paragraph_after(last, sanitize_text(t), style="List Bullet")
            if refs:
                last = insert_paragraph_after(last, "References: " + ", ".join(sanitize_text(r) for r in refs), style="Normal")
            last = insert_paragraph_after(last, "")
        last = insert_paragraph_after(last, "")
    if not paragraph_has_drawing(p_anchor):
        remove_paragraph(p_anchor)


def build_dependency_vofc_at_anchor(
    doc: Document, rows: list[tuple[str, str, str]], truncation_note: bool = False
) -> None:
    """VOFC table removed; narrative-only export. Use payload.VULN_NARRATIVE and [[VULN_NARRATIVE]]."""
    raise RuntimeError("VOFC table is removed; use VULN_NARRATIVE narrative injection only.")


# Sector page structure: Chart, Chart Synopsis, Vulnerabilities, OFCs, References (per vuln block)
VULNERABILITIES_HEADING = "Vulnerabilities"
CHART_SYNOPSIS_HEADING = "Chart Synopsis"
GATE_D_MAX_FINDINGS_PER_SECTOR = 6


def _get_sector_vulnerability_blocks(
    code: str,
    energy_dependency: dict | None,
    dependency_sections: list,
) -> list[dict]:
    """Get vulnerability_blocks for a sector. Energy from energy_dependency; others from dependency_sections."""
    if code == "ELECTRIC_POWER" and energy_dependency:
        return energy_dependency.get("vulnerability_blocks") or []
    display = CATEGORY_DISPLAY.get(code, code)
    for sec in dependency_sections or []:
        if sec.get("name") and display in str(sec.get("name", "")):
            return sec.get("vulnerability_blocks") or []
    return []


def render_sector_pages_at_anchor(
    doc: Document,
    chart_paths: dict[str, Path],
    assessment: dict,
    energy_dependency: dict | None,
    dependency_sections: list,
) -> None:
    """
    Replace [[VULN_NARRATIVE]] with per-sector pages (deprecated path).
    DEPRECATED: narrative-only export; use inject_sector_narrative_at_infra_anchors instead.
    """
    raise RuntimeError("VOFC table removed; use VULN_NARRATIVE narrative injection.")


# Narrative-only: inject at each INFRA_* anchor; Part II uses [[VULN_NARRATIVE]]. Order matches SECTOR_ORDER.
INFRA_ANCHOR_SECTORS = [
    ("[[INFRA_ENERGY]]", "ELECTRIC_POWER", "Electric Power"),
    ("[[INFRA_COMMS]]", "COMMUNICATIONS", "Communications"),
    ("[[INFRA_IT]]", "INFORMATION_TECHNOLOGY", "Information Technology"),
    ("[[INFRA_WATER]]", "WATER", "Water"),
    ("[[INFRA_WASTEWATER]]", "WASTEWATER", "Wastewater"),
]

# Map anchor -> payload key. Web computes narrative; reporter only injects these.
INFRA_MAP = {
    "[[INFRA_ENERGY]]": "INFRA_ENERGY",
    "[[INFRA_COMMS]]": "INFRA_COMMS",
    "[[INFRA_IT]]": "INFRA_IT",
    "[[INFRA_WATER]]": "INFRA_WATER",
    "[[INFRA_WASTEWATER]]": "INFRA_WASTEWATER",
}


def inject_sector_narrative_at_infra_anchors(
    doc: Document,
    chart_paths: dict[str, Path],
    assessment: dict,
    energy_dependency: dict | None,
    dependency_sections: list,
    payload: dict | None = None,
) -> dict[str, bool]:
    """
    At each [[INFRA_*]] anchor: inject page break + sector content (H1, chart, synopsis from payload, vulnerabilities).
    Web is single narrative authority: we only use payload INFRA_* values; no structured fallback.
    """
    if payload is None:
        payload = {}
    infra_filled: dict[str, bool] = {}
    for idx, (anchor, code, display) in enumerate(INFRA_ANCHOR_SECTORS):
        # Search body and table cells (doc.paragraphs + doc.tables only; no header/footer)
        p = find_paragraph_by_exact_text(doc, anchor, body_only=False)
        if p is None:
            infra_filled[anchor] = False
            continue
        infra_filled[anchor] = True
        if paragraph_has_drawing(p):
            _clear_paragraph_text_preserve_drawings(p)
        else:
            p.clear()
        insert_after = p
        title_para = insert_paragraph_after(insert_after, f"{display.upper()} — Dependency Assessment", style="Heading 1")
        set_paragraph_keep_with_next(title_para)
        insert_after = title_para
        png_path = chart_paths.get(code)
        if png_path and png_path.exists():
            caption = f"Figure {idx + 1}. {display} Dependency Curve"
            insert_after = insert_chart_after_paragraph(doc, insert_after, png_path, caption=caption)
        synopsis_h = insert_paragraph_after(insert_after, CHART_SYNOPSIS_HEADING, style="Heading 2")
        set_paragraph_keep_with_next(synopsis_h)
        # Force use of payload INFRA_* only (no structured fallback)
        key = INFRA_MAP[anchor]
        raw = payload.get(key)
        s = (raw if isinstance(raw, str) else "").replace("\u00a0", " ").strip()
        if not s:
            raise RuntimeError(f"INFRA payload empty: {key}")
        synopsis_para = insert_paragraph_after(synopsis_h, sanitize_text(s), style="Normal")
        insert_after = synopsis_para
        if code == "INFORMATION_TECHNOLOGY":
            it_section = next(
                (s for s in (dependency_sections or []) if "Information Technology" in str(s.get("name") or "")),
                None,
            )
            if it_section:
                external_services = it_section.get("external_services") or []
                cascade_narrative = (it_section.get("cascade_narrative") or "").strip()
                if external_services:
                    ext_h = insert_paragraph_after(insert_after, "External Critical Services", style="Heading 4")
                    set_paragraph_keep_with_next(ext_h)
                    headers = ("Service", "Type", "Supported Functions", "Continuity", "Likely Cascading Effect")
                    tbl = insert_table_after(doc, ext_h, 1 + len(external_services), len(headers))
                    apply_table_grid_style(tbl)
                    for c, hdr in enumerate(headers):
                        tbl.rows[0].cells[c].text = sanitize_text(str(hdr))
                    set_repeat_header_row(tbl.rows[0])
                    set_table_rows_cant_split(tbl)
                    for r_idx, svc in enumerate(external_services):
                        tbl.rows[r_idx + 1].cells[0].text = sanitize_text(svc.get("name") or "")
                        tbl.rows[r_idx + 1].cells[1].text = sanitize_text(svc.get("service_type") or "")
                        funcs = svc.get("supports_functions") or []
                        tbl.rows[r_idx + 1].cells[2].text = sanitize_text(", ".join(funcs) if isinstance(funcs, list) else str(funcs))
                        tbl.rows[r_idx + 1].cells[3].text = sanitize_text(svc.get("resilience") or "")
                        tbl.rows[r_idx + 1].cells[4].text = sanitize_text(svc.get("cascade_effect") or "")
                    insert_after = insert_paragraph_after_block(doc, tbl, "", style="Normal")
                if cascade_narrative:
                    insert_after = insert_paragraph_after(insert_after, sanitize_text(cascade_narrative), style="Normal")
                elif not external_services:
                    insert_after = insert_paragraph_after(insert_after, sanitize_text("No external critical services were identified."), style="Normal")
    return infra_filled


def inject_infra_at_anchors(
    doc: Document,
    payload: dict,
    chart_paths: dict[str, Path],
    assessment: dict,
    energy_dependency: dict | None,
    dependency_sections: list,
) -> None:
    """
    Compatibility injector for Part II sector pages.
    - Preferred path: inject at [[INFRA_*]] anchors when present.
    - Fallback path: inject at [[VULN_NARRATIVE]] when template does not expose INFRA anchors.
    """
    has_infra_anchors = any(_doc_contains_anchor(doc, anchor, body_only=False) for anchor, _, _ in INFRA_ANCHOR_SECTORS)
    if has_infra_anchors:
        inject_sector_narrative_at_infra_anchors(
            doc,
            chart_paths,
            assessment,
            energy_dependency,
            dependency_sections,
            payload=payload,
        )
        return

    def _normalize_sector_code(raw: str | None) -> str:
        s = (raw or "").strip().upper().replace("-", "_").replace(" ", "_")
        alias = {
            "ELECTRIC": "ELECTRIC_POWER",
            "POWER": "ELECTRIC_POWER",
            "COMMS": "COMMUNICATIONS",
            "IT": "INFORMATION_TECHNOLOGY",
            "INFORMATION_TECH": "INFORMATION_TECHNOLOGY",
            "WASTE_WATER": "WASTEWATER",
            "SEWER": "WASTEWATER",
        }
        return alias.get(s, s)

    def _cross_infra_by_sector(payload_obj: dict) -> dict[str, str]:
        out: dict[str, str] = {}
        txt = (payload_obj.get("cross_infra_analysis") or "").strip()
        if not txt:
            return out
        for line in txt.splitlines():
            t = (line or "").strip()
            if not t or ":" not in t:
                continue
            head, rest = t.split(":", 1)
            code = _normalize_sector_code(head)
            if code in CHART_CATEGORIES and rest.strip():
                out[code] = t
        return out

    def _primary_vuln_titles_by_sector(payload_obj: dict) -> dict[str, list[str]]:
        out: dict[str, list[str]] = {c: [] for c in CHART_CATEGORIES}
        canonical = payload_obj.get("canonicalVulnBlocks")
        if isinstance(canonical, list) and canonical:
            for v in canonical:
                code = _normalize_sector_code(v.get("domain"))
                title = (v.get("title") or "").strip()
                if code in out and title:
                    out[code].append(title)
            return out
        if _has_derived_findings(payload_obj):
            for v in load_derived_findings_from_payload(payload_obj):
                code = _normalize_sector_code(v.get("infrastructure"))
                title = (v.get("title") or "").strip()
                if code in out and title:
                    out[code].append(title)
            return out
        vm_v = (((payload_obj.get("report_vm") or {}).get("part2") or {}).get("vulnerabilities") or [])
        if isinstance(vm_v, list) and vm_v:
            for v in vm_v:
                code = _normalize_sector_code(v.get("domain") or v.get("category") or v.get("infrastructure"))
                title = (v.get("title") or "").strip()
                if code in out and title:
                    out[code].append(title)
        return out

    def _assessment_synopsis_for_sector(assessment_obj: dict, sector_code: str, sector_display: str) -> str:
        cats = (assessment_obj or {}).get("categories") or {}
        cat = cats.get(sector_code) or {}
        if not cat:
            return ""
        if sector_code == "COMMUNICATIONS":
            try:
                pace = build_pace_model_from_comm(cat)
                txt = (build_comms_pace_narrative(pace) or "").strip()
                if txt:
                    return txt
            except Exception:
                pass
        try:
            txt = (build_dependency_narrative(cat, sector_display) or "").strip()
            if txt:
                return txt
        except Exception:
            return ""
        return ""

    cross_by_sector = _cross_infra_by_sector(payload)
    primary_titles_by_sector = _primary_vuln_titles_by_sector(payload)
    vm_vulns = (((payload.get("report_vm") or {}).get("part2") or {}).get("vulnerabilities") or [])
    has_primary_vuln_source = any(len(primary_titles_by_sector.get(c, [])) > 0 for c in CHART_CATEGORIES) or bool(
        (isinstance(vm_vulns, list) and len(vm_vulns) > 0) or
        (payload.get("vulnerability_blocks") or "").strip()
    )

    p = find_paragraph_by_exact_text(doc, VULN_NARRATIVE_ANCHOR, body_only=False)
    if p is None:
        return
    if paragraph_has_drawing(p):
        _clear_paragraph_text_preserve_drawings(p)
    else:
        p.clear()

    insert_after = p
    for idx, (anchor, code, display) in enumerate(INFRA_ANCHOR_SECTORS):
        title_para = insert_paragraph_after(
            insert_after, f"{display.upper()} — Dependency Assessment", style="Heading 1"
        )
        set_paragraph_keep_with_next(title_para)
        insert_after = title_para

        png_path = chart_paths.get(code)
        if png_path and png_path.exists():
            caption = f"Figure {idx + 1}. {display} Dependency Curve"
            insert_after = insert_chart_after_paragraph(doc, insert_after, png_path, caption=caption)

        synopsis_h = insert_paragraph_after(insert_after, CHART_SYNOPSIS_HEADING, style="Heading 2")
        set_paragraph_keep_with_next(synopsis_h)
        key = INFRA_MAP[anchor]
        raw_synopsis = payload.get(key)
        synopsis = (raw_synopsis if isinstance(raw_synopsis, str) else "").replace("\u00a0", " ").strip()
        if not synopsis:
            synopsis = _assessment_synopsis_for_sector(assessment, code, display).strip()
        if not synopsis:
            synopsis = (cross_by_sector.get(code) or "").strip()
        if not synopsis:
            raise RuntimeError(
                f"Missing sector synopsis truth data for {code}: expected payload.{key}, cross_infra_analysis sector line, "
                "or derivable assessment category values."
            )
        insert_after = insert_paragraph_after(synopsis_h, sanitize_text(synopsis), style="Normal")



def _is_generic_structural_profile_summary(text: str) -> bool:
    t = (text or "").strip().lower()
    if not t:
        return True
    return (
        "structural drivers are reflected in structural constraints" in t
        and "time-to-impact exposure varies by sector" in t
    )


def _build_structural_profile_summary_from_assessment(assessment: dict) -> str:
    categories = (assessment or {}).get("categories") or {}
    rows: list[str] = []
    immediate: list[str] = []
    near_term: list[str] = []
    extended: list[str] = []
    no_backup: list[str] = []

    for code in CHART_CATEGORIES:
        cat = categories.get(code) or {}
        if not cat.get("requires_service"):
            continue
        display = CATEGORY_DISPLAY.get(code, code)
        tti = cat.get("time_to_impact_hours")
        if tti is None:
            tti = cat.get("curve_time_to_impact_hours")
        tti_text = "unknown" if tti is None else f"{tti:g}h"
        rows.append(f"{display}: severe impact at {tti_text}")
        if isinstance(tti, (int, float)):
            if tti <= 1:
                immediate.append(display)
            elif tti <= 4:
                near_term.append(display)
            else:
                extended.append(display)
        if not _effective_has_backup(cat):
            no_backup.append(display)

    if not rows:
        return "Structural sensitivity could not be summarized because no required sector dependency inputs were provided."

    lines = []
    lines.append("Structural sensitivity is driven by time-to-severe-impact and alternate capability depth across required sectors.")
    lines.append("; ".join(rows) + ".")
    if immediate:
        lines.append("Immediate-impact exposure (<=1h): " + ", ".join(immediate) + ".")
    if near_term:
        lines.append("Near-term impact exposure (>1h to <=4h): " + ", ".join(near_term) + ".")
    if extended:
        lines.append("Extended impact exposure (>4h): " + ", ".join(extended) + ".")
    if no_backup:
        lines.append("Alternate capability not confirmed: " + ", ".join(no_backup) + ".")
    return " ".join(lines)

def _insert_cyber_vofc_block_after_paragraph(
    doc: Document,
    insert_after,
    cyber_rows: list[tuple[str, str, str]],
    truncation_note: bool = False,
) -> None:
    """Insert cyber VOFC heading and table after the given paragraph. Shared by SLA and non-SLA paths."""
    heading_para = insert_paragraph_after(insert_after, CYBER_VOFC_HEADING)
    try:
        heading_para.style = "Heading 2"
    except Exception:
        pass
    set_paragraph_keep_with_next(heading_para)
    num_rows = 1 + max(1, len(cyber_rows) + (1 if truncation_note and cyber_rows else 0))
    table = insert_table_after(doc, heading_para, num_rows, 3)
    apply_table_grid_style(table)
    set_table_fixed_widths(table, VOFC_COL_WIDTHS)
    table.rows[0].cells[0].text = "Category"
    table.rows[0].cells[1].text = "Vulnerability"
    table.rows[0].cells[2].text = "Option for Consideration"
    set_repeat_header_row(table.rows[0])
    set_table_rows_cant_split(table)
    if not cyber_rows:
        table.rows[1].cells[0].text = "Information Technology"
        table.rows[1].cells[1].text = "No vulnerabilities identified"
        table.rows[1].cells[2].text = "Not applicable"
    else:
        for i, (cat, vuln, ofc) in enumerate(cyber_rows, start=1):
            table.rows[i].cells[0].text = cat
            table.rows[i].cells[1].text = vuln
            table.rows[i].cells[2].text = ofc
        if truncation_note:
            table.rows[len(cyber_rows) + 1].cells[0].text = ""
            table.rows[len(cyber_rows) + 1].cells[1].text = ""
            table.rows[len(cyber_rows) + 1].cells[2].text = sanitize_text(GATE_D_TRUNCATION_NOTE)
    insert_paragraph_after(heading_para, "")


def insert_cyber_vofc_block_after_sla_pra(
    doc: Document, cyber_rows: list[tuple[str, str, str]], truncation_note: bool = False
) -> None:
    """
    After the SLA/PRA summary block, insert 'Cybersecurity Program & Resilience...' heading and cyber VOFC table.
    When SLA/PRA module is disabled, the SLA block is not rendered; then insert after [[SLA_PRA_SUMMARY]] anchor
    and remove that anchor paragraph.
    Gate D: If truncation_note, add truncation row.
    """
    blocks = list(iter_block_items(doc))
    insert_after = None
    found_title = False
    for block in blocks:
        if isinstance(block, DocxParagraph):
            text = (block.text or "").strip()
            if SLA_PRA_SUMMARY_TITLE in text:
                found_title = True
                continue
            if found_title and not text:
                insert_after = block
                break
    if insert_after is None:
        p_anchor = find_paragraph_by_exact_text(doc, SLA_PRA_SUMMARY_ANCHOR, body_only=False)
        if p_anchor is not None:
            _insert_cyber_vofc_block_after_paragraph(doc, p_anchor, cyber_rows, truncation_note)
            remove_paragraph(p_anchor)
        return
    _insert_cyber_vofc_block_after_paragraph(doc, insert_after, cyber_rows, truncation_note)


# --- F) Optional hotel fact sheet brief at [[TABLE_SUMMARY]] ---

# Fit within 6.5" printable width (letter 8.5" - 1" margins).
# Rebalanced widths to reduce aggressive word-wrapping in Category/Provider columns.
SUMMARY_COL_WIDTHS = [1.35, 1.0, 0.8, 0.85, 0.85, 1.65]
IT_TRANSPORT_COL_WIDTHS = [1.2, 1.1, 1.6, 1.1, 1.5]
IT_HOSTED_COL_WIDTHS = [1.4, 1.0, 1.7, 1.2, 1.2]
SUMMARY_HEADERS_6 = (
    "Category",
    "Provider Identified",
    "Backup Present",
    "Time to Severe Impact (hrs)",
    "Recovery Time (hrs)",
    "Notes",
)


def _set_cell_no_wrap(cell) -> None:
    """Add w:noWrap to cell so text does not wrap (prevents mid-word splits)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    if tcPr.find(qn("w:noWrap")) is None:
        tcPr.append(OxmlElement("w:noWrap"))


def _infer_table_col_widths(block: dict, num_cols: int) -> list[float] | None:
    """Return table width profile (inches) for known dependency-summary tables."""
    title = (block.get("title") or "").strip().upper()
    headers = [str(h).strip().lower() for h in (block.get("headers") or [])]
    if num_cols == 6 and headers[:2] == ["category", "provider identified"]:
        return SUMMARY_COL_WIDTHS
    if num_cols == 5 and title == "INTERNET TRANSPORT":
        return IT_TRANSPORT_COL_WIDTHS
    if num_cols == 5 and title == "CRITICAL HOSTED SERVICES":
        return IT_HOSTED_COL_WIDTHS
    if num_cols == 5 and headers[:2] == ["role", "provider"]:
        return IT_TRANSPORT_COL_WIDTHS
    if num_cols == 5 and headers[:2] == ["service", "provider"]:
        return IT_HOSTED_COL_WIDTHS
    return None


def _normalize_table_paragraph_spacing(table) -> None:
    """Normalize cell paragraph spacing for dense, readable report tables."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                try:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                except Exception:
                    pass


def build_summary_table_at_anchor(
    doc: Document,
    summary_rows: list[dict],
    anchor: str = DEP_SUMMARY_TABLE_ANCHOR,
    assessment_json: dict | None = None,
    part2: dict | None = None,
) -> None:
    """
    Insert 6-col summary table at anchor (default: in visualization section).
    Fixed layout, noWrap on Category column. Anchor paragraph removed.
    When assessment_json provided, append ISP section, Critical Hosted Services section,
    then optional IT Transport Resilience and IT Hosted Dependencies tables (PART II – DEPENDENCY SUMMARY).
    When part2 provided (report_vm.part2), use part2 internet_transport_rows and critical_hosted_services_rows
    instead of building from assessment (aligns DOCX with in-app summary).
    DEPENDENCY SUMMARY must never be empty: fail if no known anchor is present.
    Tries anchor first, then legacy [[DEP_SUMMARY_TABLE]] for older templates.
    """
    p = find_paragraph_by_exact_text(doc, anchor, body_only=False)
    if p is None and anchor != DEP_SUMMARY_TABLE_ANCHOR:
        p = find_paragraph_by_exact_text(doc, DEP_SUMMARY_TABLE_ANCHOR, body_only=False)
    if p is None:
        raise RuntimeError(
            f"DEPENDENCY SUMMARY anchor not found: {anchor} or {DEP_SUMMARY_TABLE_ANCHOR}. "
            "Export requires one of these in the template."
        )
    if not summary_rows:
        summary_rows = [{"category": "No summary data", "primary_provider": "No", "backup_present": "No", "time_to_severe_impact_hrs": "", "recovery_time_hrs": "", "notes": SUMMARY_NOT_CONFIRMED_TEXT}]
    set_paragraph_keep_with_next(p)
    num_rows = 1 + len(summary_rows)
    table = insert_table_after(doc, p, num_rows, 6)
    apply_table_grid_style(table)
    table.autofit = False
    _set_tbl_layout_fixed(table)
    set_table_fixed_widths(table, SUMMARY_COL_WIDTHS)
    for c, h in enumerate(SUMMARY_HEADERS_6):
        table.rows[0].cells[c].text = h
        if c == 0:
            _set_cell_no_wrap(table.rows[0].cells[0])
    set_repeat_header_row(table.rows[0])
    set_table_rows_cant_split(table)
    acronyms_seen: set[str] = set()
    for r, row_data in enumerate(summary_rows, start=1):
        cat_cell = table.rows[r].cells[0]
        cat_cell.text = sanitize_text(category_no_break(str(row_data.get("category", ""))))
        _set_cell_no_wrap(cat_cell)
        provider_raw = str(row_data.get("primary_provider", "No"))
        provider_cell = expand_acronym(provider_raw, acronyms_seen)
        table.rows[r].cells[1].text = sanitize_text(provider_cell)
        table.rows[r].cells[2].text = sanitize_text(str(row_data.get("backup_present", "No")))
        table.rows[r].cells[3].text = sanitize_text(str(row_data.get("time_to_severe_impact_hrs", "N/A")))
        table.rows[r].cells[4].text = sanitize_text(str(row_data.get("recovery_time_hrs", "N/A")))
        notes_raw = str(row_data.get("notes", SUMMARY_NOT_CONFIRMED_TEXT))
        notes_cell = expand_acronym_in_text(notes_raw, acronyms_seen)
        table.rows[r].cells[5].text = sanitize_text(notes_cell)
    insert_paragraph_after(p, "")
    remove_paragraph(p)

    # DEPENDENCY SUMMARY: Prefer in-place replacement at [[IT_TRANSPORT_SECTION]] / [[IT_HOSTED_SECTION]] (template owns headings).
    # If those anchors are absent (older template), append tables after 6-col table so content is not lost (no headings emitted).
    safe_assessment = assessment_json or {}
    if part2 and isinstance(part2, dict):
        transport_rows = part2.get("internet_transport_rows") or []
        def _transport_cell(r: dict, key: str, use_not_provided_for_empty: bool = False) -> str:
            val = (r.get(key) or "").strip()
            if use_not_provided_for_empty and not val:
                return "Not provided"
            return val or "—"
        transport_block = {
            "type": "table",
            "title": "INTERNET TRANSPORT",
            "headers": ["Role", "Provider", "Demarcation", "Independence", "Notes"],
            "rows": [[r.get("role", ""), r.get("provider", ""), _transport_cell(r, "demarcation", use_not_provided_for_empty=True), _transport_cell(r, "independence", use_not_provided_for_empty=True), _transport_cell(r, "notes", use_not_provided_for_empty=True)] for r in transport_rows],
        }
        hosted_rows = part2.get("critical_hosted_services_rows") or []
        if not hosted_rows:
            hosted_rows = [{"service": "No critical hosted services identified.", "provider": "—", "service_loss_effect": "—", "continuity_strategy": "—", "notes": "—"}]
        hosted_block = {
            "type": "table",
            "title": "CRITICAL HOSTED SERVICES",
            "headers": ["Service", "Provider", "Service Loss", "Continuity Strategy", "Notes"],
            "rows": [[r.get("service", ""), r.get("provider", ""), r.get("service_loss_effect", ""), r.get("continuity_strategy", ""), r.get("notes", "—")] for r in hosted_rows],
        }
    else:
        transport_block = _internet_transport_table(safe_assessment)
        hosted_block = _it_critical_hosted_table(safe_assessment)
    replaced_transport = replace_anchor_with_table_only(doc, IT_TRANSPORT_SECTION_ANCHOR, transport_block)
    replaced_hosted = replace_anchor_with_table_only(doc, IT_HOSTED_SECTION_ANCHOR, hosted_block)
    if not replaced_transport or not replaced_hosted:
        # Fallback: append tables after 6-col table when anchors not in template (omit_title so no duplicate headings).
        insert_after = table
        if not replaced_transport:
            insert_after = _render_table_block_after(doc, transport_block, insert_after, omit_title=True)
        if not replaced_hosted:
            insert_after = _render_table_block_after(doc, hosted_block, insert_after, omit_title=True)


def _normalize_dependency_summary_rows(summary_rows: list[dict] | None) -> list[dict]:
    """
    Ensure dependency summary rows include all expected categories and stable ordering.
    Adds a default Critical Products row when missing (for payloads that provide partial rows).
    """
    rows = list(summary_rows or [])
    if not rows:
        rows = []

    by_key: dict[str, dict] = {}
    for row in rows:
        if not isinstance(row, dict):
            continue
        category = str(row.get("category") or "").strip()
        if not category:
            continue
        key = category.upper().replace(" ", "_")
        by_key[key] = row

    if "CRITICAL_PRODUCTS" not in by_key and "CRITICAL PRODUCTS" not in by_key:
        by_key["CRITICAL_PRODUCTS"] = {
            "category": "Critical Products",
            "primary_provider": "N/A",
            "backup_present": "N/A",
            "time_to_severe_impact_hrs": "N/A",
            "recovery_time_hrs": "N/A",
            "notes": "—",
        }

    order = [
        "ELECTRIC_POWER",
        "COMMUNICATIONS",
        "INFORMATION_TECHNOLOGY",
        "WATER",
        "WASTEWATER",
        "CRITICAL_PRODUCTS",
    ]
    out: list[dict] = []
    for key in order:
        if key in by_key:
            out.append(by_key[key])
    # keep any additional rows at end, preserving input order
    for row in rows:
        if not isinstance(row, dict):
            continue
        key = str(row.get("category") or "").strip().upper().replace(" ", "_")
        if key and key not in order:
            out.append(row)
    return out


def _set_tbl_layout_fixed(table) -> None:
    """Set w:tblLayout type='fixed' for fixed table layout."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblLayout.set(qn("w:type"), "fixed")
        tblPr.append(tblLayout)
    else:
        tblLayout.set(qn("w:type"), "fixed")


# --- SLA/PRA summary narrative (after VOFC section) ---

SLA_PRA_SUMMARY_TITLE = "Service Restoration Reliability Summary"


IT_SCOPE_CLARIFICATION = (
    "Scope Clarification: Information Technology evaluates externally hosted or managed digital services "
    "(SaaS, cloud applications, hosted identity, managed IT providers). Systems owned and operated by the "
    "facility are treated as critical assets and are not assessed as dependencies here. Communications "
    "evaluates carrier-based transport services (ISP circuits, fiber, wireless, satellite)."
)


def _render_table_block_after(
    doc: Document, block: dict, after_block, *, omit_title: bool = False
) -> DocxParagraph:
    """Render a table block (optionally title + table) after the given paragraph or table.
    When omit_title is True, inject only the table; template owns the section heading (e.g. INTERNET TRANSPORT, CRITICAL HOSTED SERVICES).
    Returns last inserted element."""
    title = (block.get("title") or "").strip()
    headers = block.get("headers") or []
    rows = block.get("rows") or []
    if not headers and not rows:
        return after_block
    # Insert after table: use insert_paragraph_after_block; after paragraph: use insert_paragraph_after
    if omit_title:
        # Template owns headings; insert only table (no duplicate heading).
        if hasattr(after_block, "_tbl"):
            ref_para = insert_paragraph_after_block(doc, after_block, "", style="Normal")
        else:
            ref_para = after_block
        set_paragraph_keep_with_next(ref_para)
    else:
        if hasattr(after_block, "_tbl"):
            title_para = insert_paragraph_after_block(doc, after_block, sanitize_text(title), style="Heading 4")
        else:
            title_para = insert_paragraph_after(after_block, sanitize_text(title))
            try:
                title_para.style = "Heading 4"
            except Exception:
                pass
        set_paragraph_keep_with_next(title_para)
        ref_para = title_para
    num_cols = max(len(headers), max(len(r) for r in rows) if rows else 0)
    if num_cols == 0:
        return ref_para
    tbl = insert_table_after(doc, ref_para, 1 + len(rows), num_cols)
    apply_table_grid_style(tbl)
    _set_tbl_layout_fixed(tbl)
    widths = _infer_table_col_widths(block, num_cols)
    if widths:
        set_table_fixed_widths(tbl, widths)
    for c, h in enumerate(headers):
        if c < num_cols:
            tbl.rows[0].cells[c].text = sanitize_text(str(h))
    set_repeat_header_row(tbl.rows[0])
    set_table_rows_cant_split(tbl)
    for r_idx, row in enumerate(rows):
        for c_idx, cell_val in enumerate(row):
            if c_idx < num_cols and r_idx + 1 < len(tbl.rows):
                tbl.rows[r_idx + 1].cells[c_idx].text = sanitize_text(str(cell_val))
    _normalize_table_paragraph_spacing(tbl)
    return insert_paragraph_after_block(doc, tbl, "", style="Normal")


def _render_it_table_block(doc: Document, block: dict, insert_after) -> DocxParagraph:
    """Render a table block (title + table) after the given paragraph. Returns last inserted element."""
    title = (block.get("title") or "").strip()
    headers = block.get("headers") or []
    rows = block.get("rows") or []
    if not headers and not rows:
        return insert_after
    title_para = insert_paragraph_after(insert_after, sanitize_text(title))
    try:
        title_para.style = "Heading 4"
    except Exception:
        pass
    set_paragraph_keep_with_next(title_para)
    num_cols = max(len(headers), max(len(r) for r in rows) if rows else 0)
    if num_cols == 0:
        return title_para
    tbl = insert_table_after(doc, title_para, 1 + len(rows), num_cols)
    apply_table_grid_style(tbl)
    _set_tbl_layout_fixed(tbl)
    widths = _infer_table_col_widths(block, num_cols)
    if widths:
        set_table_fixed_widths(tbl, widths)
    for c, h in enumerate(headers):
        if c < num_cols:
            tbl.rows[0].cells[c].text = sanitize_text(str(h))
    set_repeat_header_row(tbl.rows[0])
    set_table_rows_cant_split(tbl)
    for r_idx, row in enumerate(rows):
        for c_idx, cell_val in enumerate(row):
            if c_idx < num_cols and r_idx + 1 < len(tbl.rows):
                tbl.rows[r_idx + 1].cells[c_idx].text = sanitize_text(str(cell_val))
    _normalize_table_paragraph_spacing(tbl)
    return insert_paragraph_after_block(doc, tbl, "", style="Normal")


def insert_it_scope_clarification(doc: Document) -> None:
    """
    Find the first body paragraph whose text is exactly "Information Technology" (section heading).
    Insert the scope clarification paragraph immediately after it. Inserts only once.
    IT tables are now rendered in DEPENDENCY SUMMARY (after the dependency summary table).
    """
    p = find_paragraph_by_exact_text(doc, "Information Technology", body_only=True)
    if p is None:
        return
    scope_para = insert_paragraph_after(p, IT_SCOPE_CLARIFICATION)
    try:
        scope_para.style = "Normal"
    except Exception:
        pass


def insert_cross_dependency_summary_at_anchor(
    doc: Document, summary: dict | None, modules: list[dict] | None
) -> None:
    """
    Find [[CROSS_DEPENDENCY_SUMMARY]] paragraph; replace with Cross-Dependency Findings section.
    summary: {confirmed_count, top_edges, flags} or None.
    modules: list of module findings for Cross-Dependency Modules.
    """
    p = find_paragraph_by_exact_text(doc, CROSS_DEPENDENCY_SUMMARY_ANCHOR, body_only=False)
    if p is None:
        return
    title_para = insert_paragraph_after(p, "Cross-Dependency Findings")
    try:
        title_para.style = "Heading 2"
    except Exception:
        pass

    if summary and summary.get("confirmed_count", 0) > 0:
        bullets = [f"Confirmed cross-dependencies: {summary.get('confirmed_count', 0)}"]
        bullets.extend(summary.get("top_edges") or [])
        bullets.extend(summary.get("flags") or [])
        for bullet_text in bullets:
            bullet_para = insert_paragraph_after(title_para, bullet_text)
            title_para = bullet_para
            try:
                bullet_para.style = "List Bullet"
            except Exception:
                pass
    else:
        title_para = insert_paragraph_after(title_para, "No cross-dependencies captured.")

    if modules:
        mod_heading = insert_paragraph_after(title_para, "Cross-Dependency Modules")
        try:
            mod_heading.style = "Heading 3"
        except Exception:
            pass
        last_para = mod_heading
        for mod in modules:
            title = mod.get("title") or mod.get("module_code") or "Module"
            mod_title = insert_paragraph_after(last_para, str(title))
            try:
                mod_title.style = "Heading 4"
            except Exception:
                pass
            last_para = mod_title
            summary_sentences = mod.get("summary_sentences") or []
            if summary_sentences:
                summary_text = " ".join([str(s) for s in summary_sentences])
                last_para = insert_paragraph_after(last_para, summary_text)

            vulnerabilities = mod.get("vulnerabilities") or []
            if not vulnerabilities:
                last_para = insert_paragraph_after(last_para, "No significant cross-dependency control gaps identified.")
                continue

            for v in vulnerabilities:
                v_title = v.get("title") or "Vulnerability"
                v_text = v.get("text") or ""
                bullet_para = insert_paragraph_after(last_para, f"{v_title}: {v_text}".strip())
                last_para = bullet_para
                try:
                    bullet_para.style = "List Bullet"
                except Exception:
                    pass
                ofcs = v.get("ofcs") or []
                for ofc in ofcs[:4]:
                    ofc_text = ofc.get("option_for_consideration") or ""
                    if not ofc_text:
                        continue
                    ofc_para = insert_paragraph_after(last_para, f"OFC: {ofc_text}")
                    last_para = ofc_para
                    try:
                        ofc_para.style = "List Bullet"
                    except Exception:
                        pass

    insert_paragraph_after(title_para, "")
    remove_paragraph(p)


def insert_sla_pra_summary_at_anchor(doc: Document, items: list[dict]) -> bool:
    """
    Find [[SLA_PRA_SUMMARY]] paragraph; replace with heading + bullet paragraphs + spacer; remove anchor.
    items: list of {category, routine_outage_text, widespread_disaster_text}. If empty, insert heading + spacer only.
    Returns True if anchor was found and block was inserted.
    """
    p = find_paragraph_by_exact_text(doc, SLA_PRA_SUMMARY_ANCHOR, body_only=False)
    if p is None:
        return False
    title_para = insert_paragraph_after(p, SLA_PRA_SUMMARY_TITLE)
    try:
        title_para.style = "Heading 2"
    except Exception:
        pass
    if items:
        for entry in items:
            cat = entry.get("category") or ""
            routine = entry.get("routine_outage_text") or ""
            widespread = entry.get("widespread_disaster_text") or ""
            bullet_text = f"{cat}: {routine} {widespread}".strip()
            bullet_para = insert_paragraph_after(title_para, bullet_text)
            title_para = bullet_para
            try:
                bullet_para.style = "List Bullet"
            except Exception:
                pass
    insert_paragraph_after(title_para, "")
    remove_paragraph(p)
    return True


INDEPENDENCE_LABELS = {
    "UNKNOWN": "Unknown",
    "SAME_DEMARCATION": "Same demarcation point",
    "DIFFERENT_DEMARCATION_SAME_UPSTREAM": "Different demarcation, same upstream (partial resilience)",
    "DIFFERENT_LOOP_OR_PATH": "Different loop/path (resilient)",
}


def pct(value: float) -> float:
    return round(min(100, max(0, value)) * 10) / 10


def build_curve(category_input: dict, horizon_hours: int = 96, step_hours: int = 3) -> list[dict]:
    """
    Mirror of engine buildCurve: points with t_hours, capacity_without_backup, capacity_with_backup.
    Prescribed Agency math: capacity % = 100 - loss % (so capacity and loss are complements and sum to 100).
    - Without backup: before time_to_impact → 100% capacity; after → (1 - loss_fraction_no_backup) * 100.
    - With backup: from t=0 to backup_duration → (1 - loss_fraction_with_backup) * 100; after backup exhausted → same as without backup.
    """
    inp = dict(category_input) if isinstance(category_input, dict) else {}
    points = []
    for t in range(0, horizon_hours + 1, step_hours):
        if not inp.get("requires_service", False):
            cap_no, cap_with = 100.0, 100.0
        else:
            t_impact = inp.get("time_to_impact_hours", 0) or 0
            loss_no = inp.get("loss_fraction_no_backup", 0) or 0
            # Capacity without backup: 100% before impact, then (100 - loss_no)
            if t < t_impact:
                cap_no = 100.0
            else:
                cap_no = (1 - loss_no) * 100
            if not _effective_has_backup(inp):
                cap_with = cap_no
            else:
                duration_raw = inp.get("backup_duration_hours")
                try:
                    duration = float(duration_raw) if duration_raw is not None else 0.0
                except Exception:
                    duration = 0.0
                loss_with = inp.get("loss_fraction_with_backup", 0) or 0
                # Backup duration >= horizon is treated as indefinite for modeled window.
                if duration >= float(horizon_hours):
                    cap_with = (1 - loss_with) * 100
                # With backup: during backup window use (100 - loss_with); after exhausted, same as without backup.
                elif duration > 0 and t <= duration:
                    cap_with = (1 - loss_with) * 100
                else:
                    cap_with = cap_no
        points.append({
            "t_hours": t,
            "capacity_without_backup": pct(cap_no),
            "capacity_with_backup": pct(cap_with),
        })
    return points


def _effective_has_backup(inp: dict) -> bool:
    """True if has_backup_any or has_backup is True (engine compatibility)."""
    if inp.get("has_backup_any") is True:
        return True
    return inp.get("has_backup", False) is True


def _sources_summary(supply: dict | None) -> str | None:
    """Descriptive only: list explicit provider/service names. Never infer conclusions."""
    if supply is None or not isinstance(supply, dict):
        return None
    sources = supply.get("sources") or []
    if not sources:
        return None
    parts = []
    for s in sources:
        if not isinstance(s, dict):
            continue
        name = (s.get("provider_name") or s.get("service_provider") or s.get("provider") or "").strip()
        if name:
            role = (s.get("role") or "").strip()
            parts.append(f"{name}" + (f" ({role})" if role else ""))
    if not parts:
        return None
    return "; ".join(parts)


def normalize_note_text(x) -> str:
    """Force notes to a clean human string; never numeric-only, never raw int/list."""
    if x is None:
        return ""
    if isinstance(x, (int, float)):
        return ""
    if isinstance(x, (list, tuple)):
        parts = [str(p).strip() for p in x if isinstance(p, (str, int, float)) and str(p).strip()]
        if parts and all(p.isdigit() for p in parts):
            return ""
        return "; ".join(parts)
    s = str(x).replace("\u00a0", " ").strip()
    if s.isdigit():
        return ""
    return s


def _sla_summary(agreements: dict | None) -> str:
    """Legacy: used when sla_reliability_for_report is not provided."""
    if not agreements or not agreements.get("has_sla"):
        return "No"
    h = agreements.get("sla_hours")
    return f"Yes ({h}h)" if h is not None else "Yes (\u2014)"


def _format_sla_cell_stakeholder(entry: dict) -> str:
    """
    Format SLA cell for report (stakeholder tone). No 'gap' wording.
    - sla_assessed false: do not state 'No SLA in place'; show 'SLA not assessed'.
    - sla_assessed true, sla_in_place NO: 'No SLA in place'
    - sla_assessed true, sla_in_place UNKNOWN: 'SLA status unknown'
    - sla_assessed true, sla_in_place YES: 'SLA in place', MTTR-Max if set, summary_text if set.
    """
    if not entry.get("sla_assessed"):
        return "SLA not assessed"
    in_place = entry.get("sla_in_place") or "UNKNOWN"
    if in_place == "NO":
        return "No SLA in place"
    if in_place == "UNKNOWN":
        return "SLA status unknown"
    # YES
    parts = ["SLA in place"]
    mttr = entry.get("mttr_max_hours")
    if mttr is not None and isinstance(mttr, (int, float)) and mttr == mttr:  # exclude NaN
        parts.append(f"MTTR-Max: {int(mttr)} hours")
    summary = (entry.get("summary_text") or "").strip()
    if summary:
        parts.append(summary)
    return " ".join(parts)


def _apply_sla_reliability_to_summary_rows(
    summary_rows: list[dict], sla_reliability_for_report: list[dict]
) -> None:
    """Override notes SLA line in summary rows with stakeholder-safe text from sla_reliability_for_report."""
    if not sla_reliability_for_report:
        return
    label_to_entry = {e.get("topic_label"): e for e in sla_reliability_for_report if e.get("topic_label")}
    for row in summary_rows:
        cat = row.get("category")
        if cat and cat in label_to_entry:
            sla_text = _format_sla_cell_stakeholder(label_to_entry[cat])
            notes = str(row.get("notes") or "").strip()
            notes_no_sla = re.sub(r"\s*SLA: [^.]*\.?", "", notes).strip()
            row["notes"] = f"{notes_no_sla} SLA: {sla_text}".strip() if notes_no_sla else f"SLA: {sla_text}"


def _pra_summary(agreements: dict | None) -> str:
    if not agreements or not agreements.get("has_pra"):
        return "No"
    cat = agreements.get("pra_category")
    if cat is None:
        return "Yes (\u2014)"
    if cat == "OTHER":
        other = (agreements.get("pra_category_other") or "").strip()
        return f"Yes (Other: {other})" if other else "Yes (Other: \u2014)"
    return f"Yes ({cat})"


def _primary_provider_identified(inp: dict, category: str) -> str:
    """Provider Identified = YES if ANY of: E-1_utility_providers[].provider_name, supply.sources[].provider_name,
    IT-1_service_providers[].provider_name, curve_primary_provider. NO only if none exist."""
    if category == "CRITICAL_PRODUCTS":
        return "No"
    supply = inp.get("supply") or {}
    sources = supply.get("sources") or []
    has_in_sources = any(
        (s or {}).get("provider_name") and str((s or {}).get("provider_name", "")).strip()
        or (s or {}).get("service_provider") and str((s or {}).get("service_provider", "")).strip()
        or (s or {}).get("provider") and str((s or {}).get("provider", "")).strip()
        for s in sources
    )
    if has_in_sources:
        return "Yes"
    curve_primary = (inp.get("curve_primary_provider") or "").strip()
    if curve_primary:
        return "Yes"
    if category == "ELECTRIC_POWER":
        providers = inp.get("E-1_utility_providers") or []
        if isinstance(providers, list) and any(
            (p or {}).get("provider_name") and str((p or {}).get("provider_name", "")).strip()
            for p in providers
        ):
            return "Yes"
    if category == "INFORMATION_TECHNOLOGY":
        if (inp.get("IT-1_can_identify_providers") or "").strip().lower() in ("yes", "true", "1"):
            return "Yes"
        it1 = inp.get("IT-1_service_providers") or []
        if isinstance(it1, list) and any(
            (p or {}).get("provider_name") and str((p or {}).get("provider_name", "")).strip()
            for p in it1
        ):
            return "Yes"
        if isinstance(it1, dict):
            for p in (it1 or {}).values():
                if isinstance(p, dict) and (p.get("provider_name") or "").strip():
                    return "Yes"
    return "No"


def _independence_note_suffix(inp: dict, provider: str) -> str:
    """When provider is Yes but all sources have independence UNKNOWN or missing, return suffix for notes."""
    if provider != "Yes":
        return ""
    supply = inp.get("supply") or {}
    sources = supply.get("sources") or []
    if not sources:
        return ""
    all_unknown = all(
        (s or {}).get("independence") in (None, "", "UNKNOWN")
        for s in sources
    )
    if not all_unknown:
        return ""
    return " Provider identified; independence not provided."


def build_summary(assessment: dict) -> list[dict]:
    """One row per category for template summary table (D1 format).
    Critical Products is always a special-case row (— / — / —)."""
    categories = dict(assessment.get("categories") or {})
    categories.pop("CRITICAL_PRODUCTS", None)  # prevent normal processing; we add it last

    order = list(CHART_CATEGORIES)
    seen = set()
    rows = []
    for category in order:
        if category not in categories:
            continue
        seen.add(category)
        raw_inp = categories[category]
        inp = dict(raw_inp) if isinstance(raw_inp, dict) else {}
        has_backup = _effective_has_backup(inp)
        time_impact = inp.get("time_to_impact_hours")
        recovery = inp.get("recovery_time_hours")
        s = _sources_summary(inp.get("supply"))
        notes = normalize_note_text(s)
        provider = _primary_provider_identified(inp, category)
        if not notes or notes.strip() == SUMMARY_NOT_CONFIRMED_TEXT.strip():
            notes = "Provider identified." if provider == "Yes" else SUMMARY_NOT_CONFIRMED_TEXT
        elif len(notes) > 80:
            notes = notes[:77] + "..."
        notes_final = notes + _independence_note_suffix(inp, provider)
        if len(notes_final) > 80:
            notes_final = notes_final[:77] + "..."
        rows.append({
            "category": CATEGORY_DISPLAY.get(category, category),
            "primary_provider": provider,
            "backup_present": "Yes" if has_backup else "No",
            "time_to_severe_impact_hrs": "N/A" if time_impact is None else str(time_impact),
            "recovery_time_hrs": "N/A" if recovery is None else str(recovery),
            "notes": notes_final,
        })
    for category, raw_inp in categories.items():
        if category in seen:
            continue
        inp = dict(raw_inp) if isinstance(raw_inp, dict) else {}
        s = _sources_summary(inp.get("supply"))
        notes = normalize_note_text(s)
        provider = _primary_provider_identified(inp, category)
        if not notes or notes.strip() == SUMMARY_NOT_CONFIRMED_TEXT.strip():
            notes = "Provider identified." if provider == "Yes" else SUMMARY_NOT_CONFIRMED_TEXT
        notes_final = notes + _independence_note_suffix(inp, provider)
        if len(notes_final) > 80:
            notes_final = notes_final[:77] + "..."
        rows.append({
            "category": CATEGORY_DISPLAY.get(category, category),
            "primary_provider": provider,
            "backup_present": "No",
            "time_to_severe_impact_hrs": "N/A",
            "recovery_time_hrs": "N/A",
            "notes": notes_final,
        })
    # Always add Critical Products as special-case row (verifier requirement)
    rows.append({
        "category": "Critical Products",
        "primary_provider": "N/A",
        "backup_present": "N/A",
        "time_to_severe_impact_hrs": "N/A",
        "recovery_time_hrs": "N/A",
        "notes": "—",
    })
    return rows


def iter_paragraphs_and_cells(doc: Document):
    """Yield (paragraph, parent_element) for every paragraph in doc and in every table cell."""
    for p in doc.paragraphs:
        yield p, None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p, cell


def replace_anchor_in_paragraph_preserve_drawings(p, anchor: str, replacement: str) -> bool:
    """
    Replace anchor with replacement only in a run that contains the anchor and does NOT contain a drawing.
    If anchor is split across runs or only in a drawing run, return False (caller should fail loudly).
    """
    for r in p.runs:
        if not r.text or anchor not in r.text:
            continue
        if _run_has_drawing(r):
            return False
        r.text = r.text.replace(anchor, replacement)
        return True
    return False


def replace_anchor_in_paragraph(paragraph, anchor: str, replacement: str) -> bool:
    """Replace anchor with replacement in paragraph. Never clears runs that contain drawings."""
    if paragraph_has_drawing(paragraph):
        return replace_anchor_in_paragraph_preserve_drawings(paragraph, anchor, replacement)
    full = "".join(run.text for run in paragraph.runs)
    anchor_norm = " ".join(anchor.split())
    full_norm = " ".join(full.split())
    if anchor in full:
        new_full = full.replace(anchor, replacement)
    elif full_norm == anchor_norm or full_norm.strip() == anchor_norm:
        new_full = replacement
    else:
        return False
    for run in paragraph.runs:
        if not _run_has_drawing(run):
            run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_full
    else:
        paragraph.add_run(new_full)
    return True


def _iter_paragraphs_in_tables(tables: list) -> list:
    """Yield all paragraphs in tables and nested tables (python-docx doc.tables does not include nested)."""
    out = []
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    out.append(p)
                out.extend(_iter_paragraphs_in_tables(cell.tables))
    return out


def replace_anchor_in_doc(doc: Document, anchor: str, replacement: str, body_only: bool = True) -> int:
    """Replace anchor with replacement in body paragraphs and table cells (including nested). Returns number of replacements.
    When body_only=True (default), header/footer are not mutated so title page graphic and branding are preserved."""
    hits = 0
    for p, _ in iter_paragraphs_and_cells(doc):
        if replace_anchor_in_paragraph(p, anchor, replacement):
            hits += 1
    for p in _iter_paragraphs_in_tables(doc.tables):
        if replace_anchor_in_paragraph(p, anchor, replacement):
            hits += 1
    if not body_only:
        for section in doc.sections:
            for p in section.header.paragraphs:
                if replace_anchor_in_paragraph(p, anchor, replacement):
                    hits += 1
            for p in section.footer.paragraphs:
                if replace_anchor_in_paragraph(p, anchor, replacement):
                    hits += 1
    return hits


def find_placeholder_or_anchor(doc: Document, pattern: str):
    """
    Search paragraphs and table cells for a placeholder or anchor string.
    Yields (paragraph, parent_cell_or_None) where paragraph text contains pattern.
    """
    for p, parent_cell in iter_paragraphs_and_cells(doc):
        if pattern in (p.text or ""):
            yield p, parent_cell


def _paragraph_text_normalized(p) -> str:
    """Full paragraph text with normalized whitespace (template may have extra spaces/splits)."""
    return " ".join((p.text or "").split())


def find_anchor_paragraph_exact(doc: Document, anchor: str, body_only: bool = False):
    """
    Find a paragraph whose full trimmed text equals the anchor (exact match).
    Normalizes whitespace so template runs like "[[CHART_" + "ELECTRIC_POWER]]" still match.
    When body_only=True, search only doc.paragraphs (not table cells), to avoid header/title.
    Yields (paragraph, parent_cell_or_None).
    """
    anchor_norm = " ".join(anchor.split())
    if body_only:
        for p in doc.paragraphs:
            if _paragraph_text_normalized(p) == anchor_norm:
                yield p, None
        return
    for p, parent_cell in iter_paragraphs_and_cells(doc):
        if _paragraph_text_normalized(p) == anchor_norm:
            yield p, parent_cell


# Sector headings for INFRA blocks (enforced in reporter for consistent layout)
INFRA_ANCHOR_TO_SECTOR_HEADING = {
    "[[INFRA_ENERGY]]": "ELECTRIC POWER",
    "[[INFRA_COMMS]]": "COMMUNICATIONS",
    "[[INFRA_IT]]": "INFORMATION TECHNOLOGY",
    "[[INFRA_WATER]]": "WATER",
    "[[INFRA_WASTEWATER]]": "WASTEWATER",
}


def inject_infra_sector_as_paragraphs(
    doc: Document,
    anchor: str,
    sector_heading: str,
    narrative: str,
    page_break_before: bool = False,
    page_break_after: bool = False,
) -> None:
    """
    Find the INFRA anchor paragraph, clear it, then insert sector heading + narrative as styled paragraphs.
    Splits narrative on double newlines and single newlines; each non-empty line becomes one Normal paragraph.
    Page break before: inserted before sector title (title at top of new page).
    Page break after: inserted after narrative (next sector will start on new page).
    """
    candidates = list(find_anchor_paragraph_exact(doc, anchor))
    if len(candidates) != 1:
        raise RuntimeError(
            f"INFRA anchor must appear exactly once: {anchor} (found {len(candidates)} occurrences)"
        )
    p, _ = candidates[0]
    if paragraph_has_drawing(p):
        _clear_paragraph_text_preserve_drawings(p)
    else:
        p.clear()
    insert_after = p
    if page_break_before:
        break_para = insert_paragraph_after(insert_after, "")
        run = break_para.add_run()
        run.add_break(WD_BREAK.PAGE)
        insert_after = break_para
    insert_after = insert_paragraph_after(insert_after, sector_heading, style="Heading 2")
    set_paragraph_keep_with_next(insert_after)
    # Split on double newlines (paragraphs) then single newlines (lines)
    blocks = re.split(r"\n\s*\n", narrative)
    for block in blocks:
        for line in block.split("\n"):
            line = line.strip()
            if line:
                insert_after = insert_paragraph_after(
                    insert_after, sanitize_text(line), style="Normal"
                )
    insert_after = insert_paragraph_after(insert_after, "", style="Normal")  # blank line between sector blocks
    if page_break_after:
        break_para = insert_paragraph_after(insert_after, "")
        run = break_para.add_run()
        run.add_break(WD_BREAK.PAGE)


def search_placeholders_and_anchors(doc: Document) -> list[tuple[str, object, object]]:
    """
    Search all paragraphs and table cells for known placeholders and anchors.
    Returns list of (matched_text, paragraph, parent_cell_or_None) for each hit.
    """
    text_placeholders = ("{{ASSET_NAME}}", "{{VISIT_DATE}}", "{{ASSESSOR}}", "{{LOCATION}}")
    anchors = CHART_ANCHORS + [TABLE_ANCHOR, DEP_SUMMARY_TABLE_ANCHOR, VULN_NARRATIVE_ANCHOR, SLA_PRA_SUMMARY_ANCHOR]
    hits = []
    for p, parent_cell in iter_paragraphs_and_cells(doc):
        text = p.text or ""
        for token in text_placeholders + anchors:
            if token in text:
                hits.append((token, p, parent_cell))
                break
    return hits


def replace_text_in_paragraph(paragraph, old: str, new: str) -> None:
    """Replace old with new in paragraph, preserving runs where possible."""
    if old not in (paragraph.text or ""):
        return
    # Simple approach: clear and set full text (loses formatting but works)
    for run in paragraph.runs:
        run.text = run.text.replace(old, new)


def replace_all_text_placeholders(doc: Document, asset: dict) -> None:
    """Replace {{ASSET_NAME}}, {{VISIT_DATE}}, {{ASSESSOR}}, {{LOCATION}} and
    [[PSA_NAME]], [[PSA_REGION]], [[PSA_CITY]], [[PSA_PHONE]], [[PSA_EMAIL]] in doc.
    Uses paragraph-level replacement when placeholder is split across runs.
    All values are sanitized before insertion.
    PSA phone is REQUIRED; export fails if missing.
    [[PSA_CELL]] must appear exactly once; after replace, 'Cell: [' is forbidden (template typo)."""
    phone = (asset.get("psa_phone") or asset.get("psa_cell") or "").strip()
    if not phone:
        raise RuntimeError("MISSING_PSA_PHONE: PSA phone number is required for export. Set asset.psa_cell or asset.psa_phone.")
    # When template contains [[PSA_CELL]], it must appear exactly once (fail if 0 or >1)
    psa_cell_anchor = "[[PSA_CELL]]"
    psa_count = 0
    for p, _ in iter_paragraphs_and_cells(doc):
        psa_count += (p.text or "").count(psa_cell_anchor)
    if psa_count != 0 and psa_count != 1:
        raise RuntimeError(
            f"[[PSA_CELL]] must appear exactly once in the template (found {psa_count}). "
            "Fix template to use exactly one 'Cell: [[PSA_CELL]]'."
        )
    mapping = {
        "{{ASSET_NAME}}": sanitize_text((asset.get("asset_name") or "").strip()),
        "{{VISIT_DATE}}": sanitize_text((asset.get("visit_date_iso") or "").strip()),
        "{{ASSESSOR}}": sanitize_text((asset.get("assessor") or "").strip()),
        "{{LOCATION}}": sanitize_text((asset.get("location") or "").strip()),
        "[[PSA_NAME]]": sanitize_text((asset.get("psa_name") or "").strip()),
        "[[PSA_REGION]]": sanitize_text((asset.get("psa_region") or "").strip()),
        "[[PSA_CITY]]": sanitize_text((asset.get("psa_city") or "").strip()),
        "[[PSA_PHONE]]": sanitize_text(phone),
        "[[PSA_CELL]]": sanitize_text(phone),
        "[PSA_CELL]]": sanitize_text(phone),
        "[[PSA_CELL]": sanitize_text(phone),   # template typo: missing closing ]
        "[[PSA_EMAIL]]": sanitize_text((asset.get("psa_email") or "").strip()),
        "[[PSA_EMAIL]": sanitize_text((asset.get("psa_email") or "").strip()),  # template typo: missing closing ]
        "[[FACILITY_NAME]]": sanitize_text((asset.get("asset_name") or "").strip()),
        "[[ASSESSMENT_DATE]]": sanitize_text((asset.get("visit_date_iso") or "").strip()[:10]),
    }
    # Layer 1: Paragraph-level replacement (handles split runs). Never wipe paragraphs that contain drawings (e.g. title image).
    for p, _ in iter_paragraphs_and_cells(doc):
        full_text = p.text or ""
        changed = False
        new_text = full_text
        for placeholder, value in mapping.items():
            if placeholder in full_text:
                new_text = new_text.replace(placeholder, value)
                changed = True
        if changed:
            if paragraph_has_drawing(p):
                _clear_paragraph_text_preserve_drawings(p)
                first_text_run = next((r for r in p.runs if not _run_has_drawing(r)), None)
                if first_text_run is not None:
                    first_text_run.text = new_text
                else:
                    p.add_run(new_text)
            else:
                p.clear()
                p.add_run(new_text)

    # Layer 2: Handle SDT (Structured Data Tag) elements - Word content controls
    for element in doc.element.iter():
        if element.tag == qn('w:t') and element.text:
            for placeholder, value in mapping.items():
                if placeholder in element.text:
                    element.text = element.text.replace(placeholder, value)

    # Fail if PSA_CELL replacement left a malformed "Cell: [" (missing closing bracket in template)
    for p, _ in iter_paragraphs_and_cells(doc):
        if "Cell: [" in (p.text or ""):
            raise RuntimeError(
                "PSA_CELL replacement produced malformed line: 'Cell: [' (missing closing bracket in template?). "
                "Template must use exactly 'Cell: [[PSA_CELL]]'."
            )


def _narrative_values_from_assessment(assessment: dict) -> dict:
    """
    Build a dict of narrative placeholder values from assessment.
    Used to fill "____" and "_________" in template text. Keys are short names for pattern matching.
    Uses backup_capacity_pct for mitigated loss when present (mitigated = 100 - capacity);
    when missing, uses MITIGATED_UNKNOWN_CLAUSE so narrative does not say "Not identified % loss".
    """
    categories = assessment.get("categories") or {}
    hours_val = None
    capacity_pct_val = None
    backup_type_val = None
    mitigated_loss_pct_val = None
    mitigated_unknown = False
    for code in CHART_CATEGORIES:
        cat = categories.get(code)
        if not cat or not cat.get("requires_service"):
            continue
        if hours_val is None:
            h = cat.get("time_to_impact_hours")
            if h is not None:
                hours_val = int(h) if isinstance(h, (int, float)) else str(h)
        if capacity_pct_val is None:
            cap = (1 - (cat.get("loss_fraction_no_backup") or 0)) * 100
            capacity_pct_val = f"{pct(cap):.1f}"
        if mitigated_loss_pct_val is None and cat.get("has_backup_any"):
            loss_with = cat.get("loss_fraction_with_backup")
            cap_backup = cat.get("backup_capacity_pct") or cat.get("backup_capacity_percent")
            if loss_with is not None and isinstance(loss_with, (int, float)):
                mitigated_loss_pct_val = f"{pct(loss_with * 100):.1f}"
            elif cap_backup is not None and isinstance(cap_backup, (int, float)):
                mitigated_loss_pct_val = f"{pct(100 - cap_backup):.1f}"
            else:
                mitigated_loss_pct_val = MITIGATED_UNKNOWN_CLAUSE
                mitigated_unknown = True
        if backup_type_val is None and cat.get("has_backup_any"):
            backup_type_val = (cat.get("backup_type") or "").strip() or "backup system"
    return {
        "hours": hours_val or NOT_CONFIRMED,
        "capacity_pct": capacity_pct_val or NOT_CONFIRMED,
        "backup_type": backup_type_val or "backup system",
        "mitigated_loss_pct": mitigated_loss_pct_val or NOT_CONFIRMED,
        "mitigated_unknown": mitigated_unknown,
    }


def fill_narrative_blanks(doc: Document, assessment: dict) -> None:
    """
    Replace narrative blanks ("____", "_________") in paragraphs with computed values.
    Uses targeted replacements (e.g. "approximately ____ hours" -> value).
    When mitigated loss is unknown, the "only suffer a ____ % loss" clause is replaced with
    MITIGATED_UNKNOWN_CLAUSE (no "Not identified % loss"). Any other remaining blank is "not confirmed".
    """
    values = _narrative_values_from_assessment(assessment)
    # Mitigated loss: use neutral clause when unknown instead of "Not identified % loss"
    if values.get("mitigated_unknown") and values["mitigated_loss_pct"] == MITIGATED_UNKNOWN_CLAUSE:
        mitigated_repl = MITIGATED_UNKNOWN_CLAUSE
    else:
        mitigated_repl = f"only suffer a {values['mitigated_loss_pct']}% loss"
    # Pairs of (substring containing blank, replacement for that substring)
    replacements = [
        ("approximately ____ hours", f"approximately {values['hours']} hours"),
        ("approximately ____ hour", f"approximately {values['hours']} hour"),
        ("reduced to ____ % operational capacity", f"reduced to {values['capacity_pct']}% operational capacity"),
        ("reduced to ____ %", f"reduced to {values['capacity_pct']}%"),
        ("____ % operational capacity", f"{values['capacity_pct']}% operational capacity"),
        ("only suffer a  ____ % loss", mitigated_repl),
        ("only suffer a ____ % loss", mitigated_repl),
        ("mitigated with the use of a _________.", f"mitigated with the use of a {values['backup_type']}."),
        ("mitigated with the use of a ____.", f"mitigated with the use of a {values['backup_type']}."),
    ]
    # Also replace any remaining standalone blanks
    for p, _ in iter_paragraphs_and_cells(doc):
        text = p.text or ""
        if "____" not in text and "_" * 7 not in text:
            continue
        new_text = text
        for old, new in replacements:
            if old in new_text:
                new_text = new_text.replace(old, new)
        # Replace any remaining ____ or long underscores with not confirmed (unknown)
        new_text = UNDERSCORE_RE.sub(NOT_CONFIRMED, new_text)
        if new_text != text:
            p.clear()
            p.add_run(new_text)


def remove_dev_anchor_block(doc: Document) -> None:
    """
    Delete the paragraph containing 'ANCHOR BLOCK (dev only)' AND all following
    paragraphs (to end-of-doc). The dev template appends these at the end.
    """
    marker = "ANCHOR BLOCK (dev only)"
    found = False
    to_remove = []
    for p in list(doc.paragraphs):
        if marker in (p.text or ""):
            found = True
        if found:
            to_remove.append(p)
    for p in to_remove:
        remove_paragraph(p)


def modernize_executive_heading_flow(doc: Document) -> None:
    """
    Normalize Part I heading text into a concise decision-brief flow.
    This keeps content truth intact and improves scan order.
    """
    replacements = {
        "EXECUTIVE RISK POSTURE": "A. Executive Risk Posture",
        "B. DEPENDENCY SNAPSHOT TABLE": "B. At-a-Glance Dependency Profile",
        "C. CROSS-INFRASTRUCTURE SYNTHESIS": "C. Cross-Infrastructure Analysis",
        "D. PRIORITY ACTIONS": "D. Priority Actions",
    }
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        new_text = replacements.get(text)
        if not new_text:
            continue
        p.clear()
        p.add_run(sanitize_text(new_text))


def remove_forced_page_breaks_in_part1(doc: Document) -> None:
    """
    Remove hard page-break paragraphs between Part I sub-sections to prevent
    large dead whitespace in executive pages.
    """
    in_part1 = False
    for p in list(doc.paragraphs):
        text = (p.text or "").strip().upper()
        if "PART I" in text:
            in_part1 = True
            continue
        if "PART II" in text:
            break
        if not in_part1:
            continue
        if _is_page_break_paragraph(p):
            remove_paragraph(p)


def remove_orphan_drawing_paragraphs_in_part1(doc: Document) -> None:
    """
    Remove drawing-only empty paragraphs between Part I section headings where they
    introduce large whitespace and no readable content.
    """
    in_part1 = False
    for p in list(doc.paragraphs):
        text_u = (p.text or "").strip().upper()
        if "PART I" in text_u:
            in_part1 = True
            continue
        if "PART II" in text_u:
            break
        if not in_part1:
            continue
        if (p.text or "").strip():
            continue
        if paragraph_has_drawing(p) and not _is_page_break_paragraph(p):
            remove_paragraph(p)


def tighten_report_layout_spacing(doc: Document) -> None:
    """
    Reduce excessive whitespace by compacting paragraph spacing and removing
    blank spacer paragraphs before headings (except drawing containers).
    """
    def _set_style_spacing(style_name: str, before_pt: int, after_pt: int, line: float) -> None:
        s = _get_style_safe(doc, style_name)
        if s is None:
            return
        pf = s.paragraph_format
        pf.space_before = Pt(before_pt)
        pf.space_after = Pt(after_pt)
        pf.line_spacing = line

    # Compact global typography for better information density.
    _set_style_spacing("Normal", 0, 4, 1.1)
    _set_style_spacing("Heading 1", 8, 3, 1.0)
    _set_style_spacing("Heading 2", 6, 2, 1.0)
    _set_style_spacing("Heading 3", 4, 2, 1.0)
    _set_style_spacing("Heading 4", 4, 2, 1.0)

    # Drop blank paragraphs directly before headings to eliminate vertical holes.
    paras = list(doc.paragraphs)
    for i in range(1, len(paras)):
        cur = paras[i]
        prev = paras[i - 1]
        if not _is_heading_style(cur):
            continue
        prev_text = (prev.text or "").strip()
        if prev_text:
            continue
        if paragraph_has_drawing(prev):
            continue
        remove_paragraph(prev)


def compact_table_layout(doc: Document) -> None:
    """
    Reduce visual whitespace in tables by top-aligning content and removing
    paragraph-level spacing inside cells.
    """
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                try:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                except Exception:
                    pass
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(0)
                    pf.line_spacing = 1.0


def enforce_heading_content_cohesion(doc: Document) -> None:
    """
    Reduce heading/data separation across page boundaries by applying keep-with-next
    on section headings that should stay attached to immediate content.
    """
    targets = {
        "A. EXECUTIVE RISK POSTURE",
        "B. AT-A-GLANCE DEPENDENCY PROFILE",
        "C. CROSS-INFRASTRUCTURE ANALYSIS",
        "D. PRIORITY ACTIONS",
        "ANNEX OVERVIEW",
        "SECTOR REPORTS",
        "CORE INFRASTRUCTURE OVERVIEW",
        "INTERNET TRANSPORT",
        "CRITICAL HOSTED SERVICES",
        "STRUCTURAL RISK PROFILE",
        "INFRASTRUCTURE VULNERABILITIES",
    }
    for p in doc.paragraphs:
        text = (p.text or "").strip().upper()
        if not text:
            continue
        style_name = ((p.style.name if p.style is not None else "") or "").strip().lower()
        is_heading = style_name.startswith("heading")
        if is_heading or text in targets:
            try:
                p.paragraph_format.keep_with_next = True
            except Exception:
                pass


def force_page_break_before_heading(doc: Document, heading_text: str) -> None:
    """
    Ensure a section heading starts a fresh page to prevent orphan heading-at-bottom
    when followed by a large table.
    """
    target = (heading_text or "").strip().upper()
    for p in doc.paragraphs:
        if (p.text or "").strip().upper() != target:
            continue
        # avoid duplicate page breaks
        prev_el = p._element.getprevious()
        if prev_el is not None and prev_el.tag == qn("w:p"):
            prev_p = DocxParagraph(prev_el, p._parent)
            if _is_page_break_paragraph(prev_p):
                return
        pb = insert_paragraph_before(p, "")
        run = pb.add_run()
        run.add_break(WD_BREAK.PAGE)
        return


def reorder_part2_sections(doc: Document) -> None:
    """
    Reorder major Part II sections to improve narrative flow:
    1) Sector Reports
    2) Structural Risk Profile
    3) Infrastructure Vulnerabilities
    4) Core Infrastructure Overview
    5) Internet Transport
    6) Critical Hosted Services
    """
    heading_map = {
        "SECTOR REPORTS": "SECTOR_REPORTS",
        "STRUCTURAL RISK PROFILE": "STRUCTURAL_PROFILE",
        "INFRASTRUCTURE VULNERABILITIES": "VULNERABILITIES",
        "CORE INFRASTRUCTURE OVERVIEW": "CORE_OVERVIEW",
        "INTERNET TRANSPORT": "INTERNET_TRANSPORT",
        "CRITICAL HOSTED SERVICES": "HOSTED_SERVICES",
    }
    desired_order = [
        "SECTOR_REPORTS",
        "STRUCTURAL_PROFILE",
        "VULNERABILITIES",
        "CORE_OVERVIEW",
        "INTERNET_TRANSPORT",
        "HOSTED_SERVICES",
    ]

    body = doc.element.body
    body_children = list(body)

    def _element_text(el) -> str:
        parts = []
        for r in el.iter():
            if r.tag == qn("w:t") and r.text:
                parts.append(r.text)
        return " ".join("".join(parts).split()).strip()

    in_part2 = False
    heading_indices: dict[str, int] = {}
    for idx, child in enumerate(body_children):
        if child.tag != qn("w:p"):
            continue
        text = _element_text(child)
        text_u = text.upper()
        if "PART II" in text_u:
            in_part2 = True
            continue
        if not in_part2:
            continue
        key = heading_map.get(text_u)
        if key and key not in heading_indices:
            heading_indices[key] = idx

    if len(heading_indices) < 3:
        return

    ordered_starts = sorted(((idx, key) for key, idx in heading_indices.items()), key=lambda x: x[0])
    segments: dict[str, list] = {}
    for i, (start_idx, key) in enumerate(ordered_starts):
        end_idx = ordered_starts[i + 1][0] if i + 1 < len(ordered_starts) else len(body_children)
        segments[key] = body_children[start_idx:end_idx]

    insert_pos = ordered_starts[0][0]
    for _, key in ordered_starts:
        for el in segments.get(key, []):
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)

    keys_in_original_order = [key for _, key in ordered_starts]
    reordered_keys = [k for k in desired_order if k in segments] + [
        k for k in keys_in_original_order if k not in desired_order
    ]
    reinsertion = []
    for key in reordered_keys:
        reinsertion.extend(segments.get(key, []))

    for offset, el in enumerate(reinsertion):
        body.insert(insert_pos + offset, el)


def _section_has_readable_content_between(doc: Document, start_text: str, stop_texts: list[str]) -> bool:
    paras = doc.paragraphs
    start_idx = None
    stop_idx = len(paras)
    start_u = start_text.upper()
    stop_u = [s.upper() for s in stop_texts]
    for i, p in enumerate(paras):
        t = (p.text or "").strip().upper()
        if t == start_u:
            start_idx = i
            break
    if start_idx is None:
        return False
    for i in range(start_idx + 1, len(paras)):
        t = (paras[i].text or "").strip().upper()
        if t in stop_u:
            stop_idx = i
            break
    for i in range(start_idx + 1, stop_idx):
        p = paras[i]
        t = (p.text or "").strip()
        if t:
            return True
    return False


def ensure_part1_decision_brief_content(
    doc: Document,
    executive_snapshot: dict | None,
    synthesis: dict | None,
    priority_actions: dict | list | None,
    assessment: dict | None = None,
    fallback_action_titles: list[str] | None = None,
) -> None:
    """
    Ensure Part I headings always have readable body content even when template anchors
    are embedded in drawing/textbox artifacts and direct anchor injection misses.
    """
    snapshot = executive_snapshot if isinstance(executive_snapshot, dict) else {}
    posture = (snapshot.get("posture") or "").strip()
    summary = (snapshot.get("summary") or "").strip()
    matrix_rows = snapshot.get("matrixRows") or []
    if not matrix_rows:
        matrix_rows = _matrix_rows_from_assessment(assessment or {})
    if not posture:
        posture = _posture_from_matrix_rows(matrix_rows)
    if _is_generic_executive_summary_text(summary):
        summary = _build_executive_summary_from_matrix(posture, matrix_rows)

    h_a = find_paragraph_by_exact_text(doc, "A. Executive Risk Posture", body_only=False)
    if h_a and not _section_has_readable_content_between(
        doc, "A. Executive Risk Posture", ["B. At-a-Glance Dependency Profile"]
    ):
        insert_after = h_a
        if posture:
            insert_after = insert_paragraph_after(insert_after, sanitize_text(posture), style="Normal")
        if summary:
            insert_after = insert_paragraph_after(insert_after, sanitize_text(summary), style="Normal")

    h_c = find_paragraph_by_exact_text(doc, "C. Cross-Infrastructure Analysis", body_only=False)
    if h_c and not _section_has_readable_content_between(
        doc, "C. Cross-Infrastructure Analysis", ["D. Priority Actions", "PART II – TECHNICAL ANNEX"]
    ):
        insert_after = h_c
        synth = synthesis if isinstance(synthesis, dict) else {}
        paras = synth.get("paragraphs") or []
        bullets = synth.get("bullets") or []
        if paras:
            for txt in paras[:2]:
                if str(txt).strip():
                    insert_after = insert_paragraph_after(insert_after, sanitize_text(str(txt).strip()), style="Normal")
        elif matrix_rows:
            earliest = None
            for row in matrix_rows:
                try:
                    tti = float(row.get("ttiHrs"))
                except Exception:
                    continue
                sector = str(row.get("sector") or "").strip()
                if not sector:
                    continue
                if earliest is None or tti < earliest[1]:
                    earliest = (sector, tti)
            if earliest:
                insert_after = insert_paragraph_after(
                    insert_after,
                    sanitize_text(f"Earliest cross-domain degradation begins in {earliest[0]} at {earliest[1]:g} hours."),
                    style="Normal",
                )
        for b in bullets[:3]:
            label = str((b or {}).get("label") or "").strip()
            text = str((b or {}).get("text") or "").strip()
            if label and text:
                insert_after = insert_paragraph_after(
                    insert_after, sanitize_text(f"{label.rstrip(':')}: {text}"), style="List Bullet"
                )

    h_d = find_paragraph_by_exact_text(doc, "D. Priority Actions", body_only=False)
    if h_d and not _section_has_readable_content_between(
        doc, "D. Priority Actions", ["PART II – TECHNICAL ANNEX"]
    ):
        insert_after = h_d
        actions = _parse_priority_actions_payload(priority_actions if priority_actions is not None else {})
        if actions:
            for a in actions[:3]:
                lead = (a.get("leadIn") or "").strip()
                desc = (a.get("fullText") or "").strip()
                if lead:
                    insert_after = insert_paragraph_after(insert_after, sanitize_text(lead), style="List Number")
                if desc:
                    insert_after = insert_paragraph_after(insert_after, sanitize_text(desc), style="Normal")
        else:
            fallback_inserted = 0
            for title in (fallback_action_titles or [])[:3]:
                t = (title or "").strip()
                if t:
                    insert_after = insert_paragraph_after(insert_after, sanitize_text(f"Address: {t}"), style="List Number")
                    fallback_inserted += 1
            if fallback_inserted == 0 and matrix_rows:
                earliest = None
                for r in matrix_rows:
                    sector = str(r.get("sector") or "").strip()
                    if not sector:
                        continue
                    try:
                        tti = float(r.get("ttiHrs"))
                    except Exception:
                        continue
                    if earliest is None or tti < earliest[1]:
                        earliest = (sector, tti)
                if earliest:
                    insert_after = insert_paragraph_after(
                        insert_after,
                        sanitize_text(f"Prepare immediate outage-response procedures for {earliest[0]} (earliest severe impact)."),
                        style="List Number",
                    )
                no_alt = []
                for r in matrix_rows:
                    sector = str(r.get("sector") or "").strip()
                    b = r.get("backupHrs")
                    if sector and (b is None or (isinstance(b, str) and b.strip().upper() in {"", "N/A", "NONE"})):
                        no_alt.append(sector)
                if no_alt:
                    insert_after = insert_paragraph_after(
                        insert_after,
                        sanitize_text(f"Prioritize alternate-capability planning for {', '.join(no_alt)}."),
                        style="List Number",
                    )


def _find_heading_paragraph(doc: Document, heading_text: str):
    target = (heading_text or "").strip().upper()
    for p in doc.paragraphs:
        if (p.text or "").strip().upper() == target:
            return p
    return None


def _find_heading_paragraph_any(doc: Document, heading_texts: list[str]):
    targets = {(t or "").strip().upper() for t in heading_texts if (t or "").strip()}
    for p in doc.paragraphs:
        if (p.text or "").strip().upper() in targets:
            return p
    return None


def _clear_text_between_headings(doc: Document, start_heading: str, end_heading: str) -> None:
    paras = list(doc.paragraphs)
    start_idx = None
    end_idx = None
    for i, p in enumerate(paras):
        t = (p.text or "").strip().upper()
        if start_idx is None and t == start_heading.upper():
            start_idx = i
            continue
        if start_idx is not None and t == end_heading.upper():
            end_idx = i
            break
    if start_idx is None or end_idx is None:
        return
    for i in range(start_idx + 1, end_idx):
        p = paras[i]
        if paragraph_has_drawing(p) or _is_page_break_paragraph(p):
            continue
        remove_paragraph(p)


def _build_part1_metrics(matrix_rows: list[dict], vuln_count: int) -> str:
    domains = len(matrix_rows)
    earliest = None
    for r in matrix_rows:
        try:
            tti = float(r.get("ttiHrs"))
        except Exception:
            continue
        earliest = tti if earliest is None else min(earliest, tti)
    earliest_txt = f"{earliest:g}h" if earliest is not None else "Unknown"
    return f"Domains assessed: {domains} | Findings: {vuln_count} | Earliest severe impact: {earliest_txt}"


def build_part1_executive_hero(
    doc: Document,
    executive_snapshot: dict | None,
    assessment: dict | None,
    vulnerability_count: int,
    fallback_action_titles: list[str] | None = None,
) -> None:
    """
    Build a concise executive hero panel under section A:
    left = risk + summary, right = top risk drivers.
    """
    h_a = _find_heading_paragraph_any(doc, ["A. Executive Risk Posture", "EXECUTIVE RISK POSTURE"])
    h_b = _find_heading_paragraph_any(doc, ["B. At-a-Glance Dependency Profile", "B. DEPENDENCY SNAPSHOT TABLE"])
    if h_a is None or h_b is None:
        return

    snapshot = executive_snapshot if isinstance(executive_snapshot, dict) else {}
    matrix_rows = snapshot.get("matrixRows") or _matrix_rows_from_assessment(assessment or {})
    posture = (snapshot.get("posture") or "").strip() or _posture_from_matrix_rows(matrix_rows)
    summary = (snapshot.get("summary") or "").strip()
    if _is_generic_executive_summary_text(summary):
        summary = _build_executive_summary_from_matrix(posture, matrix_rows)
    drivers = [str(d).strip() for d in (snapshot.get("drivers") or []) if str(d).strip()][:3]
    if not drivers:
        drivers = [str(t).strip() for t in (fallback_action_titles or []) if str(t).strip()][:3]

    start_text = (h_a.text or "").strip()
    end_text = (h_b.text or "").strip()
    _clear_text_between_headings(doc, start_text, end_text)

    table = insert_table_after(doc, h_a, 1, 2)
    apply_table_grid_style(table)
    set_table_fixed_widths(table, [3.1, 3.7])
    _normalize_table_paragraph_spacing(table)

    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]

    left.text = ""
    p = left.paragraphs[0]
    if posture.strip():
        p.add_run(sanitize_text(posture.upper()))
        p.runs[0].bold = True
    if summary.strip():
        p2 = left.add_paragraph(sanitize_text(summary))
        try:
            p2.style = "Normal"
        except Exception:
            pass

    right.text = ""
    rh = right.paragraphs[0]
    rh.add_run("Top Risk Drivers").bold = True
    if drivers:
        for d in drivers:
            bp = right.add_paragraph(sanitize_text(d))
            try:
                bp.style = "List Bullet"
            except Exception:
                pass
    else:
        # No synthetic placeholder text; keep section title only when no truth-backed drivers exist.
        pass

    metric_line = _build_part1_metrics(matrix_rows, vulnerability_count)
    if metric_line.strip():
        metric_para = insert_paragraph_after_block(doc, table, sanitize_text(metric_line), style="Normal")
        try:
            metric_para.paragraph_format.space_before = Pt(4)
            metric_para.paragraph_format.space_after = Pt(6)
        except Exception:
            pass


def compress_part1_c_and_d(doc: Document) -> None:
    """Keep C and D concise and decision-oriented."""
    _clear_text_between_headings(doc, "C. Cross-Infrastructure Analysis", "D. Priority Actions")

def scrub_any_remaining_anchors(doc: Document) -> None:
    """
    Final scrub: remove any paragraph containing '[['; clear table cells containing '[['.
    """
    for p in list(doc.paragraphs):
        if "[[" in (p.text or ""):
            remove_paragraph(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if "[[" in (cell.text or ""):
                    cell.text = ""


def remove_anchor_block_paragraphs(doc: Document) -> None:
    """Clear any paragraph whose text contains 'ANCHOR BLOCK (dev only)' or 'ANCHOR BLOCK'. Preserve drawings."""
    for p, _ in iter_paragraphs_and_cells(doc):
        text = (p.text or "").strip()
        if "ANCHOR BLOCK (dev only)" in text or "ANCHOR BLOCK" in text:
            if paragraph_has_drawing(p):
                _clear_paragraph_text_preserve_drawings(p)
            else:
                p.clear()


# C. SECTOR ANALYSIS removed entirely — do not insert under any circumstance
NO_VULNERABILITIES_TRIGGERED = "No vulnerabilities were triggered based on provided inputs."
VULN_BLOCK_SPACER_PT = 6
VULN_NARRATIVE_SPACING_PT = 6
VULN_OFC_ITEM_SPACING_PT = 3
VULN_DIVIDER_SPACING_PT = 9  # Between vulnerability blocks

# Standalone sector names that must not appear as orphan headings in Part I
ORPHAN_SECTOR_NAMES = frozenset({
    "Electric Power", "COMMUNICATIONS", "Information Technology", "Water", "Wastewater",
    "ELECTRIC POWER", "Communications", "INFORMATION TECHNOLOGY", "WATER", "WASTEWATER",
})

# Ghost sector lines: standalone labels that must not appear in Part I (template + renderer safeguard)
SECTOR_GHOST_LINES = frozenset({
    "ELECTRIC POWER",
    "COMMUNICATIONS",
    "INFORMATION TECHNOLOGY",
    "WATER",
    "WASTEWATER",
})


def remove_ghost_sector_lines(doc: Document) -> None:
    """
    Remove exact standalone paragraphs that match sector ghost labels.
    Part II headings (e.g. "ELECTRIC POWER — Dependency Assessment") do not match
    since they are not exact. Call early after loading template to prevent regression.
    Never remove paragraphs that contain drawings (e.g. title page graphic).
    """
    for p in list(doc.paragraphs):
        if paragraph_has_drawing(p):
            continue
        t = (p.text or "").strip()
        if t.upper() in SECTOR_GHOST_LINES:
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)


def rename_annex_phantom_headings(doc: Document) -> None:
    """
    Rename Annex headings to remove phantom feel.
    - Dependency Summary → Annex Overview
    - Infrastructure Dependency Vulnerabilities and Options for Consideration → Sector Reports
    Preserve paragraphs that contain drawings (e.g. title image).
    """
    for p, _ in iter_paragraphs_and_cells(doc):
        t = (p.text or "").strip()
        tu = t.upper()
        if "DEPENDENCY SUMMARY" in tu:
            if paragraph_has_drawing(p):
                _clear_paragraph_text_preserve_drawings(p)
                p.add_run(sanitize_text(ANNEX_OVERVIEW_HEADING))
            else:
                p.clear()
                p.add_run(sanitize_text(ANNEX_OVERVIEW_HEADING))
        elif (
            "INFRASTRUCTURE DEPENDENCY VULNERABILITIES" in tu
            or "VULNERABILITIES AND OPTIONS FOR CONSIDERATION" in tu
        ):
            if paragraph_has_drawing(p):
                _clear_paragraph_text_preserve_drawings(p)
                p.add_run(sanitize_text(SECTOR_REPORTS_HEADING))
            else:
                p.clear()
                p.add_run(sanitize_text(SECTOR_REPORTS_HEADING))


def ensure_sector_reports_heading(doc: Document) -> None:
    """Ensure a visible 'Sector Reports' heading exists before sector content anchors."""
    for p in doc.paragraphs:
        if (p.text or "").strip() == SECTOR_REPORTS_HEADING:
            return
    anchor_p = (
        find_paragraph_by_exact_text(doc, VULN_NARRATIVE_ANCHOR, body_only=False)
        or find_paragraph_by_exact_text(doc, VULNERABILITY_BLOCKS_ANCHOR, body_only=False)
    )
    if anchor_p is not None:
        insert_paragraph_before(anchor_p, SECTOR_REPORTS_HEADING, style="Heading 2")


# Static sector labels for template cleanup
SECTOR_LABELS = frozenset({
    "ELECTRIC POWER",
    "COMMUNICATIONS",
    "INFORMATION TECHNOLOGY",
    "WATER",
    "WASTEWATER",
})


def remove_sector_analysis_chart_block(doc: Document) -> None:
    """
    Remove legacy C. SECTOR ANALYSIS heading only (if present). Template should use "C. OPERATIONAL CAPABILITY CURVES"; do NOT remove that.
    Charts are inserted at CHART_* anchors in Section C. Part I has no INFRA_* anchors.
    """
    # Remove C. SECTOR ANALYSIS heading entirely (legacy phantom section)
    for p, _ in iter_paragraphs_and_cells(doc):
        text_stripped = (p.text or "").strip()
        if "C. SECTOR ANALYSIS" in text_stripped.upper() or "C. Sector Analysis" in text_stripped:
            remove_paragraph(p)
            break


def remove_orphan_sector_headings_in_part1(doc: Document) -> None:
    """
    Remove orphan sector headings (Electric Power, Communications, etc.) in Part I.
    These appear under the Dependency Snapshot Table; sector content moved to Part II.
    Only removes in Part I region (stops at 'PART II').
    """
    to_remove = []
    in_part1 = True
    for block in iter_block_items(doc):
        if not isinstance(block, DocxParagraph):
            continue
        text = (block.text or "").strip()
        if "PART II" in text.upper():
            in_part1 = False
            break
        if in_part1 and text in ORPHAN_SECTOR_NAMES:
            to_remove.append(block)
    for p in to_remove:
        remove_paragraph(p)


def remove_static_sector_labels_safeguard(doc: Document) -> None:
    """
    Remove standalone sector label paragraphs (template safeguard).
    Prevents old templates from reintroducing static headings in Part I.
    Runs before save; only removes in Part I region (stops at 'PART II').
    """
    to_remove = []
    in_part1 = False
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if "DEPENDENCY SNAPSHOT" in text.upper() or "SNAPSHOT TABLE" in text.upper():
            in_part1 = True
        if "PART II" in text.upper():
            break
        if in_part1 and text.upper() in SECTOR_LABELS:
            to_remove.append(p)
    for p in to_remove:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)


# ADA Report v2: anchors that get fallback text when no engine data (not removed). Part I no longer has INFRA_*.
PLACEHOLDER_ANCHORS_WITH_FALLBACK = (
    "[[SNAPSHOT_POSTURE]]",
    "[[SNAPSHOT_SUMMARY]]",
    "[[SNAPSHOT_DRIVERS]]",
    "[[SNAPSHOT_MATRIX]]",
    "[[SNAPSHOT_CASCADE]]",
    "[[SYNTHESIS]]",
    "[[PRIORITY_ACTIONS]]",
)
FALLBACK_NO_FINDINGS = "No structural vulnerabilities identified based on provided inputs."
FALLBACK_WITH_FINDINGS = ""


def replace_placeholder_anchors_with_fallback(doc: Document, vulnerability_count: int = 0) -> None:
    """
    Replace SNAPSHOT_*, SYNTHESIS, PRIORITY_ACTIONS with fallback text.
    Never use "No structural vulnerabilities..." when vulnerability_count > 0.
    """
    fallback = FALLBACK_WITH_FINDINGS if vulnerability_count > 0 else FALLBACK_NO_FINDINGS
    for p, _ in iter_paragraphs_and_cells(doc):
        text_norm = _paragraph_text_normalized(p)
        for anchor in PLACEHOLDER_ANCHORS_WITH_FALLBACK:
            if " ".join(anchor.split()) == text_norm:
                p.clear()
                p.add_run(sanitize_text(fallback))
                break


def _is_generic_executive_summary_text(text: str) -> bool:
    t = (text or "").strip().lower()
    if not t:
        return True
    return (
        "structural drivers are reflected in structural constraints" in t
        or "inform the dependency characteristics described in this assessment" in t
    )


def _build_executive_summary_from_matrix(posture: str, matrix_rows: list[dict]) -> str:
    if not matrix_rows:
        return (posture or "").strip()
    tti_pairs: list[tuple[str, float]] = []
    for row in matrix_rows:
        sector = str(row.get("sector") or "").strip()
        raw_tti = row.get("ttiHrs")
        try:
            tti = float(raw_tti)
        except Exception:
            continue
        if sector:
            tti_pairs.append((sector, tti))
    if not tti_pairs:
        return (posture or "").strip()
    earliest_sector, earliest_tti = sorted(tti_pairs, key=lambda x: x[1])[0]
    near_term = [s for s, t in tti_pairs if t <= 4]
    near_term_list = ", ".join(near_term)
    clean_posture = (posture or "").replace(" Structural Sensitivity", "").strip().upper()
    if not clean_posture:
        return ""
    prefix = f"Overall dependency exposure is {clean_posture}."
    return (
        f"{prefix} Earliest operational degradation begins in {earliest_sector} at {earliest_tti:g} hours. "
        f"Near-term dependency pressure is concentrated in {near_term_list}."
    )


def _matrix_rows_from_assessment(assessment: dict) -> list[dict]:
    rows: list[dict] = []
    cats = (assessment or {}).get("categories") or {}
    for code in CHART_CATEGORIES:
        cat = cats.get(code) or {}
        if not cat.get("requires_service"):
            continue
        tti = cat.get("time_to_impact_hours")
        loss = cat.get("loss_fraction_no_backup")
        if isinstance(loss, (int, float)):
            loss = int(round(loss * 100))
        backup = cat.get("backup_duration_hours") if _effective_has_backup(cat) else None
        rows.append(
            {
                "sector": CATEGORY_DISPLAY.get(code, code),
                "ttiHrs": tti if tti is not None else "Unknown",
                "lossPct": loss if loss is not None else "Unknown",
                "backupHrs": backup,
            }
        )
    return rows


def _posture_from_matrix_rows(rows: list[dict]) -> str:
    if not rows:
        return ""
    min_tti = None
    for r in rows:
        try:
            tti = float(r.get("ttiHrs"))
        except Exception:
            continue
        min_tti = tti if min_tti is None else min(min_tti, tti)
    if min_tti is None:
        return ""
    if min_tti <= 1:
        return "High Structural Sensitivity"
    if min_tti <= 4:
        return "Elevated Structural Sensitivity"
    return "Moderate Structural Sensitivity"


def inject_executive_snapshot_at_anchors(doc: Document, snapshot: dict) -> None:
    """
    Inject curve-driven executive snapshot content at SNAPSHOT_* anchors.
    Replaces posture, summary, drivers, matrix table, cascade (or removes cascade if empty).
    """
    posture = (snapshot.get("posture") or "").strip()
    summary = (snapshot.get("summary") or "").strip()
    drivers = snapshot.get("drivers") or []
    matrix_rows = snapshot.get("matrixRows") or []
    cascade = snapshot.get("cascade")

    # SNAPSHOT_POSTURE: clear anchor, insert content after with explicit style
    p = find_paragraph_by_exact_text(doc, "[[SNAPSHOT_POSTURE]]", body_only=True)
    if p and posture:
        p.clear()
        insert_paragraph_after(p, sanitize_text(posture), style="Normal")

    # SNAPSHOT_SUMMARY + compact quick-view bullets derived from matrix rows
    p = find_paragraph_by_exact_text(doc, "[[SNAPSHOT_SUMMARY]]", body_only=True)
    if p and (summary or matrix_rows):
        if _is_generic_executive_summary_text(summary):
            summary = _build_executive_summary_from_matrix(posture, matrix_rows)
        p.clear()
        insert_after = p
        if summary:
            insert_after = insert_paragraph_after(p, sanitize_text(summary), style="Normal")
        quick_view_rows = []
        for row in matrix_rows[:5]:
            sector = str(row.get("sector") or "").strip()
            tti = str(row.get("ttiHrs") or "").strip()
            raw_loss = str(row.get("lossPct") or "").strip()
            loss = raw_loss
            if loss and not loss.endswith("%"):
                loss = f"{loss}%"
            backup = row.get("backupHrs")
            backup_text = "N/A" if backup is None or (isinstance(backup, str) and not backup.strip()) else f"{backup}h"
            if not sector or not tti:
                continue
            quick_view_rows.append(
                f"{sector}: severe impact {tti}h; loss {loss or 'N/A'}; alternate duration {backup_text}."
            )
        for line in quick_view_rows:
            insert_after = insert_paragraph_after(insert_after, sanitize_text(line), style="List Bullet")

    # SNAPSHOT_DRIVERS (up to 3 items as List Bullet)
    p = find_paragraph_by_exact_text(doc, "[[SNAPSHOT_DRIVERS]]", body_only=True)
    if p and drivers:
        p.clear()
        insert_after = p
        for d in drivers[:3]:
            insert_after = insert_paragraph_after(insert_after, sanitize_text(str(d)), style="List Bullet")

    # SNAPSHOT_MATRIX (table): clear anchor, insert table after
    p = find_paragraph_by_exact_text(doc, "[[SNAPSHOT_MATRIX]]", body_only=True)
    if p:
        if matrix_rows:
            p.clear()
            set_paragraph_keep_with_next(p)
            num_rows = 1 + len(matrix_rows)
            table = insert_table_after(doc, p, num_rows, 5)
            apply_table_grid_style(table)
            table.autofit = False
            widths = [1.2, 1.8, 1.2, 1.4, 2.0]
            set_table_fixed_widths(table, widths)
            for c, h in enumerate(("Sector", "Time to Severe Impact (hrs)", "Functional Loss (%)", "Backup Duration (hrs)", "Structural Posture")):
                table.rows[0].cells[c].text = h
            set_repeat_header_row(table.rows[0])
            set_table_rows_cant_split(table)
            for i, row in enumerate(matrix_rows):
                backup_hrs = row.get("backupHrs")
                if backup_hrs is None or (isinstance(backup_hrs, str) and not backup_hrs.strip()):
                    backup_hrs = "N/A"
                table.rows[i + 1].cells[0].text = sanitize_text(row.get("sector", ""))
                table.rows[i + 1].cells[1].text = sanitize_text(row.get("ttiHrs") or "Unknown")
                table.rows[i + 1].cells[2].text = sanitize_text(row.get("lossPct") or "Unknown")
                table.rows[i + 1].cells[3].text = sanitize_text(str(backup_hrs))
                table.rows[i + 1].cells[4].text = sanitize_text(row.get("structuralPosture") or "Unknown")
            insert_paragraph_after(p, "")
        else:
            remove_paragraph(p)

    # SNAPSHOT_CASCADE: clear anchor, insert content after or remove if empty
    p = find_paragraph_by_exact_text(doc, "[[SNAPSHOT_CASCADE]]", body_only=True)
    if p:
        if cascade and str(cascade).strip():
            p.clear()
            insert_paragraph_after(p, sanitize_text(str(cascade).strip()), style="Normal")
        else:
            remove_paragraph(p)


def _humanize_cross_infra_analysis_text(raw: str) -> str:
    """Rewrite colon-delimited cross-infrastructure lines into readable sentences without adding facts."""
    text = (raw or "").strip()
    if not text:
        return ""
    def _boolish(v: str) -> str | None:
        t = (v or "").strip().lower()
        if t in ("yes", "y", "true", "1"):
            return "yes"
        if t in ("no", "n", "false", "0"):
            return "no"
        return None

    def _render_fact(key: str, value: str) -> str:
        k = (key or "").strip().lower()
        v = (value or "").strip()
        b = _boolish(v)
        if "time to severe impact" in k:
            return f"reaches severe impact in {v}"
        if "restoration coordination" in k:
            if b == "yes":
                return "has restoration coordination"
            if b == "no":
                return "does not have restoration coordination"
            return f"restoration coordination is {v}"
        if "service connection count" in k:
            return f"has {v} service connection(s)"
        if "alternate water source" in k:
            if b == "yes":
                return "has an alternate water source"
            if b == "no":
                return "does not have an alternate water source"
            return f"alternate water source status is {v}"
        if "onsite pumping" in k:
            if b == "yes":
                return "has onsite pumping capability"
            if b == "no":
                return "does not have onsite pumping capability"
            return f"onsite pumping status is {v}"
        if "building entry diversity" in k:
            return f"building entry diversity is {v}"
        if "redundancy activation" in k:
            return f"redundancy activation is {v}"
        return f"{k} is {v}"

    lines: list[str] = []
    for src in text.splitlines():
        line = (src or "").strip()
        if not line:
            continue
        if ":" not in line:
            lines.append(line)
            continue
        head, rest = line.split(":", 1)
        sector = head.strip()
        fragments: list[str] = []
        for chunk in re.split(r"\.\s*", rest.strip()):
            c = chunk.strip().strip(".")
            if not c:
                continue
            if ":" in c:
                k, v = c.split(":", 1)
                val = v.strip()
                if val:
                    fragments.append(_render_fact(k, val))
            else:
                fragments.append(c)
        if fragments:
            lines.append(f"{sector}: " + "; ".join(fragments) + ".")
        else:
            lines.append(f"{sector}.")
    return "\n".join(lines).strip()


def _is_cross_dependency_enabled(data: dict, assessment: dict | None) -> bool:
    """Single source of truth for cross-dependency module gating."""
    if isinstance(data.get("cross_dependency_enabled"), bool):
        return data.get("cross_dependency_enabled") is True
    settings = (assessment or {}).get("settings") or {}
    return settings.get("cross_dependency_enabled") is True


def _remove_cross_infra_analysis_section(doc: Document) -> None:
    """
    Remove cross-infrastructure analysis section when module is disabled:
    - remove anchor paragraph
    - remove nearest preceding heading paragraph if it is the cross-infra heading
    """
    p_anchor = find_paragraph_by_exact_text(doc, CROSS_INFRA_ANALYSIS_ANCHOR, body_only=False)
    if p_anchor is None:
        return
    prev_el = p_anchor._element.getprevious()
    heading_re = re.compile(r"^cross[-\s]infrastructure analysis$", re.IGNORECASE)
    while prev_el is not None and prev_el.tag == qn("w:p"):
        prev_p = DocxParagraph(prev_el, p_anchor._parent)
        prev_text = (prev_p.text or "").strip()
        if not prev_text:
            prev_el = prev_el.getprevious()
            continue
        if _is_heading_style(prev_p) and heading_re.match(prev_text):
            remove_paragraph(prev_p)
        break
    remove_paragraph(p_anchor)


def _synthesis_sanitize_single_provider(text: str, assessment: dict) -> str:
    """
    Replace "single provider" with "single confirmed service path" when provider not identified.
    Never claim "single provider" when provider_identified is False for any referenced domain.
    """
    if not text or "single provider" not in text.lower():
        return text
    categories = assessment.get("categories") or {}
    any_provider_identified = False
    for code in CHART_CATEGORIES:
        cat = categories.get(code) or {}
        supply = cat.get("supply") or {}
        sources = supply.get("sources") or []
        if any(
            (s or {}).get("provider_name") or (s or {}).get("service_provider") or (s or {}).get("provider")
            for s in sources
        ):
            any_provider_identified = True
            break
    if not any_provider_identified:
        text = re.sub(r"\bsingle provider\b", "single confirmed service path", text, flags=re.I)
    return text


def inject_synthesis_at_anchor(doc: Document, synthesis: dict, assessment: dict | None = None) -> None:
    """
    Inject Cross-Infrastructure Synthesis at [[SYNTHESIS]] anchor.
    Template owns heading; inject only body paragraphs (Normal) and bullets (List Bullet).
    Sanitizes "single provider" -> "single confirmed service path" when provider not identified.
    """
    paragraphs = synthesis.get("paragraphs") or []
    bullets = synthesis.get("bullets") or []
    assessment = assessment or {}

    p = find_paragraph_by_exact_text(doc, "[[SYNTHESIS]]", body_only=True)
    if p is None:
        return

    p.clear()
    insert_after = p

    for para in paragraphs:
        if para and str(para).strip():
            raw = str(para).strip()
            human = ensure_synthesis_formatting(normalize_spacing(sanitize_backend_evidence(raw)))
            sanitized = _synthesis_sanitize_single_provider(human or raw, assessment)
            insert_after = insert_paragraph_after(insert_after, sanitize_text(sanitized), style="Normal")

    for b in bullets:
        label = (b.get("label") or "").strip()
        text = (b.get("text") or "").strip()
        if label and text:
            # Ensure "Category: Sentence" spacing (no "Category:Sentence")
            label = label.rstrip(":").rstrip() + ": "
            human = ensure_synthesis_formatting(normalize_spacing(sanitize_backend_evidence(text)))
            sanitized_text = _synthesis_sanitize_single_provider(human or text, assessment)
            bullet_para = insert_paragraph_after(insert_after, "", style="List Bullet")
            bold_run = bullet_para.add_run(sanitize_text(label))
            bold_run.bold = True
            bullet_para.add_run(sanitize_text(sanitized_text))
            insert_after = bullet_para
        elif text:
            human = ensure_synthesis_formatting(normalize_spacing(sanitize_backend_evidence(text)))
            insert_after = insert_paragraph_after(insert_after, sanitize_text(human or text), style="List Bullet")
        else:
            continue


def _strip_priority_action_curve_language(text: str) -> str:
    """Remove hours, percent, and 'reaches severe impact' from Priority Action text."""
    if not text or not isinstance(text, str):
        return text
    text = re.sub(r"\b\d+\s*hours?\b", "", text, flags=re.I)
    text = re.sub(r"~\s*\d+%|\b\d+%\b", "", text)
    text = re.sub(r"\b\d+\s*percent\b", "", text, flags=re.I)
    text = re.sub(r"reaches severe impact[^.]*\.?", "", text, flags=re.I)
    text = re.sub(r"this dependency[^.]*\.?", "", text, flags=re.I)
    return " ".join(text.split()).strip()


def _parse_priority_actions_payload(priority_actions: dict) -> list[dict]:
    """
    Normalize priority_actions into list of {leadIn, fullText, severe_impact_hours?, functional_loss_pct?}.
    If payload is a single string, parse by blank lines or ' — ' and return same structure.
    """
    if not priority_actions:
        return []
    if isinstance(priority_actions, str):
        # Parse by blank lines then optionally by " — " for first line = title, rest = description
        blocks = re.split(r"\n\s*\n", priority_actions.strip())
        out = []
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            lines = block.split("\n")
            first = (lines[0] or "").strip()
            rest = "\n".join(lines[1:]).strip() if len(lines) > 1 else ""
            out.append({"leadIn": first, "fullText": rest})
        return out
    if not isinstance(priority_actions, dict):
        return []
    actions = priority_actions.get("actions")
    if actions is None and isinstance(priority_actions, list):
        actions = priority_actions
    if isinstance(actions, str):
        return _parse_priority_actions_payload(actions)
    if not isinstance(actions, list):
        return []
    out = []
    for a in actions[:5]:
        if not isinstance(a, dict):
            continue
        lead_in = (a.get("leadIn") or "").strip()
        full_text = (a.get("fullText") or "").strip()
        full_text = _strip_priority_action_curve_language(full_text)
        out.append({
            "leadIn": lead_in,
            "fullText": full_text,
            "severe_impact_hours": None,  # Never include in output
            "functional_loss_pct": None,
        })
    return out


def inject_priority_actions_at_anchor(doc: Document, priority_actions: dict) -> None:
    """
    Inject Priority Actions at [[PRIORITY_ACTIONS]] anchor as grouped paragraphs (no run-on).
    Template owns heading; inject only body content.
    Per action: Para 1 (bold) "<Action Name> — <Sector>"; Para 2 description.
    """
    if not priority_actions:
        return
    actions = _parse_priority_actions_payload(priority_actions)
    if not actions:
        return

    p = find_paragraph_by_exact_text(doc, "[[PRIORITY_ACTIONS]]", body_only=True)
    if p is None:
        raise RuntimeError("[[PRIORITY_ACTIONS]] anchor not found in template; cannot inject priority actions.")

    p.clear()
    insert_after = p

    for a in actions:
        lead_in = (a.get("leadIn") or "").strip()
        full_text = (a.get("fullText") or "").strip()
        if not lead_in and not full_text:
            continue
        if not lead_in and full_text:
            # Derive a concise title from truth text; never inject placeholder labels.
            first_clause = re.split(r"[.;:]", full_text, maxsplit=1)[0].strip()
            lead_in = first_clause[:90].rstrip()
        # Para 1 (bold): Action — Sector
        title_para = insert_paragraph_after(insert_after, "", style="List Number")
        title_para.add_run(sanitize_text(lead_in))
        if title_para.runs:
            title_para.runs[0].bold = True
        insert_after = title_para
        # Para 2: description (no hours, no percent - structural-driver-based only)
        if full_text:
            desc_para = insert_paragraph_after(insert_after, sanitize_text(full_text), style="Normal")
            insert_after = desc_para
        # Never add "Severe impact: X hours" or "Functional loss: Y%" - Priority Actions are driver-based only


def clear_remaining_anchor_paragraphs(doc: Document) -> None:
    """
    Clear any paragraph that is exactly an anchor token (handles duplicate anchors
    e.g. in _dev_with_anchors when anchors are also appended at end of document).
    """
    all_anchors = list(REQUIRED_ANCHORS) + [TABLE_ANCHOR, SLA_PRA_SUMMARY_ANCHOR, CROSS_DEPENDENCY_SUMMARY_ANCHOR]
    for p, _ in iter_paragraphs_and_cells(doc):
        text_norm = _paragraph_text_normalized(p)
        for anchor in all_anchors:
            if " ".join(anchor.split()) == text_norm:
                p.clear()
                break


def assert_no_anchors_remaining(doc: Document) -> None:
    """
    Assert the document contains no anchor tokens or dev-only blocks.
    Raises SystemExit with a clear message if any are found.
    """
    all_anchors = list(REQUIRED_ANCHORS) + [TABLE_ANCHOR, SLA_PRA_SUMMARY_ANCHOR, CROSS_DEPENDENCY_SUMMARY_ANCHOR]
    found = []
    for p, _ in iter_paragraphs_and_cells(doc):
        text = p.text or ""
        for anchor in all_anchors:
            if anchor in text:
                found.append(anchor)
        if "ANCHOR BLOCK (dev only)" in text or "ANCHOR BLOCK" in text:
            found.append("ANCHOR BLOCK (dev only)")
    if found:
        print("ERROR: Generated document still contains anchor or dev-only text:", file=sys.stderr)
        for a in set(found):
            print(f"  - {a!r}", file=sys.stderr)
        sys.exit(1)


def assert_no_narrative_blanks_remaining(doc: Document) -> None:
    """Fail build if any '____' or long underscore placeholders remain in the output."""
    for p, _ in iter_paragraphs_and_cells(doc):
        text = p.text or ""
        if UNDERSCORE_RE.search(text):
            print("ERROR: Generated document still contains unreplaced narrative blank (____)", file=sys.stderr)
            sys.exit(1)


def _set_repeat_header_row(row) -> None:
    """Mark table row as repeating header on each page."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trPr.append(tbl_header)


def safety_sweep_underscore_blanks(doc: Document) -> None:
    """Replace any remaining underscore blanks (r'_{3,}') with NOT_CONFIRMED (unknown) in all paragraphs. Preserve drawings."""
    for p in _iter_paragraphs(doc):
        text = p.text or ""
        if not UNDERSCORE_RE.search(text):
            continue
        new_text = UNDERSCORE_RE.sub(NOT_CONFIRMED, text)
        if paragraph_has_drawing(p):
            _clear_paragraph_text_preserve_drawings(p)
            first_text_run = next((r for r in p.runs if not _run_has_drawing(r)), None)
            if first_text_run is not None:
                first_text_run.text = new_text
            else:
                p.add_run(new_text)
        else:
            p.clear()
            p.add_run(new_text)


def assert_exec_summary_not_followed_by_table(doc: Document) -> None:
    """Fail if 'EXECUTIVE SUMMARY' paragraph is immediately followed by a table (regression guard)."""
    blocks = list(iter_block_items(doc))
    for i, block in enumerate(blocks):
        if isinstance(block, DocxParagraph) and (block.text or "").strip() == "EXECUTIVE SUMMARY":
            if i + 1 < len(blocks) and isinstance(blocks[i + 1], DocxTable):
                print(
                    "ERROR: EXECUTIVE SUMMARY is immediately followed by a table (bug regression).",
                    file=sys.stderr,
                )
                sys.exit(1)
            break


def assert_no_export_style_tables(doc: Document) -> None:
    """Fail build if any table has export-style headers (Requires Service, Time to Impact, etc.).
    Uses exact match so template summary header 'Recovery Time (hrs)' is not flagged."""
    for table in doc.tables:
        if not table.rows:
            continue
        for row in table.rows[:2]:
            for cell in row.cells:
                c = (cell.text or "").strip()
                for bad in EXPORT_TABLE_BAD_HEADERS:
                    if c == bad:
                        print(
                            f"ERROR: Export-style table detected (header contains {bad!r}). "
                            "Use template summary/VOFC format only.",
                            file=sys.stderr,
                        )
                        sys.exit(1)


def gate_a_placeholder_anchor_leak_scanner(doc: Document) -> None:
    """
    Gate A: Placeholder/Anchor Leak Scanner (FAIL HARD).
    Scan fully-rendered text; fail export if placeholders, unresolved anchors, or deprecated terms found.
    """
    offenders: list[tuple[str, str, str]] = []  # (type, match, context)
    for p, _ in iter_paragraphs_and_cells(doc):
        text = p.text or ""
        if not text.strip():
            continue
        text_lower = text.lower()

        # 1) Placeholder phrases (case-insensitive)
        for phrase in GATE_A_PLACEHOLDER_PHRASES:
            if phrase in text_lower:
                idx = text_lower.find(phrase)
                ctx = _context_snippet(text, idx, 40)
                offenders.append(("placeholder", repr(phrase), ctx))

        # 2) Legacy FORBIDDEN_PLACEHOLDERS
        for forbidden in FORBIDDEN_PLACEHOLDERS:
            if forbidden in text:
                ctx = text[:80] + "..." if len(text) > 80 else text
                offenders.append(("placeholder", repr(forbidden), ctx))

        # 3) Unresolved anchor pattern
        for m in GATE_A_ANCHOR_RE.finditer(text):
            ctx = _context_snippet(text, m.start(), 40)
            offenders.append(("anchor", m.group(), ctx))

        # 4) Deprecated terms (case-insensitive). "safe" only as whole word (allow "unsafe", "safety")
        for term in GATE_A_DEPRECATED_TERMS:
            if term == "safe":
                m = GATE_A_SAFE_WORD_BOUNDARY_RE.search(text_lower)
                if m:
                    idx = m.start()
                    ctx = _context_snippet(text, idx, 40)
                    offenders.append(("deprecated", repr(term), ctx))
            elif term in text_lower:
                idx = text_lower.find(term)
                ctx = _context_snippet(text, idx, 40)
                offenders.append(("deprecated", repr(term), ctx))

    if offenders:
        print("ERROR: Gate A — Placeholder/Anchor/Deprecated term leakage. Export blocked.", file=sys.stderr)
        for typ, match, ctx in offenders[:10]:
            hint = "Check sector narratives" if typ == "placeholder" else "Check anchor injection" if typ == "anchor" else "Remove deprecated SAFE/security assessment references"
            print(f"  [{typ}] {match} | context: ...{ctx}... | hint: {hint}", file=sys.stderr)
        if len(offenders) > 10:
            print(f"  ... and {len(offenders) - 10} more", file=sys.stderr)
        sys.exit(1)


def _context_snippet(text: str, idx: int, half_len: int) -> str:
    """Return ±half_len chars around idx, sanitized for display."""
    start = max(0, idx - half_len)
    end = min(len(text), idx + half_len)
    return text[start:end].replace("\n", " ").strip()


def gate_b_suppress_empty_sections(doc: Document) -> None:
    """
    Gate B: Empty Section Suppression (NO DEAD HEADINGS).
    Remove headings that have no renderable content before the next section.
    SNAPSHOT_CASCADE is handled in inject_executive_snapshot_at_anchors (paragraph removed).
    """
    blocks = list(iter_block_items(doc))
    to_remove = []
    i = 0
    while i < len(blocks):
        block = blocks[i]
        if isinstance(block, DocxParagraph):
            text = (block.text or "").strip()
            # Only consider actual heading-styled paragraphs
            if not _is_heading_style(block):
                i += 1
                continue
            # Known section headings that may be empty
            empty_section_headings = (
                "Critical Products",
                "CRITICAL PRODUCTS",
                "Cross-Infrastructure Synthesis",
                "CROSS-INFRASTRUCTURE SYNTHESIS",
            )
            if text in empty_section_headings:
                # Look ahead: collect consecutive empty/blank paragraphs until next heading or content
                j = i + 1
                has_content = False
                while j < len(blocks):
                    next_b = blocks[j]
                    if isinstance(next_b, DocxTable):
                        has_content = True
                        break
                    if isinstance(next_b, DocxParagraph):
                        next_text = (next_b.text or "").strip()
                        if not next_text:
                            j += 1
                            continue
                        if next_text in empty_section_headings or _is_heading_style(next_b):
                            break
                        if next_text != FALLBACK_NO_FINDINGS and len(next_text) > 10:
                            has_content = True
                            break
                        # Fallback-only counts as minimal content; don't remove
                        if next_text == FALLBACK_NO_FINDINGS:
                            has_content = True
                            break
                        j += 1
                if not has_content and j > i + 1:
                    for k in range(i, j):
                        if isinstance(blocks[k], DocxParagraph):
                            to_remove.append(blocks[k]._element)
                i = j
                continue
        i += 1
    for el in to_remove:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def _is_heading_style(p) -> bool:
    """True if paragraph has a heading style."""
    try:
        style = (p.style and p.style.name or "").lower()
        return "heading" in style
    except Exception:
        return False


# Gate C: Allowlist for OK duplicate paragraphs (e.g. FOUO, CISA tagline, table headers)
GATE_C_ALLOWLIST = (
    "fouo",
    "for official use only",
    "cisa",
    "category",
    "vulnerability",
    "option for consideration",
    "time to severe impact",
    "functional loss",
    "backup duration",
    "structural posture",
)


def _normalize_block_for_duplicate(text: str) -> str:
    """Normalize paragraph for duplicate detection: lowercase, trim, collapse spaces, digits -> #."""
    t = (text or "").lower().strip()
    t = re.sub(r"\s+", " ", t)
    t = re.sub(r"\d+", "#", t)
    return t


def gate_c_duplicate_block_detection(doc: Document) -> None:
    """
    Gate C: Duplicate Block Detection (PREVENT COPY/PASTE BLEED).
    Detect repeated paragraphs across different sector sections; fail if found.
    """
    sector_section_markers = ("Electric Power", "Communications", "Information Technology", "Water", "Wastewater")
    seen: dict[str, list[str]] = {}  # normalized -> list of section names where it appeared
    current_section = "preamble"
    for block in iter_block_items(doc):
        if isinstance(block, DocxTable):
            continue
        if isinstance(block, DocxParagraph):
            text = (block.text or "").strip()
            if not text:
                continue
            # Update current_section for sector headings even when short (e.g. "Water", "Wastewater")
            for marker in sector_section_markers:
                if marker in text and _is_heading_style(block):
                    current_section = marker
                    break
            if len(text) < 20:
                continue
            norm = _normalize_block_for_duplicate(text)
            if not norm:
                continue
            for allowed in GATE_C_ALLOWLIST:
                if norm.startswith(allowed) or allowed in norm[:50]:
                    break
            else:
                if norm not in seen:
                    seen[norm] = []
                if current_section not in seen[norm]:
                    seen[norm].append(current_section)
    duplicates = [(norm, sections) for norm, sections in seen.items() if len(sections) > 1]
    if duplicates:
        norm, sections = duplicates[0]
        sample = norm[:120] + "..." if len(norm) > 120 else norm
        print(
            "ERROR: Gate C — Duplicate block (copy bleed) detected. Export blocked.",
            file=sys.stderr,
        )
        print(f"  Repeated text (first 120 chars): {sample!r}", file=sys.stderr)
        print(f"  Appeared in sectors: {', '.join(sections)}", file=sys.stderr)
        print("  Remediation: likely copy bleed in sector narrative builder", file=sys.stderr)
        sys.exit(1)


def sanitize_spof_language(doc):
    """Replace any SPOF/inference language with factual wording. No 'single-path' or 'likely'."""
    # Map exact phrases to factual replacements; then strip any remaining SPOF
    phrase_replacements = [
        ("(SPOF likely)", "Reported sources: 1"),
        ("SPOF likely", "Reported sources: 1"),
    ]
    banned = [
        "SPOF",
        "single point of failure",
    ]
    replacement = "Reported sources: 1"

    for p in doc.paragraphs:
        txt = p.text or ""
        for old_phrase, new_phrase in phrase_replacements:
            txt = txt.replace(old_phrase, new_phrase)
        for phrase in banned:
            if phrase in txt:
                txt = txt.replace(phrase, replacement)
        if txt != (p.text or ""):
            p.text = txt

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text or ""
                for old_phrase, new_phrase in phrase_replacements:
                    txt = txt.replace(old_phrase, new_phrase)
                for phrase in banned:
                    if phrase in txt:
                        txt = txt.replace(phrase, replacement)
                if txt != (cell.text or ""):
                    cell.text = txt


def assert_no_spof_in_output(doc: Document) -> None:
    """Fail if explicit SPOF labels appear in output. Vulnerability narratives may use 'single point of failure' as a risk term."""
    forbidden = ("(SPOF likely)", "SPOF likely")
    for p, _ in iter_paragraphs_and_cells(doc):
        text = (p.text or "").strip()
        for phrase in forbidden:
            if phrase in text:
                raise RuntimeError(
                    f"Output must not contain SPOF language: found {phrase!r} in: {text[:150]!r}..."
                )




# Chart image dimensions (2.82:1) to preserve Excel ratio
CHART_MAX_WIDTH_INCHES = 6.0
CHART_MAX_HEIGHT_INCHES = 6.0 / 2.82  # ≈ 2.13


def replace_chart_anchor(doc: Document, anchor: str, image_path: Path, narrative: str = None) -> bool:
    """
    Find paragraph whose full text equals anchor (or contains it), clear it and insert image.
    Image uses fixed width+height (2.82:1) to preserve aspect. Adds a small spacer
    paragraph after the chart (and after narrative if provided) to prevent collisions.
    Returns True if anchor was found and replaced.
    """
    # Find paragraph containing anchor. Body only (Section C) so title/header are not mutated.
    it = find_anchor_paragraph_exact(doc, anchor, body_only=True)
    try:
        p, _ = next(it)
    except StopIteration:
        it = find_placeholder_or_anchor(doc, anchor)
        try:
            p, _ = next(it)
        except StopIteration:
            return False
    if paragraph_has_drawing(p):
        return False
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(CHART_MAX_WIDTH_INCHES), height=Inches(CHART_MAX_HEIGHT_INCHES))
    body = doc.element.body
    body_children = list(body)
    para_index = index_of_element(body_children, p._element)
    if para_index is None:
        return True
    insert_at = para_index + 1
    if narrative:
        narrative_para = doc.add_paragraph(narrative)
        body.remove(narrative_para._element)
        body.insert(insert_at, narrative_para._element)
        insert_at += 1
    # Spacer paragraph after each chart to prevent collisions
    spacer = doc.add_paragraph()
    body.remove(spacer._element)
    body.insert(insert_at, spacer._element)
    return True


def _print_anchor_debug(doc: Document) -> None:
    """Print all paragraph texts that contain '[[' for debugging missing anchors."""
    for p, _ in iter_paragraphs_and_cells(doc):
        t = (p.text or "").strip()
        if "[[" in t:
            print(f"  [DEBUG] Paragraph containing '[[': {t!r}", file=sys.stderr)


def get_body_children(doc: Document):
    """Return list of body block elements (paragraphs, tables, etc.) for positioning."""
    body = doc.element.body
    return list(body)


def index_of_element(body_children, element) -> int | None:
    for i, child in enumerate(body_children):
        if child is element:
            return i
    return None


SUMMARY_TABLE_HEADERS = (
    "Category",
    "Primary provider identified?",
    "Backup present?",
    "Time to severe impact (hrs)",
    "Recovery time (hrs)",
    "Notes",
)


def insert_summary_table_at_anchor(doc: Document, summary_rows: list[dict]) -> None:
    """
    Find paragraph or table cell containing [[TABLE_SUMMARY]], remove it, insert template summary table.
    D1 format: Category | Primary provider identified? | Backup present? | Time to severe impact (hrs) | Recovery time (hrs) | Notes.
    """
    if not summary_rows:
        summary_rows = [{
            "category": "No summary data available.",
            "primary_provider": "Unknown",
            "backup_present": "No",
            "time_to_severe_impact_hrs": "",
            "recovery_time_hrs": "",
            "notes": NOT_CONFIRMED,
        }]
    body = doc.element.body
    body_children = list(body)
    target_index = None
    target_para = None
    target_row_in_table = None
    for child in body_children:
        if child.tag != qn("w:p"):
            continue
        text_parts = []
        for r in child.iter():
            if r.tag == qn("w:t") and r.text:
                text_parts.append(r.text)
        text = "".join(text_parts)
        if TABLE_ANCHOR in text:
            target_para = child
            target_index = index_of_element(body_children, child)
            break
    if target_index is None:
        for p, parent_cell in find_placeholder_or_anchor(doc, TABLE_ANCHOR):
            if parent_cell is not None:
                tr = p._element.getparent()
                if tr is not None and tr.tag == qn("w:tr"):
                    tbl = tr.getparent()
                    if tbl is not None:
                        target_index = index_of_element(body_children, tbl)
                        target_row_in_table = tr
                        break
    if target_index is None and target_para is None and target_row_in_table is None:
        print(f"ERROR: Anchor {TABLE_ANCHOR!s} not found at insertion time", file=sys.stderr)
        _print_anchor_debug(doc)
        sys.exit(1)
    if target_row_in_table is not None:
        tbl = target_row_in_table.getparent()
        tbl.remove(target_row_in_table)
        target_index = index_of_element(list(body), tbl) + 1
    elif target_para is not None:
        body.remove(target_para)
    body_children = list(body)
    num_cols = len(SUMMARY_TABLE_HEADERS)
    table = doc.add_table(rows=1 + len(summary_rows), cols=num_cols)
    table.style = "Table Grid"
    for c, h in enumerate(SUMMARY_TABLE_HEADERS):
        table.rows[0].cells[c].text = h
    _set_repeat_header_row(table.rows[0])
    for r, row_data in enumerate(summary_rows):
        row_idx = r + 1
        table.rows[row_idx].cells[0].text = sanitize_text(str(row_data.get("category", "")))
        table.rows[row_idx].cells[1].text = sanitize_text(str(row_data.get("primary_provider", "Unknown")))
        table.rows[row_idx].cells[2].text = sanitize_text(str(row_data.get("backup_present", "No")))
        table.rows[row_idx].cells[3].text = sanitize_text(str(row_data.get("time_to_severe_impact_hrs", "")))
        table.rows[row_idx].cells[4].text = sanitize_text(str(row_data.get("recovery_time_hrs", "")))
        table.rows[row_idx].cells[5].text = sanitize_text(str(row_data.get("notes", NOT_CONFIRMED)))
    tbl_element = table._tbl
    body.remove(tbl_element)
    body.insert(target_index, tbl_element)


def build_sources_narrative(
    assessment: dict,
    sla_reliability_for_report: list[dict] | None = None,
    sla_pra_module_enabled: bool = False,
) -> list[str]:
    """
    Build per-category narrative paragraphs for supply sources. When sla_pra_module_enabled is True,
    each paragraph also includes SLA/PRA status; when False, no SLA or PRA language (baseline only).
    When sla_reliability_for_report is provided and module enabled: SLA line is gated by sla_assessed.
    """
    categories = assessment.get("categories") or {}
    label_to_sla = {}
    if sla_pra_module_enabled and sla_reliability_for_report:
        label_to_sla = {
            e.get("topic_label"): e
            for e in sla_reliability_for_report
            if e.get("topic_label")
        }
    paragraphs = []
    for code in CHART_CATEGORIES:
        if code not in categories:
            continue
        inp = categories[code]
        display = CATEGORY_DISPLAY.get(code, code)
        supply = inp.get("supply")
        agreements = inp.get("agreements")
        parts = [f"{display}:"]
        sources = (supply or {}).get("sources") or []
        for i, s in enumerate(sources):
            if not s:
                continue
            provider = (s.get("provider_name") or "").strip() or "Not provided"
            parts.append(f"  Source {i + 1}: Provider {provider}.")
            desc = (s.get("demarcation_description") or "").strip()
            lat, lon = s.get("demarcation_lat"), s.get("demarcation_lon")
            if desc or (lat is not None and lon is not None):
                loc = desc or ""
                if lat is not None and lon is not None:
                    loc = f"{loc} (lat {lat}, lon {lon})".strip() if loc else f"lat {lat}, lon {lon}"
                parts.append(f"    Demarcation: {loc}.")
            ind = s.get("independence")
            if ind:
                parts.append(f"    Independence: {INDEPENDENCE_LABELS.get(ind, ind)}.")
        if sla_pra_module_enabled:
            sla_entry = label_to_sla.get(display) if label_to_sla else None
            if sla_entry is not None:
                if sla_entry.get("sla_assessed"):
                    parts.append("  SLA: " + _format_sla_cell_stakeholder(sla_entry) + ".")
            else:
                sla_text = _sla_summary(agreements)
                if sla_text == "No":
                    parts.append("  SLA: No.")
                else:
                    h = (agreements or {}).get("sla_hours")
                    parts.append(f"  SLA: Yes; target hours: {h}." if h is not None else "  SLA: Yes.")
            pra_text = _pra_summary(agreements)
            if pra_text == "No":
                parts.append("  PRA: No.")
            else:
                cat = (agreements or {}).get("pra_category")
                if cat == "OTHER":
                    other = ((agreements or {}).get("pra_category_other") or "").strip() or "\u2014"
                    parts.append(f"  PRA: Yes; category: Other ({other}).")
                else:
                    parts.append(f"  PRA: Yes; category: {cat}." if cat else "  PRA: Yes.")
        paragraphs.append(sanitize_text(" ".join(parts)))
    return paragraphs


def insert_sources_narrative_at_anchor(
    doc: Document,
    assessment: dict,
    sla_reliability_for_report: list[dict] | None = None,
    sla_pra_module_enabled: bool = False,
) -> None:
    """
    If [[NARRATIVE_SOURCES]] is present, replace it with per-category supply source narrative.
    When sla_pra_module_enabled is False, no SLA or PRA language is included (baseline only).
    """
    body = doc.element.body
    body_children = list(body)
    target_index = None
    target_para = None
    for child in body_children:
        if child.tag != qn("w:p"):
            continue
        text_parts = []
        for r in child.iter():
            if r.tag == qn("w:t") and r.text:
                text_parts.append(r.text)
        text = "".join(text_parts)
        if NARRATIVE_SOURCES_ANCHOR in text:
            target_para = child
            target_index = index_of_element(body_children, child)
            break
    if target_index is None or target_para is None:
        return
    body.remove(target_para)
    narrative_paragraphs = build_sources_narrative(
        assessment, sla_reliability_for_report, sla_pra_module_enabled
    )
    if not narrative_paragraphs:
        narrative_paragraphs = ["No supply source details were provided for any dependency category."]
    for i, text in enumerate(narrative_paragraphs):
        para = doc.add_paragraph(text)
        body.remove(para._element)
        body.insert(target_index + i, para._element)


def insert_executive_summary_at_anchor(
    doc: Document, assessment: dict, executive_summary_brief: str | None = None
) -> None:
    """
    Replace TABLE_SUMMARY anchor with Hotel Fact Sheet brief (2–5 sentences).
    When executive_summary_brief provided: remove anchor para, insert brief.
    Otherwise: remove anchor para only (template-native).
    """
    body = doc.element.body
    body_children = list(body)
    target_para = None
    for child in body_children:
        if child.tag != qn("w:p"):
            continue
        text_parts = []
        for r in child.iter():
            if r.tag == qn("w:t") and r.text:
                text_parts.append(r.text)
        text = "".join(text_parts)
        if TABLE_ANCHOR in text:
            target_para = child
            break
    if target_para is None:
        return
    target_index = index_of_element(body_children, target_para)
    in_exec_summary = False
    if target_index < 20:
        for i in range(max(0, target_index - 10), min(len(body_children), target_index + 5)):
            child = body_children[i]
            if child.tag == qn("w:p"):
                text_parts = []
                for r in child.iter():
                    if r.tag == qn("w:t") and r.text:
                        text_parts.append(r.text)
                t = "".join(text_parts).upper()
                if "EXECUTIVE SUMMARY" in t:
                    in_exec_summary = True
                    break
    if not in_exec_summary:
        return
    body.remove(target_para)
    if executive_summary_brief and str(executive_summary_brief).strip():
        new_para = doc.add_paragraph(sanitize_text(executive_summary_brief.strip()))
        body.remove(new_para._element)
        body.insert(target_index, new_para._element)


def build_dependency_narrative(category_input: dict, display_name: str) -> str:
    """Build narrative text for a dependency category with actual values."""
    requires = category_input.get("requires_service")
    if not requires:
        return f"{display_name}: This facility does not require {display_name.lower()} service for operations."
    
    impact_hours = category_input.get("time_to_impact_hours", 0)
    has_backup = category_input.get("has_backup_any", False)
    loss_no_backup = category_input.get("loss_fraction_no_backup", 0) * 100
    loss_with_backup = category_input.get("loss_fraction_with_backup", 0) * 100 if has_backup else loss_no_backup
    recovery_hours = category_input.get("recovery_time_hours", 0)
    
    narrative = f"The facility requires {display_name.lower()} service for normal operations. "
    narrative += f"Loss of {display_name.lower()} service would impact operations within {impact_hours} hours. "
    
    if has_backup:
        backup_duration = category_input.get("backup_duration_hours", 0)
        if backup_duration is None or (isinstance(backup_duration, str) and not str(backup_duration).strip()):
            narrative += "Backup systems are present; validated duration is not available. "
        else:
            narrative += f"Backup systems provide approximately {backup_duration} hours of continuity. "
        narrative += f"With backup in place, functional capacity loss is estimated at {loss_with_backup:.0f}%. "
        narrative += f"Without backup, functional loss would reach {loss_no_backup:.0f}%. "
    else:
        narrative += f"No backup systems are in place. Functional capacity loss is estimated at {loss_no_backup:.0f}%. "
    
    narrative += f"Recovery to normal operations is estimated to require {recovery_hours} hours following service restoration."
    
    return narrative


# Number of leading body paragraphs to never strip images from (title/cover graphic).
TITLE_PAGE_IMAGE_GUARD_COUNT = 3


def remove_placeholder_images(doc: Document) -> None:
    """Remove placeholder images from the document, particularly in intro/summary sections.
    Keeps chart images at their anchor positions. Never removes images from the first
    TITLE_PAGE_IMAGE_GUARD_COUNT body paragraphs so the title/cover graphic is preserved."""
    body = doc.element.body
    body_children = list(body)

    # Find position of first chart anchor
    first_chart_index = len(body_children)
    for i, child in enumerate(body_children):
        if child.tag != qn("w:p"):
            continue
        text_parts = []
        for r in child.iter():
            if r.tag == qn("w:t") and r.text:
                text_parts.append(r.text)
        text = "".join(text_parts)
        if "[[CHART_" in text:
            first_chart_index = i
            break

    # Never strip images from the first N paragraphs (title page graphic)
    start_index = TITLE_PAGE_IMAGE_GUARD_COUNT
    end_index = min(first_chart_index, 50)

    for i in range(start_index, end_index):
        if i >= len(body_children):
            break
        child = body_children[i]
        if child.tag != qn("w:p"):
            continue

        drawings = child.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing")
        pictures = child.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/picture}pic")

        if drawings or pictures:
            for run_elem in child.findall(".//" + qn("w:r")):
                has_drawing = run_elem.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing") is not None
                has_picture = run_elem.find(".//{http://schemas.openxmlformats.org/drawingml/2006/picture}pic") is not None
                if has_drawing or has_picture:
                    child.remove(run_elem)


# Required anchors for narrative-only template (ADA/report template.docx). Part I: brief + CHART_*; Part II: federal-style anchors.
REQUIRED_ANCHORS = (
    "[[SNAPSHOT_POSTURE]]",
    "[[SNAPSHOT_SUMMARY]]",
    "[[SNAPSHOT_DRIVERS]]",
    "[[SNAPSHOT_MATRIX]]",
    "[[SNAPSHOT_CASCADE]]",
    "[[CHART_ELECTRIC_POWER]]",
    "[[CHART_COMMUNICATIONS]]",
    "[[CHART_INFORMATION_TECHNOLOGY]]",
    "[[CHART_WATER]]",
    "[[CHART_WASTEWATER]]",
    "[[SYNTHESIS]]",
    "[[PRIORITY_ACTIONS]]",
    "[[TABLE_DEPENDENCY_SUMMARY]]",
    "[[STRUCTURAL_PROFILE_SUMMARY]]",
    "[[VULNERABILITY_COUNT_SUMMARY]]",
    "[[VULNERABILITY_BLOCKS]]",
    "[[CROSS_INFRA_ANALYSIS]]",
)
OPTIONAL_ANCHORS = ("[[SLA_PRA_SUMMARY]]", "[[CROSS_DEPENDENCY_SUMMARY]]", "[[NARRATIVE_SOURCES]]", "[[DESIGNATION_SERVICES]]")


def _collect_anchors_in_doc(doc: Document) -> set[str]:
    """Return set of anchor tokens found in document body (paragraphs and table cells). Body-only; no headers/footers."""
    found = set()
    for p, _ in iter_paragraphs_and_cells(doc):
        text = p.text or ""
        for anchor in REQUIRED_ANCHORS + OPTIONAL_ANCHORS:
            if anchor in text:
                found.add(anchor)
    return found


def _assert_template_has_required_anchors(doc: Document) -> None:
    """Exit 1 if any required anchor is missing. Lists missing anchors clearly."""
    found = _collect_anchors_in_doc(doc)
    missing = [a for a in REQUIRED_ANCHORS if a not in found]
    if missing:
        print("ERROR: Template missing required anchors:", file=sys.stderr)
        for a in missing:
            print(f"  - {a}", file=sys.stderr)
        sys.exit(1)


def _assert_chart_anchors_after_visualization(doc: Document) -> None:
    """
    Hard-fail if any chart anchor appears before [[VISUALIZATION_START]].
    Prevents charts from landing in the Hotel Fact Sheet section.
    Skip when template has no VISUALIZATION_START anchor.
    """
    all_text = " ".join(p.text or "" for p, _ in iter_paragraphs_and_cells(doc))
    if VISUALIZATION_START_ANCHOR not in all_text:
        return  # v2 template: no visualization boundary
    blocks = list(iter_block_items(doc))
    vis_start_index = None
    for i, block in enumerate(blocks):
        if isinstance(block, DocxParagraph) and (block.text or "").strip() == VISUALIZATION_START_ANCHOR:
            vis_start_index = i
            break
    if vis_start_index is None:
        raise RuntimeError(f"{VISUALIZATION_START_ANCHOR} not found; cannot validate chart placement")
    for i, block in enumerate(blocks):
        if not isinstance(block, DocxParagraph):
            continue
        text = (block.text or "").strip()
        for anchor in CHART_ANCHORS:
            if text == anchor:
                if i < vis_start_index:
                    raise RuntimeError(
                        f"Chart anchor appears before visualization section: {anchor}"
                    )
                break


def remove_static_chart_captions(doc: Document) -> None:
    """Remove template paragraphs matching 'Table X. ... Dependency Chart' (generator provides captions)."""
    caption_re = re.compile(r"^Table\s+\d+\.\s+.*[Dd]ependency\s+[Cc]hart", re.IGNORECASE)
    for p in list(doc.paragraphs):
        if caption_re.match((p.text or "").strip()):
            remove_paragraph(p)


def remove_empty_paragraphs_after_page_breaks(doc: Document) -> None:
    """
    Remove empty paragraphs that appear immediately after a page break.
    These create blank pages; removing them keeps content tight after each break.
    """
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:p"):
            continue
        if not _element_has_page_break(child):
            continue
        # Remove all immediately following empty paragraphs
        while True:
            nxt = child.getnext()
            if nxt is None or nxt.tag != qn("w:p"):
                break
            para = DocxParagraph(nxt, doc)
            if paragraph_has_drawing(para) or not _is_effectively_empty_paragraph(para) or _paragraph_has_page_break(para):
                break
            body.remove(nxt)


def compress_blank_paragraphs(doc: Document, max_run: int = 1) -> None:
    """
    Remove excess blank paragraphs: allow at most max_run consecutive empty paragraphs.
    Nukes 'page of nothing' between sections.
    Never remove paragraphs that contain a hard page break (w:br w:type=page).
    Never remove paragraphs that contain images/drawings (e.g. sector charts).
    """
    run = 0
    for p in list(doc.paragraphs):
        is_blank = not (p.text or "").strip()
        has_page_break = _paragraph_has_page_break(p)
        if is_blank and not has_page_break and not paragraph_has_drawing(p):
            run += 1
            if run > max_run:
                remove_paragraph(p)
        else:
            run = 0


VOFC_SECTION_HEADING = "Vulnerabilities and Options for Consideration"


def _ensure_vuln_narrative_heading(doc: Document) -> None:
    """
    Ensure vulnerability section heading exists before [[VULNERABILITY_BLOCKS]] (or legacy [[VULN_NARRATIVE]]).
    If template already has SECTOR_REPORTS_HEADING or VOFC_SECTION_HEADING, do nothing.
    """
    for p, _ in iter_paragraphs_and_cells(doc):
        txt = (p.text or "").strip()
        if SECTOR_REPORTS_HEADING in txt or VOFC_SECTION_HEADING in txt or "INFRASTRUCTURE VULNERABILITIES" in txt.upper():
            return  # Already present
    p = find_paragraph_by_exact_text(doc, VULNERABILITY_BLOCKS_ANCHOR, body_only=False)
    if p is None:
        p = find_paragraph_by_exact_text(doc, VULN_NARRATIVE_ANCHOR, body_only=False)
    if p is None:
        return  # Anchor missing (will fail in anchor assertion)
    heading_para = insert_paragraph_before(p, SECTOR_REPORTS_HEADING)
    try:
        heading_para.style = "Heading 2"
    except Exception:
        pass


def remove_placeholder_vofc_tables(doc: Document) -> None:
    """Remove all placeholder VOFC tables from the template.
    These are tables with 'Vulnerability' in the header that contain example data.
    Only the [[VULN_NARRATIVE]] anchor paragraph will remain for insertion of narrative."""
    body = doc.element.body
    tables_to_remove = []
    
    for table in doc.tables:
        # Check if this table has 'Vulnerability' in the first row (header)
        if table.rows and table.rows[0].cells:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            if 'Vulnerability' in headers or 'vulnerability' in str(headers).lower():
                # Check if this table contains the [[VULN_NARRATIVE]] anchor (keep such table; remove placeholders only)
                has_anchor = False
                for row in table.rows:
                    for cell in row.cells:
                        if VULN_NARRATIVE_ANCHOR in cell.text:
                            has_anchor = True
                            break
                
                # Only keep table if it has the anchor; all others are placeholders
                if not has_anchor:
                    tables_to_remove.append(table)
    
    # Remove placeholder tables from the body
    for table in tables_to_remove:
        tbl_element = table._tbl
        body.remove(tbl_element)


def _assert_vulnerability_block_formatting(doc: Document) -> None:
    """Regression guard: no single paragraph may contain both VULNERABILITY and Options for Consideration and Exposure Description (run-on)."""
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        if "VULNERABILITY" in txt and "Options for Consideration" in txt and "Exposure Description" in txt:
            raise RuntimeError(
                "Export failed: vulnerability block formatting regression. "
                "A single paragraph contains VULNERABILITY, Exposure Description, and Options for Consideration "
                "(should be separate paragraphs). Check that vulnerability_blocks text preserves newlines."
            )


VOFC_NO_ITEMS_SENTENCE = (
    "No significant vulnerabilities were identified from current dependency conditions."
)
VOFC_INTRO_PARAGRAPH = (
    "Vulnerabilities and Options for Consideration are generated from "
    "dependency conditions using a standardized evaluation framework."
)
# Template VOFC table: 3 columns only (match template layout)
VOFC_TEMPLATE_HEADERS = ("Category", "Vulnerability", "Option for Consideration")
# Category order for report (match dashboard order)
VOFC_CATEGORY_ORDER = (
    "ELECTRIC_POWER",
    "COMMUNICATIONS",
    "INFORMATION_TECHNOLOGY",
    "WATER",
    "WASTEWATER",
    "CRITICAL_PRODUCTS",
)
VOFC_COL_WIDTHS_INCHES = tuple(VOFC_COL_WIDTHS)  # Alias for code that expects tuple


def _set_cell_bold(cell) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True


def _vofc_items_sorted_for_report(items: list[dict]) -> list[dict]:
    """Sort VOFC items by template category order; cap at 4 per category (engine already caps)."""
    order_idx = {c: i for i, c in enumerate(VOFC_CATEGORY_ORDER)}
    def key(item):
        cat = item.get("category") or ""
        return (order_idx.get(cat, 999), cat, item.get("vofc_id", ""))
    return sorted(items, key=key)


def _vofc_rows_with_no_identified_per_category(items: list[dict]) -> list[dict]:
    """
    Build VOFC table rows: for each category in order, either vulnerability rows (up to 4)
    or one row 'No vulnerabilities identified'. Ensures every category appears.
    """
    by_cat: dict[str, list[dict]] = {}
    for item in _vofc_items_sorted_for_report(items):
        cat = item.get("category") or ""
        if cat not in by_cat:
            by_cat[cat] = []
        if len(by_cat[cat]) < 4:
            by_cat[cat].append(item)
    rows = []
    for code in VOFC_CATEGORY_ORDER:
        if code == "CRITICAL_PRODUCTS":
            continue
        if code not in by_cat or not by_cat[code]:
            rows.append({
                "category": CATEGORY_DISPLAY.get(code, code),
                "vulnerability": "No vulnerabilities identified",
                "option_for_consideration": "",
            })
        else:
            for item in by_cat[code]:
                rows.append({
                    "category": CATEGORY_DISPLAY.get(item.get("category"), item.get("category", "")),
                    "vulnerability": str(item.get("vulnerability", "")),
                    "option_for_consideration": str(item.get("option_for_consideration", "")),
                })
    return rows


def insert_vofc_table_at_anchor(doc: Document, vofc_collection: dict | None) -> None:
    """VOFC table removed; narrative-only export. Use payload.VULN_NARRATIVE and [[VULN_NARRATIVE]]."""
    raise RuntimeError("VOFC table removed; use VULN_NARRATIVE narrative injection.")


# SLA marker color (stakeholder report)
SLA_MTTR_COLOR = "#e87500"
TTI_REF_LINE_COLOR = "#666666"
ALT_REF_LINE_COLOR = "#999999"

# Chart uniformity — Excel aspect ratio ≈2.82:1 (measured from Asset Dependency Visualization.xlsm)
CHART_ASPECT_RATIO = 2.82
CHART_W_PX = 1410
CHART_H_PX = 500  # 1410/500 = 2.82
CHART_DPI = 200
CHART_FIGSIZE = (CHART_W_PX / CHART_DPI, CHART_H_PX / CHART_DPI)
CHART_XLIM = (0, 96)
CHART_YLIM = (0, 100)
# Margins: reserve right for legend, no bbox_inches='tight' (preserves aspect)
CHART_SUBPLOTS_LEFT = 0.08
CHART_SUBPLOTS_RIGHT = 0.82
CHART_SUBPLOTS_TOP = 0.86
CHART_SUBPLOTS_BOTTOM = 0.18

# Refined stroke weights — professional technical reporting (not chunky)
PRIMARY_SOLID_WIDTH = 2.6
SECONDARY_SOLID_WIDTH = 2.2
DASHED_WIDTH = 1.6
THRESHOLD_LINE_WIDTH = 1.4

# Z-order: Primary solid base, secondary solid, dashed on top
Z_SOLID_BASE = 1
Z_SOLID_SECONDARY = 2
Z_DASHED = 4

# PACE chart — Primary solid 2.6; Alternate/Contingency/Emergency solid 2.2 (different colors)
PACE_LINEWIDTHS = {
    "PRIMARY": PRIMARY_SOLID_WIDTH,
    "ALTERNATE": SECONDARY_SOLID_WIDTH,
    "CONTINGENCY": SECONDARY_SOLID_WIDTH,
    "EMERGENCY": SECONDARY_SOLID_WIDTH,
    "NOT_PRESENT": THRESHOLD_LINE_WIDTH,
}
PACE_LINESTYLES = {
    "PRIMARY": "-",
    "ALTERNATE": "-",
    "CONTINGENCY": "-",
    "EMERGENCY": "-",
    "NOT_PRESENT": ":",
}
# High-contrast, print-safe palette (no black; navy distinct from axes)
PACE_COLOR_PRIMARY = "#003f5c"
PACE_COLORS = {
    "PRIMARY": PACE_COLOR_PRIMARY,
    "ALTERNATE": "#1f77b4",
    "CONTINGENCY": "#2ca02c",
    "EMERGENCY": "#9467bd",
    "NOT_PRESENT": "#999999",
}
AXIS_SPINE_COLOR = "#444444"
AXIS_SPINE_WIDTH = 1.0
# zorder: Primary base, Alternate/Contingency/Emergency secondary
PACE_ZORDER = {"PRIMARY": Z_SOLID_BASE, "ALTERNATE": Z_SOLID_SECONDARY, "CONTINGENCY": Z_SOLID_SECONDARY, "EMERGENCY": Z_SOLID_SECONDARY, "NOT_PRESENT": 0}

# Dependency chart — Without backup solid (base), With backup dashed (on top)
DEP_CHART_LW_PRIMARY = PRIMARY_SOLID_WIDTH
DEP_CHART_LW_SECONDARY = DASHED_WIDTH
DEP_CHART_LW_THRESHOLD = THRESHOLD_LINE_WIDTH
DEP_CHART_LW_PRIMARY_THICK = 3.4
DEP_CHART_LW_SECONDARY_THICK = 2.4
_THICK_CURVE_SECTORS = {"WATER", "WASTEWATER"}


def _dep_chart_linewidths_for_sector(sector_label: str) -> tuple[float, float]:
    code = (sector_label or "").strip().upper().replace(" ", "_")
    if code in _THICK_CURVE_SECTORS:
        return (DEP_CHART_LW_PRIMARY_THICK, DEP_CHART_LW_SECONDARY_THICK)
    return (DEP_CHART_LW_PRIMARY, DEP_CHART_LW_SECONDARY)

def _assert_chart_dimensions(path: Path, allow_tight: bool = False) -> None:
    """Fail export if chart dimensions differ from canonical size. When allow_tight=True (bbox_inches='tight'), only require minimum dimensions so legend is not clipped."""
    try:
        from PIL import Image
        with Image.open(path) as img:
            w, h = img.size
    except ImportError:
        import struct
        with open(path, "rb") as f:
            f.seek(16)
            w, h = struct.unpack(">II", f.read(8))
    if allow_tight:
        if w < CHART_W_PX or h < CHART_H_PX:
            raise RuntimeError(
                f"Chart dimensions {w}x{h} below minimum {CHART_W_PX}x{CHART_H_PX}. "
                "Legend may be clipped; check figsize and bbox_inches='tight'."
            )
    elif (w, h) != (CHART_W_PX, CHART_H_PX):
        raise RuntimeError(
            f"Chart dimension mismatch: {w}x{h} expected {CHART_W_PX}x{CHART_H_PX}. "
            "Fix figsize/subplots_adjust to preserve aspect."
        )


def _ensure_curve_complete(points: list[dict], horizon: int = 96) -> list[dict]:
    """Ensure curve covers [0..96] and capacity in [0..100]. Add endpoint at x=96 if missing.
    If capacity_with_backup is missing for a point, use capacity_without_backup for that point (do not drop series)."""
    if not points:
        return [{"t_hours": 0, "capacity_without_backup": 100.0, "capacity_with_backup": 100.0}]
    out = []
    for p in points:
        t = max(0, min(horizon, int(p.get("t_hours", 0))))
        cap_no = max(0, min(100, float(p.get("capacity_without_backup", 100))))
        cap_with_raw = p.get("capacity_with_backup")
        if cap_with_raw is None or (isinstance(cap_with_raw, float) and cap_with_raw != cap_with_raw):
            cap_with = cap_no
        else:
            cap_with = max(0, min(100, float(cap_with_raw)))
        out.append({"t_hours": t, "capacity_without_backup": cap_no, "capacity_with_backup": cap_with})
    if out[-1]["t_hours"] < horizon:
        last = out[-1]
        out.append({
            "t_hours": horizon,
            "capacity_without_backup": last["capacity_without_backup"],
            "capacity_with_backup": last["capacity_with_backup"],
        })
    return out


def _series_equal(no_backup: list[float], with_backup: list[float]) -> bool:
    """True if max(abs(y1-y2)) <= 1 for all points (series effectively equal)."""
    if len(no_backup) != len(with_backup):
        return False
    return all(abs(a - b) <= 1 for a, b in zip(no_backup, with_backup))


# --- PACE model (Communications only) ---

def _to_num(v, fallback: float = 0) -> float:
    if isinstance(v, (int, float)) and v == v:
        return float(v)
    if isinstance(v, str):
        s = str(v).strip().replace("%", "")
        n = float(s) if s else fallback
        return n if n == n else fallback
    return fallback


def _is_layer_viable(layer: dict) -> bool:
    st = (layer or {}).get("system_type")
    return bool(st and st not in ("NONE", "UNKNOWN"))


def _layer_capacity_at_hour(h: int, activate_after: float, sustain_hours: float | None, capacity_pct: float, horizon: int) -> float:
    cap = max(0, min(100, capacity_pct))
    if cap <= 0:
        return 0.0
    if h < activate_after:
        return 0.0
    if sustain_hours is None:
        return cap
    # sustain_hours >= horizon means layer sustains across modeled window.
    if sustain_hours >= horizon:
        return cap
    end = activate_after + sustain_hours
    if h > end:
        return 0.0
    return cap


def _primary_capacity_at_hour(h: int, time_to_impact: float, capacity_after_impact: float) -> float:
    return 100.0 if h < time_to_impact else capacity_after_impact


def _ensure_pace_curve_complete(points: list[dict], horizon: int = 96) -> list[dict]:
    """Ensure curve spans 0..96; sort by t_hours; clamp capacity 0..100."""
    if not points:
        return []
    out = []
    for p in points:
        t = max(0, min(horizon, int(p.get("t_hours", 0))))
        cap = max(0, min(100, float(p.get("capacity_pct", p.get("capacity_without_backup", 0)))))
        out.append({"t_hours": t, "capacity_pct": cap})
    out.sort(key=lambda x: x["t_hours"])
    if out and out[0]["t_hours"] > 0:
        out.insert(0, {"t_hours": 0, "capacity_pct": out[0]["capacity_pct"]})
    if out and out[-1]["t_hours"] < horizon:
        out.append({"t_hours": horizon, "capacity_pct": out[-1]["capacity_pct"]})
    return out


def build_pace_model_from_comm(inp: dict) -> dict:
    """
    Build PACE model from comm_pace_P/A/C/E for Communications chart.
    Returns: { enabled, layers: { PRIMARY: {present, label?, curve}, ... } }
    Backward compat: if no comm_pace_* but legacy curve exists, map Primary=without, Alternate=with.
    """
    horizon = 96
    req = inp.get("requires_service") or inp.get("curve_requires_service")
    if not req:
        return {"enabled": False, "layers": {}}

    # Check for PACE data
    p_raw = inp.get("comm_pace_P") or {}
    a_raw = inp.get("comm_pace_A") or {}
    c_raw = inp.get("comm_pace_C") or {}
    e_raw = inp.get("comm_pace_E") or {}
    has_pace = any(
        _is_layer_viable(x) for x in (p_raw if isinstance(p_raw, dict) else {},
                                       a_raw if isinstance(a_raw, dict) else {},
                                       c_raw if isinstance(c_raw, dict) else {},
                                       e_raw if isinstance(e_raw, dict) else {})
    )

    if not has_pace:
        # Backward compat: build from legacy curve
        legacy_points = build_curve(inp)
        if not legacy_points:
            return {"enabled": False, "layers": {}}
        pts_no = _ensure_curve_complete(legacy_points, horizon)
        curve_primary = [{"t_hours": x["t_hours"], "capacity_pct": x["capacity_without_backup"]} for x in pts_no]
        curve_primary = _ensure_pace_curve_complete(curve_primary, horizon)
        curve_alt = [{"t_hours": x["t_hours"], "capacity_pct": x["capacity_with_backup"]} for x in pts_no]
        curve_alt = _ensure_pace_curve_complete(curve_alt, horizon)
        has_alt = _effective_has_backup(inp)
        return {
            "enabled": True,
            "layers": {
                "PRIMARY": {"present": True, "label": "Primary", "curve": curve_primary},
                "ALTERNATE": {"present": has_alt, "label": "Alternate", "curve": curve_alt if has_alt else []},
                "CONTINGENCY": {"present": False, "label": "Contingency", "curve": []},
                "EMERGENCY": {"present": False, "label": "Emergency", "curve": []},
            },
        }

    t_impact = max(0, min(horizon, _to_num(inp.get("time_to_impact_hours") or inp.get("curve_time_to_impact_hours"), 0)))
    loss_no = max(0, min(100, _to_num(inp.get("loss_fraction_no_backup") or inp.get("curve_loss_fraction_no_backup"), 0) * 100))
    cap_after = 100 - loss_no

    def build_primary():
        pts = [{"t_hours": h, "capacity_pct": _primary_capacity_at_hour(h, t_impact, cap_after)} for h in range(0, horizon + 1)]
        return _ensure_pace_curve_complete(pts, horizon)

    def build_layer(raw: dict, key: str) -> dict:
        if not isinstance(raw, dict) or not _is_layer_viable(raw):
            return {"present": False, "label": {"PRIMARY": "Primary", "ALTERNATE": "Alternate", "CONTINGENCY": "Contingency", "EMERGENCY": "Emergency"}[key], "curve": []}
        activate = max(0, min(horizon, _to_num(raw.get("activate_after_hours"), 0)))
        cap = max(0, min(100, _to_num(raw.get("effective_capacity_pct"), 0)))
        sustain_raw = raw.get("sustain_hours")
        sustain = None
        if sustain_raw is not None and sustain_raw != sustain_raw:
            pass
        elif sustain_raw is not None:
            sustain = max(0, min(horizon, _to_num(sustain_raw, 0)))
            if sustain <= 0 and key != "EMERGENCY":
                return {"present": False, "label": key.capitalize(), "curve": []}
        if key == "EMERGENCY" and (sustain_raw is None or sustain_raw == 0):
            sustain = None
        pts = [{"t_hours": h, "capacity_pct": _layer_capacity_at_hour(h, activate, sustain, cap, horizon)} for h in range(0, horizon + 1)]
        return {"present": True, "label": key.capitalize(), "curve": _ensure_pace_curve_complete(pts, horizon)}

    p_layer = build_primary()
    a_layer = build_layer(a_raw, "ALTERNATE")
    c_layer = build_layer(c_raw, "CONTINGENCY")
    e_layer = build_layer(e_raw, "EMERGENCY")

    # Primary always present when req
    if not p_layer:
        p_layer = build_primary()

    return {
        "enabled": True,
        "layers": {
            "PRIMARY": {"present": True, "label": "Primary", "curve": p_layer},
            "ALTERNATE": a_layer,
            "CONTINGENCY": c_layer,
            "EMERGENCY": e_layer,
        },
    }


def _apply_halo_path_effect(line, linewidth: float) -> None:
    """Subtle white separation when curves overlap — not outline-heavy."""
    line.set_path_effects([pe.Stroke(linewidth=linewidth + 0.8, foreground="white"), pe.Normal()])


def render_comms_pace_chart_png(pace_model: dict, out_path: Path, sector_label: str = "Communications") -> bytes:
    """
    Render Communications PACE chart: 4 series (Primary, Alternate, Contingency, Emergency).
    Present layers: plot curve. Not present: dotted 0% line with "(not present)" in legend.
    High-contrast palette, white halo for overlapping lines, Primary always top layer.
    """
    layers = pace_model.get("layers") or {}
    enabled = pace_model.get("enabled", False)

    fig, ax = plt.subplots(figsize=CHART_FIGSIZE)
    ax.set_xlim(CHART_XLIM)
    ax.set_ylim(CHART_YLIM)
    ax.set_xticks(range(0, 97, 6))
    ax.set_xticks(range(0, 97, 3), minor=True)
    ax.set_yticks(range(0, 101, 10))
    ax.set_xlabel("Time (hours)")
    ax.set_ylabel("Capacity %")
    ax.set_title(sector_label, fontsize=12)
    ax.grid(True, which="major", linewidth=0.6, alpha=0.14, color="#777777")
    ax.grid(True, which="minor", linewidth=0.4, alpha=0.07, color="#999999")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for spine in ax.spines.values():
        spine.set_color(AXIS_SPINE_COLOR)
        spine.set_linewidth(AXIS_SPINE_WIDTH)

    # Plot E, C, A first (lower zorder), then Primary last (highest zorder, top layer)
    order = ["EMERGENCY", "CONTINGENCY", "ALTERNATE", "PRIMARY"]
    if not enabled:
        (ln,) = ax.plot(
            [0, 96], [0, 0],
            linestyle=":", linewidth=PACE_LINEWIDTHS["NOT_PRESENT"],
            color=PACE_COLORS["NOT_PRESENT"], label="PACE not defined",
            zorder=PACE_ZORDER["NOT_PRESENT"],
        )
        _apply_halo_path_effect(ln, PACE_LINEWIDTHS["NOT_PRESENT"])
    else:
        for key in order:
            layer = layers.get(key) or {}
            present = layer.get("present", False)
            label = layer.get("label") or key.capitalize()
            curve = layer.get("curve") or []
            z = PACE_ZORDER.get(key, 0)
            lw = PACE_LINEWIDTHS[key]
            if present and curve:
                t = [x["t_hours"] for x in curve]
                y = [x["capacity_pct"] for x in curve]
                (ln,) = ax.plot(t, y, linestyle=PACE_LINESTYLES[key], linewidth=lw, color=PACE_COLORS[key], label=label, zorder=z)
            else:
                (ln,) = ax.plot(
                    [0, 96], [0, 0],
                    linestyle=":", linewidth=PACE_LINEWIDTHS["NOT_PRESENT"],
                    color=PACE_COLORS["NOT_PRESENT"], label=f"{label} (not present)", zorder=z,
                )
                lw = PACE_LINEWIDTHS["NOT_PRESENT"]
            _apply_halo_path_effect(ln, lw)

    ax.legend(loc="center left", bbox_to_anchor=(1.01, 0.50), frameon=False, fontsize=9)
    fig.tight_layout()
    fig.subplots_adjust(
        left=CHART_SUBPLOTS_LEFT,
        right=CHART_SUBPLOTS_RIGHT,
        top=CHART_SUBPLOTS_TOP,
        bottom=CHART_SUBPLOTS_BOTTOM,
    )
    fig.savefig(out_path, dpi=CHART_DPI, bbox_inches="tight", pad_inches=0.2)
    _assert_chart_dimensions(out_path, allow_tight=True)
    with open(out_path, "rb") as f:
        data = f.read()
    plt.close(fig)
    return data


def build_comms_pace_narrative(pace_model: dict) -> str:
    """
    PACE-specific narrative for Communications: which layers present, earliest severe degradation,
    floor capacity for Emergency, missing layers.
    """
    if not pace_model.get("enabled"):
        return "Communications: PACE plan is not defined. Operational capacity over the 0–96 hour horizon cannot be assessed."

    layers = pace_model.get("layers") or {}
    present_names = []
    missing_names = []
    earliest_severe = 96
    emergency_floor = 100

    for key in ["PRIMARY", "ALTERNATE", "CONTINGENCY", "EMERGENCY"]:
        layer = layers.get(key) or {}
        if layer.get("present") and layer.get("curve"):
            present_names.append(layer.get("label") or key.capitalize())
            curve = layer.get("curve") or []
            for p in curve:
                if p.get("capacity_pct", 100) < 50:
                    earliest_severe = min(earliest_severe, int(p.get("t_hours", 96)))
                if key == "EMERGENCY":
                    emergency_floor = min(emergency_floor, p.get("capacity_pct", 100))
        else:
            missing_names.append(layer.get("label") or key.capitalize())

    if not present_names:
        return "Communications: No PACE layers are configured. PACE plan is not defined."

    present_str = ", ".join(present_names)
    missing_str = ", ".join(missing_names) if missing_names else ""

    s1 = f"Communications PACE posture includes {present_str}."
    if missing_names:
        s1 += f" {missing_str} are not present."
    s2 = f"Earliest severe degradation across present layers occurs by {earliest_severe} hours."
    emergency_layer = layers.get("EMERGENCY") or {}
    if emergency_layer.get("present") and emergency_layer.get("curve"):
        s3 = f"Emergency layer floor capacity is approximately {round(emergency_floor)}% over the horizon."
    else:
        s3 = "Emergency layer is not present."
    if len(present_names) == 1:
        s3 += " PACE depth is limited."
    return f"{s1} {s2} {s3}"


# Dependency chart colors (high-contrast, grayscale-safe)
DEP_COLOR_WITHOUT_BACKUP = "#d62728"
DEP_COLOR_WITH_BACKUP = "#1f77b4"


def build_chart_series(points: list[dict]) -> dict:
    """
    Build ordered series for dependency chart. PRIMARY (baseline truth) = capacity_without_backup;
    secondary (mitigated/alternate) = capacity_with_backup. Used for plotting and tests.
    """
    points = _ensure_curve_complete(points or [], horizon=96)
    if not points:
        points = [{"t_hours": t, "capacity_without_backup": 100.0, "capacity_with_backup": 100.0} for t in [0, 96]]
    return {
        "t": [x["t_hours"] for x in points],
        "primary_y": [x["capacity_without_backup"] for x in points],
        "secondary_y": [x["capacity_with_backup"] for x in points],
        "labels": ("Without backup", "With backup"),
    }


def render_dependency_chart_png(
    sector_label: str,
    points: list[dict],
    meta: dict,
    out_path: Path,
) -> bytes:
    """
    Centralized chart renderer: uniform style, both series visible, deterministic.
    PRIMARY line = baseline (capacity_without_backup); secondary = capacity_with_backup.
    meta: { tti_hours, alt_sust_hours, sla_mttr_hours } for reference lines.
    Returns PNG bytes. Writes to out_path.
    """
    series = build_chart_series(points)
    t = series["t"]
    no_backup = series["primary_y"]
    with_backup = series["secondary_y"]
    meta = meta or {}
    tti = meta.get("tti_hours")
    alt_sust = meta.get("alt_sust_hours")
    sla_mttr = meta.get("sla_mttr_hours")
    lw_primary, lw_secondary = _dep_chart_linewidths_for_sector(sector_label)

    fig, ax = plt.subplots(figsize=CHART_FIGSIZE)
    ax.set_xlim(CHART_XLIM)
    ax.set_ylim(CHART_YLIM)
    ax.set_xticks(range(0, 97, 6))
    ax.set_xticks(range(0, 97, 3), minor=True)
    ax.set_yticks(range(0, 101, 10))
    ax.set_xlabel("Time (hours)")
    ax.set_ylabel("Capacity %")
    ax.set_title(sector_label, fontsize=12)
    ax.grid(True, which="major", linewidth=0.6, alpha=0.14, color="#777777")
    ax.grid(True, which="minor", linewidth=0.4, alpha=0.07, color="#999999")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for spine in ax.spines.values():
        spine.set_color(AXIS_SPINE_COLOR)
        spine.set_linewidth(AXIS_SPINE_WIDTH)

    # Solid first (base), dashed last (on top). Without backup = solid primary; With backup = dashed, on top.
    (ln_no,) = ax.plot(
        t, no_backup,
        linestyle="-",
        linewidth=lw_primary,
        color=DEP_COLOR_WITHOUT_BACKUP,
        label="Without backup",
        zorder=Z_SOLID_BASE,
    )
    _apply_halo_path_effect(ln_no, lw_primary)

    (ln_with,) = ax.plot(
        t, with_backup,
        linestyle=(0, (6, 3)),
        linewidth=lw_secondary,
        color=DEP_COLOR_WITH_BACKUP,
        label="With backup",
        zorder=Z_DASHED,
    )
    _apply_halo_path_effect(ln_with, lw_secondary)

    if tti is not None and isinstance(tti, (int, float)) and 0 <= tti <= 96:
        ax.axvline(x=tti, color=TTI_REF_LINE_COLOR, linestyle=":", linewidth=DEP_CHART_LW_THRESHOLD, alpha=0.8)
    # Don't draw a threshold line at 96h: at horizon it implies "indefinite within model".
    if alt_sust is not None and isinstance(alt_sust, (int, float)) and 0 <= alt_sust < 96:
        ax.axvline(x=alt_sust, color=ALT_REF_LINE_COLOR, linestyle="-.", linewidth=DEP_CHART_LW_THRESHOLD, alpha=0.7)
    if sla_mttr is not None and isinstance(sla_mttr, (int, float)) and sla_mttr == sla_mttr and 0 <= sla_mttr <= 96:
        ax.axvline(x=sla_mttr, color=SLA_MTTR_COLOR, linestyle="--", linewidth=DEP_CHART_LW_THRESHOLD, label="SLA MTTR-Max")

    ax.legend(loc="center left", bbox_to_anchor=(1.01, 0.50), frameon=False, fontsize=9)
    fig.tight_layout()
    fig.subplots_adjust(
        left=CHART_SUBPLOTS_LEFT,
        right=CHART_SUBPLOTS_RIGHT,
        top=CHART_SUBPLOTS_TOP,
        bottom=CHART_SUBPLOTS_BOTTOM,
    )
    fig.savefig(out_path, dpi=CHART_DPI, bbox_inches="tight", pad_inches=0.2)
    _assert_chart_dimensions(out_path, allow_tight=True)
    with open(out_path, "rb") as f:
        data = f.read()
    plt.close(fig)
    return data


def generate_chart(
    points: list[dict],
    title: str,
    out_path: Path,
    sla_mttr_hours: int | float | None = None,
    tti_hours: int | float | None = None,
    alt_sust_hours: int | float | None = None,
) -> None:
    """
    Write a PNG chart via render_dependency_chart_png.
    Uniform style: figsize per 2.82:1, dpi=200, xlim(0,96), ylim(0,100), legend center left.
    """
    meta = {"sla_mttr_hours": sla_mttr_hours, "tti_hours": tti_hours, "alt_sust_hours": alt_sust_hours}
    render_dependency_chart_png(title, points, meta, out_path)


def append_energy_dependency_section(doc: Document, energy_dependency: dict) -> None:
    """
    Append Energy dependency section when energy_dependency payload is present.
    Prefer structured findings when available; fallback to vulnerability_blocks, then themedFindings.
    Skips section entirely if no content (no empty headings).
    """
    data_blocks = energy_dependency.get("dataBlocks") or []
    vuln_blocks = energy_dependency.get("vulnerability_blocks") or []
    themed = energy_dependency.get("themedFindings") or []
    knowledge_gaps = energy_dependency.get("knowledgeGaps") or []
    has_data_blocks = any(
        b.get("type") == "narrative" and (b.get("text") or "").strip()
        or b.get("type") == "list" and (b.get("items") or [])
        or b.get("type") == "table" and (b.get("rows") or [])
        for b in data_blocks
    )
    if not has_data_blocks and not vuln_blocks and not themed and not knowledge_gaps:
        return

    doc.add_paragraph()
    h = doc.add_paragraph(sanitize_text("Energy \u2014 Dependency Data"))
    h.runs[0].bold = True
    set_paragraph_keep_with_next(h)
    for block in data_blocks:
        btype = block.get("type")
        title = sanitize_text(block.get("title") or "")
        if title:
            doc.add_paragraph(title)
        if btype == "narrative":
            doc.add_paragraph(sanitize_text(block.get("text") or ""))
        elif btype == "list":
            for item in block.get("items") or []:
                doc.add_paragraph(sanitize_text(f"\u2022 {item}"))
        elif btype == "table":
            headers = block.get("headers") or []
            rows = block.get("rows") or []
            if headers and rows is not None:
                t = doc.add_table(rows=1 + len(rows), cols=len(headers))
                apply_table_grid_style(t)
                for c, hdr in enumerate(headers):
                    t.cell(0, c).text = sanitize_text(str(hdr))
                for r, row in enumerate(rows):
                    for c, cell in enumerate(row):
                        if c < len(headers):
                            t.cell(r + 1, c).text = sanitize_text(str(cell)) if cell is not None else ""
                set_table_header_repeat(t, 0)
                set_table_rows_cant_split(t)
                doc.add_paragraph()

    # Structural findings: prefer vulnerability_blocks; fallback to themedFindings as blocks when absent.
    blocks = energy_dependency.get("vulnerability_blocks") or []
    if not blocks:
        themed = energy_dependency.get("themedFindings") or []
        blocks = [
            {"title": t.get("title") or "", "narrative": t.get("narrative") or "", "ofcs": [], "references": []}
            for t in themed[:GATE_D_MAX_FINDINGS_PER_SECTOR]
        ]
    if blocks:
        doc.add_paragraph()
        h2 = doc.add_paragraph("Structural findings")
        h2.runs[0].bold = True
        set_paragraph_keep_with_next(h2)
        for block in blocks:
            if block.get("title") or block.get("narrative") or block.get("ofcs") or block.get("references"):
                render_vulnerability_block(doc, block)
        themed = energy_dependency.get("themedFindings") or []
        if len(themed) > GATE_D_MAX_FINDINGS_PER_SECTOR and not energy_dependency.get("vulnerability_blocks"):
            doc.add_paragraph(sanitize_text(GATE_D_TRUNCATION_NOTE))

    knowledge_gaps = energy_dependency.get("knowledgeGaps") or []
    if knowledge_gaps:
        doc.add_paragraph()
        h3 = doc.add_paragraph("Knowledge gaps")
        h3.runs[0].bold = True
        for g in knowledge_gaps[:6]:
            title = sanitize_text(g.get("title") or "")
            desc = sanitize_text(g.get("description") or "")
            if title or desc:
                doc.add_paragraph(sanitize_text(f"\u2022 {title} \u2014 {desc}"))


def append_dependency_sections(doc: Document, sections: list) -> None:
    """
    Append dependency sections (Communications, IT, Water, Wastewater) with
    Structural findings and Knowledge gaps. Prefer vulnerability_blocks; fallback to themedFindings as blocks.
    """
    for section in sections or []:
        name = section.get("name") or ""
        if not name:
            continue
        blocks = section.get("vulnerability_blocks") or []
        if not blocks:
            themed = section.get("themedFindings") or []
            blocks = [
                {"title": t.get("title") or "", "narrative": t.get("narrative") or "", "ofcs": [], "references": []}
                for t in themed[:GATE_D_MAX_FINDINGS_PER_SECTOR]
            ]
        gaps = section.get("knowledgeGaps") or []
        if not blocks and not gaps:
            continue
        doc.add_paragraph()
        h = doc.add_paragraph(sanitize_text(f"{name} \u2014 Dependency Assessment"))
        h.runs[0].bold = True
        set_paragraph_keep_with_next(h)

        # IT section: External Critical Services table + cascade narrative (before vulnerabilities)
        external_services = section.get("external_services") or []
        cascade_narrative = (section.get("cascade_narrative") or "").strip()
        if external_services or cascade_narrative:
            if external_services:
                doc.add_paragraph()
                h_ext = doc.add_paragraph("External Critical Services")
                h_ext.runs[0].bold = True
                set_paragraph_keep_with_next(h_ext)
                headers = ("Service", "Type", "Supported Functions", "Continuity", "Likely Cascading Effect")
                t = doc.add_table(rows=1 + len(external_services), cols=len(headers))
                apply_table_grid_style(t)
                for c, hdr in enumerate(headers):
                    t.cell(0, c).text = sanitize_text(str(hdr))
                set_table_header_repeat(t, 0)
                set_table_rows_cant_split(t)
                for r, svc in enumerate(external_services):
                    t.cell(r + 1, 0).text = sanitize_text(svc.get("name") or "")
                    t.cell(r + 1, 1).text = sanitize_text(svc.get("service_type") or "")
                    funcs = svc.get("supports_functions") or []
                    t.cell(r + 1, 2).text = sanitize_text(", ".join(funcs) if isinstance(funcs, list) else str(funcs))
                    t.cell(r + 1, 3).text = sanitize_text(svc.get("resilience") or "")
                    t.cell(r + 1, 4).text = sanitize_text(svc.get("cascade_effect") or "")
                doc.add_paragraph()
            if cascade_narrative:
                doc.add_paragraph(sanitize_text(cascade_narrative))
        elif "Information Technology" in name and not external_services:
            doc.add_paragraph()
            doc.add_paragraph(sanitize_text("No external critical services were identified."))

        if blocks:
            doc.add_paragraph()
            h2 = doc.add_paragraph("Structural findings")
            h2.runs[0].bold = True
            set_paragraph_keep_with_next(h2)
            for block in blocks:
                if block.get("title") or block.get("narrative") or block.get("ofcs") or block.get("references"):
                    render_vulnerability_block(doc, block)
            themed = section.get("themedFindings") or []
            if len(themed) > GATE_D_MAX_FINDINGS_PER_SECTOR and not section.get("vulnerability_blocks"):
                doc.add_paragraph(sanitize_text(GATE_D_TRUNCATION_NOTE))
        if gaps:
            doc.add_paragraph()
            h3 = doc.add_paragraph("Knowledge gaps")
            h3.runs[0].bold = True
            for g in gaps[:6]:
                title = sanitize_text(g.get("title") or "")
                desc = sanitize_text(g.get("description") or "")
                if title or desc:
                    doc.add_paragraph(sanitize_text(f"\u2022 {title} \u2014 {desc}"))


def _count_rendered_vulnerability_blocks(doc: Document) -> int:
    """Count vulnerability blocks actually rendered in the document (ADA_Vuln_Header paragraphs)."""
    count = 0
    for p, _ in iter_paragraphs_and_cells(doc):
        try:
            if getattr(p, "style", None) and getattr(p.style, "name", None) == "ADA_Vuln_Header":
                count += 1
        except Exception:
            continue
    return count


def _correct_vulnerability_count_summary(doc: Document, rendered_count: int) -> None:
    """Replace 'Total findings: N' in the vulnerability count summary paragraph with the rendered block count."""
    pattern = re.compile(r"Total findings:\s*\d+", re.IGNORECASE)
    replacement = f"Total findings: {rendered_count}"
    for p, _ in iter_paragraphs_and_cells(doc):
        text = p.text or ""
        if "Total findings:" in text and pattern.search(text):
            new_text = pattern.sub(replacement, text)
            if new_text != text:
                p.clear()
                p.add_run(sanitize_text(new_text))
            break


def _count_vulnerability_blocks(data: dict) -> int:
    """Count vulnerability blocks from payload (for fallback and QC). Prefer report_vm.part2.vulnerabilities, then report_themed_findings."""
    vm = data.get("report_vm") or {}
    part2 = vm.get("part2") or {}
    vm_vulns = part2.get("vulnerabilities") or []
    if vm_vulns:
        return len(vm_vulns)
    assessment = data.get("assessment") or {}
    tf_count = 0
    for cat_id in _PART2_CATEGORY_ORDER:
        cat = (assessment.get("categories") or {}).get(cat_id) or {}
        arr = (cat.get("report_themed_findings") or [])
        for item in arr:
            if item and isinstance(item, dict):
                title = (item.get("title") or "").strip()
                narrative = (item.get("narrative") or "").strip()
                ofc = (item.get("ofcText") or "").strip()
                if title and (narrative or ofc):
                    tf_count += 1
    if tf_count > 0:
        return tf_count
    count = 0
    ed = data.get("energy_dependency") or {}
    for b in (ed.get("vulnerability_blocks") or []):
        if isinstance(b, dict) and (b.get("title") or b.get("narrative") or b.get("ofcs")):
            count += 1
    for sec in (data.get("dependency_sections") or []):
        for b in (sec.get("vulnerability_blocks") or []):
            if isinstance(b, dict) and (b.get("title") or b.get("narrative") or b.get("ofcs")):
                count += 1
    # Fallback: vulnerability_blocks as pre-rendered string may indicate content
    if count == 0 and data.get("vulnerability_blocks") and isinstance(data["vulnerability_blocks"], str):
        if len((data["vulnerability_blocks"] or "").strip()) > 100:
            count = 1  # Assume at least one block when substantial string present
    return count


def _render_report(data: dict, work_path: Path, template_path: Path, skip_canonical_template_check: bool = False) -> None:
    """
    Core report generation: validate payload, generate charts, load template, inject content, save output.docx.
    Used by main() (CLI) and run_from_payload() (serverless/API). When skip_canonical_template_check is True,
    the template path is not required to end with ADA/report template.docx (e.g. for Vercel serverless).
    """
    # Strict template path: fail if missing; always log absolute path (no silent fallback).
    template_abs = template_path.resolve()
    print(f"Template: {template_abs}", file=sys.stderr)
    if not template_path.is_file():
        raise FileNotFoundError(f"Template not found: {template_path}")
    if not skip_canonical_template_check:
        CANONICAL_TEMPLATE_SUFFIX = "ADA" + os.sep + "report template.docx"
        CANONICAL_TEMPLATE_SUFFIX_ALT = "ADA/report template.docx"
        normalized = str(template_path.resolve()).replace("\\", "/")
        if not (normalized.endswith(CANONICAL_TEMPLATE_SUFFIX_ALT) or normalized.endswith(CANONICAL_TEMPLATE_SUFFIX)):
            raise ValueError(f"Wrong template: export must use ADA/report template.docx. Got: {template_path}")
    if isinstance(data, dict) and "assessment" in data:
        assessment = data["assessment"]
        vofc_collection = data.get("vofc_collection")
        sla_pra_module_enabled = data.get("sla_pra_module_enabled") is True
        sla_reliability_for_report = data.get("sla_reliability_for_report") or [] if sla_pra_module_enabled else []
        executive_snapshot = data.get("executive_snapshot")
        synthesis = data.get("synthesis")
    else:
        assessment = data
        vofc_collection = None
        sla_pra_module_enabled = False
        sla_reliability_for_report = []
        executive_snapshot = None
        synthesis = None

    vm = data.get("report_vm") if isinstance(data, dict) else None
    use_vm = vm is not None and isinstance(vm, dict)
    part2 = (vm.get("part2") or {}) if use_vm else {}

    _debug = os.environ.get("ADA_REPORTER_DEBUG", "").strip() in ("1", "true", "yes")
    if _debug:
        print("[reporter] template OK:", template_path, file=sys.stderr)

    categories = assessment.get("categories") or {}
    # B1: Validate chart points before rendering; fail if missing
    _validate_chart_points_for_export(assessment)
    # Part D: Validate required metrics before rendering; fail export if missing
    _validate_sector_metrics_for_export(assessment)

    # Map category display name -> SLA entry for charts (SLA marker only when assessed + YES + MTTR)
    label_to_sla = {}
    if sla_reliability_for_report:
        label_to_sla = {
            e.get("topic_label"): e
            for e in sla_reliability_for_report
            if e.get("topic_label")
        }
    # Generate chart image for every category
    # COMMUNICATIONS: PACE chart (4 layers); others: render_dependency_chart_png
    chart_paths = {}
    for code in CHART_CATEGORIES:
        display = CATEGORY_DISPLAY[code]
        cat_inp = categories.get(code) or {}
        png_path = work_path / f"chart_{code}.png"
        if code == "COMMUNICATIONS":
            pace_model = build_pace_model_from_comm(cat_inp)
            render_comms_pace_chart_png(pace_model, png_path, sector_label=display)
        else:
            points = build_curve(cat_inp)
            sla_entry = label_to_sla.get(display)
            mttr_hours = None
            if sla_entry and sla_entry.get("sla_assessed") and sla_entry.get("sla_in_place") == "YES":
                m = sla_entry.get("mttr_max_hours")
                if m is not None and isinstance(m, (int, float)) and m == m:
                    mttr_hours = m
            tti = cat_inp.get("time_to_impact_hours") or cat_inp.get("curve_time_to_impact_hours")
            alt_sust = cat_inp.get("backup_duration_hours") or cat_inp.get("curve_backup_duration_hours") if _effective_has_backup(cat_inp) else None
            generate_chart(points, display, png_path, sla_mttr_hours=mttr_hours, tti_hours=tti, alt_sust_hours=alt_sust)
        chart_paths[code] = png_path
    if _debug:
        print("[reporter] charts generated", file=sys.stderr)

    # 1) Load template (render FROM template; do not build doc from scratch)
    doc = Document(str(template_path))
    if _debug:
        print("[reporter] doc loaded", file=sys.stderr)
    # Narrow body margins for space efficiency (do not modify header/footer)
    section = doc.sections[0]
    narrow = Inches(0.75)
    section.top_margin = narrow
    section.bottom_margin = narrow
    section.left_margin = narrow
    section.right_margin = narrow
    remove_dev_anchor_block(doc)
    remove_ghost_sector_lines(doc)
    rename_annex_phantom_headings(doc)
    collapse_consecutive_pagebreaks(doc)
    remove_blank_pages_after_section_b(doc)
    ensure_toc_remains_on_one_page(doc)

    # 1a) Fail fast: assert template has required anchors
    _assert_template_has_required_anchors(doc)
    _assert_chart_anchors_after_visualization(doc)

    # 1b) Defensive: remove placeholder VOFC tables, images, and static chart captions before insertions
    remove_placeholder_vofc_tables(doc)
    remove_placeholder_images(doc)
    remove_static_chart_captions(doc)

    # 1b2) Remove C. Sector Analysis heading only; keep [[CHART_*]] anchors for insertion.
    remove_sector_analysis_chart_block(doc)
    # 1b2b) Insert chart images at each [[CHART_*]] anchor when template has chart anchors (hard-fail if file missing or anchor count != 1).
    if any(_doc_contains_anchor(doc, a) for a in CHART_ANCHORS):
        insert_charts_at_anchors(doc, chart_paths)
    remove_orphan_sector_headings_in_part1(doc)
    remove_orphaned_page_breaks_before_section_d(doc)

    # 1b3) Force page break before PART II – TECHNICAL ANNEX (deterministic, no layout drift)
    ensure_part2_starts_new_page(doc)

    # 1c) Ensure Part II vulnerability heading when template has [[VULNERABILITY_BLOCKS]] (optional)
    if _doc_contains_anchor(doc, VULNERABILITY_BLOCKS_ANCHOR):
        _ensure_vuln_narrative_heading(doc)

    asset = assessment.get("asset") or {}
    try:
        replace_all_text_placeholders(doc, asset)
    except RuntimeError as e:
        if "MISSING_PSA_PHONE" in str(e):
            print(str(e), file=sys.stderr)
            sys.exit(1)
        raise

    snapshot_mode = bool(executive_snapshot and isinstance(executive_snapshot, dict))

    # 2a) Inject executive snapshot (curve-driven) when provided
    if snapshot_mode:
        inject_executive_snapshot_at_anchors(doc, executive_snapshot)
        if _debug:
            print("[reporter] executive_snapshot injected", file=sys.stderr)

    # 2a2) Inject Cross-Infrastructure Synthesis when provided.
    # In snapshot mode, suppress legacy synthesis block to avoid duplicating executive content.
    if not snapshot_mode and synthesis and isinstance(synthesis, dict):
        inject_synthesis_at_anchor(doc, synthesis, assessment=assessment)
    elif snapshot_mode:
        inject_synthesis_at_anchor(doc, {}, assessment=assessment)

    # 2a3) Inject Priority Actions when provided
    priority_actions = data.get("priority_actions")
    if priority_actions and isinstance(priority_actions, dict):
        inject_priority_actions_at_anchor(doc, priority_actions)

    # 2) Remove unexpected table directly under "EXECUTIVE SUMMARY"
    remove_unexpected_exec_summary_table(doc)

    # 2b) Hotel Fact Sheet: in snapshot mode, remove legacy TABLE_SUMMARY block; otherwise inject brief.
    exec_brief = None if snapshot_mode else ((executive_snapshot or {}).get("executive_summary_brief") if executive_snapshot else None)
    insert_executive_summary_at_anchor(doc, assessment, executive_summary_brief=exec_brief)
    if snapshot_mode:
        remove_legacy_executive_duplicate_paragraphs(doc)

    # 3) Sector narrative now in Part II per-sector Chart Synopsis (skip Part I fill)

    # 4) Insert IT scope clarification under "Information Technology" section heading (if present)
    insert_it_scope_clarification(doc)

    # 5) Part II: Summary table at [[TABLE_DEPENDENCY_SUMMARY]]
    if use_vm and part2.get("dependency_summary_rows"):
        summary_rows = _normalize_dependency_summary_rows(part2["dependency_summary_rows"])
    else:
        summary_rows = _normalize_dependency_summary_rows(build_summary(assessment))
    build_summary_table_at_anchor(doc, summary_rows, anchor=TABLE_DEPENDENCY_SUMMARY_ANCHOR, assessment_json=data, part2=part2 if use_vm else None)

    # 6) Part I: charts only at CHART_* (Section C Operational Capability Curves). No INFRA_* in Part I.
    energy_dependency = data.get("energy_dependency")
    dependency_sections = data.get("dependency_sections") or []

    ensure_sector_reports_heading(doc)

    # Keep Annex Overview table and Sector Reports contiguous (prevents template-inserted blank page).
    remove_page_breaks_between_annex_table_and_sector_reports(doc)

    # 6a) Part II sector pages at [[INFRA_*]] (preferred) or [[VULN_NARRATIVE]] (compatibility path).
    inject_infra_at_anchors(doc, data, chart_paths, assessment, energy_dependency, dependency_sections)

    # 7) Part II: federal-style anchors (structural profile, vulnerability count, vulnerability blocks, cross-infra; Annex ends with CROSS_INFRA_ANALYSIS)
    cross_dependency_enabled = _is_cross_dependency_enabled(data, assessment)
    for anchor, key in [
        (STRUCTURAL_PROFILE_SUMMARY_ANCHOR, "structural_profile_summary"),
        (VULNERABILITY_COUNT_SUMMARY_ANCHOR, "vulnerability_count_summary"),
        (VULNERABILITY_BLOCKS_ANCHOR, "vulnerability_blocks"),
        (CROSS_INFRA_ANALYSIS_ANCHOR, "cross_infra_analysis"),
    ]:
        if _doc_contains_anchor(doc, anchor):
            val = data.get(key) or ""
            if key == "vulnerability_blocks":
                n = _count_anchor_occurrences(doc, VULNERABILITY_BLOCKS_ANCHOR, body_only=True)
                if n != 1:
                    raise RuntimeError(
                        f"Template must contain exactly one [[VULNERABILITY_BLOCKS]] anchor (body paragraphs); found {n}. "
                        "Duplicate anchors cause insertion in the wrong place and boundary verification to fail."
                    )
                vm_vulns = part2.get("vulnerabilities") if use_vm else None

                canonical = data.get("canonicalVulnBlocks")
                if isinstance(canonical, list) and len(canonical) > 0:
                    # Single source of truth: do NOT re-derive or rephrase. Render exactly from
                    # canonicalVulnBlocks (title, narrative, ofcs) so DOCX matches Online Summary 1:1.
                    blocks = [_canonical_vuln_block_to_block(c) for c in canonical]
                    blocks = [b for b in blocks if b and (b.get("title") or "").strip()]
                    if blocks:
                        _render_structured_vulnerabilities(doc, anchor, blocks, body_only=True)
                        if os.environ.get("EMIT_VULN_MANIFEST"):
                            data["_vuln_manifest_blocks"] = [
                                {"title": (b.get("title") or "").strip(), "narrative": (b.get("narrative") or "").strip(), "ofcText": "\n".join(b.get("ofcs") or [])}
                                for b in blocks
                            ]
                    else:
                        inject_text_at_anchor(doc, anchor, NO_VULNERABILITIES_TRIGGERED, body_only=True)
                elif _has_derived_findings(data):
                    # Fallback: render from assessment.sessions.<domain>.derived; never recompute.
                    derived_list = load_derived_findings_from_payload(data)
                    blocks = [_derived_finding_to_block(f) for f in derived_list]
                    blocks = [b for b in blocks if b and (b.get("title") or "").strip()]
                    if blocks:
                        _render_structured_vulnerabilities(doc, anchor, blocks, body_only=True)
                        if os.environ.get("EMIT_VULN_MANIFEST"):
                            data["_vuln_manifest_blocks"] = [
                                {"title": (b.get("title") or "").strip(), "narrative": (b.get("narrative") or "").strip(), "ofcText": "\n".join(b.get("ofcs") or [])}
                                for b in blocks
                            ]
                    else:
                        inject_text_at_anchor(doc, anchor, NO_VULNERABILITIES_TRIGGERED, body_only=True)
                elif use_vm and isinstance(vm_vulns, list) and len(vm_vulns) > 0:
                    # Back-compat: payload has no sessions.derived but has part2.vulnerabilities.
                    _render_structured_vulnerabilities(doc, anchor, vm_vulns, body_only=True)
                else:
                    try:
                        themed = get_part2_findings(assessment)
                        inject_themed_findings_at_anchor(doc, anchor, themed, body_only=True)
                    except Exception:
                        # Truthful no-findings path: avoid hard crash when legacy payloads omit
                        # report_themed_findings and no canonical/session/part2 vulnerabilities exist.
                        inject_text_at_anchor(doc, anchor, NO_VULNERABILITIES_TRIGGERED, body_only=True)
            else:
                if key == "cross_infra_analysis":
                    if cross_dependency_enabled:
                        inject_text_at_anchor(
                            doc,
                            anchor,
                            _humanize_cross_infra_analysis_text((val or "").replace("\u00a0", " ").strip()),
                            body_only=True,
                        )
                    else:
                        _remove_cross_infra_analysis_section(doc)
                elif key == "structural_profile_summary":
                    profile = (val or "").replace("\u00a0", " ").strip()
                    if _is_generic_structural_profile_summary(profile):
                        profile = _build_structural_profile_summary_from_assessment(assessment)
                    inject_text_at_anchor(doc, anchor, profile, body_only=True)
                else:
                    inject_text_at_anchor(doc, anchor, (val or "").replace("\u00a0", " ").strip(), body_only=True)

    # Guard: after injecting [[VULNERABILITY_BLOCKS]], next section heading must still appear after last vuln content (no bleed)
    if cross_dependency_enabled:
        _verify_vulnerability_section_boundary(doc)

    # Final pass: prevent accidental back-to-back page breaks introduced by anchor/template drift.
    collapse_consecutive_pagebreaks(doc)

    # Total findings count must match rendered vulnerability blocks (not raw payload length)
    rendered_vuln_count = _count_rendered_vulnerability_blocks(doc)
    _correct_vulnerability_count_summary(doc, rendered_vuln_count)
    fallback_titles: list[str] = []
    canonical_vulns = data.get("canonicalVulnBlocks")
    if isinstance(canonical_vulns, list):
        fallback_titles = [str(v.get("title") or "").strip() for v in canonical_vulns if isinstance(v, dict)]
    if not fallback_titles:
        vm_vulns = (((data.get("report_vm") or {}).get("part2") or {}).get("vulnerabilities") or [])
        if isinstance(vm_vulns, list):
            fallback_titles = [str(v.get("title") or "").strip() for v in vm_vulns if isinstance(v, dict)]
    if not fallback_titles and isinstance(data.get("vulnerability_blocks"), str):
        for line in data["vulnerability_blocks"].splitlines():
            m = re.match(r"^\s*VULNERABILITY\s+\d+\s+(.+)$", (line or "").strip(), flags=re.I)
            if m and m.group(1).strip():
                fallback_titles.append(m.group(1).strip())
        fallback_titles = fallback_titles[:12]
    if not fallback_titles:
        try:
            derived_f = load_derived_findings_from_payload(data)
            fallback_titles = [str(f.get("title") or "").strip() for f in derived_f if str(f.get("title") or "").strip()]
        except Exception:
            fallback_titles = []
    build_part1_executive_hero(
        doc,
        executive_snapshot,
        assessment,
        rendered_vuln_count,
        fallback_action_titles=fallback_titles,
    )

    # Designation block: list services this facility provides (from assessment.asset.services_provided)
    if _doc_contains_anchor(doc, DESIGNATION_SERVICES_ANCHOR):
        assessment = data.get("assessment") or {}
        asset = assessment.get("asset") or {}
        services = [s.strip() for s in (asset.get("services_provided") or []) if (s or "").strip()]
        if services:
            designation_text = "This facility provides: " + ", ".join(services) + "."
        else:
            designation_text = "None specified."
        inject_text_at_anchor(doc, DESIGNATION_SERVICES_ANCHOR, designation_text, body_only=True)

    items = (vofc_collection or {}).get("items") or []
    _, cyber_rows = _clean_ofcs_and_build_rows(items)
    cyber_capped, cyber_truncated = _cap_vofc_rows_per_sector(
        cyber_rows, max_per=REPORT_MAX_FINDINGS_PER_SECTOR_MAIN
    )

    # SLA/PRA narrative block only when module is enabled (payload has sla_pra_summary with items)
    sla_pra_summary = data.get("sla_pra_summary")
    sla_pra_items = (sla_pra_summary or {}).get("items") or []
    sla_pra_rendered = False
    if sla_pra_summary is not None and sla_pra_items:
        sla_pra_rendered = insert_sla_pra_summary_at_anchor(doc, sla_pra_items)
    # When SLA/PRA disabled: [[SLA_PRA_SUMMARY]] anchor is removed by insert_cyber_vofc_block_after_sla_pra (fallback path)
    print(f"SLA_PRA_RENDERED={'1' if sla_pra_rendered else '0'}", file=sys.stderr)

    # Cross-dependency summary (optional anchor; payload from Cross-Dependencies tab)
    cross_dependency_summary = data.get("cross_dependency_summary")
    cross_dependency_modules = data.get("cross_dependency_modules")
    insert_cross_dependency_summary_at_anchor(doc, cross_dependency_summary, cross_dependency_modules)

    # Cyber VOFC subsection after SLA/PRA block
    insert_cyber_vofc_block_after_sla_pra(doc, cyber_capped, truncation_note=cyber_truncated)

    # Optional: sources narrative (energy/dependency content now in per-sector pages)
    insert_sources_narrative_at_anchor(
        doc, assessment, sla_reliability_for_report, sla_pra_module_enabled
    )

    # 8) Safety sweep: replace any remaining underscore blanks with "not confirmed"
    safety_sweep_underscore_blanks(doc)

    # 8) Cleanup and hard validation
    remove_anchor_block_paragraphs(doc)
    replace_placeholder_anchors_with_fallback(doc, vulnerability_count=rendered_vuln_count)
    clear_remaining_anchor_paragraphs(doc)
    scrub_any_remaining_anchors(doc)
    assert_no_anchors_remaining(doc)
    assert_no_narrative_blanks_remaining(doc)
    assert_no_export_style_tables(doc)
    assert_exec_summary_not_followed_by_table(doc)
    gate_b_suppress_empty_sections(doc)
    gate_c_duplicate_block_detection(doc)
    gate_a_placeholder_anchor_leak_scanner(doc)
    sanitize_spof_language(doc)
    assert_no_spof_in_output(doc)

    # Replace encoding artifacts (U+FFFD, mojibake) so QC passes
    _replace_encoding_artifacts_in_doc(doc)

    # Export QC gate runs AFTER INFRA injection (encoding, anchors, limits, required sections)
    if _debug:
        print("[reporter] running export QC", file=sys.stderr)
    if template_path:
        try:
            from qc_export import REQUIRED_SECTION_ANCHORS, run_export_qc
            try:
                from qc.validators import run_all_qc_validators
            except ImportError:
                run_all_qc_validators = None  # optional package not present
            vuln_blocks_flat = []
            if energy_dependency and energy_dependency.get("vulnerability_blocks"):
                vuln_blocks_flat.extend(energy_dependency["vulnerability_blocks"])
            for sec in dependency_sections or []:
                vuln_blocks_flat.extend(sec.get("vulnerability_blocks") or [])
            snapshot = executive_snapshot if isinstance(executive_snapshot, dict) else {}
            synthesis_has_content = bool(
                synthesis and isinstance(synthesis, dict)
                and (synthesis.get("paragraphs") or synthesis.get("bullets") or (synthesis.get("title") or "").strip())
            )
            actions_list = (priority_actions or {}).get("actions") if isinstance(priority_actions, dict) else (priority_actions if isinstance(priority_actions, list) else None) or []
            vuln_blocks_str = data.get("vulnerability_blocks") or ""
            required_sections_filled = {
                "[[SNAPSHOT_POSTURE]]": bool(snapshot.get("posture")),
                "[[SNAPSHOT_MATRIX]]": bool(snapshot.get("matrixRows")),
                "[[SYNTHESIS]]": synthesis_has_content,
                "[[PRIORITY_ACTIONS]]": len(actions_list) > 0,
                "[[VULNERABILITY_BLOCKS]]": bool((vuln_blocks_str or "").strip()),
            }
            # Only require sections that are in REQUIRED_SECTION_ANCHORS (subset of keys we pass)
            required_sections_filled = {k: required_sections_filled.get(k, False) for k in REQUIRED_SECTION_ANCHORS}
            # Only enforce required-sections when payload has full report data (all key sections present)
            has_required_data = (
                bool(snapshot.get("posture"))
                and bool(snapshot.get("matrixRows"))
                and synthesis_has_content
                and len(actions_list) > 0
            )
            run_export_qc(
                doc,
                str(template_path),
                priority_actions=actions_list,
                vulnerability_blocks=vuln_blocks_flat,
                required_sections_filled=required_sections_filled if has_required_data else None,
                payload=data,
                assessment_json=data,
            )
            if run_all_qc_validators is not None:
                run_all_qc_validators(
                    doc,
                    vulnerability_count=rendered_vuln_count,
                    vulnerability_blocks=vuln_blocks_flat,
                )
        except ValueError as e:
            print(f"ERROR: Export QC failed: {e}", file=sys.stderr)
            sys.exit(1)
        except Exception as e:
            if isinstance(e, ImportError) and "qc" in str(e).lower():
                pass  # qc package optional; continue without run_all_qc_validators
            else:
                try:
                    from qc.validators import QCValidationError
                    if isinstance(e, QCValidationError):
                        print(f"ERROR: Export QC validation failed: {e}", file=sys.stderr)
                        sys.exit(1)
                except ImportError:
                    pass
                raise

    # Template safeguard: remove any static sector labels in Part I before save
    reorder_part2_sections(doc)
    remove_static_sector_labels_safeguard(doc)
    modernize_executive_heading_flow(doc)
    ensure_part1_decision_brief_content(
        doc,
        executive_snapshot,
        synthesis,
        priority_actions,
        assessment=assessment,
        fallback_action_titles=fallback_titles,
    )
    remove_forced_page_breaks_in_part1(doc)
    remove_orphan_drawing_paragraphs_in_part1(doc)
    tighten_report_layout_spacing(doc)
    compact_table_layout(doc)
    enforce_heading_content_cohesion(doc)
    force_page_break_before_heading(doc, "CORE INFRASTRUCTURE OVERVIEW")

    # Final hygiene: collapse consecutive breaks, remove blank-page-causing empties, trim trailing
    collapse_consecutive_pagebreaks(doc)
    remove_empty_paragraphs_after_page_breaks(doc)
    _trim_trailing_empty_paragraphs(doc)
    compress_blank_paragraphs(doc, max_run=0)
    collapse_consecutive_pagebreaks(doc)

    output_path = work_path / "output.docx"
    assert_no_backend_leak(doc)
    _assert_vulnerability_block_formatting(doc)
    doc.save(str(output_path))
    if _debug:
        print("[reporter] saved:", output_path, file=sys.stderr)


def run_from_payload(data: dict, work_dir: str | Path, template_path: str | Path) -> bytes:
    """
    Generate report from payload dict and return DOCX bytes. For serverless/API (e.g. Vercel).
    Uses skip_canonical_template_check so the template can live at any path (e.g. /tmp or deployment path).
    """
    work_path = Path(work_dir)
    work_path.mkdir(parents=True, exist_ok=True)
    _render_report(data, work_path, Path(template_path), skip_canonical_template_check=True)
    return (work_path / "output.docx").read_bytes()


def _render_generic_report(data: dict, work_path: Path) -> None:
    """
    Universal DOCX renderer.
    Expects payload.generic_report with:
      - title: string
      - subtitle: optional string
      - header_left / header_right / footer_left / footer_right: optional strings
      - sections: array of { heading: string, paragraphs?: string[], bullets?: string[] }
    """
    generic = data.get("generic_report") if isinstance(data, dict) else None
    if not isinstance(generic, dict):
        raise ValueError("generic_report payload missing or invalid")

    title = sanitize_text((generic.get("title") or "Report").strip())
    subtitle = sanitize_text((generic.get("subtitle") or "").strip())
    sections = generic.get("sections") if isinstance(generic.get("sections"), list) else []
    header_left = sanitize_text((generic.get("header_left") or "").strip())
    header_right = sanitize_text((generic.get("header_right") or "").strip())
    footer_left = sanitize_text((generic.get("footer_left") or "").strip())
    footer_right = sanitize_text((generic.get("footer_right") or "").strip())

    script_dir = Path(__file__).resolve().parent
    repo_root = script_dir.parent.parent
    template_path_env = os.environ.get("TEMPLATE_PATH")
    template_path = Path(template_path_env) if template_path_env else repo_root / "ADA" / "report template.docx"
    if template_path.is_file():
        doc = Document(str(template_path))
        replace_anchor_in_doc(doc, "[[REPORT_TITLE]]", title, body_only=False)
        replace_anchor_in_doc(doc, "[[REPORT_SUBTITLE]]", subtitle, body_only=False)
        replace_anchor_in_doc(doc, "[[REPORT_HEADER_LEFT]]", header_left, body_only=False)
        replace_anchor_in_doc(doc, "[[REPORT_HEADER_RIGHT]]", header_right, body_only=False)
        replace_anchor_in_doc(doc, "[[REPORT_FOOTER_LEFT]]", footer_left, body_only=False)
        replace_anchor_in_doc(doc, "[[REPORT_FOOTER_RIGHT]]", footer_right, body_only=False)
    else:
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        try:
            p.style = "Heading 1"
        except Exception:
            pass

        if subtitle:
            sub = doc.add_paragraph(sanitize_text(subtitle))
            try:
                sub.style = "Heading 2"
            except Exception:
                pass

    for section in sections:
        if not isinstance(section, dict):
            continue
        heading = sanitize_text((section.get("heading") or "").strip())
        if not heading:
            continue
        h = doc.add_paragraph()
        h_run = h.add_run(heading)
        h_run.bold = True
        try:
            h.style = "Heading 2"
        except Exception:
            pass

        paragraphs = section.get("paragraphs") if isinstance(section.get("paragraphs"), list) else []
        bullets = section.get("bullets") if isinstance(section.get("bullets"), list) else []

        for para in paragraphs:
            text = sanitize_text(str(para).strip())
            if text:
                doc.add_paragraph(text)

        for bullet in bullets:
            text = sanitize_text(str(bullet).strip())
            if text:
                try:
                    doc.add_paragraph(text, style="List Bullet")
                except Exception:
                    doc.add_paragraph(text)

    output_path = work_path / "output.docx"
    doc.save(str(output_path))


def main() -> None:
    work_dir = os.environ.get("WORK_DIR")
    if not work_dir:
        script_dir = Path(__file__).resolve().parent
        repo_root = script_dir.parent.parent
        work_dir = str(repo_root / "data" / "temp" / "reporter")
    work_path = Path(work_dir)
    work_path.mkdir(parents=True, exist_ok=True)

    raw = sys.stdin.read()
    if not raw.strip():
        print("No JSON input on stdin", file=sys.stderr)
        sys.exit(1)
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"Invalid JSON: {e}", file=sys.stderr)
        sys.exit(1)

    if isinstance(data, dict) and isinstance(data.get("generic_report"), dict):
        _render_generic_report(data, work_path)
        print(work_path / "output.docx")
        return

    CANONICAL_TEMPLATE_SUFFIX = "ADA" + os.sep + "report template.docx"
    CANONICAL_TEMPLATE_SUFFIX_ALT = "ADA/report template.docx"
    script_dir = Path(__file__).resolve().parent
    repo_root = script_dir.parent.parent
    template_path = os.environ.get("TEMPLATE_PATH")
    if template_path:
        template_path = Path(template_path)
    else:
        template_path = repo_root / "ADA" / "report template.docx"
    template_path = template_path.resolve()
    normalized = str(template_path).replace("\\", "/")
    if not (normalized.endswith(CANONICAL_TEMPLATE_SUFFIX_ALT) or normalized.endswith(CANONICAL_TEMPLATE_SUFFIX)):
        print("ERROR: Wrong template: export must use /ADA/report template.docx", file=sys.stderr)
        print(f"  Got: {template_path}", file=sys.stderr)
        sys.exit(1)
    if not template_path.is_file():
        print(f"Template not found: {template_path}", file=sys.stderr)
        sys.exit(1)
    print(f"Template: {template_path}", file=sys.stderr)

    _render_report(data, work_path, template_path, skip_canonical_template_check=False)
    if os.environ.get("EMIT_VULN_MANIFEST") and data.get("_vuln_manifest_blocks") is not None:
        manifest_path = work_path / "vuln_manifest.json"
        with open(manifest_path, "w", encoding="utf-8") as f:
            json.dump(data["_vuln_manifest_blocks"], f, indent=2)
    output_path = work_path / "output.docx"
    print(output_path)


# Re-exports for tests (import from new modules so "from main import X" still works)
from docx_ops import (
    paragraph_has_drawing,
    replace_anchor_in_paragraph_preserve_drawings,
    replace_anchor_in_paragraph,
    remove_paragraph,
    find_paragraph_by_exact_text,
    find_anchor_paragraph_exact,
)
from render_vulns import (
    render_vulnerability_block,
    render_vulnerability_block_after,
    _render_structured_vulnerabilities,
)
from io_payload import get_part2_findings
from render_part2 import (
    _internet_transport_table,
    _it_critical_hosted_table,
    _is_transport_provider,
    _get_it_isp_names_from_curve,
    _format_it_internet_connectivity_narrative,
    _service_loss_description,
    _hosted_continuity_label,
    _hosted_continuity_label_for_summary,
)
from qc_pipeline import (
    gate_a_placeholder_anchor_leak_scanner,
    gate_c_duplicate_block_detection,
    gate_b_suppress_empty_sections,
)

def _emit_vuln_model(json_path: str, out_path: str) -> None:
    """
    Emit the vuln blocks the reporter would render (pre-DOCX model) as JSON.
    Used by truth_diff harness. No template or DOCX; deterministic output.
    """
    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)
    out_list: list[dict] = []
    canonical = data.get("canonicalVulnBlocks")
    if isinstance(canonical, list) and len(canonical) > 0:
        for c in canonical:
            domain = (c.get("domain") or "").strip()
            vuln_id = (c.get("vuln_id") or "").strip() or (c.get("title") or "").strip()
            title = (c.get("title") or "").strip()
            if not title:
                continue
            narrative = (c.get("narrative") or "").strip()
            ofc_text = c.get("ofcText") or ""
            ofcs = [s.strip() for s in (ofc_text.strip().split("\n")) if s.strip()] if ofc_text else []
            out_list.append({
                "domain": domain,
                "vuln_id": vuln_id,
                "title": title,
                "narrative": narrative,
                "ofcs": ofcs,
            })
    else:
        derived_list = load_derived_findings_from_payload(data)
        for f in derived_list:
            domain = (f.get("infrastructure") or "").strip()
            vuln_id = (f.get("id") or "").strip() or (f.get("title") or "").strip()
            title = (f.get("title") or "").strip()
            if not title:
                continue
            narrative = (f.get("narrative") or "").strip()
            ofc_text = f.get("ofcText") or ""
            ofcs = [s.strip() for s in (ofc_text.strip().split("\n")) if s.strip()] if ofc_text else []
            out_list.append({
                "domain": domain,
                "vuln_id": vuln_id,
                "title": title,
                "narrative": narrative,
                "ofcs": ofcs,
            })
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(out_list, f, indent=2)


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Reporter: stdin JSON -> DOCX, or --emit-vuln-model for truth diff.")
    parser.add_argument("--emit-vuln-model", action="store_true", help="Emit vuln model JSON only (no DOCX)")
    parser.add_argument("--json", dest="json_path", help="Path to export JSON (for --emit-vuln-model)")
    parser.add_argument("--out", dest="out_path", help="Output path for vuln model JSON")
    args, _ = parser.parse_known_args()
    if args.emit_vuln_model and args.json_path and args.out_path:
        _emit_vuln_model(args.json_path, args.out_path)
        sys.exit(0)
    main()
