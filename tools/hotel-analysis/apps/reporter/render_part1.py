"""
Part I rendering: executive snapshot, synthesis, priority actions, charts at CHART_* anchors,
replace_all_text_placeholders. Chart generation (build_curve, render_dependency_chart_png, etc.)
remains in main and is re-exported for tests; render_part1 only inserts pre-generated charts.
"""
from __future__ import annotations

import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from constants import (
    CHART_CATEGORIES,
    CHART_W_INCHES,
    CHART_H_INCHES,
    CHART_SPACING_PT,
    CHART_HEADING_SPACE_BEFORE_PT,
    CHART_HEADING_SPACE_AFTER_PT,
    CHART_IMAGE_SPACING_PT,
    SECTION_C_SECTOR_HEADINGS,
)
from docx_ops import (
    find_paragraph_by_exact_text,
    find_anchor_paragraph_exact,
    insert_paragraph_after,
    remove_paragraph,
    set_paragraph_keep_with_next,
    insert_table_after,
    apply_table_grid_style,
    set_table_fixed_widths,
    set_repeat_header_row,
    set_table_rows_cant_split,
    iter_paragraphs_and_cells,
    paragraph_has_drawing,
    _clear_paragraph_text_preserve_drawings,
)
from sanitize import sanitize_text
from report_humanize import (
    ensure_synthesis_formatting,
    normalize_spacing,
    sanitize_backend_evidence,
)

# For replace_all_text_placeholders we need _run_has_drawing from docx_ops (internal)
try:
    from docx_ops import _run_has_drawing
except ImportError:
    def _run_has_drawing(run) -> bool:
        from docx_ops import _element_has_drawing_or_pict_child
        return _element_has_drawing_or_pict_child(run._element)


def insert_charts_at_anchors(doc: Document, chart_paths: dict) -> None:
    """Insert chart image at each [[CHART_*]] anchor. Heading 3 (sector name) above image."""
    for code in CHART_CATEGORIES:
        anchor = f"[[CHART_{code}]]"
        path = chart_paths.get(code)
        if path is None:
            raise RuntimeError(f"Chart path not set for sector: {code}")
        path = Path(path)
        if not path.exists():
            raise RuntimeError(f"Chart missing: {code} expected at {path}")
        candidates = list(find_anchor_paragraph_exact(doc, anchor, body_only=True))
        if len(candidates) != 1:
            raise RuntimeError(
                f"Chart anchor must appear exactly once: {anchor} (found {len(candidates)} occurrences)"
            )
        sector_heading = SECTION_C_SECTOR_HEADINGS.get(code, code)
        insert_chart_at_anchor(doc, anchor, path, sector_label=sector_heading)


def insert_chart_at_anchor(
    doc: Document,
    anchor_text: str,
    png_path: Path | str,
    width_in: float | None = None,
    height_in: float | None = None,
    caption: str | None = None,
    sector_label: str | None = None,
) -> None:
    """Find anchor paragraph (body only). Insert Heading 3 (sector label) above image, then image."""
    w = width_in if width_in is not None else CHART_W_INCHES
    h = height_in if height_in is not None else CHART_H_INCHES
    p = find_paragraph_by_exact_text(doc, anchor_text, body_only=True)
    if p is None:
        raise RuntimeError(f"Chart anchor not found: {anchor_text}")
    insert_after = p
    if sector_label:
        label_para = insert_paragraph_after(insert_after, sanitize_text(sector_label))
        try:
            label_para.style = "Heading 3"
        except Exception:
            pass
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            label_para.paragraph_format.space_before = Pt(CHART_HEADING_SPACE_BEFORE_PT)
            label_para.paragraph_format.space_after = Pt(CHART_HEADING_SPACE_AFTER_PT)
        except Exception:
            pass
        set_paragraph_keep_with_next(label_para)
        insert_after = label_para
    img_para = insert_paragraph_after(insert_after, "")
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        img_para.paragraph_format.space_before = Pt(CHART_IMAGE_SPACING_PT)
        img_para.paragraph_format.space_after = Pt(CHART_IMAGE_SPACING_PT)
    except Exception:
        pass
    set_paragraph_keep_with_next(img_para)
    run = img_para.add_run()
    run.add_picture(str(png_path), width=Inches(w), height=Inches(h))
    if caption:
        caption_para = insert_paragraph_after(img_para, sanitize_text(caption))
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            caption_para.paragraph_format.space_before = Pt(CHART_IMAGE_SPACING_PT)
            caption_para.paragraph_format.space_after = Pt(CHART_IMAGE_SPACING_PT)
        except Exception:
            pass
        set_paragraph_keep_with_next(caption_para)
    remove_paragraph(p)


def insert_chart_after_paragraph(
    doc: Document,
    after_para,
    png_path: Path | str,
    caption: str | None = None,
):
    """Insert chart image (and optional caption) after the given paragraph. Returns last inserted paragraph."""
    w, h = CHART_W_INCHES, CHART_H_INCHES
    insert_after = after_para
    img_para = doc.add_paragraph("")
    run = img_para.add_run()
    run.add_picture(str(png_path), width=Inches(w), height=Inches(h))
    img_para._element.getparent().remove(img_para._element)
    insert_after._element.addnext(img_para._element)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        img_para.paragraph_format.space_before = Pt(CHART_SPACING_PT)
        img_para.paragraph_format.space_after = Pt(CHART_SPACING_PT)
    except Exception:
        pass
    set_paragraph_keep_with_next(img_para)
    insert_after = img_para
    if caption:
        caption_para = insert_paragraph_after(insert_after, sanitize_text(caption))
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            caption_para.paragraph_format.space_before = Pt(CHART_SPACING_PT)
            caption_para.paragraph_format.space_after = Pt(CHART_SPACING_PT)
        except Exception:
            pass
        set_paragraph_keep_with_next(caption_para)
        insert_after = caption_para
    spacer = insert_paragraph_after(insert_after, "")
    return spacer


def inject_executive_snapshot_at_anchors(doc: Document, snapshot: dict) -> None:
    """Inject curve-driven executive snapshot content at SNAPSHOT_* anchors."""
    posture = (snapshot.get("posture") or "").strip()
    summary = (snapshot.get("summary") or "").strip()
    vulnerabilities = snapshot.get("vulnerabilities") or []
    drivers = snapshot.get("drivers") or []
    matrix_rows = snapshot.get("matrixRows") or []
    cascade = snapshot.get("cascade")

    p = find_paragraph_by_exact_text(doc, "[[SNAPSHOT_POSTURE]]", body_only=True)
    if p and posture:
        p.clear()
        insert_paragraph_after(p, sanitize_text(posture), style="Normal")

    p = find_paragraph_by_exact_text(doc, "[[SNAPSHOT_SUMMARY]]", body_only=True)
    if p and summary:
        p.clear()
        insert_after = insert_paragraph_after(p, sanitize_text(summary), style="Normal")
        valid_vulns = [
            v for v in vulnerabilities
            if isinstance(v, dict) and (str(v.get("title") or "").strip() or str(v.get("narrative") or "").strip())
        ]
        if valid_vulns:
            heading = insert_paragraph_after(insert_after, "Executive Vulnerabilities", style="Heading 3")
            set_paragraph_keep_with_next(heading)
            insert_after = heading
            for vuln in valid_vulns[:5]:
                title = str(vuln.get("title") or "").strip()
                narrative = str(vuln.get("narrative") or "").strip()
                sector = str(vuln.get("sector") or "").strip()
                if title and sector:
                    label = f"{sector} - {title}"
                else:
                    label = title or sector
                if label:
                    insert_after = insert_paragraph_after(insert_after, sanitize_text(label), style="List Bullet")
                    if narrative:
                        insert_after = insert_paragraph_after(insert_after, sanitize_text(narrative), style="Normal")

    p = find_paragraph_by_exact_text(doc, "[[SNAPSHOT_DRIVERS]]", body_only=True)
    if p and drivers:
        p.clear()
        insert_after = p
        for d in drivers[:3]:
            insert_after = insert_paragraph_after(insert_after, sanitize_text(str(d)), style="List Bullet")

    p = find_paragraph_by_exact_text(doc, "[[SNAPSHOT_MATRIX]]", body_only=True)
    if p:
        if matrix_rows:
            p.clear()
            set_paragraph_keep_with_next(p)
            num_rows = 1 + len(matrix_rows)
            table = insert_table_after(doc, p, num_rows, 5)
            apply_table_grid_style(table)
            table.autofit = False
            widths = [1.2, 1.8, 1.2, 1.4, 2.0]
            set_table_fixed_widths(table, widths)
            for c, h in enumerate(("Sector", "Time to Severe Impact (hrs)", "Functional Loss (%)", "Backup Duration (hrs)", "Structural Posture")):
                table.rows[0].cells[c].text = h
            set_repeat_header_row(table.rows[0])
            set_table_rows_cant_split(table)
            for i, row in enumerate(matrix_rows):
                backup_hrs = row.get("backupHrs")
                if backup_hrs is None or (isinstance(backup_hrs, str) and not backup_hrs.strip()):
                    backup_hrs = "N/A"
                table.rows[i + 1].cells[0].text = sanitize_text(row.get("sector", ""))
                table.rows[i + 1].cells[1].text = sanitize_text(row.get("ttiHrs") or "Unknown")
                table.rows[i + 1].cells[2].text = sanitize_text(row.get("lossPct") or "Unknown")
                table.rows[i + 1].cells[3].text = sanitize_text(str(backup_hrs))
                table.rows[i + 1].cells[4].text = sanitize_text(row.get("structuralPosture") or "Unknown")
            insert_paragraph_after(p, "")
        else:
            remove_paragraph(p)

    p = find_paragraph_by_exact_text(doc, "[[SNAPSHOT_CASCADE]]", body_only=True)
    if p:
        if cascade and str(cascade).strip():
            p.clear()
            insert_paragraph_after(p, sanitize_text(str(cascade).strip()), style="Normal")
        else:
            remove_paragraph(p)


def _synthesis_sanitize_single_provider(text: str, assessment: dict) -> str:
    """Replace 'single provider' with 'single confirmed service path' when provider not identified."""
    if not text or "single provider" not in text.lower():
        return text
    categories = assessment.get("categories") or {}
    any_provider_identified = False
    for code in CHART_CATEGORIES:
        cat = categories.get(code) or {}
        supply = cat.get("supply") or {}
        sources = supply.get("sources") or []
        if any(
            (s or {}).get("provider_name") or (s or {}).get("service_provider") or (s or {}).get("provider")
            for s in sources
        ):
            any_provider_identified = True
            break
    if not any_provider_identified:
        text = re.sub(r"\bsingle provider\b", "single confirmed service path", text, flags=re.I)
    return text


def inject_synthesis_at_anchor(doc: Document, synthesis: dict, assessment: dict | None = None) -> None:
    """Inject Cross-Infrastructure Synthesis at [[SYNTHESIS]] anchor. Template owns heading; inject only body and bullets."""
    paragraphs = synthesis.get("paragraphs") or []
    bullets = synthesis.get("bullets") or []
    assessment = assessment or {}

    p = find_paragraph_by_exact_text(doc, "[[SYNTHESIS]]", body_only=True)
    if p is None:
        return

    p.clear()
    insert_after = p

    for para in paragraphs:
        if para and str(para).strip():
            raw = str(para).strip()
            human = ensure_synthesis_formatting(normalize_spacing(sanitize_backend_evidence(raw)))
            sanitized = _synthesis_sanitize_single_provider(human or raw, assessment)
            insert_after = insert_paragraph_after(insert_after, sanitize_text(sanitized), style="Normal")

    for b in bullets:
        label = (b.get("label") or "").strip()
        text = (b.get("text") or "").strip()
        if label and text:
            label = label.rstrip(":").rstrip() + ": "
            human = ensure_synthesis_formatting(normalize_spacing(sanitize_backend_evidence(text)))
            sanitized_text = _synthesis_sanitize_single_provider(human or text, assessment)
            bullet_para = insert_paragraph_after(insert_after, "", style="List Bullet")
            bold_run = bullet_para.add_run(sanitize_text(label))
            bold_run.bold = True
            bullet_para.add_run(sanitize_text(sanitized_text))
            insert_after = bullet_para
        elif text:
            human = ensure_synthesis_formatting(normalize_spacing(sanitize_backend_evidence(text)))
            insert_after = insert_paragraph_after(insert_after, sanitize_text(human or text), style="List Bullet")


def _strip_priority_action_curve_language(text: str) -> str:
    """Remove hours, percent, and 'reaches severe impact' from Priority Action text."""
    if not text or not isinstance(text, str):
        return text
    text = re.sub(r"\b\d+\s*hours?\b", "", text, flags=re.I)
    text = re.sub(r"~\s*\d+%|\b\d+%\b", "", text)
    text = re.sub(r"\b\d+\s*percent\b", "", text, flags=re.I)
    text = re.sub(r"reaches severe impact[^.]*\.?", "", text, flags=re.I)
    text = re.sub(r"this dependency[^.]*\.?", "", text, flags=re.I)
    return " ".join(text.split()).strip()


def _parse_priority_actions_payload(priority_actions: dict) -> list:
    """Normalize priority_actions into list of {leadIn, fullText}."""
    if not priority_actions:
        return []
    if isinstance(priority_actions, str):
        blocks = re.split(r"\n\s*\n", priority_actions.strip())
        out = []
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            lines = block.split("\n")
            first = (lines[0] or "").strip()
            rest = "\n".join(lines[1:]).strip() if len(lines) > 1 else ""
            out.append({"leadIn": first, "fullText": rest})
        return out
    if not isinstance(priority_actions, dict):
        return []
    actions = priority_actions.get("actions")
    if actions is None and isinstance(priority_actions, list):
        actions = priority_actions
    if isinstance(actions, str):
        return _parse_priority_actions_payload(actions)
    if not isinstance(actions, list):
        return []
    out = []
    for a in actions[:5]:
        if not isinstance(a, dict):
            continue
        lead_in = (a.get("leadIn") or "").strip()
        full_text = (a.get("fullText") or "").strip()
        full_text = _strip_priority_action_curve_language(full_text)
        out.append({"leadIn": lead_in, "fullText": full_text})
    return out


def inject_priority_actions_at_anchor(doc: Document, priority_actions: dict) -> None:
    """Inject Priority Actions at [[PRIORITY_ACTIONS]] anchor. Template owns heading; inject only body content."""
    if not priority_actions:
        return
    actions = _parse_priority_actions_payload(priority_actions)
    if not actions:
        return

    p = find_paragraph_by_exact_text(doc, "[[PRIORITY_ACTIONS]]", body_only=True)
    if p is None:
        raise RuntimeError("[[PRIORITY_ACTIONS]] anchor not found in template; cannot inject priority actions.")

    p.clear()
    insert_after = p

    for a in actions:
        lead_in = (a.get("leadIn") or "").strip()
        full_text = (a.get("fullText") or "").strip()
        if not lead_in and not full_text:
            continue
        title_para = insert_paragraph_after(insert_after, "", style="List Number")
        title_para.add_run(sanitize_text(lead_in or "(No title)"))
        if title_para.runs:
            title_para.runs[0].bold = True
        insert_after = title_para
        if full_text:
            desc_para = insert_paragraph_after(insert_after, sanitize_text(full_text), style="Normal")
            insert_after = desc_para
        insert_after = insert_paragraph_after(insert_after, "", style="Normal")


def replace_all_text_placeholders(doc: Document, asset: dict) -> None:
    """Replace {{ASSET_NAME}}, {{VISIT_DATE}}, [[PSA_*]], etc. in doc. PSA phone is REQUIRED."""
    phone = (asset.get("psa_phone") or asset.get("psa_cell") or "").strip()
    if not phone:
        raise RuntimeError("MISSING_PSA_PHONE: PSA phone number is required for export. Set asset.psa_cell or asset.psa_phone.")
    psa_cell_anchor = "[[PSA_CELL]]"
    psa_count = 0
    for p, _ in iter_paragraphs_and_cells(doc):
        psa_count += (p.text or "").count(psa_cell_anchor)
    if psa_count != 0 and psa_count != 1:
        raise RuntimeError(
            f"[[PSA_CELL]] must appear exactly once in the template (found {psa_count}). "
            "Fix template to use exactly one 'Cell: [[PSA_CELL]]'."
        )
    mapping = {
        "{{ASSET_NAME}}": sanitize_text((asset.get("asset_name") or "").strip()),
        "{{VISIT_DATE}}": sanitize_text((asset.get("visit_date_iso") or "").strip()),
        "{{ASSESSOR}}": sanitize_text((asset.get("assessor") or "").strip()),
        "{{LOCATION}}": sanitize_text((asset.get("location") or "").strip()),
        "[[PSA_NAME]]": sanitize_text((asset.get("psa_name") or "").strip()),
        "[[PSA_REGION]]": sanitize_text((asset.get("psa_region") or "").strip()),
        "[[PSA_CITY]]": sanitize_text((asset.get("psa_city") or "").strip()),
        "[[PSA_PHONE]]": sanitize_text(phone),
        "[[PSA_CELL]]": sanitize_text(phone),
        "[PSA_CELL]]": sanitize_text(phone),
        "[[PSA_CELL]": sanitize_text(phone),
        "[[PSA_EMAIL]]": sanitize_text((asset.get("psa_email") or "").strip()),
        "[[PSA_EMAIL]": sanitize_text((asset.get("psa_email") or "").strip()),
        "[[FACILITY_NAME]]": sanitize_text((asset.get("asset_name") or "").strip()),
        "[[ASSESSMENT_DATE]]": sanitize_text((asset.get("visit_date_iso") or "").strip()[:10]),
    }
    for p, _ in iter_paragraphs_and_cells(doc):
        full_text = p.text or ""
        changed = False
        new_text = full_text
        for placeholder, value in mapping.items():
            if placeholder in full_text:
                new_text = new_text.replace(placeholder, value)
                changed = True
        if changed:
            if paragraph_has_drawing(p):
                _clear_paragraph_text_preserve_drawings(p)
                first_text_run = next((r for r in p.runs if not _run_has_drawing(r)), None)
                if first_text_run is not None:
                    first_text_run.text = new_text
                else:
                    p.add_run(new_text)
            else:
                p.clear()
                p.add_run(new_text)

    for element in doc.element.iter():
        if element.tag == qn("w:t") and element.text:
            for placeholder, value in mapping.items():
                if placeholder in element.text:
                    element.text = element.text.replace(placeholder, value)

    for p, _ in iter_paragraphs_and_cells(doc):
        if "Cell: [" in (p.text or ""):
            raise RuntimeError(
                "PSA_CELL replacement produced malformed line: 'Cell: [' (missing closing bracket in template?). "
                "Template must use exactly 'Cell: [[PSA_CELL]]'."
            )


def render_part1(doc: Document, data: dict, chart_paths: dict, template_path: str | None = None) -> None:
    """
    Part I: charts at CHART_* anchors, replace_all_text_placeholders, executive snapshot,
    synthesis, priority actions. Chart paths must be pre-generated by caller (main).
    """
    if chart_paths:
        insert_charts_at_anchors(doc, chart_paths)

    assessment = data.get("assessment") or {}
    asset = assessment.get("asset") or {}
    try:
        replace_all_text_placeholders(doc, asset)
    except RuntimeError as e:
        if "MISSING_PSA_PHONE" in str(e):
            import sys
            print(str(e), file=sys.stderr)
            sys.exit(1)
        raise

    executive_snapshot = data.get("executive_snapshot")
    if executive_snapshot and isinstance(executive_snapshot, dict):
        inject_executive_snapshot_at_anchors(doc, executive_snapshot)

    synthesis = data.get("synthesis")
    if synthesis and isinstance(synthesis, dict):
        inject_synthesis_at_anchor(doc, synthesis, assessment=assessment)

    priority_actions = data.get("priority_actions")
    if priority_actions and isinstance(priority_actions, dict):
        inject_priority_actions_at_anchor(doc, priority_actions)
