"""
Vulnerability block rendering and ADA_Vuln_* styles.
Extracted from main.py; imports from constants, docx_ops, sanitize, report_humanize.
"""
from __future__ import annotations

import re
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches
from docx.text.paragraph import Paragraph as DocxParagraph

from constants import (
    ADA_VULN_STYLE_NAMES,
    REQUIRED_ADA_STYLES,
    FRANKLIN_FONT_DEMI,
    FRANKLIN_FONT_MEDIUM,
    FRANKLIN_FONT_BOOK,
    VULN_NARRATIVE_SPACING_PT,
    VULN_OFC_ITEM_SPACING_PT,
    VULN_DIVIDER_SPACING_PT,
    _SEVERITY_ORDER,
)
from docx_ops import (
    insert_paragraph_after,
    find_anchor_paragraph_exact,
    paragraph_has_drawing,
    _clear_paragraph_text_preserve_drawings,
    remove_paragraph,
    set_paragraph_keep_with_next,
)
from sanitize import sanitize_text, sanitize_vulnerability_text
from report_humanize import debot_vulnerability_narrative, sanitize_backend_evidence, dedupe_sentences

# Severity line: "HIGH – Structural Concentration" (dash can be -, –, —; case-insensitive)
_SEVERITY_LINE_PATTERN = re.compile(
    r"(?i)^(HIGH|ELEVATED|MODERATE)\s*[\-\u2013\u2014]\s*.+"
)
_SEVERITY_LINE_FALLBACK = re.compile(
    r"(?i)^(HIGH|ELEVATED|MODERATE)[\s\-–—:].+"
)
_EN_DASH = "\u2013"


def render_vulnerability_block(doc: Document, block: dict) -> None:
    """
    Render a single RenderedVulnerabilityBlock with strict paragraph boundaries.
    Order: Title (bold) -> [Category] -> Condition Identified -> Operational Exposure -> Why This Matters
          -> OFC heading -> numbered OFCs -> References heading -> refs.
    When condition_identified/operational_exposure/why_this_matters present, use those; else narrative.
    """
    title = sanitize_vulnerability_text(block.get("title") or "")
    narrative = sanitize_vulnerability_text(block.get("narrative") or "")
    cond_id = sanitize_vulnerability_text(block.get("condition_identified") or "")
    op_exp = sanitize_vulnerability_text(block.get("operational_exposure") or "")
    why = sanitize_vulnerability_text(block.get("why_this_matters") or "")
    driver_cat = block.get("driverCategory") or ""
    ofcs = block.get("ofcs") or []
    references = block.get("references") or []

    # 1) Title paragraph (bold)
    p_title = doc.add_paragraph()
    r = p_title.add_run(title)
    r.bold = True
    set_paragraph_keep_with_next(p_title)

    # 2) Category (when present)
    if driver_cat:
        doc.add_paragraph(f"Category: {driver_cat}")

    # 3) Condition Identified / Operational Exposure / Why This Matters (canonical) or fallback narrative
    if cond_id or op_exp or why:
        if cond_id:
            p_h = doc.add_paragraph()
            r = p_h.add_run("Condition Identified")
            r.bold = True
            doc.add_paragraph(cond_id)
        if op_exp:
            p_h = doc.add_paragraph()
            r = p_h.add_run("Operational Exposure")
            r.bold = True
            doc.add_paragraph(op_exp)
        if why:
            p_h = doc.add_paragraph()
            r = p_h.add_run("Why This Matters")
            r.bold = True
            doc.add_paragraph(why)
    elif narrative:
        doc.add_paragraph(narrative)

    # 4) Options for Consideration heading + numbered OFCs
    if ofcs:
        p_ofc_h = doc.add_paragraph()
        r = p_ofc_h.add_run("Options for Consideration")
        r.bold = True
        for i, ofc in enumerate(ofcs[:4], 1):
            doc.add_paragraph(sanitize_vulnerability_text(f"{i}. {ofc}"))

    # 5) References heading + bullet refs (always after OFCs)
    if references:
        p_ref_h = doc.add_paragraph()
        r = p_ref_h.add_run("References")
        r.bold = True
        for ref in references:
            doc.add_paragraph(sanitize_text(f"\u2022 {ref}"))


def render_vulnerability_block_after(
    doc: Document, block: dict, after_paragraph, add_spacer_after: bool = True
) -> DocxParagraph:
    """
    Strict narrative vulnerability block: Heading 3 (title), Normal (category/body), Heading 4 (section labels),
    Normal (body), List Number (OFCs, max 4), List Bullet (references). No manual "1." or "•" numbering.
    """
    title = sanitize_vulnerability_text(block.get("title") or "")
    narrative = sanitize_vulnerability_text(block.get("narrative") or "")
    cond_id = sanitize_vulnerability_text(block.get("condition_identified") or "")
    op_exp = sanitize_vulnerability_text(block.get("operational_exposure") or "")
    why = sanitize_vulnerability_text(block.get("why_this_matters") or "")
    driver_cat = block.get("driverCategory") or ""
    ofcs = block.get("ofcs") or []
    references = block.get("references") or []
    last = after_paragraph
    p_title = insert_paragraph_after(last, title, style="Heading 3")
    set_paragraph_keep_with_next(p_title)
    last = p_title
    if driver_cat:
        last = insert_paragraph_after(last, f"Category: {driver_cat}", style="Normal")
    if cond_id or op_exp or why:
        if cond_id:
            p_h = insert_paragraph_after(last, "Condition Identified", style="Heading 4")
            set_paragraph_keep_with_next(p_h)
            last = insert_paragraph_after(p_h, cond_id, style="Normal")
        if op_exp:
            p_h = insert_paragraph_after(last, "Operational Exposure", style="Heading 4")
            set_paragraph_keep_with_next(p_h)
            last = insert_paragraph_after(p_h, op_exp, style="Normal")
        if why:
            p_h = insert_paragraph_after(last, "Why This Matters", style="Heading 4")
            set_paragraph_keep_with_next(p_h)
            last = insert_paragraph_after(p_h, why, style="Normal")
    elif narrative:
        last = insert_paragraph_after(last, narrative, style="Normal")
    try:
        if last:
            last.paragraph_format.space_after = Pt(VULN_NARRATIVE_SPACING_PT)
    except Exception:
        pass
    if ofcs:
        p_ofc_h = insert_paragraph_after(last, "Options for Consideration", style="Heading 4")
        set_paragraph_keep_with_next(p_ofc_h)
        last = p_ofc_h
        for ofc in (ofcs[:4] if isinstance(ofcs, list) else [])[:4]:
            last = insert_paragraph_after(last, sanitize_vulnerability_text(str(ofc)), style="List Number")
            try:
                last.paragraph_format.space_after = Pt(VULN_OFC_ITEM_SPACING_PT)
            except Exception:
                pass
    if references:
        p_ref_h = insert_paragraph_after(last, "References", style="Heading 4")
        set_paragraph_keep_with_next(p_ref_h)
        last = p_ref_h
        refs_deduped = sorted(set(str(r).strip() for r in references if r), key=str.lower)
        for ref in refs_deduped:
            last = insert_paragraph_after(last, sanitize_text(ref), style="List Bullet")
    if add_spacer_after:
        try:
            last.paragraph_format.space_after = Pt(VULN_DIVIDER_SPACING_PT)
        except Exception:
            pass
    return last


def _get_style_safe(doc: Document, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return None


def _get_paragraph_style_names(doc: Document) -> list[str]:
    """Return list of paragraph style names in the document (exact names as in Word)."""
    names = []
    try:
        for s in doc.styles:
            if getattr(s, "type", None) == WD_STYLE_TYPE.PARAGRAPH:
                names.append(getattr(s, "name", "") or "")
    except Exception:
        pass
    return sorted(names)


def _norm_style(x: str | None) -> str:
    """Normalize for style name/styleId comparison (strip only; case-sensitive)."""
    return (x or "").strip()


def preflight_assert_ada_styles(doc: Document) -> None:
    """
    Assert all required ADA_Vuln_* paragraph styles exist in the template.
    Match by style name OR style_id (normalized strip); only PARAGRAPH styles count.
    Call after ensure_ada_vuln_styles() so on-the-fly creation has run.
    """
    present = set()
    for s in doc.styles:
        try:
            if getattr(s, "type", None) != WD_STYLE_TYPE.PARAGRAPH:
                continue
            present.add(_norm_style(getattr(s, "name", None)))
            present.add(_norm_style(getattr(s, "style_id", None)))
        except Exception:
            pass
    missing = [r for r in REQUIRED_ADA_STYLES if _norm_style(r) not in present]
    if missing:
        print("[QC] Available PARAGRAPH styles (name | style_id):")
        for s in doc.styles:
            try:
                if getattr(s, "type", None) == WD_STYLE_TYPE.PARAGRAPH:
                    print(f"[QC] {getattr(s,'name',None)!r} | {getattr(s,'style_id',None)!r}")
            except Exception:
                pass
        available = sorted(present)
        raise ValueError(
            f"Template is missing required paragraph styles: {missing!r}. "
            "Open report template.docx and add the styles with exact names (e.g. ADA_Vuln_Severity, ADA_Vuln_Label), "
            "or run: python apps/reporter/add_ada_vuln_styles_to_template.py. "
            f"Available paragraph styles (name/styleId): {available!r}"
        )


def ensure_ada_vuln_styles(doc: Document) -> None:
    """Ensure ADA_Vuln_* paragraph styles exist (Franklin typography). Idempotent per style."""
    styles = doc.styles
    if _get_style_safe(doc, "ADA_Vuln_Header") is None:
        s = styles.add_style("ADA_Vuln_Header", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = FRANKLIN_FONT_DEMI
        s.font.size = Pt(13)
        s.font.bold = True
        s.font.all_caps = True
        pf = s.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.0
        pf.keep_with_next = True
    if _get_style_safe(doc, "ADA_Vuln_Severity") is None:
        s = styles.add_style("ADA_Vuln_Severity", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = FRANKLIN_FONT_MEDIUM
        s.font.size = Pt(11)
        s.font.bold = True
        s.font.all_caps = True
        pf = s.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.keep_with_next = True
    if _get_style_safe(doc, "ADA_Vuln_Meta") is None:
        s = styles.add_style("ADA_Vuln_Meta", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = FRANKLIN_FONT_BOOK
        s.font.size = Pt(10.5)
        s.font.bold = False
        pf = s.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15
        pf.keep_with_next = True
    if _get_style_safe(doc, "ADA_Vuln_Label") is None:
        s = styles.add_style("ADA_Vuln_Label", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = FRANKLIN_FONT_MEDIUM
        s.font.size = Pt(10.5)
        s.font.bold = True
        pf = s.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.0
    if _get_style_safe(doc, "ADA_Vuln_Body") is None:
        s = styles.add_style("ADA_Vuln_Body", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = FRANKLIN_FONT_BOOK
        s.font.size = Pt(10.5)
        s.font.bold = False
        pf = s.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15
    if _get_style_safe(doc, "ADA_Vuln_Bullets") is None:
        s = styles.add_style("ADA_Vuln_Bullets", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = FRANKLIN_FONT_BOOK
        s.font.size = Pt(10.5)
        pf = s.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.first_line_indent = Inches(-0.15)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.15
    if _get_style_safe(doc, "ADA_Vuln_Numbered") is None:
        s = styles.add_style("ADA_Vuln_Numbered", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = FRANKLIN_FONT_BOOK
        s.font.size = Pt(10.5)
        pf = s.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.first_line_indent = Inches(-0.15)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.15


def _normalize_severity_line(line: str) -> str:
    """Remove stray Unicode from severity line; produce exactly 'SEVERITY – DRIVER' with U+2013 only."""
    if not line or not line.strip():
        return line
    s = (line or "").strip()
    s = s.replace("\u201c", "").replace("\u201d", "").replace('"', "")
    s = re.sub(r"[\u2014\u2013\-]+", _EN_DASH, s)
    s = re.sub(r"\s+", " ", s).strip()
    parts = re.split(r"\s*" + re.escape(_EN_DASH) + r"\s*", s, maxsplit=1)
    if len(parts) == 2:
        severity_part = parts[0].strip()
        driver_part = parts[1].strip()
        return f"{severity_part} {_EN_DASH} {driver_part}"
    return s


def _classify_ada_vuln_line(line: str) -> str:
    """Return ADA_Vuln_* style name for a single line of federal block content."""
    line = (line or "").strip()
    if not line:
        return "ADA_Vuln_Body"
    if re.match(r"^VULNERABILITY \d+\s+(?:HIGH|ELEVATED|MODERATE)\s*[\u2013\-]\s*.+", line, re.IGNORECASE):
        return "ADA_Vuln_Header"
    if re.match(r"^VULNERABILITY \d+\s+(?:HIGH|ELEVATED|MODERATE)\s*$", line, re.IGNORECASE):
        return "ADA_Vuln_Header"
    if re.match(r"^VULNERABILITY \d+$", line, re.IGNORECASE):
        return "ADA_Vuln_Header"
    if _SEVERITY_LINE_PATTERN.match(line) or _SEVERITY_LINE_FALLBACK.match(line):
        return "ADA_Vuln_Severity"
    if line.startswith("Infrastructure Domain:") or line.startswith("Risk Type:"):
        return "ADA_Vuln_Meta"
    if line in (
        "Operational Consequence",
        "Exposure Description",
        "Structural Indicator Profile",
        "Standards Alignment",
        "Options for Consideration",
    ):
        return "ADA_Vuln_Label"
    if re.match(r"^\d+\.\s+", line):
        return "ADA_Vuln_Numbered"
    if line.startswith("- "):
        return "ADA_Vuln_Bullets"
    return "ADA_Vuln_Body"


def _add_severity_paragraph_after(last, raw_line: str, doc: Document) -> DocxParagraph:
    """Create a dedicated paragraph for the severity line. Style is set on the paragraph (not run)."""
    severity_line = _normalize_severity_line(raw_line)
    p = insert_paragraph_after(last, "", style=None)
    p.clear()
    p.add_run(sanitize_text(severity_line))
    if _get_style_safe(doc, "ADA_Vuln_Severity") is None:
        available = [getattr(s, "name", "") or "" for s in doc.styles]
        raise ValueError(
            "Template is missing required paragraph style ADA_Vuln_Severity. "
            "Open report template.docx and add the style with exact name, or run: python apps/reporter/add_ada_vuln_styles_to_template.py. "
            f"Available styles: {available!r}"
        )
    p.style = doc.styles["ADA_Vuln_Severity"]
    return p


def _add_paragraph_with_style_after(last, text: str, style_name: str, doc: Document):
    """Insert a paragraph after last and set its style (ADA_Vuln_Header/Label/Body/etc.)."""
    p = insert_paragraph_after(last, "", style=None)
    p.clear()
    run = p.add_run(sanitize_text(text))
    if style_name == "ADA_Vuln_Header":
        run.bold = True
    if style_name in ADA_VULN_STYLE_NAMES:
        if _get_style_safe(doc, style_name) is None:
            available = _get_paragraph_style_names(doc)
            raise RuntimeError(
                f"Template is missing required paragraph style {style_name!r}. "
                "Open report template.docx and add the styles with exact names, or run: python apps/reporter/add_ada_vuln_styles_to_template.py. "
                f"Available paragraph styles: {available!r}"
            )
        p.style = doc.styles[style_name]
    else:
        try:
            p.style = doc.styles[style_name]
        except (KeyError, ValueError, AttributeError):
            try:
                p.style = "Normal"
            except Exception:
                pass
    return p


def _add_meta_paragraph_after(last, line: str, doc: Document) -> DocxParagraph:
    """Insert ADA_Vuln_Meta paragraph with bold label run and normal value run."""
    style_name = "ADA_Vuln_Meta"
    if ": " in line:
        label, _, value = line.partition(": ")
        label = label.strip() + ": "
        value = value.strip()
    else:
        label, value = "", line
    p = insert_paragraph_after(last, "", style=None)
    p.clear()
    if label:
        r1 = p.add_run(sanitize_text(label))
        r1.bold = True
    if value:
        p.add_run(sanitize_text(value))
    if _get_style_safe(doc, style_name) is None:
        available = _get_paragraph_style_names(doc)
        raise RuntimeError(
            f"Required paragraph style {style_name!r} not in template. "
            f"Available paragraph styles: {available!r}"
        )
    p.style = doc.styles[style_name]
    return p


def _split_ofc_text(ofc_text: str) -> list[str]:
    """Split ofcText into 1–4 bullets (newline, bullet, or numbered)."""
    if not ofc_text or not isinstance(ofc_text, str):
        return []
    t = ofc_text.replace("\u00a0", " ").strip()
    if not t:
        return []
    parts = re.split(r"\n+|[\u2022\u2023]\s*|(?:\d+\.)\s*", t)
    out = [p.strip() for p in parts if p.strip()]
    return out[:4]


def inject_themed_findings_at_anchor(doc: Document, anchor: str, findings: list, body_only: bool = True) -> int:
    """
    Replace anchor with Part II content from report_themed_findings (title, narrative, ofcText).
    Uses ADA_Vuln_* styles.
    """
    ensure_ada_vuln_styles(doc)
    preflight_assert_ada_styles(doc)
    candidates = list(find_anchor_paragraph_exact(doc, anchor, body_only=body_only))
    if len(candidates) != 1:
        raise RuntimeError(
            f"Anchor {anchor} must appear exactly once (found {len(candidates)} occurrences)"
        )
    p, _ = candidates[0]
    if paragraph_has_drawing(p):
        _clear_paragraph_text_preserve_drawings(p)
    else:
        p.clear()
    last = p
    count = 0
    for idx, f in enumerate(findings, start=1):
        title = (f.get("title") or "").strip()
        narrative = (f.get("narrative") or "").strip()
        ofc_text = (f.get("ofcText") or "").strip()
        severity = (f.get("severity") or "").strip()
        if not title:
            continue
        if narrative and narrative[-1] not in ".!?":
            narrative = narrative + "."
        header_text = f"VULNERABILITY {idx} {title}"
        last = _add_paragraph_with_style_after(last, header_text, "ADA_Vuln_Header", doc)
        count += 1
        if severity and severity.upper() in _SEVERITY_ORDER:
            sev_line = f"{severity.upper()} \u2013 Structural"
            last = _add_severity_paragraph_after(last, sev_line, doc)
            count += 1
        last = _add_paragraph_with_style_after(last, "Exposure Description", "ADA_Vuln_Label", doc)
        count += 1
        if narrative:
            last = _add_paragraph_with_style_after(last, narrative, "ADA_Vuln_Body", doc)
            count += 1
        bullets = _split_ofc_text(ofc_text)
        if bullets:
            last = _add_paragraph_with_style_after(last, "Options for Consideration", "ADA_Vuln_Label", doc)
            count += 1
            for bullet in bullets:
                last = _add_paragraph_with_style_after(last, bullet, "ADA_Vuln_Bullets", doc)
                count += 1
        last = insert_paragraph_after(last, "", style=None)
        count += 1
    if not paragraph_has_drawing(p):
        remove_paragraph(p)
    return count


def _part2_vuln_to_block(v: dict) -> dict | None:
    """Adapter: convert payload.report_vm.part2.vulnerabilities[] item to reporter block dict."""
    raw_title = (v.get("title") or "").strip()
    if not raw_title:
        return None
    title = sanitize_vulnerability_text(raw_title)
    raw_narrative = (v.get("narrative") or "").strip()
    narrative = debot_vulnerability_narrative(sanitize_backend_evidence(raw_narrative)) if raw_narrative else ""
    ofcs = v.get("ofcs")
    if isinstance(ofcs, list):
        ofcs = [sanitize_vulnerability_text(item if isinstance(item, str) else (item.get("text") or item.get("title") or "")) for item in ofcs[:4]]
    else:
        ofcs = _split_ofc_text((v.get("ofcText") or "").strip())
        ofcs = [sanitize_vulnerability_text(x) for x in ofcs if x][:4]
    ofcs = [x for x in ofcs if x]
    refs = v.get("references") or []
    if not isinstance(refs, list):
        refs = []
    refs = [sanitize_text(str(r).strip()) for r in refs[:10] if str(r).strip()]
    return {
        "severity": (v.get("severity") or "").strip(),
        "title": title,
        "narrative": narrative,
        "ofcs": ofcs,
        "references": refs,
    }


def _render_structured_vulnerabilities(doc: Document, anchor: str, vulns: list, body_only: bool = True) -> int:
    """
    Replace anchor with structured Part II vulnerabilities from report_vm.part2.vulnerabilities.
    Single source: same findings as in-app Summary. Uses ADA_Vuln_* styles.
    """
    ensure_ada_vuln_styles(doc)
    preflight_assert_ada_styles(doc)
    candidates = list(find_anchor_paragraph_exact(doc, anchor, body_only=body_only))
    if len(candidates) != 1:
        raise RuntimeError(
            f"Anchor {anchor} must appear exactly once (found {len(candidates)} occurrences)"
        )
    p, _ = candidates[0]
    if paragraph_has_drawing(p):
        _clear_paragraph_text_preserve_drawings(p)
    else:
        p.clear()
    last = p
    count = 0
    blocks: list[dict] = []
    for v in (vulns or []):
        block = _part2_vuln_to_block(v)
        if block and block.get("title"):
            blocks.append(block)
    if blocks:
        lead_in = (
            f"The following {len(blocks)} vulnerabilities represent the highest-priority "
            "dependency exposures identified from assessment responses."
        )
        last = _add_paragraph_with_style_after(last, lead_in, "ADA_Vuln_Body", doc)
        count += 1
    for idx, block in enumerate(blocks, start=1):
        title = block["title"]
        narrative = block.get("narrative") or ""
        severity = (block.get("severity") or "").strip()
        ofcs = block.get("ofcs") or []
        refs = block.get("references") or []
        if severity and severity.upper() in _SEVERITY_ORDER:
            header_text = f"VULNERABILITY {idx} {severity.upper()} \u2013 {title}"
        else:
            header_text = f"VULNERABILITY {idx} {title}"
        last = _add_paragraph_with_style_after(last, header_text, "ADA_Vuln_Header", doc)
        count += 1
        last = _add_severity_paragraph_after(last, "Structural", doc)
        count += 1
        last = _add_paragraph_with_style_after(last, "Exposure Description", "ADA_Vuln_Label", doc)
        count += 1
        if narrative:
            last = _add_paragraph_with_style_after(last, narrative, "ADA_Vuln_Body", doc)
            count += 1
        if ofcs:
            last = _add_paragraph_with_style_after(last, "Options for Consideration", "ADA_Vuln_Label", doc)
            count += 1
            for bullet in ofcs[:4]:
                last = _add_paragraph_with_style_after(last, bullet, "ADA_Vuln_Bullets", doc)
                count += 1
        if refs:
            last = _add_paragraph_with_style_after(last, "References", "ADA_Vuln_Label", doc)
            count += 1
            for r in refs:
                last = _add_paragraph_with_style_after(last, r, "ADA_Vuln_Bullets", doc)
                count += 1
        # Separation handled by style spacing; avoid explicit blank spacer paragraphs.
    if not paragraph_has_drawing(p):
        remove_paragraph(p)
    return count


def inject_part2_vulnerabilities_at_anchor(doc: Document, anchor: str, vulns: list, body_only: bool = True) -> int:
    """Replace anchor with Part II vulnerabilities from report_vm.part2.vulnerabilities."""
    return _render_structured_vulnerabilities(doc, anchor, vulns, body_only)


_VULN_FLAT_SPLIT = re.compile(
    r"(?=(?:VULNERABILITY \d+|Operational Consequence|Exposure Description|"
    r"Structural Indicator Profile|Standards Alignment|Options for Consideration|"
    r"Infrastructure Domain:\s*|Risk Type:\s*))",
    re.IGNORECASE,
)


def _vulnerability_block_to_lines(text: str) -> list[str]:
    """Split vulnerability block string into lines for one-paragraph-per-line rendering."""
    if not text or not isinstance(text, str):
        return []
    t = text.replace("\u00a0", " ").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not t:
        return []
    lines = [ln.rstrip() for ln in t.split("\n")]
    if len(lines) == 1 and len(lines[0]) > 400:
        one = lines[0]
        if "VULNERABILITY" in one and "Exposure Description" in one:
            parts = _VULN_FLAT_SPLIT.split(one)
            lines = [p.strip() for p in parts if p.strip()]
    filtered = _filter_vuln_lines_for_print(lines)
    return _merge_vuln_header_severity_lines(filtered)


def _merge_vuln_header_severity_lines(lines: list[str]) -> list[str]:
    """Keep both header and severity as separate lines for ADA_Vuln_Header and ADA_Vuln_Severity."""
    return lines


def _filter_vuln_lines_for_print(lines: list[str]) -> list[str]:
    """Remove Structural Indicator Profile and Standards Alignment sections from vulnerability block lines."""
    out: list[str] = []
    skip = False
    for line in lines:
        s = line.strip()
        if re.match(r"^VULNERABILITY \d+", s, re.IGNORECASE):
            skip = False
            out.append(line)
            continue
        if s.startswith("Structural Indicator Profile"):
            skip = True
            continue
        if s.startswith("Standards Alignment"):
            skip = True
            continue
        if s.startswith("Options for Consideration"):
            skip = False
            out.append(line)
            continue
        if skip:
            continue
        out.append(line)
    return out


def _add_vulnerability_block_paragraphs(doc: Document, last_para, text: str):
    """Insert vulnerability block content as multiple paragraphs (one per line). Returns (last_paragraph, count)."""
    ensure_ada_vuln_styles(doc)
    preflight_assert_ada_styles(doc)
    lines = _vulnerability_block_to_lines(text)
    count = 0
    prev_was_exposure_label = False
    for raw_line in lines:
        line_stripped = raw_line.strip()
        if not line_stripped:
            last_para = insert_paragraph_after(last_para, "", style=None)
            count += 1
            continue
        if prev_was_exposure_label and _classify_ada_vuln_line(line_stripped) == "ADA_Vuln_Body":
            line_stripped = dedupe_sentences(line_stripped)
            if not line_stripped:
                continue
        prev_was_exposure_label = line_stripped == "Exposure Description"
        style_name = _classify_ada_vuln_line(line_stripped)
        if style_name == "ADA_Vuln_Meta":
            last_para = _add_meta_paragraph_after(last_para, line_stripped, doc)
        elif style_name == "ADA_Vuln_Severity":
            last_para = _add_severity_paragraph_after(last_para, line_stripped, doc)
        else:
            last_para = _add_paragraph_with_style_after(last_para, line_stripped, style_name, doc)
        count += 1
    return last_para, count


def inject_vulnerability_blocks_at_anchor(doc: Document, anchor: str, text: str, body_only: bool = True) -> int:
    """
    LEGACY: Replace anchor with vulnerability block text, applying ADA_Vuln_* styles per line.
    Only used when REPORTER_USE_LEGACY_VULN_STRING=1.
    """
    text_trimmed = (text or "").replace("\u00a0", " ").strip()
    candidates = list(find_anchor_paragraph_exact(doc, anchor, body_only=body_only))
    if len(candidates) != 1:
        raise RuntimeError(
            f"Anchor {anchor} must appear exactly once (found {len(candidates)} occurrences)"
        )
    p, _ = candidates[0]
    if paragraph_has_drawing(p):
        _clear_paragraph_text_preserve_drawings(p)
    else:
        p.clear()
    last, count = _add_vulnerability_block_paragraphs(doc, p, text_trimmed)
    if not paragraph_has_drawing(p):
        remove_paragraph(p)
    return count
