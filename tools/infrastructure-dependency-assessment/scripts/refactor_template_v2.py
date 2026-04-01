#!/usr/bin/env python3
"""
Refactor Asset Dependency Assessment Report_BLANK.docx to ADA Report v2 structure.
STRUCTURAL CONTENT REFACTOR ONLY - preserves header, footer, FOUO, CISA branding, formatting.

Run from asset-dependency-tool: python scripts/refactor_template_v2.py
"""
from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError as e:
    print(f"ERROR: {e}", file=sys.stderr)
    sys.exit(2)

REPO_ROOT = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = REPO_ROOT / "assets" / "templates" / "Asset Dependency Assessment Report_BLANK.docx"
BACKUP_PATH = REPO_ROOT / "assets" / "templates" / "Asset Dependency Assessment Report_BLANK.v2backup.docx"


def _get_paragraph_text(elem) -> str:
    """Extract all text from a paragraph element."""
    texts = []
    for t in elem.iter():
        if hasattr(t, "text") and t.text:
            texts.append(t.text)
    return "".join(texts)


def _find_purpose_index(body) -> int:
    """Find index of first body child that contains PURPOSE (start of replaceable content)."""
    for i, child in enumerate(body):
        if child.tag == qn("w:p"):
            txt = _get_paragraph_text(child)
            if "PURPOSE" in txt and len(txt.strip()) < 50:
                return i
    return -1


def _find_sectpr_index(body) -> int:
    """Find sectPr (section properties) - must stay at end."""
    for i, child in enumerate(body):
        if child.tag == qn("w:sectPr"):
            return i
    return len(body)


def _create_anchor_paragraph(doc, anchor: str):
    """Create a new paragraph containing only the anchor text."""
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(anchor)
    return p


def _create_heading_paragraph(doc, text: str, style: str = "Heading 1"):
    """Create a heading paragraph."""
    p = doc.add_paragraph(text, style=style)
    return p


def _insert_body_content(doc, body, insert_after_idx: int, sectpr_idx: int):
    """
    Remove content from PURPOSE through sectPr (exclusive), then insert new structure.
    Preserves: elements before PURPOSE, sectPr at end.
    """
    # Collect elements to remove (from PURPOSE to sectPr exclusive)
    to_remove = list(body)[insert_after_idx:sectpr_idx]

    # Remove in reverse order to preserve indices
    for elem in reversed(to_remove):
        body.remove(elem)

    # Reference element to insert after (last preserved content before removed block)
    insert_after = list(body)[insert_after_idx - 1] if insert_after_idx > 0 else None

    # Build new structure as list of (element, insert_after) - we'll insert in order
    new_paras = []

    def add_para(text: str, style=None):
        p = doc.add_paragraph()
        p.clear()
        run = p.add_run(text)
        if style:
            try:
                p.style = style
            except Exception:
                pass
        new_paras.append((p._element, insert_after))

    def add_anchor(anchor: str):
        p = doc.add_paragraph()
        p.clear()
        p.add_run(anchor)
        new_paras.append((p._element, insert_after))

    # We need to insert after insert_after, then each subsequent after the previous
    # Use a different approach: create elements and insert one by one
    pass  # Will do insertion in main logic


def refactor_document(doc: Document) -> None:
    """
    Refactor document body to ADA Report v2 structure.
    Preserves: elements before PURPOSE (cover), sectPr.
    """
    body = doc.element.body
    children = list(body)

    purpose_idx = _find_purpose_index(body)
    sectpr_idx = _find_sectpr_index(body)

    if purpose_idx < 0:
        print("ERROR: PURPOSE section not found", file=sys.stderr)
        sys.exit(1)

    # Elements to remove: from purpose_idx through sectpr_idx-1
    to_remove = list(body)[purpose_idx:sectpr_idx]
    for elem in reversed(to_remove):
        body.remove(elem)

    # Insert point: after the last preserved element (purpose_idx - 1)
    insert_after = list(body)[purpose_idx - 1]

    # Build new structure - insert each element after insert_after, then update insert_after
    def add_para(text: str, style=None):
        nonlocal insert_after
        p = doc.add_paragraph()
        p.clear()
        p.add_run(text)
        if style:
            try:
                p.style = style
            except Exception:
                pass
        el = p._element
        el.getparent().remove(el)
        insert_after.addnext(el)
        insert_after = el

    def add_anchor(anchor: str):
        nonlocal insert_after
        p = doc.add_paragraph()
        p.clear()
        p.add_run(anchor)
        el = p._element
        el.getparent().remove(el)
        insert_after.addnext(el)
        insert_after = el

    def add_empty():
        nonlocal insert_after
        p = doc.add_paragraph()
        el = p._element
        el.getparent().remove(el)
        insert_after.addnext(el)
        insert_after = el

    # =========================================================================
    # ADA REPORT v2 PART I – EXECUTIVE DEPENDENCY RISK BRIEF
    # =========================================================================

    # Title (retitled per spec)
    add_para("Asset Dependency Risk Analysis", "Title")
    add_empty()

    # A. EXECUTIVE RISK POSTURE
    add_para("PART I – EXECUTIVE DEPENDENCY RISK BRIEF", "Heading 1")
    add_empty()
    add_para("A. EXECUTIVE RISK POSTURE", "Heading 2")
    add_anchor("[[SNAPSHOT_POSTURE]]")
    add_anchor("[[SNAPSHOT_SUMMARY]]")
    add_anchor("[[SNAPSHOT_DRIVERS]]")
    add_empty()

    # B. DEPENDENCY SNAPSHOT TABLE
    add_para("B. DEPENDENCY SNAPSHOT TABLE", "Heading 2")
    add_anchor("[[SNAPSHOT_MATRIX]]")
    add_anchor("[[SNAPSHOT_CASCADE]]")
    add_empty()

    # C. SECTOR ANALYSIS
    add_para("C. SECTOR ANALYSIS", "Heading 2")

    sectors = [
        ("ELECTRIC POWER", "[[CHART_ELECTRIC_POWER]]", "[[INFRA_ENERGY]]"),
        ("COMMUNICATIONS", "[[CHART_COMMUNICATIONS]]", "[[INFRA_COMMS]]"),
        ("INFORMATION TECHNOLOGY", "[[CHART_INFORMATION_TECHNOLOGY]]", "[[INFRA_IT]]"),
        ("WATER", "[[CHART_WATER]]", "[[INFRA_WATER]]"),
        ("WASTEWATER", "[[CHART_WASTEWATER]]", "[[INFRA_WASTEWATER]]"),
    ]
    for title, chart_anchor, infra_anchor in sectors:
        add_para(title, "Heading 3")
        add_anchor(chart_anchor)
        add_anchor(infra_anchor)
        add_empty()

    # D. CROSS-INFRASTRUCTURE SYNTHESIS
    add_para("D. CROSS-INFRASTRUCTURE SYNTHESIS", "Heading 2")
    add_anchor("[[SYNTHESIS]]")
    add_empty()

    # E. PRIORITY ACTIONS
    add_para("E. PRIORITY ACTIONS", "Heading 2")
    add_anchor("[[PRIORITY_ACTIONS]]")
    add_empty()

    # =========================================================================
    # PART II – TECHNICAL ANNEX
    # =========================================================================
    add_para("PART II – TECHNICAL ANNEX", "Heading 1")
    add_empty()
    add_para("Dependency Summary", "Heading 2")
    add_anchor("[[DEP_SUMMARY_TABLE]]")
    add_empty()
    add_para("Infrastructure Dependency Vulnerabilities and Options for Consideration", "Heading 2")
    add_anchor("[[TABLE_VOFC]]")


def main() -> int:
    if not TEMPLATE_PATH.is_file():
        print(f"ERROR: Template not found: {TEMPLATE_PATH}", file=sys.stderr)
        return 1

    # Backup
    import shutil

    shutil.copy(TEMPLATE_PATH, BACKUP_PATH)
    print(f"Backup: {BACKUP_PATH}", file=sys.stderr)

    doc = Document(str(TEMPLATE_PATH))
    refactor_document(doc)
    doc.save(str(TEMPLATE_PATH))
    print(f"Refactored: {TEMPLATE_PATH}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    sys.exit(main())
