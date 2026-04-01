#!/usr/bin/env python3
"""
Reporter CLI: read Assessment JSON from stdin, produce DOCX with charts and summary table.
Uses WORK_DIR env for working directory; prints only output path to stdout.
"""
import json
import os
import re
import sys
from pathlib import Path

# Startup dependency check: fail early with clear message (exit 2)
try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Inches
except Exception as e:
    print("ERROR: python-docx is not available", file=sys.stderr)
    print(str(e), file=sys.stderr)
    sys.exit(2)
try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
except ImportError:
    print("Install matplotlib: pip install matplotlib", file=sys.stderr)
    sys.exit(1)

# Category codes and display names for chart titles (charts only for these five)
CHART_CATEGORIES = (
    "ELECTRIC_POWER",
    "COMMUNICATIONS",
    "INFORMATION_TECHNOLOGY",
    "WATER",
    "WASTEWATER",
)
CATEGORY_DISPLAY = {
    "ELECTRIC_POWER": "Electric Power",
    "COMMUNICATIONS": "Communications",
    "INFORMATION_TECHNOLOGY": "Information Technology",
    "WATER": "Water",
    "WASTEWATER": "Wastewater",
    "CRITICAL_PRODUCTS": "Critical Products",  # no chart; contributes to VOFC only
}
CHART_ANCHORS = [f"[[CHART_{c}]]" for c in CHART_CATEGORIES]
TABLE_ANCHOR = "[[TABLE_SUMMARY]]"
TABLE_VOFC_ANCHOR = "[[TABLE_VOFC]]"
NARRATIVE_SOURCES_ANCHOR = "[[NARRATIVE_SOURCES]]"

INDEPENDENCE_LABELS = {
    "UNKNOWN": "Unknown",
    "SAME_DEMARCATION": "Same demarcation point (single point of failure)",
    "DIFFERENT_DEMARCATION_SAME_UPSTREAM": "Different demarcation, same upstream (partial resilience)",
    "DIFFERENT_LOOP_OR_PATH": "Different loop/path (resilient)",
}


def pct(value: float) -> float:
    return round(min(100, max(0, value)) * 10) / 10


def build_curve(category_input: dict, horizon_hours: int = 72, step_hours: int = 3) -> list[dict]:
    """Mirror of engine buildCurve: points with t_hours, capacity_without_backup, capacity_with_backup."""
    inp = dict(category_input) if isinstance(category_input, dict) else {}
    points = []
    for t in range(0, horizon_hours + 1, step_hours):
        if not inp.get("requires_service", False):
            cap_no, cap_with = 100.0, 100.0
        else:
            t_impact = inp.get("time_to_impact_hours", 0)
            if t < t_impact:
                cap_no = 100.0
            else:
                cap_no = (1 - inp.get("loss_fraction_no_backup", 0)) * 100
            if not _effective_has_backup(inp):
                cap_with = cap_no
            else:
                duration = inp.get("backup_duration_hours") or 0
                loss_with = inp.get("loss_fraction_with_backup", 0)
                if t < duration:
                    cap_with = 100.0
                else:
                    cap_with = (1 - loss_with) * 100
        points.append({
            "t_hours": t,
            "capacity_without_backup": pct(cap_no),
            "capacity_with_backup": pct(cap_with),
        })
    return points


def _effective_has_backup(inp: dict) -> bool:
    """True if has_backup_any or has_backup is True (engine compatibility)."""
    if inp.get("has_backup_any") is True:
        return True
    return inp.get("has_backup", False) is True


def _sources_summary(supply: dict | None) -> str | None:
    """SPOF/sources summary: 1 (SPOF likely) | 2+ (verify independence) | 2+ (independent) | None for CP."""
    if not supply:
        return "1 (SPOF likely)"
    if not supply.get("has_alternate_source"):
        return "1 (SPOF likely)"
    sources = supply.get("sources") or []
    has_independent = any(
        (s or {}).get("independence") == "DIFFERENT_LOOP_OR_PATH" for s in sources
    )
    return "2+ (independent)" if has_independent else "2+ (verify independence)"


def _sla_summary(agreements: dict | None) -> str:
    """Legacy: used when sla_reliability_for_report is not provided."""
    if not agreements or not agreements.get("has_sla"):
        return "No"
    h = agreements.get("sla_hours")
    return f"Yes ({h}h)" if h is not None else "Yes (ΓÇö)"


def _format_sla_cell_stakeholder(entry: dict) -> str:
    """
    Format SLA cell for report (stakeholder tone). No 'gap' wording.
    - sla_assessed false: do not state 'No SLA documented'; show 'SLA not assessed'.
    - sla_assessed true, sla_in_place NO: 'No SLA documented'
    - sla_assessed true, sla_in_place UNKNOWN: 'SLA status unknown'
    - sla_assessed true, sla_in_place YES: 'SLA documented', MTTR-Max if set, summary_text if set.
    """
    if not entry.get("sla_assessed"):
        return "SLA not assessed"
    in_place = entry.get("sla_in_place") or "UNKNOWN"
    if in_place == "NO":
        return "No SLA documented"
    if in_place == "UNKNOWN":
        return "SLA status unknown"
    # YES
    parts = ["SLA documented"]
    mttr = entry.get("mttr_max_hours")
    if mttr is not None and isinstance(mttr, (int, float)) and mttr == mttr:  # exclude NaN
        parts.append(f"MTTR-Max: {int(mttr)} hours")
    summary = (entry.get("summary_text") or "").strip()
    if summary:
        parts.append(summary)
    return " ".join(parts)


def _apply_sla_reliability_to_summary_rows(
    summary_rows: list[dict], sla_reliability_for_report: list[dict]
) -> None:
    """Override SLA column in summary rows with stakeholder-safe text from sla_reliability_for_report."""
    if not sla_reliability_for_report:
        return
    label_to_entry = {e.get("topic_label"): e for e in sla_reliability_for_report if e.get("topic_label")}
    for row in summary_rows:
        cat = row.get("category")
        label = CATEGORY_DISPLAY.get(cat) if cat else None
        if label and label in label_to_entry:
            row["sla"] = _format_sla_cell_stakeholder(label_to_entry[label])


def _pra_summary(agreements: dict | None) -> str:
    if not agreements or not agreements.get("has_pra"):
        return "No"
    cat = agreements.get("pra_category")
    if cat is None:
        return "Yes (ΓÇö)"
    if cat == "OTHER":
        other = (agreements.get("pra_category_other") or "").strip()
        return f"Yes (Other: {other})" if other else "Yes (Other: ΓÇö)"
    return f"Yes ({cat})"


def build_summary(assessment: dict) -> list[dict]:
    """Mirror of engine buildSummary: one row per category with required columns."""
    categories = assessment.get("categories") or {}
    rows = []
    for category, raw_inp in categories.items():
        inp = dict(raw_inp) if isinstance(raw_inp, dict) else {}
        if category == "CRITICAL_PRODUCTS" and "critical_products" in inp and inp.get("loss_fraction_no_backup") is None:
            rows.append({
                "category": category,
                "requires_service": False,
                "time_to_impact_hours": 0,
                "capacity_after_impact_no_backup": 100.0,
                "has_backup": False,
                "backup_duration_hours": None,
                "capacity_after_backup_exhausted": None,
                "recovery_time_hours": 0,
                "sources": None,
                "sla": None,
                "pra": None,
            })
            continue
        cap_no = pct((1 - inp.get("loss_fraction_no_backup", 0)) * 100)
        cap_exhausted = None
        if _effective_has_backup(inp) and inp.get("loss_fraction_with_backup") is not None:
            cap_exhausted = pct((1 - inp.get("loss_fraction_with_backup", 0)) * 100)
        supply = inp.get("supply")
        agreements = inp.get("agreements")
        sources_val = None if category == "CRITICAL_PRODUCTS" else _sources_summary(supply)
        sla_val = None if category == "CRITICAL_PRODUCTS" else _sla_summary(agreements)
        pra_val = None if category == "CRITICAL_PRODUCTS" else _pra_summary(agreements)
        rows.append({
            "category": category,
            "requires_service": inp.get("requires_service", False),
            "time_to_impact_hours": inp.get("time_to_impact_hours", 0),
            "capacity_after_impact_no_backup": cap_no,
            "has_backup": _effective_has_backup(inp),
            "backup_duration_hours": inp.get("backup_duration_hours"),
            "capacity_after_backup_exhausted": cap_exhausted,
            "recovery_time_hours": inp.get("recovery_time_hours", 0),
            "sources": sources_val,
            "sla": sla_val,
            "pra": pra_val,
        })
    return rows


def iter_paragraphs_and_cells(doc: Document):
    """Yield (paragraph, parent_element) for every paragraph in doc and in every table cell."""
    for p in doc.paragraphs:
        yield p, None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p, cell


def find_placeholder_or_anchor(doc: Document, pattern: str):
    """
    Search paragraphs and table cells for a placeholder or anchor string.
    Yields (paragraph, parent_cell_or_None) where paragraph text contains pattern.
    """
    for p, parent_cell in iter_paragraphs_and_cells(doc):
        if pattern in (p.text or ""):
            yield p, parent_cell


def search_placeholders_and_anchors(doc: Document) -> list[tuple[str, object, object]]:
    """
    Search all paragraphs and table cells for known placeholders and anchors.
    Returns list of (matched_text, paragraph, parent_cell_or_None) for each hit.
    """
    text_placeholders = ("{{ASSET_NAME}}", "{{VISIT_DATE}}", "{{ASSESSOR}}", "{{LOCATION}}")
    anchors = CHART_ANCHORS + [TABLE_ANCHOR, TABLE_VOFC_ANCHOR]
    hits = []
    for p, parent_cell in iter_paragraphs_and_cells(doc):
        text = p.text or ""
        for token in text_placeholders + anchors:
            if token in text:
                hits.append((token, p, parent_cell))
                break
    return hits


def replace_text_in_paragraph(paragraph, old: str, new: str) -> None:
    """Replace old with new in paragraph, preserving runs where possible."""
    if old not in (paragraph.text or ""):
        return
    # Simple approach: clear and set full text (loses formatting but works)
    for run in paragraph.runs:
        run.text = run.text.replace(old, new)


def replace_all_text_placeholders(doc: Document, asset: dict) -> None:
    """Replace {{ASSET_NAME}}, {{VISIT_DATE}}, {{ASSESSOR}}, {{LOCATION}} in doc."""
    mapping = {
        "{{ASSET_NAME}}": (asset.get("asset_name") or "").strip(),
        "{{VISIT_DATE}}": (asset.get("visit_date_iso") or "").strip(),
        "{{ASSESSOR}}": (asset.get("assessor") or "").strip(),
        "{{LOCATION}}": (asset.get("location") or "").strip(),
    }
    for p, _ in iter_paragraphs_and_cells(doc):
        for run in p.runs:
            for placeholder, value in mapping.items():
                if placeholder in (run.text or ""):
                    run.text = (run.text or "").replace(placeholder, value)


def replace_chart_anchor(doc: Document, anchor: str, image_path: Path) -> bool:
    """
    Find paragraph containing anchor, clear it and insert image in its place.
    Returns True if anchor was found and replaced.
    """
    for p, _ in find_placeholder_or_anchor(doc, anchor):
        p.clear()
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(5.5))
        return True
    return False


def _print_anchor_debug(doc: Document) -> None:
    """Print all paragraph texts that contain '[[' for debugging missing anchors."""
    for p, _ in iter_paragraphs_and_cells(doc):
        t = (p.text or "").strip()
        if "[[" in t:
            print(f"  [DEBUG] Paragraph containing '[[': {t!r}", file=sys.stderr)


def get_body_children(doc: Document):
    """Return list of body block elements (paragraphs, tables, etc.) for positioning."""
    body = doc.element.body
    return list(body)


def index_of_element(body_children, element) -> int | None:
    for i, child in enumerate(body_children):
        if child is element:
            return i
    return None


def insert_summary_table_at_anchor(doc: Document, summary_rows: list[dict]) -> None:
    """
    Find paragraph or table cell containing [[TABLE_SUMMARY]], remove it, insert a table with summary rows.
    If summary_rows is empty, insert a one-row table saying "No summary data available."
    """
    # Tolerant of empty sections: one placeholder row
    if not summary_rows:
        summary_rows = [{
            "category": "No summary data available.",
            "requires_service": False,
            "time_to_impact_hours": "",
            "capacity_after_impact_no_backup": "",
            "has_backup": False,
            "backup_duration_hours": "",
            "capacity_after_backup_exhausted": "",
            "recovery_time_hours": "",
            "sources": "",
            "sla": "",
            "pra": "",
        }]
    body = doc.element.body
    body_children = list(body)
    target_index = None
    target_para = None
    target_row_in_table = None  # when anchor is in a table cell
    # Search body paragraphs
    for child in body_children:
        if child.tag != qn("w:p"):
            continue
        text_parts = []
        for r in child.iter():
            if r.tag == qn("w:t") and r.text:
                text_parts.append(r.text)
        text = "".join(text_parts)
        if TABLE_ANCHOR in text:
            target_para = child
            target_index = index_of_element(body_children, child)
            break
    # Search table cells if not found in body
    if target_index is None:
        for p, parent_cell in find_placeholder_or_anchor(doc, TABLE_ANCHOR):
            if parent_cell is not None:
                # Anchor in a table cell: get row (tr) and table index in body
                tr = p._element.getparent()
                if tr is not None and tr.tag == qn("w:tr"):
                    tbl = tr.getparent()
                    if tbl is not None:
                        target_index = index_of_element(body_children, tbl)
                        target_row_in_table = tr
                        break
    if target_index is None and target_para is None and target_row_in_table is None:
        print(f"ERROR: Anchor {TABLE_ANCHOR!s} not found at insertion time", file=sys.stderr)
        _print_anchor_debug(doc)
        sys.exit(1)
    if target_row_in_table is not None:
        # Remove the row that contained the anchor, insert table after the table
        tbl = target_row_in_table.getparent()
        tbl.remove(target_row_in_table)
        target_index = index_of_element(list(body), tbl) + 1
    elif target_para is not None:
        body.remove(target_para)
    body_children = list(body)
    num_cols = 11
    table = doc.add_table(rows=1 + len(summary_rows), cols=num_cols)
    table.style = "Table Grid"
    headers = (
        "Category", "Requires Service", "Time to Impact (hrs)",
        "Capacity After Impact (No Backup)", "Has Backup", "Backup Duration (hrs)",
        "Capacity After Backup Exhausted", "Recovery Time (hrs)", "Sources",
        "SLA", "PRA",
    )
    for c, h in enumerate(headers):
        table.rows[0].cells[c].text = h
    for r, row_data in enumerate(summary_rows):
        row_idx = r + 1
        table.rows[row_idx].cells[0].text = row_data.get("category", "")
        table.rows[row_idx].cells[1].text = "Yes" if row_data.get("requires_service") else "No"
        table.rows[row_idx].cells[2].text = str(row_data.get("time_to_impact_hours", ""))
        table.rows[row_idx].cells[3].text = str(row_data.get("capacity_after_impact_no_backup", ""))
        table.rows[row_idx].cells[4].text = "Yes" if row_data.get("has_backup") else "No"
        bd = row_data.get("backup_duration_hours")
        table.rows[row_idx].cells[5].text = str(bd) if bd is not None else ""
        ce = row_data.get("capacity_after_backup_exhausted")
        table.rows[row_idx].cells[6].text = str(ce) if ce is not None else ""
        table.rows[row_idx].cells[7].text = str(row_data.get("recovery_time_hours", ""))
        table.rows[row_idx].cells[8].text = str(row_data.get("sources") or "")
        table.rows[row_idx].cells[9].text = str(row_data.get("sla") or "")
        table.rows[row_idx].cells[10].text = str(row_data.get("pra") or "")
    # add_table appends to body; we need to move this table to target_index.
    tbl_element = table._tbl
    body.remove(tbl_element)
    body.insert(target_index, tbl_element)


def build_sources_narrative(
    assessment: dict,
    sla_reliability_for_report: list[dict] | None = None,
) -> list[str]:
    """
    Build per-category narrative paragraphs for supply sources and SLA/PRA.
    Only non-CP categories are included. Each paragraph includes sources (if any) and SLA/PRA status.
    When sla_reliability_for_report is provided: SLA line is gated by sla_assessedΓÇönever "No SLA documented"
    when not assessed; when assessed use stakeholder wording (SLA documented / No SLA documented / etc).
    """
    categories = assessment.get("categories") or {}
    label_to_sla = {}
    if sla_reliability_for_report:
        label_to_sla = {
            e.get("topic_label"): e
            for e in sla_reliability_for_report
            if e.get("topic_label")
        }
    paragraphs = []
    for code in CHART_CATEGORIES:
        if code not in categories:
            continue
        inp = categories[code]
        display = CATEGORY_DISPLAY.get(code, code)
        supply = inp.get("supply")
        agreements = inp.get("agreements")
        parts = [f"{display}:"]
        sources = (supply or {}).get("sources") or []
        for i, s in enumerate(sources):
            if not s:
                continue
            provider = (s.get("provider_name") or "").strip() or "ΓÇö"
            parts.append(f"  Source {i + 1}: Provider {provider}.")
            desc = (s.get("demarcation_description") or "").strip()
            lat, lon = s.get("demarcation_lat"), s.get("demarcation_lon")
            if desc or (lat is not None and lon is not None):
                loc = desc or ""
                if lat is not None and lon is not None:
                    loc = f"{loc} (lat {lat}, lon {lon})".strip() if loc else f"lat {lat}, lon {lon}"
                parts.append(f"    Demarcation: {loc}.")
            ind = s.get("independence")
            if ind:
                parts.append(f"    Independence: {INDEPENDENCE_LABELS.get(ind, ind)}.")
        # SLA: use sla_reliability_for_report when present so we never say "No SLA documented" when not assessed
        sla_entry = label_to_sla.get(display) if label_to_sla else None
        if sla_entry is not None:
            if sla_entry.get("sla_assessed"):
                parts.append("  SLA: " + _format_sla_cell_stakeholder(sla_entry) + ".")
            # when not assessed: omit SLA line entirely (no phantom "No SLA documented")
        else:
            sla_text = _sla_summary(agreements)
            if sla_text == "No":
                parts.append("  SLA: No.")
            else:
                h = (agreements or {}).get("sla_hours")
                parts.append(f"  SLA: Yes; target hours: {h}." if h is not None else "  SLA: Yes.")
        pra_text = _pra_summary(agreements)
        if pra_text == "No":
            parts.append("  PRA: No.")
        else:
            cat = (agreements or {}).get("pra_category")
            if cat == "OTHER":
                other = ((agreements or {}).get("pra_category_other") or "").strip() or "ΓÇö"
                parts.append(f"  PRA: Yes; category: Other ({other}).")
            else:
                parts.append(f"  PRA: Yes; category: {cat}." if cat else "  PRA: Yes.")
        paragraphs.append(" ".join(parts))
    return paragraphs


def insert_sources_narrative_at_anchor(
    doc: Document,
    assessment: dict,
    sla_reliability_for_report: list[dict] | None = None,
) -> None:
    """
    If [[NARRATIVE_SOURCES]] is present, replace it with per-category source/SPOF narrative.
    If anchor is absent, do nothing (optional section).
    """
    body = doc.element.body
    body_children = list(body)
    target_index = None
    target_para = None
    for child in body_children:
        if child.tag != qn("w:p"):
            continue
        text_parts = []
        for r in child.iter():
            if r.tag == qn("w:t") and r.text:
                text_parts.append(r.text)
        text = "".join(text_parts)
        if NARRATIVE_SOURCES_ANCHOR in text:
            target_para = child
            target_index = index_of_element(body_children, child)
            break
    if target_index is None or target_para is None:
        return
    body.remove(target_para)
    narrative_paragraphs = build_sources_narrative(assessment, sla_reliability_for_report)
    if not narrative_paragraphs:
        narrative_paragraphs = ["No supply source details were provided for any dependency category."]
    for i, text in enumerate(narrative_paragraphs):
        para = doc.add_paragraph(text)
        body.remove(para._element)
        body.insert(target_index + i, para._element)


VOFC_NO_ITEMS_SENTENCE = (
    "No significant vulnerabilities were identified based on assessed dependencies."
)
VOFC_INTRO_PARAGRAPH = (
    "Vulnerabilities and Options for Consideration are derived from observed or inferred "
    "dependency conditions using a standardized evaluation framework."
)
VOFC_HEADERS = (
    "Category",
    "Vulnerability",
    "Impact",
    "Option for Consideration",
    "Applicability",
    "Severity (Calibrated)",
    "Severity (Base)",
    "Notes",
)


def _set_cell_bold(cell) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True


def insert_vofc_table_at_anchor(doc: Document, vofc_collection: dict | None) -> None:
    """
    Find paragraph or table cell containing [[TABLE_VOFC]], remove it, insert VOFC table or no-items sentence.
    vofc_collection: { "items": [ ... ] }. Empty items -> "No significant vulnerabilities..." sentence.
    """
    body = doc.element.body
    body_children = list(body)
    target_index = None
    target_para = None
    target_row_in_table = None
    for child in body_children:
        if child.tag != qn("w:p"):
            continue
        text_parts = []
        for r in child.iter():
            if r.tag == qn("w:t") and r.text:
                text_parts.append(r.text)
        text = "".join(text_parts)
        if TABLE_VOFC_ANCHOR in text:
            target_para = child
            target_index = index_of_element(body_children, child)
            break
    if target_index is None:
        for p, parent_cell in find_placeholder_or_anchor(doc, TABLE_VOFC_ANCHOR):
            if parent_cell is not None:
                tr = p._element.getparent()
                if tr is not None and tr.tag == qn("w:tr"):
                    tbl = tr.getparent()
                    if tbl is not None:
                        target_index = index_of_element(body_children, tbl) + 1
                        target_row_in_table = tr
                        break
    if target_index is None or (target_para is None and target_row_in_table is None):
        print(f"ERROR: Anchor {TABLE_VOFC_ANCHOR!s} not found at insertion time", file=sys.stderr)
        _print_anchor_debug(doc)
        sys.exit(1)
    if target_row_in_table is not None:
        tbl = target_row_in_table.getparent()
        tbl.remove(target_row_in_table)
        target_index = index_of_element(list(body), tbl) + 1
    else:
        body.remove(target_para)
    items = (vofc_collection or {}).get("items") or []
    if not items:
        para = doc.add_paragraph(VOFC_NO_ITEMS_SENTENCE)
        body.remove(para._element)
        body.insert(target_index, para._element)
        return
    intro_para = doc.add_paragraph(VOFC_INTRO_PARAGRAPH)
    body.remove(intro_para._element)
    body.insert(target_index, intro_para._element)
    num_cols = len(VOFC_HEADERS)
    table = doc.add_table(rows=1 + len(items), cols=num_cols)
    table.style = "Table Grid"
    for c, h in enumerate(VOFC_HEADERS):
        cell = table.rows[0].cells[c]
        cell.text = h
        _set_cell_bold(cell)
    for r, row_data in enumerate(items):
        row_idx = r + 1
        table.rows[row_idx].cells[0].text = str(row_data.get("category", ""))
        table.rows[row_idx].cells[1].text = str(row_data.get("vulnerability", ""))
        impact = row_data.get("impact")
        table.rows[row_idx].cells[2].text = str(impact) if impact is not None else ""
        table.rows[row_idx].cells[3].text = str(row_data.get("option_for_consideration", ""))
        table.rows[row_idx].cells[4].text = str(row_data.get("applicability", ""))
        cal_sev = row_data.get("calibrated_severity", "")
        table.rows[row_idx].cells[5].text = str(cal_sev) if cal_sev else ""
        base_sev = row_data.get("base_severity", "")
        table.rows[row_idx].cells[6].text = str(base_sev) if base_sev else ""
        notes = "" if cal_sev == base_sev else (row_data.get("calibration_reason") or "")
        table.rows[row_idx].cells[7].text = str(notes).strip()
    tbl_element = table._tbl
    body.remove(tbl_element)
    body.insert(target_index + 1, tbl_element)


# SLA marker color (stakeholder report)
SLA_MTTR_COLOR = "#e87500"


def generate_chart(
    points: list[dict],
    title: str,
    out_path: Path,
    sla_mttr_hours: int | float | None = None,
) -> None:
    """
    Write a PNG chart: t_hours vs capacity %, two lines (without/with backup).
    When sla_mttr_hours is set (and finite), draw a vertical line at that hour (SLA MTTR marker).
    Only call with sla_mttr_hours when SLA was assessed and in place with a max MTTR.
    """
    t = [x["t_hours"] for x in points]
    no_backup = [x["capacity_without_backup"] for x in points]
    with_backup = [x["capacity_with_backup"] for x in points]
    plt.figure(figsize=(7, 4))
    plt.plot(t, no_backup, label="Without backup", marker=".", markersize=4)
    plt.plot(t, with_backup, label="With backup", marker=".", markersize=4)
    if sla_mttr_hours is not None and isinstance(sla_mttr_hours, (int, float)):
        try:
            h = float(sla_mttr_hours)
            if h == h and h >= 0:
                plt.axvline(x=h, color=SLA_MTTR_COLOR, linestyle="--", linewidth=1.5, label="SLA MTTR-Max")
        except (TypeError, ValueError):
            pass
    plt.xlabel("Time (hours)")
    plt.ylabel("Capacity %")
    plt.title(title)
    plt.ylim(0, 105)
    plt.legend()
    plt.tight_layout()
    plt.savefig(out_path, dpi=120)
    plt.close()


def append_energy_dependency_section(doc: Document, energy_dependency: dict) -> None:
    """
    Append Energy dependency section when energy_dependency payload is present.
    Expects: dataBlocks (list), vulnerabilities (list), ofcs (list).
    """
    data_blocks = energy_dependency.get("dataBlocks") or []
    vulnerabilities = energy_dependency.get("vulnerabilities") or []
    ofcs = energy_dependency.get("ofcs") or []
    ofc_by_vuln = {o.get("vulnerability_id"): o for o in ofcs if o.get("vulnerability_id")}

    doc.add_paragraph()
    h = doc.add_paragraph("Energy ΓÇö Dependency Data")
    h.runs[0].bold = True
    for block in data_blocks:
        btype = block.get("type")
        title = block.get("title") or ""
        if title:
            doc.add_paragraph(title)
        if btype == "narrative":
            doc.add_paragraph(block.get("text") or "")
        elif btype == "list":
            for item in block.get("items") or []:
                doc.add_paragraph(f"ΓÇó {item}")
        elif btype == "table":
            headers = block.get("headers") or []
            rows = block.get("rows") or []
            if headers and rows is not None:
                t = doc.add_table(rows=1 + len(rows), cols=len(headers))
                t.style = "Table Grid"
                for c, hdr in enumerate(headers):
                    t.cell(0, c).text = str(hdr)
                for r, row in enumerate(rows):
                    for c, cell in enumerate(row):
                        if c < len(headers):
                            t.cell(r + 1, c).text = str(cell) if cell is not None else ""
                doc.add_paragraph()

    doc.add_paragraph()
    h2 = doc.add_paragraph("Energy ΓÇö Vulnerabilities and Options for Consideration")
    h2.runs[0].bold = True
    for v in vulnerabilities:
        vuln_id = v.get("id")
        text = v.get("text") or ""
        doc.add_paragraph(f"Vulnerability: {text}")
        ofc = ofc_by_vuln.get(vuln_id) if vuln_id else None
        if ofc and ofc.get("text"):
            doc.add_paragraph(f"Option for consideration: {ofc.get('text')}")
        doc.add_paragraph()


def main() -> None:
    work_dir = os.environ.get("WORK_DIR")
    if not work_dir:
        # Default when run manually: repo/data/temp/reporter
        script_dir = Path(__file__).resolve().parent
        repo_root = script_dir.parent.parent
        work_dir = str(repo_root / "data" / "temp" / "reporter")
    work_path = Path(work_dir)
    work_path.mkdir(parents=True, exist_ok=True)

    # Read JSON from stdin: either { "assessment": {...}, "vofc_collection": {...} } or bare assessment
    raw = sys.stdin.read()
    if not raw.strip():
        print("No JSON input on stdin", file=sys.stderr)
        sys.exit(1)
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"Invalid JSON: {e}", file=sys.stderr)
        sys.exit(1)
    if isinstance(data, dict) and "assessment" in data:
        assessment = data["assessment"]
        vofc_collection = data.get("vofc_collection")
        sla_reliability_for_report = data.get("sla_reliability_for_report") or []
    else:
        assessment = data
        vofc_collection = None
        sla_reliability_for_report = []

    # Template path: TEMPLATE_PATH env (deployment) or repo-relative (dev)
    template_path = os.environ.get("TEMPLATE_PATH")
    if template_path:
        template_path = Path(template_path)
    else:
        script_dir = Path(__file__).resolve().parent
        repo_root = script_dir.parent.parent
        template_path = repo_root / "assets" / "templates" / "Asset Dependency Assessment Report_BLANK.docx"
    if not template_path.is_file():
        print(f"Template not found: {template_path}", file=sys.stderr)
        sys.exit(1)

    categories = assessment.get("categories") or {}
    # Map category display name -> SLA entry for charts (SLA marker only when assessed + YES + MTTR)
    label_to_sla = {}
    if sla_reliability_for_report:
        label_to_sla = {
            e.get("topic_label"): e
            for e in sla_reliability_for_report
            if e.get("topic_label")
        }
    chart_paths = {}
    for code in CHART_CATEGORIES:
        if code not in categories:
            continue
        display = CATEGORY_DISPLAY[code]
        points = build_curve(categories[code])
        png_path = work_path / f"chart_{code}.png"
        sla_entry = label_to_sla.get(display)
        mttr_hours = None
        if sla_entry and sla_entry.get("sla_assessed") and sla_entry.get("sla_in_place") == "YES":
            m = sla_entry.get("mttr_max_hours")
            if m is not None and isinstance(m, (int, float)) and m == m:
                mttr_hours = m
        generate_chart(points, display, png_path, sla_mttr_hours=mttr_hours)
        chart_paths[code] = png_path

    # Open template and replace placeholders / anchors
    doc = Document(str(template_path))

    asset = assessment.get("asset") or {}
    replace_all_text_placeholders(doc, asset)

    for anchor in CHART_ANCHORS:
        match = re.match(r"\[\[CHART_(.+)\]\]", anchor)
        if match and chart_paths.get(match.group(1)):
            if not replace_chart_anchor(doc, anchor, chart_paths[match.group(1)]):
                print(f"ERROR: Anchor {anchor!s} not found at insertion time", file=sys.stderr)
                _print_anchor_debug(doc)
                sys.exit(1)

    summary_rows = build_summary(assessment)
    _apply_sla_reliability_to_summary_rows(summary_rows, sla_reliability_for_report)
    insert_summary_table_at_anchor(doc, summary_rows)
    insert_sources_narrative_at_anchor(doc, assessment, sla_reliability_for_report)
    insert_vofc_table_at_anchor(doc, vofc_collection)

    energy_dependency = data.get("energy_dependency")
    if energy_dependency:
        append_energy_dependency_section(doc, energy_dependency)

    output_path = work_path / "output.docx"
    doc.save(str(output_path))
    print(output_path)


if __name__ == "__main__":
    main()
