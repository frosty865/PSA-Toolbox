#!/usr/bin/env python3
"""
One-time fix: insert [[TABLE_SUMMARY]] anchor on its own paragraph immediately after
the Summary section heading in the production template. Run from repo root:
  python scripts/add_table_summary_anchor.py
  python scripts/add_table_summary_anchor.py "C:\path\to\template.docx"
Requires: pip install python-docx
"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
except ImportError:
    print("Install python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(2)

ANCHOR = "[[TABLE_SUMMARY]]"
DEFAULT_TEMPLATE = Path(__file__).resolve().parent.parent.parent / "assets" / "templates" / "Asset Dependency Assessment Report_BLANK.docx"


def insert_paragraph_after(paragraph, text: str):
    """Insert a new paragraph after the given paragraph with the given text."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    return new_para


def main() -> int:
    template_path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_TEMPLATE
    template_path = template_path.resolve()

    if not template_path.is_file():
        print(f"Template not found: {template_path}", file=sys.stderr)
        return 1

    doc = Document(str(template_path))
    body_text = "\n".join(p.text for p in doc.paragraphs)
    if ANCHOR in body_text:
        print(f"Anchor {ANCHOR!r} already present. No change.")
        return 0

    # Find first body paragraph that looks like the Summary section heading
    summary_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t and "summary" in t.lower() and len(t) < 80:
            summary_idx = i
            break

    if summary_idx is None:
        print("Could not find a 'Summary' section heading in the document body.", file=sys.stderr)
        print("Add the anchor manually under the Summary heading. See docs/TEMPLATE_ANCHORS.md", file=sys.stderr)
        return 1

    target = doc.paragraphs[summary_idx]
    insert_paragraph_after(target, ANCHOR)
    doc.save(str(template_path))
    print(f"Inserted {ANCHOR!r} after paragraph {summary_idx + 1} ({target.text[:50]!r}...).")
    print(f"Saved: {template_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
