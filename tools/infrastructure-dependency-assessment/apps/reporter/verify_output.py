"""
Verify generated output.docx: no unreplaced anchors, required structure present,
no narrative blanks, and no raw export-style tables (parity check).
Exit 0 on success, non-zero on failure.
Usage: python verify_output.py [path_to_output.docx]
Default path: output.docx in current directory (or WORK_DIR/output.docx if set).
Env: SLA_PRA_EXPECTED=1 — require SLA/PRA summary section (set when template had [[SLA_PRA_SUMMARY]]).
"""
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    sys.stderr.write("python-docx required: pip install python-docx\n")
    sys.exit(2)


# At least one of these must appear (template may use "EXECUTIVE SUMMARY" instead of "Summary")
REQUIRED_HEADING_ANY = ("Summary", "EXECUTIVE SUMMARY")
REQUIRED_HEADINGS = (
    "Vulnerabilities and Options for Consideration",
)
# Summary table must include all categories (display or code)
REQUIRED_SUMMARY_CATEGORY_ANY = ("CRITICAL_PRODUCTS", "Critical Products")

# VOFC table deprecated; narrative-only ([[VULN_NARRATIVE]]). Output must NOT contain VOFC-style table.
VOFC_HEADER_EXPECTED = ("Category", "Vulnerability", "Option for Consideration")

# SLA/PRA summary section title (must appear when reporter inserts the block)
SLA_PRA_SUMMARY_TITLE = "Service Restoration Reliability Summary"

# Must match main.EXPORT_TABLE_BAD_HEADERS: export-style headers that must NOT appear (exact match;
# template summary uses "Recovery Time (hrs)" which must not be flagged)
EXPORT_TABLE_BAD_HEADERS = (
    "Requires Service",
    "Time to Impact",
    "Loss of Function",
    "Recovery Time",
    "Percent",
    "Time to Impact (hrs)",
    "Capacity After Impact (No Backup)",
)
UNDERSCORE_RE = re.compile(r"_{3,}")
CHOOSE_AN_ITEM = "Choose an item."
NO_VULNERABILITIES_IDENTIFIED = "No vulnerabilities identified"
VISUALIZATION_SECTION_START = "Asset Dependency Visualization"

# D1: Mojibake - UTF-8 misinterpreted as Latin-1
MOJIBAKE_PATTERN = "\u0393\u00c7"  # "ΓÇ" prefix of ΓÇó, ΓÇö, ΓÇô
# D1: Forbidden placeholders
FORBIDDEN_PLACEHOLDERS = ("TBD", "Insert ", "Region__")
# D1: Deprecated reference
SAFE_FORBIDDEN = "SAFE"
# D1: Generator language outside Energy
BACKUP_GENERATOR_PHRASE = "backup generator"
# C3: Physical security keywords in VOFC
PHYSICAL_SECURITY_KEYWORDS = (
    "cctv", "ids", "badging", "keycard", "access levels", "terminated personnel",
)
VOFC_SECTION_HEADER = "Vulnerabilities and Options for Consideration"
CYBER_TERMS_IN_VISUALIZATION = ("cyber", "cybersecurity", "NIST", "US-CERT", "ICS-CERT", "ISAC")
ABOVE_CHART_PHRASE = "In the above chart"
MAX_CHART_NARRATIVE_PARAS = 5


def get_all_text(doc: "Document") -> str:
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def get_tables_with_headers(doc: "Document") -> list[tuple[list[str], int]]:
    """Return list of (header_cell_texts, num_columns) for each table."""
    result = []
    for table in doc.tables:
        if not table.rows:
            result.append(([], 0))
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        result.append((headers, len(table.rows[0].cells)))
    return result


def iter_block_items_verify(doc: "Document"):
    """Yield ('paragraph', text) or ('table', None) in document order."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            text_parts = []
            for el in child.iter():
                if el.tag == qn("w:t") and el.text:
                    text_parts.append(el.text)
            yield "paragraph", "".join(text_parts)
        elif child.tag == qn("w:tbl"):
            yield "table", None


def main() -> int:
    if len(sys.argv) > 1:
        path = Path(sys.argv[1])
    else:
        work_dir = Path(__file__).resolve().parent
        if "WORK_DIR" in __import__("os").environ:
            work_dir = Path(__import__("os").environ["WORK_DIR"])
        path = work_dir / "output.docx"

    if not path.exists():
        sys.stderr.write(f"File not found: {path}\n")
        return 1

    try:
        doc = Document(str(path))
    except Exception as e:
        sys.stderr.write(f"Failed to open DOCX: {e}\n")
        return 1

    text = get_all_text(doc)
    errors = []

    # No anchor tokens or dev text in any paragraph
    if "[[" in text:
        errors.append("Output still contains unreplaced anchor token(s) containing '[['")
    if "ANCHOR BLOCK" in text:
        errors.append("Output still contains 'ANCHOR BLOCK (dev only)' or similar")

    # No narrative blanks left (____ or longer underscores)
    if UNDERSCORE_RE.search(text):
        errors.append("Output still contains unreplaced narrative blank (____)")

    # E2: No placeholder text
    if CHOOSE_AN_ITEM in text:
        errors.append(f"Output must not contain placeholder text {CHOOSE_AN_ITEM!r}")

    # D1: No mojibake (ΓÇ prefix)
    if MOJIBAKE_PATTERN in text:
        errors.append("Output must not contain mojibake sequences (ΓÇ); use proper Unicode bullet/dash.")

    # D1: No forbidden placeholders
    for forbidden in FORBIDDEN_PLACEHOLDERS:
        if forbidden in text:
            errors.append(f"Output must not contain placeholder {forbidden!r}")

    # D1: No SAFE reference
    if SAFE_FORBIDDEN in text:
        errors.append("Output must not contain deprecated SAFE reference.")

    if not any(h in text for h in REQUIRED_HEADING_ANY):
        errors.append("Required heading missing: one of " + str(REQUIRED_HEADING_ANY))
    for heading in REQUIRED_HEADINGS:
        if heading not in text:
            errors.append(f"Required heading missing: {heading!r}")
    if not any(cat in text for cat in REQUIRED_SUMMARY_CATEGORY_ANY):
        errors.append(f"Summary table must include category: one of {REQUIRED_SUMMARY_CATEGORY_ANY!r}")

    # Parity: no VOFC table (deprecated); vulnerability content is narrative-only ([[VULN_NARRATIVE]]). No export-style tables.
    tables_info = get_tables_with_headers(doc)
    vofc_style_found = any(
        num_cols >= 3
        and (headers[0] or "").strip() == "Category"
        and (headers[1] or "").strip() == "Vulnerability"
        and "Option" in ((headers[2] or "").strip())
        for headers, num_cols in tables_info
    )
    if vofc_style_found:
        errors.append(
            "Output must not contain VOFC table (Category | Vulnerability | Option for Consideration). "
            "Use [[VULN_NARRATIVE]] narrative-only export."
        )
    export_bad = None
    for table in doc.tables:
        for row in table.rows[:2]:
            for cell in row.cells:
                c = (cell.text or "").strip()
                for bad in EXPORT_TABLE_BAD_HEADERS:
                    if c == bad:  # exact match so "Recovery Time (hrs)" is allowed
                        export_bad = bad
                        break
            if export_bad:
                break
        if export_bad:
            break
    if export_bad:
        errors.append(f"Export-style table detected (header={export_bad!r}). Use template format only.")

    # SLA/PRA summary block: only require when template had [[SLA_PRA_SUMMARY]] (SLA_PRA_EXPECTED=1)
    if os.environ.get("SLA_PRA_EXPECTED") == "1" and SLA_PRA_SUMMARY_TITLE not in text:
        errors.append(
            f"Expected SLA/PRA summary section with title {SLA_PRA_SUMMARY_TITLE!r} in output."
        )

    # E5: "single source" must never appear (no inference in Notes)
    if "single source" in text.lower():
        errors.append("Output must not contain 'single source' (no inference; use explicit provider/source data only).")

    # E6: Internal category IDs must not leak into DOCX (use canonical display names only)
    INTERNAL_CATEGORY_IDS = ("ELECTRIC_POWER", "INFORMATION_TECHNOLOGY", "CRITICAL_PRODUCTS")
    for internal_id in INTERNAL_CATEGORY_IDS:
        if internal_id in text:
            errors.append(f"Output must not contain internal category ID {internal_id!r}; use canonical display name.")

    # E7: Exec Summary section only must not contain per-dependency injected lines
    # Sector narratives (Communications:, etc.) belong in CRITICAL INFRASTRUCTURE, not Exec Summary
    EXEC_SUMMARY_FORBIDDEN_SIGNATURES = (
        "Electric Power:",
        "Communications:",
        "Information Technology:",
        "Water:",
        "Wastewater:",
        "Loss of service impacts operations",
    )
    exec_summary_text_parts = []
    in_exec_summary = False
    for kind, para_text in iter_block_items_verify(doc):
        pt = (para_text or "").strip()
        if kind == "paragraph" and pt == "EXECUTIVE SUMMARY":
            in_exec_summary = True
        elif kind == "paragraph" and (
            VISUALIZATION_SECTION_START in pt
            or "CRITICAL INFRASTRUCTURE" in pt
            or "SECTOR ANALYSIS" in pt
            or "ELECTRIC POWER" == pt
        ):
            in_exec_summary = False
        elif in_exec_summary and kind == "paragraph" and pt:
            exec_summary_text_parts.append(pt)
    exec_summary_text = " ".join(exec_summary_text_parts).lower()
    for sig in EXEC_SUMMARY_FORBIDDEN_SIGNATURES:
        if sig.lower() in exec_summary_text:
            errors.append(
                f"Executive Summary must not contain injected per-dependency text ({sig!r}); "
                "keep template-native narrative only."
            )

    # EXECUTIVE SUMMARY: no table between it and Asset Dependency Visualization
    blocks = list(iter_block_items_verify(doc))
    in_exec_summary = False
    tables_in_exec_summary = 0
    for i, (kind, para_text) in enumerate(blocks):
        pt = (para_text or "").strip()
        if kind == "paragraph" and pt == "EXECUTIVE SUMMARY":
            in_exec_summary = True
        elif kind == "paragraph" and VISUALIZATION_SECTION_START in pt:
            in_exec_summary = False
        elif in_exec_summary and kind == "table":
            tables_in_exec_summary += 1
    if tables_in_exec_summary > 0:
        errors.append(
            f"Executive Summary must not contain any table; found {tables_in_exec_summary} "
            "table(s) between EXECUTIVE SUMMARY and Asset Dependency Visualization."
        )
    # Also check immediate follow (legacy regression guard)
    for i, (kind, para_text) in enumerate(blocks):
        if kind == "paragraph" and (para_text or "").strip() == "EXECUTIVE SUMMARY":
            if i + 1 < len(blocks) and blocks[i + 1][0] == "table":
                errors.append("EXECUTIVE SUMMARY is immediately followed by a table (bug regression).")
            break

    # E2: "No vulnerabilities identified" must not appear in a table that also has other vulnerability rows
    for table in doc.tables:
        if not table.rows or len(table.rows) < 2:
            continue
        has_no_vuln_row = False
        for row in table.rows[1:]:
            cell_text = " ".join((c.text or "").strip() for c in row.cells)
            if NO_VULNERABILITIES_IDENTIFIED in cell_text:
                has_no_vuln_row = True
                break
        if has_no_vuln_row and len(table.rows) > 2:
            errors.append(
                "Table contains 'No vulnerabilities identified' and other vulnerability rows in the same subsection."
            )

    # E3: At most 5 chart narrative paragraphs (no duplicated "In the above chart..." blocks)
    above_chart_count = text.count(ABOVE_CHART_PHRASE)
    if above_chart_count > MAX_CHART_NARRATIVE_PARAS:
        errors.append(
            f"Too many chart narrative paragraphs ({above_chart_count} containing {ABOVE_CHART_PHRASE!r}); expected at most {MAX_CHART_NARRATIVE_PARAS}."
        )

    # C3: VOFC rows - dependency category must not have physical security keywords; D1: no "backup generator" in non-Energy
    for table in doc.tables:
        if not table.rows or len(table.rows) < 2:
            continue
        headers = [c.text.strip().lower() for c in table.rows[0].cells]
        cat_idx = next((i for i, h in enumerate(headers) if "category" in h), -1)
        vuln_idx = next((i for i, h in enumerate(headers) if "vulnerability" in h), -1)
        ofc_idx = next((i for i, h in enumerate(headers) if "option" in h), vuln_idx + 1 if vuln_idx >= 0 else -1)
        if cat_idx < 0 or vuln_idx < 0:
            continue
        dep_cats = ("energy", "electric", "communications", "information technology", "water", "wastewater")
        non_energy = ("communications", "information technology", "water", "wastewater")
        for row in table.rows[1:]:
            cells = row.cells
            if cat_idx >= len(cells) or vuln_idx >= len(cells):
                continue
            cat = (cells[cat_idx].text or "").strip().lower()
            vuln = (cells[vuln_idx].text or "").strip().lower()
            ofc = (cells[ofc_idx].text or "").strip().lower() if 0 <= ofc_idx < len(cells) else ""
            if any(dc in cat for dc in dep_cats):
                for kw in PHYSICAL_SECURITY_KEYWORDS:
                    if kw in vuln:
                        errors.append(
                            f"VOFC row: dependency category '{cat}' has physical security keyword '{kw}' in vulnerability."
                        )
                        break
            if any(ne in cat for ne in non_energy):
                if BACKUP_GENERATOR_PHRASE.lower() in vuln or BACKUP_GENERATOR_PHRASE.lower() in ofc:
                    errors.append(
                        f"VOFC row: non-Energy category '{cat}' must not contain 'backup generator'."
                    )

    # E4: No cyber text in Asset Dependency Visualization section (between section start and VOFC header)
    in_visualization = False
    for kind, para_text in blocks:
        if kind != "paragraph" or not para_text:
            continue
        pt = (para_text or "").strip()
        if VISUALIZATION_SECTION_START in pt:
            in_visualization = True
            continue
        if in_visualization and VOFC_SECTION_HEADER in pt:
            break
        if in_visualization:
            for term in CYBER_TERMS_IN_VISUALIZATION:
                if term.lower() in pt.lower():
                    errors.append(
                        f"Cyber-related text ({term!r}) must not appear in Asset Dependency Visualization section."
                    )
                    break

    if errors:
        for e in errors:
            sys.stderr.write(e + "\n")
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
