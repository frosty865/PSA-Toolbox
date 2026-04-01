"""
DOCX document operations: paragraph/table iteration, anchor search/replace, page breaks, table helpers.
Extracted from main.py; imports constants only where needed.
"""
from __future__ import annotations

import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_BREAK
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from sanitize import sanitize_text

from constants import ANNEX_OVERVIEW_HEADING, SECTOR_REPORTS_HEADING


def _iter_paragraphs(doc: Document):
    """Yield every paragraph in doc (body and inside all table cells)."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def iter_block_items(doc: Document):
    """
    Yield Paragraph or Table in true document order (body block items only).
    Uses doc._element.body.iterchildren(); wraps w:p -> Paragraph, w:tbl -> Table.
    """
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, doc)


def _doc_contains_anchor(doc: Document, token: str, body_only: bool = True) -> bool:
    """Return True if any paragraph (or cell) contains exact text token."""
    needle = (token or "").strip()
    if body_only:
        for p in doc.paragraphs:
            if (p.text or "").strip() == needle:
                return True
        return False
    for p in _iter_paragraphs(doc):
        if (p.text or "").strip() == needle:
            return True
    return False


def _normalize_paragraph_text(t: str) -> str:
    """Collapse any run of whitespace to a single space and strip. Helps match anchors split across runs."""
    if not t:
        return ""
    return re.sub(r"\s+", " ", (t or "").strip())


def find_paragraph_by_exact_text(doc: Document, needle: str, body_only: bool = False):
    """
    Find paragraph whose trimmed text equals needle. If body_only=True, search only body paragraphs (not table cells).
    Tries exact match first, then normalized (whitespace collapsed) so anchors split across runs still match.
    Returns the first matching Paragraph or None.
    """
    needle = (needle or "").strip()
    needle_norm = _normalize_paragraph_text(needle)
    if body_only:
        for p in doc.paragraphs:
            raw = (p.text or "").strip()
            if raw == needle or _normalize_paragraph_text(raw) == needle_norm:
                return p
        return None
    for p in _iter_paragraphs(doc):
        raw = (p.text or "").strip()
        if raw == needle or _normalize_paragraph_text(raw) == needle_norm:
            return p
    return None


def _element_has_drawing_or_pict_child(el) -> bool:
    """True if element or any descendant has local tag name 'drawing', 'pict', or 'blip' (OOXML inline/floating image)."""
    for child in el.iter():
        local = (child.tag or "").split("}")[-1] if "}" in str(child.tag) else (child.tag or "")
        if local in ("drawing", "pict", "blip"):
            return True
    return False


def _run_has_drawing(run) -> bool:
    """True if run contains inline picture or drawing (w:drawing / w:pict)."""
    return _element_has_drawing_or_pict_child(run._element)


def paragraph_has_drawing(p) -> bool:
    """Detect inline pictures and floating drawings anchored to this paragraph/run."""
    return _element_has_drawing_or_pict_child(p._element)


def _clear_paragraph_text_preserve_drawings(p) -> None:
    """Clear text only in runs that do not contain drawings. Leaves drawing runs untouched."""
    for r in p.runs:
        if not _run_has_drawing(r):
            r.text = ""


def remove_paragraph(paragraph) -> None:
    """Remove paragraph from document (safe removal from XML tree). Never remove paragraphs that contain drawings."""
    if paragraph_has_drawing(paragraph):
        _clear_paragraph_text_preserve_drawings(paragraph)
        return
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def remove_table(table) -> None:
    """Remove table from document."""
    parent = table._tbl.getparent()
    if parent is not None:
        parent.remove(table._tbl)


def insert_paragraph_after(paragraph, text: str = "", style=None):
    """
    Insert a new paragraph immediately after the given paragraph (body or cell).
    Returns the new Paragraph. Uses add_paragraph then repositions element.
    Style is applied when the document has that style (e.g. List Bullet, Normal); otherwise skipped.
    """
    parent = paragraph._parent
    new_p = parent.add_paragraph(text)
    new_el = new_p._element
    new_el.getparent().remove(new_el)
    paragraph._element.addnext(new_el)
    if style is not None:
        try:
            new_p.style = style
        except (KeyError, ValueError, AttributeError):
            pass  # Template may not define List Bullet / List Number / etc.
    return new_p


def insert_paragraph_before(paragraph, text: str = "", style=None):
    """
    Insert a new paragraph immediately before the given paragraph (body or cell).
    Returns the new Paragraph. Style applied if present in document; otherwise skipped.
    """
    parent = paragraph._parent
    new_p = parent.add_paragraph(text)
    new_el = new_p._element
    new_el.getparent().remove(new_el)
    paragraph._element.addprevious(new_el)
    if style is not None:
        try:
            new_p.style = style
        except (KeyError, ValueError, AttributeError):
            pass
    return new_p


def add_page_break_paragraph(doc: Document, after_paragraph=None):
    """Insert a page break paragraph. If after_paragraph given, insert after it; else append to body."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    if after_paragraph is not None:
        new_el = p._element
        new_el.getparent().remove(new_el)
        after_paragraph._element.addnext(new_el)
    return p


def _paragraph_has_page_break(para) -> bool:
    """Check if paragraph contains a hard page break (w:br w:type=page)."""
    for el in para._element.iter():
        if el.tag == qn("w:br") and el.get(qn("w:type")) == "page":
            return True
    return False


def _is_page_break_paragraph(p) -> bool:
    """True if paragraph contains a hard page break."""
    if p is None:
        return False
    return _paragraph_has_page_break(p)


def _is_effectively_empty_paragraph(p) -> bool:
    """True if paragraph has no meaningful content (empty text, no images/tables)."""
    if p is None:
        return True
    t = (p.text or "").strip()
    if t:
        return False
    if paragraph_has_drawing(p):
        return False
    return True


def iter_paragraphs_and_cells(doc: Document):
    """Yield (paragraph, parent_element) for every paragraph in doc and in every table cell."""
    for p in doc.paragraphs:
        yield p, None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p, cell


def replace_anchor_in_paragraph_preserve_drawings(p, anchor: str, replacement: str) -> bool:
    """
    Replace anchor with replacement only in a run that contains the anchor and does NOT contain a drawing.
    If anchor is split across runs or only in a drawing run, return False (caller should fail loudly).
    """
    for r in p.runs:
        if not r.text or anchor not in r.text:
            continue
        if _run_has_drawing(r):
            return False
        r.text = r.text.replace(anchor, replacement)
        return True
    return False


def replace_anchor_in_paragraph(paragraph, anchor: str, replacement: str) -> bool:
    """Replace anchor with replacement in paragraph. Never clears runs that contain drawings."""
    if paragraph_has_drawing(paragraph):
        return replace_anchor_in_paragraph_preserve_drawings(paragraph, anchor, replacement)
    full = "".join(run.text for run in paragraph.runs)
    anchor_norm = " ".join(anchor.split())
    full_norm = " ".join(full.split())
    if anchor in full:
        new_full = full.replace(anchor, replacement)
    elif full_norm == anchor_norm or full_norm.strip() == anchor_norm:
        new_full = replacement
    else:
        return False
    for run in paragraph.runs:
        if not _run_has_drawing(run):
            run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_full
    else:
        paragraph.add_run(new_full)
    return True


def _iter_paragraphs_in_tables(tables: list) -> list:
    """Yield all paragraphs in tables and nested tables (python-docx doc.tables does not include nested)."""
    out = []
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    out.append(p)
                out.extend(_iter_paragraphs_in_tables(cell.tables))
    return out


def replace_anchor_in_doc(doc: Document, anchor: str, replacement: str, body_only: bool = True) -> int:
    """Replace anchor with replacement in body paragraphs and table cells (including nested). Returns number of replacements.
    When body_only=True (default), header/footer are not mutated so title page graphic and branding are preserved."""
    hits = 0
    for p, _ in iter_paragraphs_and_cells(doc):
        if replace_anchor_in_paragraph(p, anchor, replacement):
            hits += 1
    for p in _iter_paragraphs_in_tables(doc.tables):
        if replace_anchor_in_paragraph(p, anchor, replacement):
            hits += 1
    if not body_only:
        for section in doc.sections:
            for p in section.header.paragraphs:
                if replace_anchor_in_paragraph(p, anchor, replacement):
                    hits += 1
            for p in section.footer.paragraphs:
                if replace_anchor_in_paragraph(p, anchor, replacement):
                    hits += 1
    return hits


def find_placeholder_or_anchor(doc: Document, pattern: str):
    """
    Search paragraphs and table cells for a placeholder or anchor string.
    Yields (paragraph, parent_cell_or_None) where paragraph text contains pattern.
    """
    for p, parent_cell in iter_paragraphs_and_cells(doc):
        if pattern in (p.text or ""):
            yield p, parent_cell


def _paragraph_text_normalized(p) -> str:
    """Full paragraph text with normalized whitespace (template may have extra spaces/splits)."""
    return " ".join((p.text or "").split())


def find_anchor_paragraph_exact(doc: Document, anchor: str, body_only: bool = False):
    """
    Find a paragraph whose full trimmed text equals the anchor (exact match).
    Normalizes whitespace so template runs like "[[CHART_" + "ELECTRIC_POWER]]" still match.
    When body_only=True, search only doc.paragraphs (not table cells), to avoid header/title.
    Yields (paragraph, parent_cell_or_None).
    """
    anchor_norm = " ".join(anchor.split())
    if body_only:
        for p in doc.paragraphs:
            if _paragraph_text_normalized(p) == anchor_norm:
                yield p, None
        return
    for p, parent_cell in iter_paragraphs_and_cells(doc):
        if _paragraph_text_normalized(p) == anchor_norm:
            yield p, parent_cell


def inject_text_at_anchor(doc: Document, anchor: str, text: str, body_only: bool = True) -> int:
    """
    Replace anchor with text as multiple paragraphs. Text is split on newlines (single or double).
    Returns number of paragraphs inserted. Anchor must appear exactly once.
    """
    text_trimmed = (text or "").replace("\u00a0", " ").strip()
    candidates = list(find_anchor_paragraph_exact(doc, anchor, body_only=body_only))
    if len(candidates) != 1:
        raise RuntimeError(
            f"Anchor {anchor} must appear exactly once (found {len(candidates)} occurrences)"
        )
    p, _ = candidates[0]
    if paragraph_has_drawing(p):
        _clear_paragraph_text_preserve_drawings(p)
    else:
        p.clear()
    last = p
    count = 0
    blocks = re.split(r"\n\s*\n", text_trimmed) if text_trimmed else []
    for block in blocks:
        for line in block.split("\n"):
            line = line.strip()
            if line:
                last = insert_paragraph_after(last, sanitize_text(line), style="Normal")
                count += 1
    if not paragraph_has_drawing(p):
        remove_paragraph(p)
    return count


def insert_table_after(doc: Document, paragraph, rows: int, cols: int):
    """
    Insert a new table immediately after the given paragraph. doc.add_table() creates at end;
    we detach and addnext to place after paragraph. Returns the Table.
    """
    table = doc.add_table(rows=rows, cols=cols)
    tbl_el = table._tbl
    tbl_el.getparent().remove(tbl_el)
    paragraph._element.addnext(tbl_el)
    return table


def insert_paragraph_after_block(doc: Document, block, text: str = "", style=None) -> DocxParagraph:
    """Insert a paragraph after a block (paragraph or table). Returns the new paragraph."""
    el = block._element if hasattr(block, "_element") else block._tbl
    new_p = doc.add_paragraph(text)
    new_el = new_p._element
    new_el.getparent().remove(new_el)
    el.addnext(new_el)
    if style is not None:
        try:
            new_p.style = style
        except (KeyError, ValueError, AttributeError):
            pass
    return new_p


def set_table_fixed_widths(table, col_widths_inches: list[float]) -> None:
    """Set fixed column widths (inches); disable autofit."""
    table.autofit = False
    for idx, w_in in enumerate(col_widths_inches):
        if idx >= len(table.columns):
            break
        w = Inches(w_in)
        table.columns[idx].width = w
        for cell in table.columns[idx].cells:
            cell.width = w


def set_repeat_header_row(row) -> None:
    """Mark row as repeating header on each page (w:tblHeader)."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trPr.append(tbl_header)


def set_table_rows_cant_split(table) -> None:
    """Add w:cantSplit to each row so rows don't split across pages."""
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(OxmlElement("w:cantSplit"))


def _infer_table_col_widths(block: dict, num_cols: int) -> list[float] | None:
    """Return table width profile (inches) for known dependency-summary tables."""
    title = (block.get("title") or "").strip().upper()
    headers = [str(h).strip().lower() for h in (block.get("headers") or [])]
    if num_cols == 6 and headers[:2] == ["category", "provider identified"]:
        return [1.35, 1.0, 0.8, 0.85, 0.85, 1.65]
    if num_cols == 5 and title == "INTERNET TRANSPORT":
        return [1.2, 1.1, 1.6, 1.1, 1.5]
    if num_cols == 5 and title == "CRITICAL HOSTED SERVICES":
        return [1.4, 1.0, 1.7, 1.2, 1.2]
    if num_cols == 5 and headers[:2] == ["role", "provider"]:
        return [1.2, 1.1, 1.6, 1.1, 1.5]
    if num_cols == 5 and headers[:2] == ["service", "provider"]:
        return [1.4, 1.0, 1.7, 1.2, 1.2]
    return None


def _normalize_table_paragraph_spacing(table) -> None:
    """Normalize cell paragraph spacing for dense, readable report tables."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                try:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                except Exception:
                    pass


def replace_anchor_with_table_only(doc: Document, anchor: str, block: dict) -> bool:
    """
    Replace a single anchor paragraph with a table only (no heading). Template owns the section heading.
    Returns True if anchor was found and replaced, False if anchor not present.
    """
    p = find_paragraph_by_exact_text(doc, anchor, body_only=True)
    if p is None:
        return False
    headers = block.get("headers") or []
    rows = block.get("rows") or []
    if not headers and not rows:
        rows = [["Not provided", "—", "—", "—"]]
        headers = ["", "", "", ""]
    num_cols = max(len(headers), max(len(r) for r in rows) if rows else 0)
    if num_cols == 0:
        remove_paragraph(p)
        return True
    tbl = insert_table_after(doc, p, 1 + len(rows), num_cols)
    apply_table_grid_style(tbl)
    widths = _infer_table_col_widths(block, num_cols)
    if widths:
        set_table_fixed_widths(tbl, widths)
    for c, h in enumerate(headers):
        if c < num_cols:
            tbl.rows[0].cells[c].text = sanitize_text(str(h))
    set_repeat_header_row(tbl.rows[0])
    set_table_rows_cant_split(tbl)
    for r_idx, row in enumerate(rows):
        for c_idx, cell_val in enumerate(row):
            if c_idx < num_cols and r_idx + 1 < len(tbl.rows):
                tbl.rows[r_idx + 1].cells[c_idx].text = sanitize_text(str(cell_val))
    _normalize_table_paragraph_spacing(tbl)
    remove_paragraph(p)
    return True


def apply_table_grid_style(table) -> None:
    """Set Word style to Table Grid for visible gridlines."""
    try:
        table.style = "Table Grid"
    except Exception:
        pass


def set_paragraph_keep_with_next(paragraph) -> None:
    """Keep paragraph with next (prevent heading/table split across pages)."""
    try:
        paragraph.paragraph_format.keep_with_next = True
    except Exception:
        pass


def _trim_trailing_empty_paragraphs(doc: Document, max_trim: int = 25) -> None:
    """Remove trailing empty paragraphs to prevent Word from creating blank pages. Never remove drawing paragraphs."""
    trimmed = 0
    while trimmed < max_trim and doc.paragraphs:
        p = doc.paragraphs[-1]
        if _is_page_break_paragraph(p) or paragraph_has_drawing(p):
            break
        if not _is_effectively_empty_paragraph(p):
            break
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)
            trimmed += 1
        else:
            break


def _trim_empty_paragraphs_before(para, max_trim: int = 25) -> None:
    """Remove empty paragraphs immediately before the given paragraph. Never remove drawing paragraphs."""
    trimmed = 0
    while trimmed < max_trim:
        prev = para._element.getprevious()
        if prev is None or prev.tag != qn("w:p"):
            break
        prev_para = DocxParagraph(prev, para._parent)
        if _is_page_break_paragraph(prev_para) or paragraph_has_drawing(prev_para):
            break
        if not _is_effectively_empty_paragraph(prev_para):
            break
        parent = prev.getparent()
        if parent is not None:
            parent.remove(prev)
            trimmed += 1
        else:
            break


def remove_blank_pages_after_section_b(doc: Document) -> None:
    """
    Remove hard-coded blank pages after Section B (DEPENDENCY SNAPSHOT TABLE).
    Strips empty paragraphs and page-break paragraphs between SNAPSHOT_CASCADE
    and the next content (C. SECTOR ANALYSIS, D. CROSS, etc.), up to 5 paras.
    """
    p_cascade = find_paragraph_by_exact_text(doc, "[[SNAPSHOT_CASCADE]]", body_only=True)
    if p_cascade is None:
        return
    body = doc.element.body
    el = p_cascade._element
    removed = 0
    for _ in range(5):
        nxt = el.getnext()
        if nxt is None or nxt.tag != qn("w:p"):
            break
        para = DocxParagraph(nxt, doc)
        text = (para.text or "").strip()
        # Stop at next section or anchor
        if any(
            x in text.upper()
            for x in (
                "C. SECTOR",
                "D. CROSS",
                "E. PRIORITY",
                "PART II",
                "[[",
            )
        ):
            break
        if _is_page_break_paragraph(para) or _is_effectively_empty_paragraph(para):
            body.remove(nxt)
            removed += 1
        else:
            break


def remove_page_breaks_between_annex_table_and_sector_reports(doc: Document) -> None:
    """
    Remove template page break between Annex Overview table and Sector Reports.
    Prevents blank page when table ends and only 'Sector Reports' heading would appear.
    """
    # Find Sector Reports / Infrastructure Dependency heading (immediately before VULN_NARRATIVE)
    p_heading = None
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if SECTOR_REPORTS_HEADING in t or "Infrastructure Dependency" in t:
            p_heading = p
            break
    if p_heading is None:
        return
    body = doc.element.body
    for _ in range(5):
        prev = p_heading._element.getprevious()
        if prev is None:
            break
        if prev.tag == qn("w:tbl"):
            break
        if prev.tag != qn("w:p"):
            continue
        para = DocxParagraph(prev, doc)
        text = (para.text or "").strip()
        if text and (ANNEX_OVERVIEW_HEADING in text or "Dependency Summary" in text):
            break
        if _is_page_break_paragraph(para) or _is_effectively_empty_paragraph(para):
            body.remove(prev)
        else:
            break


def remove_orphaned_page_breaks_before_section_d(doc: Document) -> None:
    """
    Remove orphaned page-break paragraphs left after removing C. SECTOR ANALYSIS
    and sector content. Each sector had a page break before it; removing the
    sector headings/charts leaves 5 orphaned page breaks before D. CROSS.
    """
    p_d = None
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if "D. CROSS" in t.upper() and "SYNTHESIS" in t.upper():
            p_d = p
            break
    if p_d is None:
        return
    body = doc.element.body
    for _ in range(10):
        prev = p_d._element.getprevious()
        if prev is None:
            break
        if prev.tag != qn("w:p"):
            break
        para = DocxParagraph(prev, doc)
        if _is_page_break_paragraph(para) or _is_effectively_empty_paragraph(para):
            body.remove(prev)
        else:
            break


def collapse_consecutive_pagebreaks(doc: Document) -> None:
    """Remove page break paragraphs when two occur back-to-back."""
    body = doc.element.body
    children = list(body)
    i = 1
    while i < len(children):
        prev_el = children[i - 1]
        cur_el = children[i]
        if prev_el.tag != qn("w:p") or cur_el.tag != qn("w:p"):
            i += 1
            continue
        prev_para = DocxParagraph(prev_el, doc)
        cur_para = DocxParagraph(cur_el, doc)
        if _is_page_break_paragraph(prev_para) and _is_page_break_paragraph(cur_para):
            body.remove(cur_el)
            children = list(body)
            continue
        i += 1


def ensure_single_page_break_before(doc: Document, para) -> None:
    """
    Ensure exactly one page break exists immediately before the given paragraph.
    Trims trailing empty paragraphs, avoids double breaks.
    """
    _trim_empty_paragraphs_before(para)
    prev = para._element.getprevious()
    if prev is not None and prev.tag == qn("w:p") and _element_has_page_break(prev):
        return
    new_p = insert_paragraph_before(para, "")
    run = new_p.add_run()
    run.add_break(WD_BREAK.PAGE)


def ensure_single_page_break_after(doc: Document, after_para):
    """
    Ensure exactly one page break exists after the given paragraph.
    Returns the paragraph to insert after (either after_para or the new break para).
    """
    if _is_page_break_paragraph(after_para):
        return after_para
    return add_page_break_paragraph(doc, after_para)


def _element_has_page_break(elm) -> bool:
    """Check if an XML element (e.g. w:p) contains a hard page break."""
    for child in elm.iter():
        if child.tag == qn("w:br") and child.get(qn("w:type")) == "page":
            return True
    return False


def ensure_part2_starts_new_page(doc: Document) -> None:
    """
    Force exactly one page break before PART II – TECHNICAL ANNEX.
    Uses page_break_before on the paragraph (no extra paragraph = no blank page).
    """
    part2_para = None
    for p in doc.paragraphs:
        t = (p.text or "").strip().upper()
        if "PART II" in t and "TECHNICAL ANNEX" in t:
            part2_para = p
            break
    if part2_para is None:
        return
    try:
        part2_para.paragraph_format.page_break_before = True
    except Exception:
        ensure_single_page_break_before(doc, part2_para)
    set_paragraph_keep_with_next(part2_para)
