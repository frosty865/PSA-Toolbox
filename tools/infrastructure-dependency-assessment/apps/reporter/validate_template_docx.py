#!/usr/bin/env python3
"""
Validate the report template DOCX: required anchors, optional anchors, styles, and PSA_CELL.
Exit 0 if all required fields are addressed; non-zero and list issues otherwise.
Usage: python validate_template_docx.py [path_to_template.docx]
Default: ADA/report template.docx (relative to repo root).
"""
from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Install python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(2)

# Must stay in sync with main.REQUIRED_ANCHORS / main.OPTIONAL_ANCHORS
REQUIRED_ANCHORS = (
    "[[SNAPSHOT_POSTURE]]",
    "[[SNAPSHOT_SUMMARY]]",
    "[[SNAPSHOT_DRIVERS]]",
    "[[SNAPSHOT_MATRIX]]",
    "[[SNAPSHOT_CASCADE]]",
    "[[CHART_ELECTRIC_POWER]]",
    "[[CHART_COMMUNICATIONS]]",
    "[[CHART_INFORMATION_TECHNOLOGY]]",
    "[[CHART_WATER]]",
    "[[CHART_WASTEWATER]]",
    "[[SYNTHESIS]]",
    "[[PRIORITY_ACTIONS]]",
    "[[TABLE_DEPENDENCY_SUMMARY]]",
    "[[STRUCTURAL_PROFILE_SUMMARY]]",
    "[[VULNERABILITY_COUNT_SUMMARY]]",
    "[[VULNERABILITY_BLOCKS]]",
    "[[CROSS_INFRA_ANALYSIS]]",
)
OPTIONAL_ANCHORS = (
    "[[SLA_PRA_SUMMARY]]",
    "[[CROSS_DEPENDENCY_SUMMARY]]",
    "[[NARRATIVE_SOURCES]]",
    "[[DESIGNATION_SERVICES]]",
    "[[IT_TRANSPORT_SECTION]]",
    "[[IT_HOSTED_SECTION]]",
)
CHART_ANCHORS = [
    "[[CHART_ELECTRIC_POWER]]",
    "[[CHART_COMMUNICATIONS]]",
    "[[CHART_INFORMATION_TECHNOLOGY]]",
    "[[CHART_WATER]]",
    "[[CHART_WASTEWATER]]",
]
PSA_CELL_ANCHOR = "[[PSA_CELL]]"

# Cover-page strings that must not appear as consecutive duplicates in the first page
COVER_DUPLICATE_TEXTS = (
    "Asset Dependency Assessment",
    "UNCLASSIFIED//FOR OFFICIAL USE ONLY",
)
COVER_LOOKAHEAD_PARAS = 15

# Must stay in sync with main.REQUIRED_ADA_STYLES (vulnerability block typography)
REQUIRED_ADA_STYLES = [
    "ADA_Vuln_Header",
    "ADA_Vuln_Severity",
    "ADA_Vuln_Meta",
    "ADA_Vuln_Label",
    "ADA_Vuln_Body",
    "ADA_Vuln_Bullets",
    "ADA_Vuln_Numbered",
]


def iter_paragraphs_and_cells(doc: Document):
    """Yield (paragraph, parent_cell) for every paragraph in body and table cells."""
    for p in doc.paragraphs:
        yield p, None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p, cell


def collect_anchors(doc: Document) -> set[str]:
    found = set()
    for p, _ in iter_paragraphs_and_cells(doc):
        text = p.text or ""
        for anchor in REQUIRED_ANCHORS + OPTIONAL_ANCHORS:
            if anchor in text:
                found.add(anchor)
    return found


def count_anchor_occurrences(doc: Document, anchor: str) -> int:
    count = 0
    for p, _ in iter_paragraphs_and_cells(doc):
        count += (p.text or "").count(anchor)
    return count


def get_paragraph_style_names(doc: Document) -> list[str]:
    names = []
    for s in doc.styles:
        if getattr(s, "type", None) == WD_STYLE_TYPE.PARAGRAPH:
            names.append(getattr(s, "name", "") or "")
    return sorted(names)


def check_cover_page_no_duplicate_strings(doc: Document) -> list[str]:
    """Return errors if first page has consecutive duplicate cover strings."""
    errors: list[str] = []
    paras = [p for p in doc.paragraphs[:COVER_LOOKAHEAD_PARAS] if p is not None]
    for i in range(1, len(paras)):
        prev_text = (paras[i - 1].text or "").strip()
        curr_text = (paras[i].text or "").strip()
        if curr_text in COVER_DUPLICATE_TEXTS and prev_text == curr_text:
            errors.append(
                f"Cover page has consecutive duplicate paragraph: {curr_text!r}. "
                "Run add_ada_vuln_styles_to_template.py to normalize the template."
            )
    return errors


def validate_template(path: Path) -> tuple[list[str], list[str]]:
    """Return (errors, warnings). Errors mean required fields are not addressed."""
    errors: list[str] = []
    warnings: list[str] = []

    if not path.is_file():
        return ([f"Template file not found: {path}"], [])
    doc = Document(str(path))

    # 0) Cover page: no consecutive duplicate title/classification
    errors.extend(check_cover_page_no_duplicate_strings(doc))

    # 1) Required anchors
    found = collect_anchors(doc)
    missing_required = [a for a in REQUIRED_ANCHORS if a not in found]
    if missing_required:
        errors.append("Missing required anchors:")
        for a in missing_required:
            errors.append(f"  - {a}")

    # 2) Optional anchors (informational)
    present_optional = [a for a in OPTIONAL_ANCHORS if a in found]
    if present_optional:
        warnings.append("Optional anchors present: " + ", ".join(present_optional))

    # 3) Chart anchors must appear exactly once each (reporter inserts one image per anchor)
    for anchor in CHART_ANCHORS:
        if anchor not in found:
            continue
        n = count_anchor_occurrences(doc, anchor)
        if n != 1:
            errors.append(f"Anchor {anchor} must appear exactly once (found {n})")

    # 4) [[PSA_CELL]] must appear 0 or 1 time
    psa_count = count_anchor_occurrences(doc, PSA_CELL_ANCHOR)
    if psa_count not in (0, 1):
        errors.append(f"{PSA_CELL_ANCHOR} must appear 0 or 1 time (found {psa_count})")

    # 5) ADA_Vuln_* styles: reporter creates them at export if missing; warn if not in template
    available_styles = get_paragraph_style_names(doc)
    missing_styles = [s for s in REQUIRED_ADA_STYLES if s not in available_styles]
    if missing_styles:
        warnings.append("Paragraph styles not in template (reporter will create at export):")
        for s in missing_styles:
            warnings.append(f"  - {s}")

    return (errors, warnings)


def main() -> int:
    script_dir = Path(__file__).resolve().parent
    repo_root = script_dir.parent.parent
    default_template = repo_root / "ADA" / "report template.docx"

    if len(sys.argv) > 1:
        template_path = Path(sys.argv[1]).resolve()
    else:
        template_path = default_template

    errors, warnings = validate_template(template_path)
    for line in errors:
        print(line, file=sys.stderr)
    for line in warnings:
        print(line, file=sys.stderr)
    if not errors:
        print("OK: Template has all required fields addressed.", file=sys.stderr)
        return 0
    return 1


if __name__ == "__main__":
    sys.exit(main())
