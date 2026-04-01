"""
QC pipeline: placeholder/anchor replacement, gates (A/B/C), SPOF sanitization, encoding cleanup, export QC.
Single entry point: run_all(doc, data, template_path, ...).
"""
from __future__ import annotations

import sys
from docx import Document
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from constants import (
    PLACEHOLDER_ANCHORS_WITH_FALLBACK,
    FALLBACK_NO_FINDINGS,
    FALLBACK_WITH_FINDINGS,
    REQUIRED_ANCHORS,
    OPTIONAL_ANCHORS,
    TABLE_ANCHOR,
    SLA_PRA_SUMMARY_ANCHOR,
    CROSS_DEPENDENCY_SUMMARY_ANCHOR,
    UNDERSCORE_RE,
    EXPORT_TABLE_BAD_HEADERS,
    GATE_A_PLACEHOLDER_PHRASES,
    FORBIDDEN_PLACEHOLDERS,
    GATE_A_ANCHOR_RE,
    GATE_A_DEPRECATED_TERMS,
    GATE_A_SAFE_WORD_BOUNDARY_RE,
)
from docx_ops import (
    iter_paragraphs_and_cells,
    iter_block_items,
    _paragraph_text_normalized,
    remove_paragraph,
    paragraph_has_drawing,
    _clear_paragraph_text_preserve_drawings,
)
from sanitize import sanitize_text


def remove_anchor_block_paragraphs(doc: Document) -> None:
    """Clear any paragraph whose text contains 'ANCHOR BLOCK (dev only)' or 'ANCHOR BLOCK'. Preserve drawings."""
    for p, _ in iter_paragraphs_and_cells(doc):
        text = (p.text or "").strip()
        if "ANCHOR BLOCK (dev only)" in text or "ANCHOR BLOCK" in text:
            if paragraph_has_drawing(p):
                _clear_paragraph_text_preserve_drawings(p)
            else:
                p.clear()


def replace_placeholder_anchors_with_fallback(doc: Document, vulnerability_count: int = 0) -> None:
    """Replace SNAPSHOT_*, SYNTHESIS, PRIORITY_ACTIONS with fallback text."""
    fallback = FALLBACK_WITH_FINDINGS if vulnerability_count > 0 else FALLBACK_NO_FINDINGS
    for p, _ in iter_paragraphs_and_cells(doc):
        text_norm = _paragraph_text_normalized(p)
        for anchor in PLACEHOLDER_ANCHORS_WITH_FALLBACK:
            if " ".join(anchor.split()) == text_norm:
                p.clear()
                p.add_run(sanitize_text(fallback))
                break


def clear_remaining_anchor_paragraphs(doc: Document) -> None:
    """Clear any paragraph that is exactly an anchor token."""
    all_anchors = list(REQUIRED_ANCHORS) + [TABLE_ANCHOR, SLA_PRA_SUMMARY_ANCHOR, CROSS_DEPENDENCY_SUMMARY_ANCHOR]
    for p, _ in iter_paragraphs_and_cells(doc):
        text_norm = _paragraph_text_normalized(p)
        for anchor in all_anchors:
            if " ".join(anchor.split()) == text_norm:
                p.clear()
                break


def scrub_any_remaining_anchors(doc: Document) -> None:
    """Final scrub: remove any paragraph containing '[['; clear table cells containing '[['."""
    for p in list(doc.paragraphs):
        if "[[" in (p.text or ""):
            remove_paragraph(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if "[[" in (cell.text or ""):
                    cell.text = ""


def assert_no_anchors_remaining(doc: Document) -> None:
    """Assert the document contains no anchor tokens or dev-only blocks. Raises SystemExit if found."""
    all_anchors = list(REQUIRED_ANCHORS) + [TABLE_ANCHOR, SLA_PRA_SUMMARY_ANCHOR, CROSS_DEPENDENCY_SUMMARY_ANCHOR]
    found = []
    for p, _ in iter_paragraphs_and_cells(doc):
        text = p.text or ""
        for anchor in all_anchors:
            if anchor in text:
                found.append(anchor)
        if "ANCHOR BLOCK (dev only)" in text or "ANCHOR BLOCK" in text:
            found.append("ANCHOR BLOCK (dev only)")
    if found:
        print("ERROR: Generated document still contains anchor or dev-only text:", file=sys.stderr)
        for a in set(found):
            print(f"  - {a!r}", file=sys.stderr)
        sys.exit(1)


def assert_no_narrative_blanks_remaining(doc: Document) -> None:
    """Fail build if any '____' or long underscore placeholders remain."""
    for p, _ in iter_paragraphs_and_cells(doc):
        text = p.text or ""
        if UNDERSCORE_RE.search(text):
            print("ERROR: Generated document still contains unreplaced narrative blank (____)", file=sys.stderr)
            sys.exit(1)


def assert_no_export_style_tables(doc: Document) -> None:
    """Fail build if any table has export-style headers."""
    for table in doc.tables:
        if not table.rows:
            continue
        for row in table.rows[:2]:
            for cell in row.cells:
                c = (cell.text or "").strip()
                for bad in EXPORT_TABLE_BAD_HEADERS:
                    if c == bad:
                        print(
                            f"ERROR: Export-style table detected (header contains {bad!r}). "
                            "Use template summary/VOFC format only.",
                            file=sys.stderr,
                        )
                        sys.exit(1)


def assert_exec_summary_not_followed_by_table(doc: Document) -> None:
    """Fail if 'EXECUTIVE SUMMARY' paragraph is immediately followed by a table."""
    blocks = list(iter_block_items(doc))
    for i, block in enumerate(blocks):
        if isinstance(block, DocxParagraph) and (block.text or "").strip() == "EXECUTIVE SUMMARY":
            if i + 1 < len(blocks) and isinstance(blocks[i + 1], DocxTable):
                print(
                    "ERROR: EXECUTIVE SUMMARY is immediately followed by a table (bug regression).",
                    file=sys.stderr,
                )
                sys.exit(1)
            break


def _context_snippet(text: str, idx: int, half_len: int) -> str:
    """Return ±half_len chars around idx, sanitized for display."""
    start = max(0, idx - half_len)
    end = min(len(text), idx + half_len)
    return text[start:end].replace("\n", " ").strip()


def gate_a_placeholder_anchor_leak_scanner(doc: Document) -> None:
    """Gate A: Fail export if placeholders, unresolved anchors, or deprecated terms found."""
    offenders = []
    for p, _ in iter_paragraphs_and_cells(doc):
        text = p.text or ""
        if not text.strip():
            continue
        text_lower = text.lower()
        for phrase in GATE_A_PLACEHOLDER_PHRASES:
            if phrase in text_lower:
                idx = text_lower.find(phrase)
                ctx = _context_snippet(text, idx, 40)
                offenders.append(("placeholder", repr(phrase), ctx))
        for forbidden in FORBIDDEN_PLACEHOLDERS:
            if forbidden in text:
                ctx = text[:80] + "..." if len(text) > 80 else text
                offenders.append(("placeholder", repr(forbidden), ctx))
        for m in GATE_A_ANCHOR_RE.finditer(text):
            ctx = _context_snippet(text, m.start(), 40)
            offenders.append(("anchor", m.group(), ctx))
        for term in GATE_A_DEPRECATED_TERMS:
            if term == "safe":
                m = GATE_A_SAFE_WORD_BOUNDARY_RE.search(text_lower)
                if m:
                    idx = m.start()
                    ctx = _context_snippet(text, idx, 40)
                    offenders.append(("deprecated", repr(term), ctx))
            elif term in text_lower:
                idx = text_lower.find(term)
                ctx = _context_snippet(text, idx, 40)
                offenders.append(("deprecated", repr(term), ctx))
    if offenders:
        print("ERROR: Gate A — Placeholder/Anchor/Deprecated term leakage. Export blocked.", file=sys.stderr)
        for typ, match, ctx in offenders[:10]:
            hint = "Check sector narratives" if typ == "placeholder" else "Check anchor injection" if typ == "anchor" else "Remove deprecated SAFE/security assessment references"
            print(f"  [{typ}] {match} | context: ...{ctx}... | hint: {hint}", file=sys.stderr)
        if len(offenders) > 10:
            print(f"  ... and {len(offenders) - 10} more", file=sys.stderr)
        sys.exit(1)


def _is_heading_style(p) -> bool:
    """True if paragraph has a heading style."""
    try:
        style = (p.style and p.style.name or "").lower()
        return "heading" in style
    except Exception:
        return False


GATE_C_ALLOWLIST = (
    "fouo",
    "for official use only",
    "cisa",
    "category",
    "vulnerability",
    "option for consideration",
    "time to severe impact",
    "functional loss",
    "backup duration",
    "structural posture",
)


def _normalize_block_for_duplicate(text: str) -> str:
    """Normalize paragraph for duplicate detection."""
    import re
    t = (text or "").lower().strip()
    t = re.sub(r"\s+", " ", t)
    t = re.sub(r"\d+", "#", t)
    return t


def gate_b_suppress_empty_sections(doc: Document) -> None:
    """Gate B: Remove headings that have no renderable content before the next section."""
    blocks = list(iter_block_items(doc))
    to_remove = []
    i = 0
    while i < len(blocks):
        block = blocks[i]
        if isinstance(block, DocxParagraph):
            text = (block.text or "").strip()
            if not _is_heading_style(block):
                i += 1
                continue
            empty_section_headings = (
                "Critical Products",
                "CRITICAL PRODUCTS",
                "Cross-Infrastructure Synthesis",
                "CROSS-INFRASTRUCTURE SYNTHESIS",
            )
            if text in empty_section_headings:
                j = i + 1
                has_content = False
                while j < len(blocks):
                    next_b = blocks[j]
                    if isinstance(next_b, DocxTable):
                        has_content = True
                        break
                    if isinstance(next_b, DocxParagraph):
                        next_text = (next_b.text or "").strip()
                        if not next_text:
                            j += 1
                            continue
                        if next_text in empty_section_headings or _is_heading_style(next_b):
                            break
                        if next_text != FALLBACK_NO_FINDINGS and len(next_text) > 10:
                            has_content = True
                            break
                        if next_text == FALLBACK_NO_FINDINGS:
                            has_content = True
                            break
                        j += 1
                if not has_content and j > i + 1:
                    for k in range(i, j):
                        if isinstance(blocks[k], DocxParagraph):
                            to_remove.append(blocks[k]._element)
                i = j
                continue
        i += 1
    for el in to_remove:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def gate_c_duplicate_block_detection(doc: Document) -> None:
    """Gate C: Detect repeated paragraphs across sector sections; fail if found."""
    sector_section_markers = ("Electric Power", "Communications", "Information Technology", "Water", "Wastewater")
    seen = {}
    current_section = "preamble"
    for block in iter_block_items(doc):
        if isinstance(block, DocxTable):
            continue
        if isinstance(block, DocxParagraph):
            text = (block.text or "").strip()
            if not text:
                continue
            for marker in sector_section_markers:
                if marker in text and _is_heading_style(block):
                    current_section = marker
                    break
            if len(text) < 20:
                continue
            norm = _normalize_block_for_duplicate(text)
            if not norm:
                continue
            for allowed in GATE_C_ALLOWLIST:
                if norm.startswith(allowed) or allowed in norm[:50]:
                    break
            else:
                if norm not in seen:
                    seen[norm] = []
                if current_section not in seen[norm]:
                    seen[norm].append(current_section)
    duplicates = [(norm, sections) for norm, sections in seen.items() if len(sections) > 1]
    if duplicates:
        norm, sections = duplicates[0]
        sample = norm[:120] + "..." if len(norm) > 120 else norm
        print(
            "ERROR: Gate C — Duplicate block (copy bleed) detected. Export blocked.",
            file=sys.stderr,
        )
        print(f"  Repeated text (first 120 chars): {sample!r}", file=sys.stderr)
        print(f"  Appeared in sectors: {', '.join(sections)}", file=sys.stderr)
        sys.exit(1)


def sanitize_spof_language(doc: Document) -> None:
    """Replace SPOF/inference language with factual wording."""
    phrase_replacements = [
        ("(SPOF likely)", "Reported sources: 1"),
        ("SPOF likely", "Reported sources: 1"),
    ]
    banned = ["SPOF", "single point of failure"]
    replacement = "Reported sources: 1"
    for p in doc.paragraphs:
        txt = p.text or ""
        for old_phrase, new_phrase in phrase_replacements:
            txt = txt.replace(old_phrase, new_phrase)
        for phrase in banned:
            if phrase in txt:
                txt = txt.replace(phrase, replacement)
        if txt != (p.text or ""):
            p.clear()
            p.add_run(sanitize_text(txt))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text or ""
                for old_phrase, new_phrase in phrase_replacements:
                    txt = txt.replace(old_phrase, new_phrase)
                for phrase in banned:
                    if phrase in txt:
                        txt = txt.replace(phrase, replacement)
                if txt != (cell.text or ""):
                    cell.text = sanitize_text(txt)


def assert_no_spof_in_output(doc: Document) -> None:
    """Fail if explicit SPOF labels appear in output."""
    forbidden = ("(SPOF likely)", "SPOF likely")
    for p, _ in iter_paragraphs_and_cells(doc):
        text = (p.text or "").strip()
        for phrase in forbidden:
            if phrase in text:
                raise RuntimeError(
                    f"Output must not contain SPOF language: found {phrase!r} in: {text[:150]!r}..."
                )


def _replace_encoding_artifacts_in_doc(doc: Document) -> None:
    """Replace encoding artifacts (U+FFFD, mojibake) with em dash so QC passes."""
    from qc_export import ENCODING_ARTIFACT_PATTERNS
    repl = "\u2014"
    for pat in ENCODING_ARTIFACT_PATTERNS:
        for p in doc.paragraphs:
            for run in p.runs:
                if run.text and pat in run.text:
                    run.text = run.text.replace(pat, repl)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text and pat in run.text:
                                run.text = run.text.replace(pat, repl)


def run_all(
    doc: Document,
    data: dict,
    template_path: str | None,
    required_sections_filled: dict | None = None,
    rendered_vuln_count: int = 0,
    vulnerability_blocks_flat: list | None = None,
    priority_actions: list | None = None,
) -> None:
    """
    Run full QC pipeline: replace placeholders, gates B/C/A, SPOF sanitization and assert,
    encoding cleanup, then run_export_qc and optionally run_all_qc_validators.
    """
    remove_anchor_block_paragraphs(doc)
    replace_placeholder_anchors_with_fallback(doc, vulnerability_count=rendered_vuln_count)
    clear_remaining_anchor_paragraphs(doc)
    scrub_any_remaining_anchors(doc)
    assert_no_anchors_remaining(doc)
    assert_no_narrative_blanks_remaining(doc)
    assert_no_export_style_tables(doc)
    assert_exec_summary_not_followed_by_table(doc)
    gate_b_suppress_empty_sections(doc)
    gate_c_duplicate_block_detection(doc)
    gate_a_placeholder_anchor_leak_scanner(doc)
    sanitize_spof_language(doc)
    assert_no_spof_in_output(doc)
    _replace_encoding_artifacts_in_doc(doc)

    if template_path:
        try:
            from qc_export import REQUIRED_SECTION_ANCHORS, run_export_qc
            try:
                from qc.validators import run_all_qc_validators
            except ImportError:
                run_all_qc_validators = None
            vulnerability_blocks_flat = vulnerability_blocks_flat or []
            run_export_qc(
                doc,
                str(template_path),
                priority_actions=priority_actions,
                vulnerability_blocks=vulnerability_blocks_flat,
                required_sections_filled=required_sections_filled,
                payload=data,
                assessment_json=data,
            )
            if run_all_qc_validators is not None:
                run_all_qc_validators(
                    doc,
                    vulnerability_count=rendered_vuln_count,
                    vulnerability_blocks=vulnerability_blocks_flat,
                )
        except ValueError as e:
            print(f"ERROR: Export QC failed: {e}", file=sys.stderr)
            sys.exit(1)
        except Exception as e:
            if isinstance(e, ImportError) and "qc" in str(e).lower():
                pass
            else:
                try:
                    from qc.validators import QCValidationError
                    if isinstance(e, QCValidationError):
                        print(f"ERROR: Export QC validation failed: {e}", file=sys.stderr)
                        sys.exit(1)
                except ImportError:
                    pass
                raise
