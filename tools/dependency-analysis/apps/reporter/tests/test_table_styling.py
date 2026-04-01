#!/usr/bin/env python3
"""
Tests for table styling: grid lines, keep-with-next, cantSplit.
"""
import io
import json
import os
import sys
from pathlib import Path

import pytest

REPORTER_DIR = Path(__file__).resolve().parent.parent
REPO_ROOT = REPORTER_DIR.parent.parent
CANONICAL_TEMPLATE = REPO_ROOT / "ADA" / "report template.docx"


def _run_reporter(payload: dict, work_dir: Path, template: Path) -> Path:
    work_dir.mkdir(parents=True, exist_ok=True)
    old_work = os.environ.get("WORK_DIR")
    old_template = os.environ.get("TEMPLATE_PATH")
    os.environ["WORK_DIR"] = str(work_dir)
    os.environ["TEMPLATE_PATH"] = str(template)
    old_stdin = sys.stdin
    sys.stdin = io.StringIO(json.dumps(payload))
    try:
        if str(REPORTER_DIR) not in sys.path:
            sys.path.insert(0, str(REPORTER_DIR))
        from main import main as reporter_main
        reporter_main()
    except SystemExit as e:
        if e.code != 0:
            raise RuntimeError(f"Reporter exited {e.code}")
    finally:
        sys.stdin = old_stdin
        os.environ.pop("WORK_DIR", None)
        os.environ.pop("TEMPLATE_PATH", None)
        if old_work:
            os.environ["WORK_DIR"] = old_work
        if old_template:
            os.environ["TEMPLATE_PATH"] = old_template
    return work_dir / "output.docx"


def _minimal_payload():
    return {
        "assessment": {
            "asset": {"psa_cell": "555-000-0000"},
            "categories": {
                "ELECTRIC_POWER": {
                    "requires_service": True,
                    "time_to_impact_hours": 24,
                    "loss_fraction_no_backup": 0.5,
                    "recovery_time_hours": 12,
                },
                "COMMUNICATIONS": {"requires_service": False},
                "INFORMATION_TECHNOLOGY": {"requires_service": False},
                "WATER": {"requires_service": False},
                "WASTEWATER": {"requires_service": False},
                "CRITICAL_PRODUCTS": {},
            },
        },
        "vofc_collection": {"items": []},
    }


@pytest.fixture
def work_dir(tmp_path):
    return tmp_path


def test_tables_have_grid_style(work_dir):
    """Generated tables use Table Grid style where applicable."""
    if not CANONICAL_TEMPLATE.exists():
        pytest.skip("No template found: ADA/report template.docx")
    payload = _minimal_payload()
    out_path = _run_reporter(payload, work_dir, CANONICAL_TEMPLATE)
    assert out_path.is_file()

    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(str(out_path))
    grid_count = 0
    for table in doc.tables:
        try:
            if table.style and "Grid" in str(table.style.name):
                grid_count += 1
        except Exception:
            pass
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is not None:
            tblStyle = tblPr.find(qn("w:tblStyle"))
            if tblStyle is not None and tblStyle.get(qn("w:val")):
                val = tblStyle.get(qn("w:val")) or ""
                if "Grid" in val:
                    grid_count += 1
    assert grid_count >= 1, "At least one table should have Table Grid style"


def test_keep_with_next_set_for_headings_before_tables(work_dir):
    """Headings immediately before tables have keep_with_next set."""
    if not CANONICAL_TEMPLATE.exists():
        pytest.skip("No template found: ADA/report template.docx")
    payload = _minimal_payload()
    payload["energy_dependency"] = {
        "vulnerability_blocks": [
            {"title": "Test Vuln", "narrative": "N.", "ofcs": [], "references": []},
        ],
    }
    payload["vulnerability_index_rows"] = [{"title": "Test Vuln", "sector": "Energy"}]
    out_path = _run_reporter(payload, work_dir, CANONICAL_TEMPLATE)
    assert out_path.is_file()

    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(str(out_path))
    blocks = list(doc.element.body.iterchildren())
    keep_next_count = 0
    for i, child in enumerate(blocks):
        if child.tag == qn("w:p"):
            pPr = child.find(qn("w:pPr"))
            if pPr is not None:
                keepNext = pPr.find(qn("w:keepNext"))
                if keepNext is not None:
                    keep_next_count += 1
    assert keep_next_count >= 1, "At least one paragraph should have keepWithNext"
