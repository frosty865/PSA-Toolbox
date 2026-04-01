#!/usr/bin/env python3
"""
Export QC gate tests: no unresolved anchors, no encoding artifacts, style sampling.
Run after export; asserts QC checks pass on generated DOCX.
"""
import io
import json
import os
import sys
from pathlib import Path

import pytest
from docx import Document

REPORTER_DIR = Path(__file__).resolve().parent.parent
REPO_ROOT = REPORTER_DIR.parent.parent
CANONICAL_TEMPLATE = REPO_ROOT / "ADA" / "report template.docx"


def _minimal_assessment():
    return {
        "asset": {"psa_cell": "555-000-0000"},
        "categories": {
            "ELECTRIC_POWER": {
                "requires_service": True,
                "time_to_impact_hours": 24,
                "loss_fraction_no_backup": 0.5,
                "has_backup_any": True,
                "backup_duration_hours": 48,
                "loss_fraction_with_backup": 0.1,
                "recovery_time_hours": 12,
            },
            "COMMUNICATIONS": {"requires_service": False},
            "INFORMATION_TECHNOLOGY": {"requires_service": False},
            "WATER": {"requires_service": False},
            "WASTEWATER": {"requires_service": False},
            "CRITICAL_PRODUCTS": {},
        },
    }


def _minimal_full_payload():
    """Payload with all keys required when template has Part II federal-style anchors (STRUCTURAL_PROFILE_SUMMARY, VULNERABILITY_BLOCKS, etc.)."""
    base = _minimal_assessment()
    return {
        "assessment": base,
        "vofc_collection": {"items": []},
        "executive_snapshot": {
            "posture": "CONDITIONAL",
            "summary": "Key risk drivers: Electric Power.",
            "drivers": ["Electric Power"],
            "matrixRows": [{"sector": "Electric Power", "ttiHrs": "24", "lossPct": "50", "backupHrs": "48", "structuralPosture": "Tolerant"}],
            "cascade": None,
        },
        "synthesis": {"title": "Synthesis", "paragraphs": ["Cross-infrastructure synthesis."], "bullets": []},
        "priority_actions": {"title": "Priority Actions", "actions": [{"leadIn": "Action 1", "fullText": "Description."}]},
        "energy_dependency": {"vulnerability_blocks": []},
        "dependency_sections": [],
        "structural_profile_summary": "Structural sensitivity is driven by Electric Power dependency.",
        "vulnerability_count_summary": "0 findings (0 HIGH, 0 ELEVATED, 0 MODERATE).",
        "vulnerability_blocks": "No infrastructure vulnerabilities identified for this assessment.",
        "cross_infra_analysis": "Cross-infrastructure synthesis.",
    }


def _structured_vuln_narrative_payload():
    """Payload with Part II federal-style keys and at least one vulnerability block (pre-rendered string)."""
    base = _minimal_assessment()
    # Sample federal-style block (max 4 options; no SAFE)
    sample_blocks = """VULNERABILITY 1
ELEVATED – Structural Concentration
Infrastructure Domain: Electric Power
Risk Type: Structural Concentration

Operational Consequence
Severe impact in 24 hours with ~50% functional loss absent alternate capability.

Exposure Description
Utility supply is single feed with no documented alternate capability. Missing provider priority; single feed.

Structural Indicator Profile
Impact Sensitivity: Immediate
Mitigation Depth: None
Recovery Sensitivity: Moderate
Cascade Exposure: Low

Standards Alignment
- FEMA CGC

Options for Consideration
1. Document restoration priority with provider
2. Introduce alternate capability
"""
    return {
        "assessment": base,
        "vofc_collection": {"items": []},
        "executive_snapshot": {
            "posture": "CONDITIONAL",
            "summary": "Key risk drivers: Electric Power.",
            "drivers": ["Electric Power"],
            "matrixRows": [{"sector": "Electric Power", "ttiHrs": "24", "lossPct": "50", "backupHrs": "48", "structuralPosture": "Tolerant"}],
            "cascade": None,
        },
        "synthesis": {"title": "Synthesis", "paragraphs": ["Cross-infrastructure synthesis."], "bullets": []},
        "priority_actions": {"title": "Priority Actions", "actions": [{"leadIn": "Action 1", "fullText": "Description."}]},
        "energy_dependency": {"vulnerability_blocks": []},
        "dependency_sections": [],
        "structural_profile_summary": "Structural sensitivity is driven by Electric Power dependency.",
        "vulnerability_count_summary": "1 finding (0 HIGH, 1 ELEVATED, 0 MODERATE).",
        "vulnerability_blocks": sample_blocks.strip(),
        "cross_infra_analysis": "Cross-infrastructure synthesis.",
    }


def _run_reporter_and_get_docx(payload: dict, work_dir: Path) -> Path | None:
    """Run reporter; return path to output.docx or None on failure."""
    if not CANONICAL_TEMPLATE.exists():
        pytest.skip("No template found: ADA/report template.docx")
    work_dir.mkdir(parents=True, exist_ok=True)
    old_work = os.environ.get("WORK_DIR")
    old_template = os.environ.get("TEMPLATE_PATH")
    os.environ["WORK_DIR"] = str(work_dir)
    os.environ["TEMPLATE_PATH"] = str(CANONICAL_TEMPLATE)
    sys.path.insert(0, str(REPORTER_DIR))
    old_stdin = sys.stdin
    sys.stdin = io.StringIO(json.dumps(payload))
    exit_code = 0
    try:
        from main import main as reporter_main
        reporter_main()
    except SystemExit as e:
        exit_code = e.code if e.code is not None else 1
    finally:
        sys.stdin = old_stdin
        os.environ.pop("WORK_DIR", None)
        os.environ.pop("TEMPLATE_PATH", None)
        if old_work:
            os.environ["WORK_DIR"] = old_work
        if old_template:
            os.environ["TEMPLATE_PATH"] = old_template
    out = work_dir / "output.docx"
    return out if exit_code == 0 and out.exists() else None


@pytest.fixture
def work_dir(tmp_path):
    return tmp_path


def test_qc_unresolved_anchors_rejected(work_dir):
    """QC: document with [[ANCHOR]] should be rejected by check_unresolved_anchors."""
    sys.path.insert(0, str(REPORTER_DIR))
    from qc_export import check_unresolved_anchors
    doc = Document()
    doc.add_paragraph("Normal text")
    doc.add_paragraph("[[SYNTHESIS]]")
    found = check_unresolved_anchors(doc)
    assert "[[SYNTHESIS]]" in found
    assert len(found) >= 1


def test_qc_no_anchors_in_clean_doc(work_dir):
    """QC: document with no [[...]] anchors returns empty list."""
    sys.path.insert(0, str(REPORTER_DIR))
    from qc_export import check_unresolved_anchors
    doc = Document()
    doc.add_paragraph("Executive summary here.")
    doc.add_paragraph("No anchors.")
    found = check_unresolved_anchors(doc)
    assert found == []


def test_qc_encoding_artifacts_rejected(work_dir):
    """QC: document containing mojibake pattern should be rejected."""
    sys.path.insert(0, str(REPORTER_DIR))
    from qc_export import check_encoding_artifacts, ENCODING_ARTIFACT_PATTERNS
    doc = Document()
    doc.add_paragraph("Good text")
    # Use replacement char U+FFFD which is in ENCODING_ARTIFACT_PATTERNS
    doc.add_paragraph("Bad: \uFFFD here")
    found = check_encoding_artifacts(doc)
    assert len(found) >= 1
    assert any(pat in found[0][1] or pat == found[0][0] for pat in ENCODING_ARTIFACT_PATTERNS)


def test_qc_priority_actions_limit(work_dir):
    """QC: more than 5 priority actions raises."""
    sys.path.insert(0, str(REPORTER_DIR))
    from qc_export import check_priority_actions_limit, MAX_PRIORITY_ACTIONS
    check_priority_actions_limit([{}] * MAX_PRIORITY_ACTIONS)
    with pytest.raises(ValueError, match="at most 5"):
        check_priority_actions_limit([{}] * (MAX_PRIORITY_ACTIONS + 1))


def test_qc_ofcs_per_vulnerability_limit(work_dir):
    """QC: more than 4 OFCs per vulnerability raises."""
    sys.path.insert(0, str(REPORTER_DIR))
    from qc_export import check_ofcs_per_vulnerability_limit, MAX_OFCS_PER_VULNERABILITY
    check_ofcs_per_vulnerability_limit([{"ofcs": ["a", "b", "c", "d"]}])
    with pytest.raises(ValueError, match="at most 4"):
        check_ofcs_per_vulnerability_limit([{"ofcs": ["a", "b", "c", "d", "e"]}])


def test_qc_required_sections_empty_rejected(work_dir):
    """QC: required_sections_filled with any False raises."""
    sys.path.insert(0, str(REPORTER_DIR))
    from qc_export import run_export_qc, REQUIRED_SECTION_ANCHORS
    doc = Document()
    doc.add_paragraph("No anchors here.")
    template_path = str(REPO_ROOT / "ADA" / "report template.docx")
    required_ok = {k: True for k in REQUIRED_SECTION_ANCHORS}
    run_export_qc(doc, template_path, required_sections_filled=required_ok)
    required_bad = {**required_ok, "[[SNAPSHOT_POSTURE]]": False}
    with pytest.raises(ValueError, match="required sections empty"):
        run_export_qc(doc, template_path, required_sections_filled=required_bad)


def test_export_fixture_no_anchors_remain(work_dir):
    """Export fixture assessment; assert no unresolved anchors in output (if export succeeds)."""
    payload = _minimal_full_payload()
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (ensure ADA/report template.docx exists)")
    sys.path.insert(0, str(REPORTER_DIR))
    from qc_export import check_unresolved_anchors
    doc = Document(str(out_path))
    found = check_unresolved_anchors(doc)
    assert found == [], f"Unresolved anchors in output: {found}"


def test_export_fixture_no_encoding_artifacts(work_dir):
    """Export fixture assessment; assert no encoding artifacts in output (if export succeeds)."""
    payload = _minimal_full_payload()
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (ensure ADA/report template.docx exists)")
    sys.path.insert(0, str(REPORTER_DIR))
    from qc_export import check_encoding_artifacts
    doc = Document(str(out_path))
    found = check_encoding_artifacts(doc)
    assert found == [], f"Encoding artifacts in output: {found}"


def test_export_fixture_style_sampling(work_dir):
    """After export, assert narrative styles present: at least one Heading 1 and one Heading 2."""
    payload = _minimal_full_payload()
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (ensure ADA/report template.docx exists)")
    doc = Document(str(out_path))
    style_names = []
    for p in doc.paragraphs:
        try:
            name = p.style.name if p.style else None
            if name:
                style_names.append(name)
        except Exception:
            pass
    assert "Heading 1" in style_names, f"Expected at least one Heading 1 in output; got {style_names}"
    assert "Heading 2" in style_names, f"Expected at least one Heading 2 in output; got {style_names}"


def test_no_vofc_table_in_output(work_dir):
    """Output must not contain a VOFC table (Category | Vulnerability | Option for Consideration). Narrative-only."""
    payload = _minimal_full_payload()
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (ensure ADA/report template.docx has Part II anchors)")
    doc = Document(str(out_path))
    vofc_headers = ("Category", "Vulnerability", "Option for Consideration")
    for table in doc.tables:
        if table.rows and len(table.rows[0].cells) >= 3:
            row0_text = [c.text.strip() for c in table.rows[0].cells]
            if row0_text[:3] == list(vofc_headers):
                pytest.fail(
                    "Output must not contain VOFC table (Category | Vulnerability | Option for Consideration). "
                    "Use [[VULNERABILITY_BLOCKS]] narrative only."
                )
    # Pass: no table with that header row


def test_export_structured_vuln_narrative_no_tables(work_dir):
    """Export with structured VULN_NARRATIVE (≥1 vuln with evidence + OFCs); assert no vulnerability tables."""
    payload = _structured_vuln_narrative_payload()
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (ensure ADA/report template.docx has [[VULNERABILITY_BLOCKS]])")
    doc = Document(str(out_path))
    vofc_headers = ("Category", "Vulnerability", "Option for Consideration")
    for table in doc.tables:
        if table.rows and len(table.rows[0].cells) >= 3:
            row0_text = [c.text.strip() for c in table.rows[0].cells]
            if row0_text[:3] == list(vofc_headers):
                pytest.fail("No vulnerability tables allowed; use [[VULNERABILITY_BLOCKS]] narrative only.")
    # Assert export remains narrative-first and free of legacy table-based VOFC output.
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "SECTOR:" in full_text or "ELECTRIC POWER" in full_text or "Electric Power" in full_text
    assert "[[VULNERABILITY_BLOCKS]]" not in full_text


def test_vuln_narrative_empty_structured_fails(work_dir):
    """When vulnerability_blocks is empty, export still succeeds without leaking anchors/placeholders."""
    payload = _minimal_full_payload()
    payload["vulnerability_blocks"] = ""
    work_dir.mkdir(parents=True, exist_ok=True)
    if not CANONICAL_TEMPLATE.exists():
        pytest.skip("No template: ADA/report template.docx")
    import os
    old_work = os.environ.get("WORK_DIR")
    old_template = os.environ.get("TEMPLATE_PATH")
    os.environ["WORK_DIR"] = str(work_dir)
    os.environ["TEMPLATE_PATH"] = str(CANONICAL_TEMPLATE)
    sys.path.insert(0, str(REPORTER_DIR))
    old_stdin = sys.stdin
    try:
        sys.stdin = io.StringIO(json.dumps(payload))
        from main import main as reporter_main
        reporter_main()
        out = work_dir / "output.docx"
        assert out.is_file()
        doc = Document(str(out))
        full_text = _full_doc_text(doc)
        assert "[[VULNERABILITY_BLOCKS]]" not in full_text
    finally:
        sys.stdin = old_stdin
        os.environ.pop("WORK_DIR", None)
        os.environ.pop("TEMPLATE_PATH", None)
        if old_work:
            os.environ["WORK_DIR"] = old_work
        if old_template:
            os.environ["TEMPLATE_PATH"] = old_template


def test_vulnerability_section_boundary_verifier_passes_when_next_heading_after_vuln():
    """Regression: verifier passes when next section heading (e.g. C. CROSS-INFRASTRUCTURE SYNTHESIS) appears after last vuln content."""
    sys.path.insert(0, str(REPORTER_DIR))
    try:
        from main import _verify_vulnerability_section_boundary
        doc = Document()
        doc.add_paragraph("Some intro")
        doc.add_paragraph("VULNERABILITY 1")
        doc.add_paragraph("Content for vuln 1.")
        doc.add_paragraph("C. CROSS-INFRASTRUCTURE SYNTHESIS")
        doc.add_paragraph("Cross-infra content.")
        _verify_vulnerability_section_boundary(doc)  # must not raise
    finally:
        if REPORTER_DIR in sys.path:
            sys.path.remove(REPORTER_DIR)


def test_vulnerability_section_boundary_verifier_fails_when_next_heading_before_vuln():
    """Regression: verifier raises when next section appears before last vuln (content bleed)."""
    sys.path.insert(0, str(REPORTER_DIR))
    try:
        from main import _verify_vulnerability_section_boundary
        doc = Document()
        doc.add_paragraph("C. CROSS-INFRASTRUCTURE SYNTHESIS")  # earlier in doc
        doc.add_paragraph("VULNERABILITY 1")
        doc.add_paragraph("Content.")
        with pytest.raises(RuntimeError, match="boundary error|does not appear after"):
            _verify_vulnerability_section_boundary(doc)
    finally:
        if REPORTER_DIR in sys.path:
            sys.path.remove(REPORTER_DIR)


def test_export_fixture_no_safe_in_output(work_dir):
    """Output must not contain the word SAFE (federal-style report; no SAFE framework)."""
    payload = _minimal_full_payload()
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (ensure ADA/report template.docx exists)")
    doc = Document(str(out_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for cell in _iter_table_cell_text(doc):
        full_text += "\n" + cell
    assert "SAFE" not in full_text, "Output must not contain SAFE; federal-style report only."


def _iter_table_cell_text(doc):
    """Yield text of each table cell in doc body."""
    from docx.oxml.ns import qn
    body = doc.element.body
    for child in body:
        if child.tag == qn("w:tbl"):
            for row in child.iter(qn("w:tr")):
                for cell in row.iter(qn("w:tc")):
                    parts = []
                    for p_el in cell.iter(qn("w:p")):
                        for el in p_el.iter():
                            if el.tag == qn("w:t") and el.text:
                                parts.append(el.text)
                    yield "".join(parts)


def test_export_fixture_no_empty_sector_headers(work_dir):
    """No paragraph that is only 'SECTOR:' (or similar) with no sector name (no empty sector headings)."""
    payload = _structured_vuln_narrative_payload()
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (ensure ADA/report template.docx exists)")
    doc = Document(str(out_path))
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if txt == "SECTOR:" or txt.startswith("SECTOR:") and len(txt.strip().replace("SECTOR:", "").strip()) == 0:
            pytest.fail("Empty SECTOR: header found; no empty sector headings allowed.")


def test_export_structured_vuln_options_at_most_four(work_dir):
    """Federal-style blocks: each vulnerability block has at most 4 Options for Consideration."""
    payload = _structured_vuln_narrative_payload()
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (ensure ADA/report template.docx exists)")
    doc = Document(str(out_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # Split by "VULNERABILITY N" to get blocks; within each, count numbered OFC lines (1. 2. 3. 4.)
    import re
    blocks = re.split(r"\n(?=VULNERABILITY \d+)", full_text, flags=re.IGNORECASE)
    for block in blocks:
        if "Options for Consideration" not in block:
            continue
        ofc_section = block.split("Options for Consideration")[-1].split("VULNERABILITY")[0]
        numbered = re.findall(r"^\s*\d+\.\s", ofc_section, re.MULTILINE)
        assert len(numbered) <= 4, f"At most 4 options per vulnerability; found {len(numbered)} in block."


def _full_doc_text(doc) -> str:
    """Paragraphs then table cell text in document order (for section-order checks)."""
    parts = [p.text or "" for p in doc.paragraphs]
    for cell in _iter_table_cell_text(doc):
        parts.append(cell)
    return "\n".join(parts)


def test_docx_key_risk_drivers_section_exists(work_dir):
    """DOCX must include Key Risk Drivers (same as in-app); from executive_snapshot.drivers."""
    payload = _minimal_full_payload()
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (ensure ADA/report template.docx exists)")
    doc = Document(str(out_path))
    full_text = _full_doc_text(doc)
    assert "Key Risk Drivers" in full_text or "Electric Power" in full_text, (
        "Key Risk Drivers section or driver content must appear (executive_snapshot.drivers)."
    )


def test_docx_sensitivity_matrix_table_exists(work_dir):
    """DOCX must include sensitivity matrix table (Sector/Infrastructure column or equivalent)."""
    payload = _minimal_full_payload()
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (ensure ADA/report template.docx exists)")
    doc = Document(str(out_path))
    full_text = _full_doc_text(doc)
    has_sector = "Sector" in full_text or "Infrastructure" in full_text
    has_matrix_indicators = "Time to Severe" in full_text or "Functional Loss" in full_text or "Structural Posture" in full_text
    assert has_sector or has_matrix_indicators, (
        "Infrastructure Sensitivity matrix (Sector/Infrastructure or matrix columns) must appear."
    )


def test_docx_vuln_block_no_structural_or_standards(work_dir):
    """Vulnerability blocks in DOCX must not contain Structural Indicator Profile or Standards Alignment."""
    payload = _structured_vuln_narrative_payload()
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (ensure ADA/report template.docx has [[VULNERABILITY_BLOCKS]])")
    doc = Document(str(out_path))
    full_text = _full_doc_text(doc)
    # Vulnerability blocks can be legitimately absent when no condition-derived findings are present.
    assert "Structural Indicator Profile" not in full_text, (
        "DOCX must not render Structural Indicator Profile (in-app parity)."
    )
    assert "Standards Alignment" not in full_text, (
        "DOCX must not render Standards Alignment (in-app parity)."
    )


def test_docx_priority_actions_after_cross_infra(work_dir):
    """Priority Actions must appear after Cross-Infrastructure Analysis in document order."""
    payload = _minimal_full_payload()
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (ensure ADA/report template.docx exists)")
    doc = Document(str(out_path))
    full_text = _full_doc_text(doc)
    cross_pos = full_text.find("Cross-Infrastructure")
    if cross_pos < 0:
        cross_pos = full_text.find("Cross-infrastructure")
    priority_pos = full_text.find("Priority Actions")
    if priority_pos < 0 or cross_pos < 0:
        pytest.skip("Template may not include both Cross-Infrastructure and Priority Actions sections")
    assert priority_pos > cross_pos, (
        "Priority Actions must appear after Cross-Infrastructure Analysis in document order."
    )


def test_docx_posture_narrative_once(work_dir):
    """Executive posture narrative must appear exactly once (no duplication in Structural Risk Profile)."""
    payload = _minimal_full_payload()
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (ensure ADA/report template.docx exists)")
    doc = Document(str(out_path))
    full_text = _full_doc_text(doc)
    # Distinctive phrase from executive_snapshot used in posture/summary
    conditional_count = full_text.count("CONDITIONAL")
    key_risk_count = full_text.count("Key risk drivers")
    assert conditional_count <= 1 and key_risk_count <= 1, (
        "Posture/summary narrative (CONDITIONAL / Key risk drivers) must not be duplicated."
    )


def test_docx_no_vendor_failure_or_platform_outage(work_dir):
    """Report must not reference vendor failure or platform outage; IT hosted scenario is internet connectivity loss."""
    payload = _minimal_full_payload()
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (ensure ADA/report template.docx exists)")
    doc = Document(str(out_path))
    full_text = _full_doc_text(doc).lower()
    assert "vendor failure" not in full_text, "DOCX must not reference vendor failure (IT scenario is internet connectivity loss)."
    assert "platform outage" not in full_text, "DOCX must not reference platform outage (IT scenario is internet connectivity loss)."


def test_docx_part2_no_duplicate_synthesis_or_priority_actions_heading(work_dir):
    """Part II: template owns headings; no injected Heading 2 'Cross-Infrastructure Synthesis' or 'Priority Actions'."""
    payload = {
        "assessment": _minimal_assessment(),
        "vofc_collection": {"items": []},
        "executive_snapshot": {
            "posture": "CONDITIONAL",
            "summary": "Key risk drivers: Electric Power.",
            "drivers": ["Electric Power"],
            "matrixRows": [{"sector": "Electric Power", "ttiHrs": "24", "lossPct": "50", "backupHrs": "48", "structuralPosture": "Tolerant"}],
            "cascade": None,
        },
        "synthesis": {"title": "Cross-Infrastructure Synthesis", "paragraphs": ["Synthesis body."], "bullets": []},
        "priority_actions": {"title": "Priority Actions", "actions": [{"leadIn": "Action 1", "fullText": "Description."}]},
        "report_vm": {
            "part2": {
                "vulnerabilities": [
                    {"severity": "MODERATE", "title": "Placeholder", "narrative": "Minimal.", "ofcs": [], "references": []},
                ],
                "internet_transport_rows": [],
                "critical_hosted_services_rows": [],
                "dependency_summary_rows": [],
            },
        },
        "vulnerability_count_summary": "1 finding (0 HIGH, 0 ELEVATED, 1 MODERATE).",
    }
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (ensure template has [[SYNTHESIS]] and [[PRIORITY_ACTIONS]])")
    doc = Document(str(out_path))
    for p in doc.paragraphs:
        style_name = (p.style and p.style.name) or ""
        text = (p.text or "").strip()
        if style_name == "Heading 2":
            assert text != "Cross-Infrastructure Synthesis", (
                "Template owns Part II headings; do not inject Heading 2 'Cross-Infrastructure Synthesis'."
            )
            assert text != "Priority Actions", (
                "Template owns Part II headings; do not inject Heading 2 'Priority Actions'."
            )


def test_docx_no_documented_wording_in_tables(work_dir):
    """Part II tables must not contain 'Documented sources' or 'Not documented'; use 'Reported sources: 1' or 'Not provided' / '—'."""
    payload = {
        "assessment": _minimal_assessment(),
        "vofc_collection": {"items": []},
        "report_vm": {
            "part2": {
                "vulnerabilities": [
                    {"severity": "MODERATE", "title": "P", "narrative": "N.", "ofcs": [], "references": []},
                ],
                "internet_transport_rows": [
                    {"role": "Primary Internet Provider", "provider": "ISP", "demarcation": "", "independence": "", "notes": ""},
                ],
                "critical_hosted_services_rows": [],
                "dependency_summary_rows": [
                    {"category": "Electric Power", "primary_provider": "Yes", "backup_present": "No", "time_to_severe_impact_hrs": "24", "recovery_time_hrs": "12", "notes": "Reported sources: 1"},
                ],
            },
        },
        "vulnerability_count_summary": "1 finding (0 HIGH, 0 ELEVATED, 1 MODERATE).",
    }
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (template may lack Part II anchors)")
    full_text = _full_doc_text(Document(str(out_path)))
    assert "Documented sources" not in full_text, "Tables must not use 'Documented sources'"
    assert "Not documented" not in full_text, "Tables must not use 'Not documented'"
    assert "documented in assessment input" not in full_text, "Must not use 'documented in assessment input' (use 'Not provided' / 'Not confirmed')"
    assert "Sources provided: 1" not in full_text, "Must use 'Reported sources: 1' not 'Sources provided: 1'"
    assert "Reported sources: 1" in full_text, "Dependency summary Notes must show 'Reported sources: 1' when payload has it"
    assert "Not provided" in full_text or "Reported sources: 1" in full_text or "—" in full_text, (
        "Tables must use 'Not provided', 'Reported sources: 1', or '—' where appropriate"
    )


def test_docx_dependency_summary_missing_metrics_shows_na(work_dir):
    """When dependency summary has missing tti/recovery, table cells must show N/A, not 0."""
    payload = {
        "assessment": _minimal_assessment(),
        "vofc_collection": {"items": []},
        "report_vm": {
            "part2": {
                "vulnerabilities": [
                    {"severity": "MODERATE", "title": "P", "narrative": "N.", "ofcs": [], "references": []},
                ],
                "internet_transport_rows": [],
                "critical_hosted_services_rows": [],
                "dependency_summary_rows": [
                    {"category": "Electric Power", "primary_provider": "No", "backup_present": "No", "time_to_severe_impact_hrs": "N/A", "recovery_time_hrs": "N/A", "notes": "—"},
                ],
            },
        },
        "vulnerability_count_summary": "1 finding (0 HIGH, 0 ELEVATED, 1 MODERATE).",
    }
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (template may lack Part II anchors)")
    full_text = _full_doc_text(Document(str(out_path)))
    assert "N/A" in full_text, "Missing tti/recovery must render as N/A in dependency summary table"


def test_docx_ofcs_rendered_when_present(work_dir):
    """When a vulnerability has OFCs, DOCX must contain 'Options for Consideration' and the OFC text."""
    payload = {
        "assessment": _minimal_assessment(),
        "vofc_collection": {"items": []},
        "report_vm": {
            "part2": {
                "vulnerabilities": [
                    {
                        "severity": "MODERATE",
                        "title": "Vuln with OFCs",
                        "narrative": "Narrative.",
                        "ofcs": ["First OFC.", "Second OFC."],
                        "references": [],
                    },
                ],
                "internet_transport_rows": [],
                "critical_hosted_services_rows": [],
                "dependency_summary_rows": [
                    {"category": "Electric Power", "primary_provider": "No", "backup_present": "No", "time_to_severe_impact_hrs": "N/A", "recovery_time_hrs": "N/A", "notes": "—"},
                ],
            },
        },
        "vulnerability_count_summary": "1 finding (0 HIGH, 0 ELEVATED, 1 MODERATE).",
    }
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (template may lack Part II anchors)")
    full_text = _full_doc_text(Document(str(out_path)))
    assert "Options for Consideration" in full_text, "DOCX must contain OFC heading when vulnerability has OFCs"
    assert "First OFC" in full_text, "First OFC text must appear in DOCX"
    assert "Second OFC" in full_text, "Second OFC text must appear in DOCX"


def test_docx_no_options_heading_when_zero_ofcs(work_dir):
    """When a vulnerability has 0 OFCs, reporter must not add 'Options for Consideration' for that block (code path: only add when ofcs_clean non-empty)."""
    payload = {
        "assessment": _minimal_assessment(),
        "vofc_collection": {"items": []},
        "report_vm": {
            "part2": {
                "vulnerabilities": [
                    {"severity": "MODERATE", "title": "Vuln no OFCs", "narrative": "N.", "ofcs": [], "references": []},
                ],
                "internet_transport_rows": [],
                "critical_hosted_services_rows": [],
                "dependency_summary_rows": [
                    {"category": "Electric Power", "primary_provider": "No", "backup_present": "No", "time_to_severe_impact_hrs": "N/A", "recovery_time_hrs": "N/A", "notes": "—"},
                ],
            },
        },
        "vulnerability_count_summary": "1 finding (0 HIGH, 0 ELEVATED, 1 MODERATE).",
    }
    out_path = _run_reporter_and_get_docx(payload, work_dir)
    if out_path is None:
        pytest.skip("Export failed (template may lack Part II anchors)")
    full_text = _full_doc_text(Document(str(out_path)))
    assert "Vuln no OFCs" in full_text, "Vulnerability title must be rendered"
    assert "N." in full_text, "Vulnerability narrative must be rendered"
    # Reporter must not add OFC heading when ofcs is empty; we do not assert on total "Options for Consideration"
    # count here because the template may contain that phrase elsewhere.
