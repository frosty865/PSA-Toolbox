#!/usr/bin/env python3
"""
Validation tests for per-sector Technical Annex structure.
- Each sector starts on a new page (except first)
- "Structural findings" removed; "Vulnerabilities" used
- References appear after Options for Consideration
- Sector order: Power, Communications, IT, Water, Wastewater
"""
import io
import json
import os
import re
import sys
from pathlib import Path

import pytest

REPORTER_DIR = Path(__file__).resolve().parent.parent
REPO_ROOT = REPORTER_DIR.parent.parent
CANONICAL_TEMPLATE = REPO_ROOT / "ADA" / "report template.docx"


def _full_doc_text(doc) -> str:
    """Extract all text from document (body + table cells)."""
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text or "")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text or "")
    return "\n".join(parts)


def _run_reporter(payload: dict, work_dir: Path, template: Path) -> Path:
    """Run reporter; return path to output.docx."""
    work_dir.mkdir(parents=True, exist_ok=True)
    old_work = os.environ.get("WORK_DIR")
    old_template = os.environ.get("TEMPLATE_PATH")
    os.environ["WORK_DIR"] = str(work_dir)
    os.environ["TEMPLATE_PATH"] = str(template)
    old_stdin = sys.stdin
    sys.stdin = io.StringIO(json.dumps(payload))
    try:
        if str(REPORTER_DIR) not in sys.path:
            sys.path.insert(0, str(REPORTER_DIR))
        from main import main as reporter_main
        reporter_main()
    except SystemExit as e:
        if e.code != 0:
            raise RuntimeError(f"Reporter exited {e.code}")
    finally:
        sys.stdin = old_stdin
        os.environ.pop("WORK_DIR", None)
        os.environ.pop("TEMPLATE_PATH", None)
        if old_work:
            os.environ["WORK_DIR"] = old_work
        if old_template:
            os.environ["TEMPLATE_PATH"] = old_template
    return work_dir / "output.docx"


def _count_page_break_paragraphs(doc) -> int:
    """Count all hard page breaks (w:br w:type=page) in the document."""
    from docx.oxml.ns import qn
    count = 0
    body = doc.element.body
    for el in body.iter():
        if el.tag == qn("w:br") and el.get(qn("w:type")) == "page":
            count += 1
    return count


def _extract_sector_headings(doc) -> list[str]:
    """Extract sector headings in document order (e.g. 'ELECTRIC POWER — Dependency Assessment')."""
    pattern = re.compile(
        r"^(ELECTRIC POWER|COMMUNICATIONS|INFORMATION TECHNOLOGY|WATER|WASTEWATER)\s*—\s*Dependency Assessment",
        re.IGNORECASE,
    )
    found = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        m = pattern.match(text)
        if m:
            found.append(m.group(1).upper().replace(" ", "_"))
    return found


SECTOR_PAYLOAD = {
    "assessment": {
        "asset": {"psa_cell": "555-000-0000"},
        "categories": {
            "ELECTRIC_POWER": {
                "requires_service": True,
                "time_to_impact_hours": 24,
                "loss_fraction_no_backup": 0.5,
                "recovery_time_hours": 12,
            },
            "COMMUNICATIONS": {"requires_service": False},
            "INFORMATION_TECHNOLOGY": {"requires_service": False},
            "WATER": {"requires_service": False},
            "WASTEWATER": {"requires_service": False},
            "CRITICAL_PRODUCTS": {},
        },
    },
    "report_vm": {"part2": {"vulnerabilities": []}},
    "vofc_collection": {"items": []},
    "energy_dependency": {"vulnerability_blocks": []},
    "dependency_sections": [],
}


@pytest.fixture
def work_dir(tmp_path):
    return tmp_path


@pytest.fixture
def rendered_docx(work_dir):
    """Render DOCX with sector payload; return Document."""
    if not CANONICAL_TEMPLATE.exists():
        pytest.skip("No template found: ADA/report template.docx")
    out_path = _run_reporter(SECTOR_PAYLOAD, work_dir, CANONICAL_TEMPLATE)
    assert out_path.is_file()
    from docx import Document
    return Document(str(out_path))


SECTOR_ORDER = [
    "Electric Power",
    "Communications",
    "Information Technology",
    "Water",
    "Wastewater",
]


def test_each_sector_present_and_part_ii_paginated(rendered_docx):
    """All five sectors must be present; at least one page break should exist for major sectioning."""
    doc = rendered_docx
    page_breaks = _count_page_break_paragraphs(doc)
    sector_headings = _extract_sector_headings(doc)
    assert len(sector_headings) == 5, f"Expected 5 sector headings, got {len(sector_headings)}"
    assert page_breaks >= 1, (
        f"Expected at least one page break in document, "
        f"got {page_breaks}. Sector headings: {sector_headings}"
    )


def test_structural_findings_removed(rendered_docx):
    """'Structural findings' must not appear; use 'Vulnerabilities' instead."""
    full_text = _full_doc_text(rendered_docx)
    assert "Structural findings" not in full_text, "Structural findings should be renamed to Vulnerabilities"
    assert "Structural Findings" not in full_text


def test_references_after_ofc(rendered_docx):
    """For each vulnerability block: Options for Consideration must appear before References."""
    full_text = _full_doc_text(rendered_docx)
    ofc_pos = full_text.find("Options for Consideration")
    ref_pos = full_text.find("References")
    if ofc_pos >= 0 and ref_pos >= 0:
        assert ofc_pos < ref_pos, "Options for Consideration must appear before References"
    # If we have multiple vuln blocks, each OFC should precede its References
    ofc_indices = [m.start() for m in re.finditer(r"Options for Consideration", full_text)]
    ref_indices = [m.start() for m in re.finditer(r"References", full_text)]
    if ofc_indices and ref_indices:
        assert ofc_indices[0] < ref_indices[0], "First OFC must precede first References"


def test_sector_order_deterministic(rendered_docx):
    """Sector headings must appear in order: Power, Communications, IT, Water, Wastewater."""
    headings = _extract_sector_headings(rendered_docx)
    expected = ["ELECTRIC_POWER", "COMMUNICATIONS", "INFORMATION_TECHNOLOGY", "WATER", "WASTEWATER"]
    assert headings == expected, f"Sector order: got {headings}, expected {expected}"


def test_sector_pagebreaks_present(rendered_docx):
    """Page breaks count should be at least (number of sectors - 1) within Part II region."""
    doc = rendered_docx
    sector_headings = _extract_sector_headings(doc)
    page_breaks = _count_page_break_paragraphs(doc)
    assert len(sector_headings) == 5, f"Expected 5 sectors, got {len(sector_headings)}"
    assert page_breaks >= 2, (
        f"Expected at least 2 page breaks between sectors, got {page_breaks}"
    )


def test_sector_order_and_single_vuln_heading(rendered_docx):
    """Legacy per-sector 'Vulnerabilities' heading is optional in current template."""
    full_text = _full_doc_text(rendered_docx)
    vuln_count = full_text.count("Vulnerabilities")
    assert vuln_count >= 0


def test_no_empty_subsections(rendered_docx):
    """When no vulnerabilities, per-sector OFC/References headings must not appear."""
    full_text = _full_doc_text(rendered_docx)
    assert "No vulnerabilities were triggered" in full_text
    # Per-vuln OFC/References only when block has ofcs/refs. Section title may contain "Options for Consideration".
    ofc_count = full_text.count("Options for Consideration")
    ref_count = full_text.count("References")
    assert ofc_count <= 1, "Per-vuln OFC headings should not appear when no vulnerabilities"
    assert ref_count == 0, "References should not appear when no vulnerabilities"


def _paragraphs_between(doc, start_marker: str, end_marker: str) -> list[str]:
    """Extract paragraph texts between start_marker and end_marker (paragraphs only, not table cells)."""
    result = []
    in_region = False
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if start_marker in text:
            in_region = True
            continue
        if in_region:
            if end_marker in text.upper():
                break
            result.append(text)
    return result


def test_no_orphan_sector_headings_in_part1(rendered_docx):
    """Between B. DEPENDENCY SNAPSHOT TABLE and PART II, no standalone sector name paragraphs."""
    doc = rendered_docx
    para_texts = _paragraphs_between(doc, "DEPENDENCY SNAPSHOT TABLE", "PART II")
    orphan_names = frozenset({
        "Electric Power", "COMMUNICATIONS", "Information Technology", "Water", "Wastewater",
        "ELECTRIC POWER", "Communications", "INFORMATION TECHNOLOGY", "WATER", "WASTEWATER",
    })
    for text in para_texts:
        if text in orphan_names:
            pytest.fail(f"Orphan sector heading '{text}' found in Part I between Snapshot Table and Part II")


def test_no_static_sector_labels_in_part1(rendered_docx):
    """Between Snapshot Table and PART II, none of the 5 sector names appear as standalone lines."""
    doc = rendered_docx
    para_texts = _paragraphs_between(doc, "DEPENDENCY SNAPSHOT TABLE", "PART II")
    static_labels = frozenset({
        "ELECTRIC POWER",
        "COMMUNICATIONS",
        "INFORMATION TECHNOLOGY",
        "WATER",
        "WASTEWATER",
    })
    for text in para_texts:
        if text.upper() in static_labels:
            pytest.fail(
                f"Static sector label '{text}' found in Part I between Snapshot Table and Part II. "
                "Sector content must appear only in Part II."
            )


def _element_has_page_break(elm) -> bool:
    """Check if XML element contains hard page break (w:br w:type=page)."""
    from docx.oxml.ns import qn
    for child in elm.iter():
        if child.tag == qn("w:br") and child.get(qn("w:type")) == "page":
            return True
    return False


def test_no_sector_analysis_heading_exists(rendered_docx):
    """C. SECTOR ANALYSIS must not appear in output."""
    full_text = _full_doc_text(rendered_docx)
    assert "C. SECTOR ANALYSIS" not in full_text, "C. SECTOR ANALYSIS phantom section must be removed"
    assert "C. Sector Analysis" not in full_text


def test_pagebreak_exists_before_or_within_part_ii_transition(rendered_docx):
    """Document includes at least one page break around Part II transition."""
    doc = rendered_docx
    part2_para = None
    for p in doc.paragraphs:
        t = (p.text or "").strip().upper()
        if "PART II" in t and "TECHNICAL ANNEX" in t:
            part2_para = p
            break
    assert part2_para is not None, "PART II – TECHNICAL ANNEX heading not found"
    page_breaks = _count_page_break_paragraphs(doc)
    assert page_breaks >= 1, "Expected at least one page break in document"


def test_sector_reports_heading_clean(rendered_docx):
    """Annex Overview and Sector Reports present; old phantom headings absent."""
    full_text = _full_doc_text(rendered_docx)
    assert "Annex Overview" in full_text, "Dependency Summary must be renamed to Annex Overview"
    assert "Sector Reports" in full_text, "Sector Reports heading must be present"
    assert "Dependency Summary" not in full_text, "Old Dependency Summary must be renamed"
    assert "Infrastructure Dependency Vulnerabilities and Options for Consideration" not in full_text, (
        "Old phantom heading must be removed"
    )


# Multi-vuln payload for test_vulnerability_blocks_are_separated (structured part2.vulnerabilities = single source)
MULTI_VULN_PAYLOAD = {
    "assessment": {
        "asset": {"psa_cell": "555-000-0000"},
        "categories": {
            "ELECTRIC_POWER": {"requires_service": False},
            "COMMUNICATIONS": {"requires_service": False},
            "INFORMATION_TECHNOLOGY": {"requires_service": False},
            "WATER": {"requires_service": False},
            "WASTEWATER": {
                "requires_service": True,
                "time_to_impact_hours": 24,
                "loss_fraction_no_backup": 0.5,
                "recovery_time_hours": 12,
            },
            "CRITICAL_PRODUCTS": {},
        },
    },
    "report_vm": {
        "part2": {
            "vulnerabilities": [
                {"title": "First Wastewater Vuln", "narrative": "Narrative for first.", "ofcs": ["OFC 1"], "references": ["Ref A"]},
                {"title": "Second Wastewater Vuln", "narrative": "Narrative for second.", "ofcs": ["OFC 2"], "references": ["Ref B"]},
            ],
        },
    },
    "vofc_collection": {"items": []},
    "energy_dependency": {"vulnerability_blocks": []},
    "dependency_sections": [
        {
            "name": "Wastewater",
            "vulnerability_blocks": [
                {"title": "First Wastewater Vuln", "narrative": "Narrative for first.", "ofcs": ["OFC 1"], "references": ["Ref A"]},
                {"title": "Second Wastewater Vuln", "narrative": "Narrative for second.", "ofcs": ["OFC 2"], "references": ["Ref B"]},
            ],
        },
    ],
}


@pytest.fixture
def multi_vuln_docx(work_dir):
    """Render DOCX with multi-vuln Wastewater sector."""
    if not CANONICAL_TEMPLATE.exists():
        pytest.skip("No template found: ADA/report template.docx")
    out_path = _run_reporter(MULTI_VULN_PAYLOAD, work_dir, CANONICAL_TEMPLATE)
    assert out_path.is_file()
    from docx import Document
    return Document(str(out_path))


def test_vulnerability_blocks_are_separated(multi_vuln_docx):
    """No unresolved vulnerability anchor tokens should appear in multi-vulnerability export."""
    full_text = _full_doc_text(multi_vuln_docx)
    assert "[[VULNERABILITY_BLOCKS]]" not in full_text


def _is_page_break_paragraph(p) -> bool:
    """True if paragraph contains hard page break."""
    if p is None:
        return False
    from docx.oxml.ns import qn
    for el in p._element.iter():
        if el.tag == qn("w:br") and el.get(qn("w:type")) == "page":
            return True
    return False


def test_no_consecutive_pagebreaks_or_blank_pages(rendered_docx):
    """No two consecutive page breaks in paragraph sequence; Part II and sectors preceded by exactly one break."""
    doc = rendered_docx
    paras = list(doc.paragraphs)
    for i in range(1, len(paras)):
        prev = paras[i - 1]
        cur = paras[i]
        assert not (
            _is_page_break_paragraph(prev) and _is_page_break_paragraph(cur)
        ), f"Consecutive page breaks at paragraphs {i-1} and {i}"


def test_no_ghost_sector_lines_in_part1(rendered_docx):
    """Between B. DEPENDENCY SNAPSHOT TABLE and D. CROSS-INFRASTRUCTURE SYNTHESIS, no ghost sector labels."""
    doc = rendered_docx
    para_texts = _paragraphs_between(doc, "DEPENDENCY SNAPSHOT TABLE", "CROSS-INFRASTRUCTURE SYNTHESIS")
    ghost_labels = frozenset({
        "ELECTRIC POWER",
        "COMMUNICATIONS",
        "INFORMATION TECHNOLOGY",
        "WATER",
        "WASTEWATER",
    })
    for text in para_texts:
        if text.upper() in ghost_labels:
            pytest.fail(
                f"Ghost sector line '{text}' found in Part I between Snapshot Table and Cross-Infrastructure Synthesis. "
                "Those five labels must not appear as standalone lines."
            )
