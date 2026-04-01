#!/usr/bin/env python3
"""
Output Gates (ACTION ITEM 4) — unit and integration tests.
Tests: placeholder leak, anchor leak, deprecated term, empty section, duplicate block, max 4 cap.
"""
import io
import json
import os
import sys
from pathlib import Path

import pytest

REPORTER_DIR = Path(__file__).resolve().parent.parent
REPO_ROOT = REPORTER_DIR.parent.parent
CANONICAL_TEMPLATE = REPO_ROOT / "ADA" / "report template.docx"


def _minimal_assessment():
    return {
        "asset": {"psa_cell": "555-000-0000"},
        "categories": {
            "ELECTRIC_POWER": {
                "requires_service": True,
                "time_to_impact_hours": 24,
                "loss_fraction_no_backup": 0.5,
                "has_backup_any": True,
                "backup_duration_hours": 48,
                "loss_fraction_with_backup": 0.1,
                "recovery_time_hours": 12,
            },
            "COMMUNICATIONS": {"requires_service": False},
            "INFORMATION_TECHNOLOGY": {"requires_service": False},
            "WATER": {"requires_service": False},
            "WASTEWATER": {"requires_service": False},
            "CRITICAL_PRODUCTS": {},
        },
    }


def _run_reporter(payload: dict, work_dir: Path) -> tuple[int, str]:
    """Run reporter with payload; return (exit_code, stderr)."""
    if not CANONICAL_TEMPLATE.exists():
        pytest.skip("No template found: ADA/report template.docx")
    template = CANONICAL_TEMPLATE
    work_dir.mkdir(parents=True, exist_ok=True)
    old_work = os.environ.get("WORK_DIR")
    old_template = os.environ.get("TEMPLATE_PATH")
    os.environ["WORK_DIR"] = str(work_dir)
    os.environ["TEMPLATE_PATH"] = str(template)
    old_stdin = sys.stdin
    sys.stdin = io.StringIO(json.dumps(payload))
    stderr_capture = io.StringIO()
    old_stderr = sys.stderr
    sys.stderr = stderr_capture
    exit_code = 0
    try:
        if str(REPORTER_DIR) not in sys.path:
            sys.path.insert(0, str(REPORTER_DIR))
        from main import main as reporter_main
        reporter_main()
    except SystemExit as e:
        exit_code = e.code if e.code is not None else 1
    finally:
        sys.stdin = old_stdin
        sys.stderr = old_stderr
        os.environ.pop("WORK_DIR", None)
        os.environ.pop("TEMPLATE_PATH", None)
        if old_work:
            os.environ["WORK_DIR"] = old_work
        if old_template:
            os.environ["TEMPLATE_PATH"] = old_template
    return exit_code, stderr_capture.getvalue()


@pytest.fixture
def work_dir(tmp_path):
    return tmp_path


def test_placeholder_leak_fails_export(work_dir):
    """Gate A scanner catches placeholder text in rendered content."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Choose an item.")
    sys.path.insert(0, str(REPORTER_DIR))
    from main import gate_a_placeholder_anchor_leak_scanner
    with pytest.raises(SystemExit):
        gate_a_placeholder_anchor_leak_scanner(doc)


def test_deprecated_safe_fails_export(work_dir):
    """Gate A scanner catches deprecated SAFE term in rendered content."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("SAFE assessment required")
    sys.path.insert(0, str(REPORTER_DIR))
    from main import gate_a_placeholder_anchor_leak_scanner
    with pytest.raises(SystemExit):
        gate_a_placeholder_anchor_leak_scanner(doc)


def test_max_4_cap_applied(work_dir):
    """Legacy vofc_collection items are ignored by canonical pipeline (no raw VOFC text leakage)."""
    items = [
        {
            "category": "Electric Power",
            "vulnerability": f"Vulnerability {i}",
            "option_for_consideration": f"OFC {i}",
        }
        for i in range(6)
    ]
    payload = {
        "assessment": _minimal_assessment(),
        "vofc_collection": {"items": items},
    }
    exit_code, stderr = _run_reporter(payload, work_dir)
    assert exit_code == 0, f"Export should succeed: {stderr}"
    output_docx = work_dir / "output.docx"
    assert output_docx.exists()
    from docx import Document
    doc = Document(str(output_docx))
    full_text = " ".join(p.text or "" for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                full_text += " " + (cell.text or "")
    assert "Vulnerability 0" not in full_text
    assert "Vulnerability 5" not in full_text


def test_anchor_leak_fails_export(work_dir):
    """Gate A: Unresolved [[ANCHOR]] pattern should cause export to fail.
    We test by creating a doc with anchor and calling gate_a directly (scrub runs before gate_a in pipeline)."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Normal content")
    doc.add_paragraph("[[SYNTHESIS]]")
    doc.add_paragraph("More content")
    sys.path.insert(0, str(REPORTER_DIR))
    from main import gate_a_placeholder_anchor_leak_scanner
    with pytest.raises(SystemExit):
        gate_a_placeholder_anchor_leak_scanner(doc)


def test_duplicate_block_fails_export(work_dir):
    """Gate C: Identical paragraph in Water and Wastewater sections should cause export to fail."""
    from docx import Document
    doc = Document()
    h1 = doc.add_paragraph("Electric Power")
    h2 = doc.add_paragraph("Water")
    try:
        h1.style = "Heading 2"
        h2.style = "Heading 2"
    except Exception:
        pass
    doc.add_paragraph("Water service connections share collocated utility corridors with other utilities.")
    h3 = doc.add_paragraph("Wastewater")
    try:
        h3.style = "Heading 2"
    except Exception:
        pass
    doc.add_paragraph("Water service connections share collocated utility corridors with other utilities.")
    sys.path.insert(0, str(REPORTER_DIR))
    from main import gate_c_duplicate_block_detection
    with pytest.raises(SystemExit):
        gate_c_duplicate_block_detection(doc)


def test_empty_section_suppression(work_dir):
    """Gate B: Empty Critical Products heading with no content should be removed."""
    from docx import Document
    doc = Document()
    h1 = doc.add_paragraph("Critical Products")
    try:
        h1.style = "Heading 2"
    except Exception:
        pass
    doc.add_paragraph("")  # empty
    doc.add_paragraph("   ")  # blank
    h2 = doc.add_paragraph("Next Section")
    try:
        h2.style = "Heading 2"
    except Exception:
        pass
    doc.add_paragraph("Content here.")
    sys.path.insert(0, str(REPORTER_DIR))
    from main import gate_b_suppress_empty_sections
    gate_b_suppress_empty_sections(doc)
    texts = [p.text or "" for p in doc.paragraphs]
    assert "Critical Products" not in texts
    assert "Next Section" in texts
    assert "Content here." in texts


def test_priority_actions_injection(work_dir):
    """Priority Actions: when provided, exactly 5 actions render at [[PRIORITY_ACTIONS]].
    Assert: token removed, Priority Actions present, exactly five numbered actions (1. .. 5.)."""
    priority_actions = {
        "title": "Priority Actions",
        "actions": [
            {"number": 1, "leadIn": "Increase Structural Redundancy — Electric Power", "fullText": "Establish geographically separated paths. This dependency reaches severe impact in 4 hours with ~65% functional loss."},
            {"number": 2, "leadIn": "Introduce Alternate Capability — Electric Power", "fullText": "Develop alternate capability. This dependency reaches severe impact in 4 hours with ~65% functional loss."},
            {"number": 3, "leadIn": "Formalize Restoration Priority — Electric Power", "fullText": "Formalize priority restoration agreements. This dependency reaches severe impact in 4 hours with ~65% functional loss."},
            {"number": 4, "leadIn": "Increase Structural Redundancy — Communications", "fullText": "Establish geographically separated paths. This dependency reaches severe impact in 72 hours with ~0% functional loss."},
            {"number": 5, "leadIn": "Increase Structural Redundancy — Water", "fullText": "Establish geographically separated paths. This dependency reaches severe impact in 72 hours with ~0% functional loss."},
        ],
    }
    payload = {
        "assessment": _minimal_assessment(),
        "vofc_collection": {"items": []},
        "priority_actions": priority_actions,
    }
    exit_code, stderr = _run_reporter(payload, work_dir)
    assert exit_code == 0, f"Export should succeed: {stderr}"
    output_docx = work_dir / "output.docx"
    assert output_docx.exists()
    from docx import Document
    doc = Document(str(output_docx))
    full_text = " ".join(p.text or "" for p in doc.paragraphs)
    assert "[[PRIORITY_ACTIONS]]" not in full_text, "Anchor token should be removed"
    assert "Priority Actions" in full_text
    assert "Increase Structural Redundancy" in full_text
    for i in range(1, 6):
        assert f"{i}." in full_text or f"{i} " in full_text, f"Numbered action {i} should appear"
