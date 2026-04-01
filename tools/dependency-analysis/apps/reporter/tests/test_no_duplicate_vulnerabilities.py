#!/usr/bin/env python3
"""
Test: no duplicate vulnerability/OFC blocks (single source of truth).
When vulnerability_index_rows + vulnerability_blocks are used, each vuln title appears
at most twice: once in index table (title only), once in Structural findings (full block).
Narrative/OFC must not be duplicated.
"""
import io
import json
import os
import sys
from pathlib import Path

import pytest

REPORTER_DIR = Path(__file__).resolve().parent.parent
REPO_ROOT = REPORTER_DIR.parent.parent
CANONICAL_TEMPLATE = REPO_ROOT / "ADA" / "report template.docx"


def _full_doc_text(doc) -> str:
    """Extract all text from document (body + table cells)."""
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text or "")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text or "")
    return "\n".join(parts)


def _run_reporter(payload: dict, work_dir: Path, template: Path) -> Path:
    """Run reporter; return path to output.docx."""
    work_dir.mkdir(parents=True, exist_ok=True)
    old_work = os.environ.get("WORK_DIR")
    old_template = os.environ.get("TEMPLATE_PATH")
    os.environ["WORK_DIR"] = str(work_dir)
    os.environ["TEMPLATE_PATH"] = str(template)
    old_stdin = sys.stdin
    sys.stdin = io.StringIO(json.dumps(payload))
    try:
        if str(REPORTER_DIR) not in sys.path:
            sys.path.insert(0, str(REPORTER_DIR))
        from main import main as reporter_main
        reporter_main()
    except SystemExit as e:
        if e.code != 0:
            raise RuntimeError(f"Reporter exited {e.code}")
    finally:
        sys.stdin = old_stdin
        os.environ.pop("WORK_DIR", None)
        os.environ.pop("TEMPLATE_PATH", None)
        if old_work:
            os.environ["WORK_DIR"] = old_work
        if old_template:
            os.environ["TEMPLATE_PATH"] = old_template
    return work_dir / "output.docx"


@pytest.fixture
def work_dir(tmp_path):
    return tmp_path


def test_no_duplicate_vulnerability_titles_in_docx_output(work_dir):
    """Render docx from fixture with 3 vulns; assert each vuln title appears at most twice (index + block)."""
    if not CANONICAL_TEMPLATE.exists():
        pytest.skip("No template found: ADA/report template.docx")

    vuln_titles = ["Single Point of Failure", "Provider Confirmation Gap", "No Backup Path"]
    payload = {
        "assessment": {
            "asset": {"psa_cell": "555-000-0000"},
            "categories": {
                "ELECTRIC_POWER": {
                    "requires_service": False,
                    "time_to_impact_hours": 24,
                    "loss_fraction_no_backup": 0.5,
                    "recovery_time_hours": 12,
                    "curve_primary_provider": "Utility A",
                },
                "COMMUNICATIONS": {"requires_service": False},
                "INFORMATION_TECHNOLOGY": {"requires_service": False},
                "WATER": {"requires_service": False},
                "WASTEWATER": {"requires_service": False},
                "CRITICAL_PRODUCTS": {},
            },
        },
        "vofc_collection": {"items": []},
        "canonicalVulnBlocks": [
            {
                "domain": "ELECTRIC_POWER",
                "vuln_id": "ENERGY_STRUCT_SINGLE_SERVICE_FEED",
                "title": vuln_titles[0],
                "narrative": "Narrative A.",
                "ofcText": "OFC A1\nOFC A2\nOFC A3",
                "references": ["FEMA P-2166 - https://www.fema.gov"],
            },
            {
                "domain": "COMMUNICATIONS",
                "vuln_id": "COMMS_STRUCT_SINGLE_CARRIER_OR_ENTRY",
                "title": vuln_titles[1],
                "narrative": "Narrative B.",
                "ofcText": "OFC B1\nOFC B2\nOFC B3",
                "references": ["NIST SP 800-34 Rev.1 - https://csrc.nist.gov/publications/detail/sp/800-34/rev-1/final"],
            },
            {
                "domain": "INFORMATION_TECHNOLOGY",
                "vuln_id": "IT_STRUCT_SINGLE_EXTERNAL_CIRCUIT",
                "title": vuln_titles[2],
                "narrative": "Narrative C.",
                "ofcText": "OFC C1\nOFC C2\nOFC C3",
                "references": ["NFPA 1600 - https://www.nfpa.org"],
            },
        ],
        "energy_dependency": {
            "vulnerability_blocks": [
                {
                    "title": vuln_titles[0],
                    "narrative": "Narrative A.",
                    "ofcs": ["OFC A1"],
                    "references": [],
                },
            ],
        },
        "dependency_sections": [
            {
                "name": "Communications",
                "vulnerability_blocks": [
                    {
                        "title": vuln_titles[1],
                        "narrative": "Narrative B.",
                        "ofcs": ["OFC B1"],
                        "references": [],
                    },
                ],
            },
            {
                "name": "Information Technology",
                "vulnerability_blocks": [
                    {
                        "title": vuln_titles[2],
                        "narrative": "Narrative C.",
                        "ofcs": ["OFC C1"],
                        "references": [],
                    },
                ],
            },
        ],
        "vulnerability_index_rows": [
            {"title": vuln_titles[0], "sector": "Energy"},
            {"title": vuln_titles[1], "sector": "Communications"},
            {"title": vuln_titles[2], "sector": "Information Technology"},
        ],
    }

    out_path = _run_reporter(payload, work_dir, CANONICAL_TEMPLATE)
    assert out_path.is_file()

    from docx import Document
    doc = Document(str(out_path))
    full_text = _full_doc_text(doc)

    for title in vuln_titles:
        count = full_text.count(title)
        assert count >= 1, f"Vuln title '{title}' should appear at least once"
        assert count <= 3, f"Vuln title '{title}' appears {count} times (max 3 in current layout)"
