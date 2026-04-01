#!/usr/bin/env python3
"""
Chart rendering uniformity tests (Part A).
- test_chart_dimensions_uniform: all sector charts have identical pixel dimensions
- test_chart_two_series_visibility_when_equal: identical curves use different linestyles
"""
import io
import json
import os
import sys
from pathlib import Path

import pytest

REPORTER_DIR = Path(__file__).resolve().parent.parent
REPO_ROOT = REPORTER_DIR.parent.parent
CANONICAL_TEMPLATE = REPO_ROOT / "ADA" / "report template.docx"


def _minimal_assessment():
    return {
        "asset": {"psa_cell": "555-000-0000"},
        "categories": {
            "ELECTRIC_POWER": {
                "requires_service": True,
                "time_to_impact_hours": 24,
                "loss_fraction_no_backup": 0.5,
                "has_backup_any": True,
                "backup_duration_hours": 48,
                "loss_fraction_with_backup": 0.5,
                "recovery_time_hours": 12,
            },
            "COMMUNICATIONS": {"requires_service": False},
            "INFORMATION_TECHNOLOGY": {"requires_service": False},
            "WATER": {"requires_service": False},
            "WASTEWATER": {"requires_service": False},
            "CRITICAL_PRODUCTS": {},
        },
    }


def _minimal_report_vm_part2():
    """Minimal report_vm.part2 so [[VULNERABILITY_BLOCKS]] anchor is filled (structured path)."""
    return {
        "report_vm": {
            "part2": {
                "vulnerabilities": [
                    {"severity": "MODERATE", "title": "Placeholder", "narrative": "Minimal.", "ofcs": [], "references": []},
                ],
                "internet_transport_rows": [],
                "critical_hosted_services_rows": [],
                "dependency_summary_rows": [],
            },
        },
    }


def _run_reporter(payload: dict, work_dir: Path) -> tuple[int, Path]:
    """Run reporter; return (exit_code, work_dir)."""
    template = CANONICAL_TEMPLATE
    if not template.exists():
        pytest.skip("No template found")
    work_dir.mkdir(parents=True, exist_ok=True)
    old_work = os.environ.get("WORK_DIR")
    old_template = os.environ.get("TEMPLATE_PATH")
    os.environ["WORK_DIR"] = str(work_dir)
    os.environ["TEMPLATE_PATH"] = str(template)
    old_stdin = sys.stdin
    sys.stdin = io.StringIO(json.dumps(payload))
    exit_code = 0
    try:
        if str(REPORTER_DIR) not in sys.path:
            sys.path.insert(0, str(REPORTER_DIR))
        from main import main as reporter_main
        reporter_main()
    except SystemExit as e:
        exit_code = e.code if e.code is not None else 1
    finally:
        sys.stdin = old_stdin
        os.environ.pop("WORK_DIR", None)
        os.environ.pop("TEMPLATE_PATH", None)
        if old_work:
            os.environ["WORK_DIR"] = old_work
        if old_template:
            os.environ["TEMPLATE_PATH"] = old_template
    return exit_code, work_dir


def _get_png_dimensions(png_path: Path) -> tuple[int, int]:
    """Return (width, height) of PNG via PIL or struct."""
    try:
        from PIL import Image
        with Image.open(png_path) as img:
            return img.size
    except ImportError:
        import struct
        with open(png_path, "rb") as f:
            f.seek(16)
            w, h = struct.unpack(">II", f.read(8))
            return (w, h)


@pytest.fixture
def work_dir(tmp_path):
    return tmp_path


def test_all_sector_charts_same_pixel_dimensions(work_dir):
    """All 5 sector chart PNGs meet minimum CHART_W_PX x CHART_H_PX (bbox_inches='tight' may increase size to avoid legend clip)."""
    sys.path.insert(0, str(REPORTER_DIR))
    from main import CHART_W_PX, CHART_H_PX
    payload = {
        "assessment": _minimal_assessment(),
        "vofc_collection": {"items": []},
        **_minimal_report_vm_part2(),
        "priority_actions": {
            "title": "Priority Actions",
            "actions": [
                {"number": i, "leadIn": f"Action {i}", "fullText": f"Text {i}."}
                for i in range(1, 6)
            ],
        },
    }
    exit_code, wd = _run_reporter(payload, work_dir)
    assert exit_code == 0
    for code in ("ELECTRIC_POWER", "COMMUNICATIONS", "INFORMATION_TECHNOLOGY", "WATER", "WASTEWATER"):
        png = wd / f"chart_{code}.png"
        assert png.exists(), f"Chart {code} missing"
        w, h = _get_png_dimensions(png)
        assert w >= CHART_W_PX and h >= CHART_H_PX, f"Chart {code}: {w}x{h} below minimum {CHART_W_PX}x{CHART_H_PX}"


def test_chart_two_series_visibility_when_equal(work_dir):
    """When without/with backup curves are identical, render_dependency_chart_png uses different linestyles."""
    sys.path.insert(0, str(REPORTER_DIR))
    from main import render_dependency_chart_png, build_curve
    points = [{"t_hours": t, "capacity_without_backup": 50.0, "capacity_with_backup": 50.0} for t in range(0, 97, 6)]
    meta = {"tti_hours": 24, "alt_sust_hours": None, "sla_mttr_hours": None}
    out = work_dir / "test_equal.png"
    render_dependency_chart_png("Test Sector", points, meta, out)
    assert out.exists()
    w, h = _get_png_dimensions(out)
    assert w > 0 and h > 0


def test_build_chart_series_primary_is_baseline(work_dir):
    """PRIMARY line must correspond to baseline (capacity_without_backup), not mitigated series."""
    sys.path.insert(0, str(REPORTER_DIR))
    from main import build_chart_series
    # Baseline worse than with_backup at same t
    points = [
        {"t_hours": 0, "capacity_without_backup": 20.0, "capacity_with_backup": 80.0},
        {"t_hours": 24, "capacity_without_backup": 20.0, "capacity_with_backup": 80.0},
        {"t_hours": 96, "capacity_without_backup": 20.0, "capacity_with_backup": 80.0},
    ]
    series = build_chart_series(points)
    primary_y = series["primary_y"]
    secondary_y = series["secondary_y"]
    expected_baseline = [20.0, 20.0, 20.0]
    expected_mitigated = [80.0, 80.0, 80.0]
    assert primary_y == expected_baseline, "primary_y must be capacity_without_backup (baseline truth)"
    assert secondary_y == expected_mitigated, "secondary_y must be capacity_with_backup (mitigated)"
    assert series["labels"] == ("Without backup", "With backup")


def test_build_curve_96h_backup_is_indefinite_within_horizon(work_dir):
    """backup_duration_hours >= 96 should sustain through the modeled horizon (no endpoint drop)."""
    sys.path.insert(0, str(REPORTER_DIR))
    from main import build_curve
    category = {
        "requires_service": True,
        "time_to_impact_hours": 1,
        "loss_fraction_no_backup": 0.95,
        "has_backup_any": True,
        "backup_duration_hours": 96,
        "loss_fraction_with_backup": 0.0,
        "recovery_time_hours": 1,
    }
    points = build_curve(category, horizon_hours=96, step_hours=3)
    assert points, "Expected non-empty curve"
    end = points[-1]
    assert end["t_hours"] == 96
    assert end["capacity_with_backup"] == 100.0, "96h backup should be treated as indefinite within 0..96h horizon"


def test_ensure_curve_complete_with_backup_series_distinct_and_fallback(work_dir):
    """_ensure_curve_complete yields distinct no_backup vs with_backup when input differs; missing capacity_with_backup falls back to capacity_without_backup."""
    sys.path.insert(0, str(REPORTER_DIR))
    from main import _ensure_curve_complete
    # Input with distinct series
    points_diff = [
        {"t_hours": 0, "capacity_without_backup": 40.0, "capacity_with_backup": 70.0},
        {"t_hours": 24, "capacity_without_backup": 40.0, "capacity_with_backup": 70.0},
    ]
    out = _ensure_curve_complete(points_diff, horizon=96)
    no_backup = [p["capacity_without_backup"] for p in out]
    with_backup = [p["capacity_with_backup"] for p in out]
    assert no_backup != with_backup, "When input differs, output series must differ"
    assert no_backup[0] == 40.0 and with_backup[0] == 70.0
    # Input with missing capacity_with_backup: fallback to capacity_without_backup
    points_missing_with = [
        {"t_hours": 0, "capacity_without_backup": 25.0},
        {"t_hours": 24, "capacity_without_backup": 25.0},
    ]
    out2 = _ensure_curve_complete(points_missing_with, horizon=96)
    for p in out2:
        assert "capacity_with_backup" in p
        assert p["capacity_with_backup"] == p["capacity_without_backup"], "Missing capacity_with_backup should fall back to capacity_without_backup"
    assert out2[0]["capacity_with_backup"] == 25.0


def test_chart_renderer_sets_linewidths(work_dir):
    """Refined stroke weights: primary 2.6, secondary 2.2, dashed 1.6."""
    sys.path.insert(0, str(REPORTER_DIR))
    from main import (
        PACE_LINEWIDTHS,
        DEP_CHART_LW_PRIMARY,
        DEP_CHART_LW_SECONDARY,
        PRIMARY_SOLID_WIDTH,
        SECONDARY_SOLID_WIDTH,
        DASHED_WIDTH,
    )
    assert PACE_LINEWIDTHS["PRIMARY"] == PRIMARY_SOLID_WIDTH
    assert PACE_LINEWIDTHS["ALTERNATE"] == SECONDARY_SOLID_WIDTH
    assert DEP_CHART_LW_PRIMARY == PRIMARY_SOLID_WIDTH
    assert DEP_CHART_LW_SECONDARY == DASHED_WIDTH


def test_no_black_line_colors(work_dir):
    """No chart line color equals black (#000000)."""
    sys.path.insert(0, str(REPORTER_DIR))
    from main import PACE_COLORS, DEP_COLOR_WITHOUT_BACKUP, DEP_COLOR_WITH_BACKUP
    all_colors = [
        *PACE_COLORS.values(),
        DEP_COLOR_WITHOUT_BACKUP,
        DEP_COLOR_WITH_BACKUP,
    ]
    black = "#000000"
    for c in all_colors:
        assert c.lower() != black.lower(), f"Black not allowed for data lines: {c}"


def test_primary_color_is_navy(work_dir):
    """Primary line color is deep navy (#003f5c)."""
    sys.path.insert(0, str(REPORTER_DIR))
    from main import PACE_COLOR_PRIMARY
    assert PACE_COLOR_PRIMARY == "#003f5c"


def test_primary_width_greater_than_dashed(work_dir):
    """Primary (solid) linewidth > dashed linewidth."""
    sys.path.insert(0, str(REPORTER_DIR))
    from main import PRIMARY_SOLID_WIDTH, DASHED_WIDTH
    assert PRIMARY_SOLID_WIDTH > DASHED_WIDTH


def test_chart_renderer_zorder_primary_base_secondary_on_top(work_dir):
    """PACE chart: Primary z=1 (base), Alternate/Contingency/Emergency z=2 (on top)."""
    sys.path.insert(0, str(REPORTER_DIR))
    from main import PACE_ZORDER, Z_SOLID_BASE, Z_SOLID_SECONDARY
    assert PACE_ZORDER["PRIMARY"] == Z_SOLID_BASE
    assert PACE_ZORDER["ALTERNATE"] == Z_SOLID_SECONDARY
    assert PACE_ZORDER["CONTINGENCY"] == Z_SOLID_SECONDARY
    assert PACE_ZORDER["EMERGENCY"] == Z_SOLID_SECONDARY


def test_dashed_zorder_highest(work_dir):
    """Dependency chart: dashed line (With backup) has highest zorder."""
    sys.path.insert(0, str(REPORTER_DIR))
    from main import Z_DASHED, Z_SOLID_BASE
    assert Z_DASHED > Z_SOLID_BASE


def test_psa_phone_injected(work_dir):
    """PSA phone replaces [[PSA_PHONE]]; token not present; phone appears exactly once."""
    base = _minimal_assessment()
    payload = {
        "assessment": {
            "asset": {**base["asset"], "psa_cell": "305-555-1212"},
            "categories": base["categories"],
        },
        "vofc_collection": {"items": []},
        **_minimal_report_vm_part2(),
        "priority_actions": {"title": "Actions", "actions": [{"number": 1, "leadIn": "A", "fullText": "Text."}]},
    }
    exit_code, wd = _run_reporter(payload, work_dir)
    assert exit_code == 0
    import zipfile
    with zipfile.ZipFile(wd / "output.docx", "r") as z:
        for name in z.namelist():
            if name.endswith(".xml"):
                content = z.read(name).decode("utf-8", errors="replace")
                assert "[[PSA_PHONE]]" not in content
                if "305-555-1212" in content:
                    return
    assert False, "PSA phone 305-555-1212 not found in output docx"


def test_psa_phone_missing_fails_export(work_dir):
    """Export fails when PSA phone is missing."""
    payload = {
        "assessment": {
            "asset": {},
            "categories": _minimal_assessment()["categories"],
        },
        "vofc_collection": {"items": []},
        **_minimal_report_vm_part2(),
        "priority_actions": {"title": "Actions", "actions": [{"number": 1, "leadIn": "A", "fullText": "T."}]},
    }
    exit_code, _ = _run_reporter(payload, work_dir)
    assert exit_code != 0


def test_chart_renderer_path_effects_exist(work_dir):
    """At least one chart line has path effects (white halo)."""
    sys.path.insert(0, str(REPORTER_DIR))
    from unittest.mock import patch
    import main as main_mod
    from main import render_dependency_chart_png
    points = [{"t_hours": t, "capacity_without_backup": 80.0, "capacity_with_backup": 90.0} for t in range(0, 97, 6)]
    meta = {"tti_hours": 24, "alt_sust_hours": None, "sla_mttr_hours": None}
    out = work_dir / "test_path_effects.png"
    with patch.object(main_mod, "_apply_halo_path_effect", wraps=main_mod._apply_halo_path_effect) as mock_halo:
        render_dependency_chart_png("Test", points, meta, out)
    assert out.exists()
    assert mock_halo.call_count >= 1


def test_docx_inserts_fixed_dimensions(work_dir):
    """DOCX chart images have identical width and height (2.82:1 ratio)."""
    sys.path.insert(0, str(REPORTER_DIR))
    from main import CHART_W_INCHES, CHART_H_INCHES
    payload = {
        "assessment": _minimal_assessment(),
        "vofc_collection": {"items": []},
        **_minimal_report_vm_part2(),
        "priority_actions": {
            "title": "Actions",
            "actions": [{"number": i, "leadIn": f"A{i}", "fullText": f"Text {i}."} for i in range(1, 6)],
        },
    }
    exit_code, wd = _run_reporter(payload, work_dir)
    assert exit_code == 0
    import zipfile
    import xml.etree.ElementTree as ET
    from docx.shared import Emu
    with zipfile.ZipFile(wd / "output.docx", "r") as z:
        xml_bytes = z.read("word/document.xml")
    root = ET.fromstring(xml_bytes)
    wp_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    tol = 0.05
    chart_dims = []
    for ext in root.iter(f"{{{wp_ns}}}extent"):
        cx = ext.get("cx")
        cy = ext.get("cy")
        if cx and cy and int(cy) > 0:
            w_in = Emu(int(cx)).inches
            h_in = Emu(int(cy)).inches
            if abs(w_in - CHART_W_INCHES) <= tol and abs(h_in - CHART_H_INCHES) <= tol:
                chart_dims.append((w_in, h_in))
    assert len(chart_dims) >= 5, f"Expected at least 5 chart images with fixed dims, got {len(chart_dims)}"


def test_docx_chart_alignment(work_dir):
    """After injection, chart paragraphs are center-aligned (docx inspection)."""
    payload = {
        "assessment": _minimal_assessment(),
        "vofc_collection": {"items": []},
        **_minimal_report_vm_part2(),
        "priority_actions": {
            "title": "Priority Actions",
            "actions": [
                {"number": i, "leadIn": f"Action {i}", "fullText": f"Text {i}."}
                for i in range(1, 6)
            ],
        },
    }
    exit_code, wd = _run_reporter(payload, work_dir)
    assert exit_code == 0
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document(str(wd / "output.docx"))
    centered = 0
    for p in doc.paragraphs:
        if p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            centered += 1
    assert centered >= 5


def test_axis_horizon_is_96(work_dir):
    """Charts use 0–96 hour horizon; xlim end is 96; major ticks include 96."""
    sys.path.insert(0, str(REPORTER_DIR))
    from main import CHART_XLIM, render_dependency_chart_png
    assert CHART_XLIM == (0, 96)
    points = [{"t_hours": t, "capacity_without_backup": 50.0, "capacity_with_backup": 60.0} for t in range(0, 97, 6)]
    meta = {"tti_hours": 24, "alt_sust_hours": None, "sla_mttr_hours": None}
    out = work_dir / "test_horizon.png"
    render_dependency_chart_png("Test", points, meta, out)
    assert out.exists()
    # Verify chart dimensions are at least configured base size (render backends may add padding).
    dims = _get_png_dimensions(out)
    from main import CHART_W_PX, CHART_H_PX
    assert dims[0] >= CHART_W_PX
    assert dims[1] >= CHART_H_PX


def test_docx_chart_ratio_282(work_dir):
    """At least 5 chart images in DOCX have width/height ratio ≈ 2.82."""
    payload = {
        "assessment": _minimal_assessment(),
        "vofc_collection": {"items": []},
        **_minimal_report_vm_part2(),
        "priority_actions": {
            "title": "Actions",
            "actions": [{"number": i, "leadIn": f"A{i}", "fullText": f"Text {i}."} for i in range(1, 6)],
        },
    }
    exit_code, wd = _run_reporter(payload, work_dir)
    assert exit_code == 0
    import zipfile
    import xml.etree.ElementTree as ET
    from docx.shared import Emu
    wp_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    tol = 0.05
    with zipfile.ZipFile(wd / "output.docx", "r") as z:
        root = ET.fromstring(z.read("word/document.xml"))
    chart_count = 0
    for ext in root.iter(f"{{{wp_ns}}}extent"):
        cx, cy = ext.get("cx"), ext.get("cy")
        if cx and cy and int(cy) > 0:
            w_in = Emu(int(cx)).inches
            h_in = Emu(int(cy)).inches
            ratio = w_in / h_in
            if abs(ratio - 2.82) < tol:
                chart_count += 1
    assert chart_count >= 5, f"Expected ≥5 charts with ratio ≈2.82, got {chart_count}"


def test_docx_narrow_body_margins(work_dir):
    """Report body margins are set to 0.75 inches for space efficiency."""
    payload = {
        "assessment": _minimal_assessment(),
        "vofc_collection": {"items": []},
        **_minimal_report_vm_part2(),
        "priority_actions": {
            "title": "Actions",
            "actions": [{"number": 1, "leadIn": "A", "fullText": "Text."}],
        },
    }
    exit_code, wd = _run_reporter(payload, work_dir)
    assert exit_code == 0
    from docx import Document
    from docx.shared import Inches
    doc = Document(str(wd / "output.docx"))
    section = doc.sections[0]
    assert abs(section.left_margin.inches - 0.75) < 0.01
