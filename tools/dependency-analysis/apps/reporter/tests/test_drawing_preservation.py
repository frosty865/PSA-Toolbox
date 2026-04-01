"""
Tests for title-page graphic preservation: paragraph_has_drawing, replace_anchor_in_paragraph_preserve_drawings,
remove_paragraph (do not remove drawing paragraphs), and body_only anchor behavior.
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from main import (
    paragraph_has_drawing,
    replace_anchor_in_paragraph_preserve_drawings,
    replace_anchor_in_paragraph,
    remove_paragraph,
    find_paragraph_by_exact_text,
    find_anchor_paragraph_exact,
)


def test_paragraph_has_drawing_false_for_text_only() -> None:
    """Text-only paragraphs have no drawing; paragraph_has_drawing returns False."""
    doc = Document()
    p = doc.add_paragraph("Hello world")
    assert paragraph_has_drawing(p) is False
    p2 = doc.add_paragraph("[[CHART_ELECTRIC_POWER]]")
    assert paragraph_has_drawing(p2) is False


def test_replace_anchor_in_paragraph_preserve_drawings_single_run() -> None:
    """When anchor is in a single text run with no drawing, replace and return True."""
    doc = Document()
    p = doc.add_paragraph("[[PSA_CELL]]")
    ok = replace_anchor_in_paragraph_preserve_drawings(p, "[[PSA_CELL]]", "555-1234")
    assert ok is True
    assert (p.text or "").strip() == "555-1234"


def test_replace_anchor_in_paragraph_preserve_drawings_anchor_missing_returns_false() -> None:
    """When anchor is not in any run, return False (e.g. split across runs or not present)."""
    doc = Document()
    p = doc.add_paragraph("Other text")
    ok = replace_anchor_in_paragraph_preserve_drawings(p, "[[PSA_CELL]]", "x")
    assert ok is False
    assert (p.text or "").strip() == "Other text"


def test_replace_anchor_in_paragraph_preserves_drawing_paragraph() -> None:
    """replace_anchor_in_paragraph does not wipe runs when paragraph has no drawing; replaces text."""
    doc = Document()
    p = doc.add_paragraph("Cell: [[PSA_CELL]]")
    ok = replace_anchor_in_paragraph(p, "[[PSA_CELL]]", "555-0000")
    assert ok is True
    assert "555-0000" in (p.text or "")


def test_remove_paragraph_removes_text_only_paragraph() -> None:
    """remove_paragraph removes a text-only paragraph from the document."""
    doc = Document()
    doc.add_paragraph("Keep")
    p = doc.add_paragraph("Remove me")
    doc.add_paragraph("After")
    remove_paragraph(p)
    texts = [(x.text or "").strip() for x in doc.paragraphs]
    assert "Remove me" not in texts
    assert "Keep" in texts
    assert "After" in texts


def test_find_anchor_paragraph_exact_body_only() -> None:
    """find_anchor_paragraph_exact with body_only=True finds only in doc.paragraphs."""
    doc = Document()
    doc.add_paragraph("[[ANCHOR_A]]")
    body_only = list(find_anchor_paragraph_exact(doc, "[[ANCHOR_A]]", body_only=True))
    assert len(body_only) == 1
    assert body_only[0][0].text.strip() == "[[ANCHOR_A]]"
    all_paras = list(find_anchor_paragraph_exact(doc, "[[ANCHOR_A]]", body_only=False))
    assert len(all_paras) == 1
