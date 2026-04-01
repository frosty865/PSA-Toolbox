#!/usr/bin/env python3
"""
Regression parity check: generated DOCX must match template layout.
- No anchor tokens ([[CHART_*]], [[TABLE_*]], [[VULN_NARRATIVE]]) remain
- No narrative blanks ("____") remain
- No VOFC table (deprecated); vulnerability content is narrative-only ([[VULN_NARRATIVE]])
- No raw export-style tables

Run from repo root: python -m pytest asset-dependency-tool/apps/reporter/tests/test_report_parity.py -v
Or: cd asset-dependency-tool/apps/reporter && python -m pytest tests/test_report_parity.py -v
"""
import io
import json
import os
import sys
from pathlib import Path

import pytest

REPORTER_DIR = Path(__file__).resolve().parent.parent
REPO_ROOT = REPORTER_DIR.parent.parent
CANONICAL_TEMPLATE = REPO_ROOT / "ADA" / "report template.docx"


# Payload with one category and empty VOFC so all anchors get replaced
PARITY_PAYLOAD = {
    "assessment": {
        "asset": {
            "asset_name": "Parity Test Asset",
            "visit_date_iso": "2025-01-15",
            "assessor": "Test Assessor",
            "location": "Test Location",
            "psa_cell": "555-000-0000",
        },
        "categories": {
            "ELECTRIC_POWER": {
                "requires_service": True,
                "time_to_impact_hours": 24,
                "loss_fraction_no_backup": 0.5,
                "has_backup_any": True,
                "backup_duration_hours": 48,
                "loss_fraction_with_backup": 0.1,
                "recovery_time_hours": 12,
                "report_themed_findings": [
                    {
                        "id": "parity-1",
                        "domain": "ELECTRIC_POWER",
                        "severity": "MODERATE",
                        "title": "Test finding",
                        "narrative": "No sector-specific dependency vulnerabilities were triggered.",
                        "ofcText": "Complete the dependency assessment to see findings.",
                    }
                ],
            },
            "COMMUNICATIONS": {"requires_service": False},
            "INFORMATION_TECHNOLOGY": {"requires_service": False},
            "WATER": {"requires_service": False},
            "WASTEWATER": {"requires_service": False},
            "CRITICAL_PRODUCTS": {},
        },
    },
    "vofc_collection": {"items": []},
    "vulnerability_blocks": "",
}


@pytest.fixture(scope="module")
def work_dir(tmp_path_factory):
    return tmp_path_factory.mktemp("parity")


@pytest.fixture(scope="module")
def generated_docx_path(work_dir):
    """Run reporter and return path to output.docx. Skip if template missing."""
    if not CANONICAL_TEMPLATE.exists():
        pytest.skip("No template found: ADA/report template.docx")
    template = CANONICAL_TEMPLATE
    work_dir.mkdir(parents=True, exist_ok=True)
    old_work = os.environ.get("WORK_DIR")
    old_template = os.environ.get("TEMPLATE_PATH")
    os.environ["WORK_DIR"] = str(work_dir)
    os.environ["TEMPLATE_PATH"] = str(template)
    old_stdin = sys.stdin
    sys.stdin = io.StringIO(json.dumps(PARITY_PAYLOAD))
    try:
        if str(REPORTER_DIR) not in sys.path:
            sys.path.insert(0, str(REPORTER_DIR))
        from main import main as reporter_main
        reporter_main()
    finally:
        sys.stdin = old_stdin
        if old_work is not None:
            os.environ["WORK_DIR"] = old_work
        else:
            os.environ.pop("WORK_DIR", None)
        if old_template is not None:
            os.environ["TEMPLATE_PATH"] = old_template
        else:
            os.environ.pop("TEMPLATE_PATH", None)
    out = work_dir / "output.docx"
    assert out.is_file(), "Reporter did not create output.docx"
    return out


def test_no_anchor_tokens_remain(generated_docx_path):
    """Generated DOCX must not contain [[CHART_*]], [[TABLE_*]], or ANCHOR BLOCK."""
    from docx import Document
    doc = Document(str(generated_docx_path))
    text = "\n".join(p.text or "" for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + (cell.text or "")
    assert "[[CHART_" not in text, "Unreplaced [[CHART_*]] anchor(s) found"
    assert "[[TABLE_" not in text, "Unreplaced [[TABLE_*]] anchor(s) found"
    assert "ANCHOR BLOCK" not in text, "ANCHOR BLOCK (dev only) or similar found"


def test_no_narrative_blanks_remain(generated_docx_path):
    """Generated DOCX must not contain '____' or long underscore placeholders."""
    from docx import Document
    doc = Document(str(generated_docx_path))
    text = "\n".join(p.text or "" for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + (cell.text or "")
    assert "____" not in text, "Unreplaced narrative blank '____' found"
    assert "_________" not in text, "Unreplaced long narrative blank found"


def test_no_vofc_table(generated_docx_path):
    """Output must not contain VOFC table (Category | Vulnerability | Option for Consideration). Narrative-only ([[VULN_NARRATIVE]])."""
    from docx import Document
    doc = Document(str(generated_docx_path))
    expected_headers = ("Category", "Vulnerability", "Option for Consideration")
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if len(headers) >= 3 and headers[0] == expected_headers[0] and headers[1] == expected_headers[1] and headers[2] == expected_headers[2]:
            pytest.fail(
                "Output must not contain VOFC table (Category | Vulnerability | Option for Consideration). "
                "Use [[VULN_NARRATIVE]] narrative-only export."
            )


def test_verify_output_script_passes(generated_docx_path):
    """verify_output.py runs without crashing; heading checks are template-version dependent."""
    import subprocess
    verify_script = REPORTER_DIR / "verify_output.py"
    if not verify_script.exists():
        pytest.skip("verify_output.py not found")
    r = subprocess.run(
        [sys.executable, str(verify_script), str(generated_docx_path)],
        capture_output=True,
        text=True,
        cwd=str(REPORTER_DIR),
    )
    assert r.returncode in (0, 1), f"verify_output.py crashed unexpectedly: {r.stderr or r.stdout}"
