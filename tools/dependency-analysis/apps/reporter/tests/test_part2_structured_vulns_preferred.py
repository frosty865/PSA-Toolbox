#!/usr/bin/env python3
"""
Tripwire: when payload has both report_vm.part2.vulnerabilities and vulnerability_blocks,
the reporter must use the structured path only. Legacy string injection must not run.
"""
import io
import json
import os
import sys
from pathlib import Path

import pytest

REPORTER_DIR = Path(__file__).resolve().parent.parent
REPO_ROOT = REPORTER_DIR.parent.parent
CANONICAL_TEMPLATE = REPO_ROOT / "ADA" / "report template.docx"

LEGACY_MARKER = "LEGACY_MARKER_SHOULD_NOT_APPEAR"


def _full_doc_text(doc) -> str:
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text or "")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text or "")
    return "\n".join(parts)


def _run_reporter(payload: dict, work_dir: Path, template: Path) -> Path:
    work_dir.mkdir(parents=True, exist_ok=True)
    old_work = os.environ.get("WORK_DIR")
    old_template = os.environ.get("TEMPLATE_PATH")
    os.environ["WORK_DIR"] = str(work_dir)
    os.environ["TEMPLATE_PATH"] = str(template)
    old_stdin = sys.stdin
    sys.stdin = io.StringIO(json.dumps(payload))
    try:
        if str(REPORTER_DIR) not in sys.path:
            sys.path.insert(0, str(REPORTER_DIR))
        from main import main as reporter_main
        reporter_main()
    finally:
        sys.stdin = old_stdin
        os.environ.pop("WORK_DIR", None)
        os.environ.pop("TEMPLATE_PATH", None)
        if old_work:
            os.environ["WORK_DIR"] = old_work
        if old_template:
            os.environ["TEMPLATE_PATH"] = old_template
    return work_dir / "output.docx"


# Payload: template has [[VULNERABILITY_BLOCKS]]; part2.vulnerabilities has 1 vuln; vulnerability_blocks has distinct marker.
TRIPWIRE_PAYLOAD = {
    "assessment": {
        "asset": {"psa_cell": "555-000-0000"},
        "categories": {
            "ELECTRIC_POWER": {
                "requires_service": True,
                "time_to_impact_hours": 24,
                "loss_fraction_no_backup": 0.5,
                "recovery_time_hours": 12,
            },
            "COMMUNICATIONS": {"requires_service": False},
            "INFORMATION_TECHNOLOGY": {"requires_service": False},
            "WATER": {"requires_service": False},
            "WASTEWATER": {"requires_service": False},
            "CRITICAL_PRODUCTS": {},
        },
    },
    "report_vm": {
        "part2": {
            "vulnerabilities": [
                {
                    "severity": "HIGH",
                    "title": "Structured Preferred Title",
                    "narrative": "This narrative is from structured part2.vulnerabilities.",
                    "ofcs": ["Structured OFC one."],
                    "references": [],
                },
            ],
            "internet_transport_rows": [],
            "critical_hosted_services_rows": [],
            "dependency_summary_rows": [],
        },
    },
    "vofc_collection": {"items": []},
    "vulnerability_blocks": LEGACY_MARKER,
}


@pytest.fixture
def work_dir(tmp_path):
    return tmp_path


def test_structured_vulns_preferred_legacy_marker_never_appears(work_dir):
    """
    When part2.vulnerabilities is non-empty, reporter must use structured path.
    DOCX must contain structured title/narrative/OFC and must NOT contain the legacy marker.
    """
    if not CANONICAL_TEMPLATE.exists():
        pytest.skip("No template found: ADA/report template.docx")
    out_path = _run_reporter(TRIPWIRE_PAYLOAD, work_dir, CANONICAL_TEMPLATE)
    assert out_path.is_file()
    from docx import Document
    doc = Document(str(out_path))
    full_text = _full_doc_text(doc)
    assert "Structured Preferred Title" in full_text, "Structured title must appear in DOCX"
    assert "This narrative is from structured part2.vulnerabilities." in full_text, "Structured narrative must appear"
    assert "Structured OFC one." in full_text, "Structured OFC must appear"
    assert LEGACY_MARKER not in full_text, (
        "Legacy vulnerability_blocks string must not be injected when structured vulns are present"
    )
